"""
Trademark Sphere — Bulk Report Orchestrator.

Goal: Bulk reports must look IDENTICAL to individual reports.

Strategy:
    1. Generate every per-mark PDF using the EXISTING `build_report_pdf()`
       (same renderer => guaranteed pixel-parity with single reports).
    2. Generate a one-page Executive Summary cover PDF using the same
       reportlab palette/styles already used in pdf_renderer.py.
    3. Merge cover + per-mark PDFs with pypdf into one downloadable file.

Also exports DOCX (python-docx) and XLSX (openpyxl) for the same dataset,
plus portfolio analytics (success probability, recommendation badge,
conflict severity, smart class, alt-name availability, device similarity,
monitoring suggestion).

Wire-in (see quickcompany_trademark_router.py for the endpoint).
"""
from __future__ import annotations

import asyncio
import io
import logging
import time
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as RLImage,
    Paragraph,
    PageBreak,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from backend.pdf_renderer import (
    BLUE,
    BORDER,
    CONTENT_W,
    L_MARGIN,
    LIGHT_BLUE,
    MUTED,
    NAVY,
    PAGE_H,
    PAGE_W,
    R_MARGIN,
    SUBTLE,
    VERDICT_PALETTE,
    WHITE,
    _S,
    _decode_logo,
    _make_page_cb,
    _section_rule,
    build_report_pdf,
)

log = logging.getLogger("trademark-bulk")

# ──────────────────────────────────────────────────────────────────────────────
# In-process result cache (short TTL) to avoid duplicate scrapes inside one
# bulk run and across rapid retries from the same user.
# ──────────────────────────────────────────────────────────────────────────────
_CACHE: Dict[str, Tuple[float, dict]] = {}
_CACHE_TTL_SECONDS = 15 * 60  # 15 min


def _cache_key(name: str, class_filter: Optional[int], device_only: bool) -> str:
    return f"{(name or '').strip().lower()}|{class_filter or ''}|{int(bool(device_only))}"


def cache_get(name: str, class_filter: Optional[int], device_only: bool) -> Optional[dict]:
    k = _cache_key(name, class_filter, device_only)
    entry = _CACHE.get(k)
    if not entry:
        return None
    ts, value = entry
    if (time.time() - ts) > _CACHE_TTL_SECONDS:
        _CACHE.pop(k, None)
        return None
    return value


def cache_put(name: str, class_filter: Optional[int], device_only: bool, value: dict) -> None:
    _CACHE[_cache_key(name, class_filter, device_only)] = (time.time(), value)


# ──────────────────────────────────────────────────────────────────────────────
# Branding validation
# ──────────────────────────────────────────────────────────────────────────────
def validate_branding(branding: Dict[str, Any]) -> List[str]:
    """Return a list of human-readable problems. Empty list == valid."""
    problems: List[str] = []
    if not branding:
        return ["No branding configuration supplied."]
    # We allow logo to be optional but warn elsewhere; tagline+footer should exist.
    if not (branding.get("tagline") or "").strip():
        problems.append("Header tagline is missing.")
    if not (branding.get("footer") or "").strip():
        problems.append("Footer text is missing.")
    return problems


# ──────────────────────────────────────────────────────────────────────────────
# Per-mark analytics enrichment
# ──────────────────────────────────────────────────────────────────────────────
_RECOMMENDATION_LEVELS = [
    # (risk_threshold, label, color_hex)
    (20,  "SAFE",      "#166534"),
    (45,  "CAUTION",   "#92400E"),
    (70,  "HIGH RISK", "#B45309"),
    (101, "AVOID",     "#7F1D1D"),
]


def _recommendation_badge(risk_score: int) -> Tuple[str, str]:
    for thr, label, hex_color in _RECOMMENDATION_LEVELS:
        if (risk_score or 0) < thr:
            return label, hex_color
    return "AVOID", "#7F1D1D"


def _success_probability(report: dict) -> int:
    """0-100% chance of successful registration. Inverse of risk with smoothing."""
    risk = int(report.get("risk_score") or 0)
    counts = report.get("summary_counts") or {}
    exact = int(counts.get("exact_matches") or 0)
    phonetic = int(counts.get("phonetic_matches") or 0)
    # Hard floor when exact matches exist.
    if exact > 0:
        return max(0, 15 - exact * 3)
    base = max(0, 100 - risk)
    # Phonetic matches lower confidence a little extra.
    base -= min(20, phonetic * 4)
    return max(0, min(100, int(round(base))))


def _conflict_severity(report: dict) -> str:
    counts = report.get("summary_counts") or {}
    if (counts.get("exact_matches") or 0) > 0:
        return "CRITICAL"
    if (counts.get("phonetic_matches") or 0) >= 3:
        return "HIGH"
    if (counts.get("phonetic_matches") or 0) >= 1 or (counts.get("contains_matches") or 0) >= 3:
        return "MEDIUM"
    if (counts.get("total_results") or 0) > 0:
        return "LOW"
    return "NONE"


def _smart_class_suggestion(report: dict) -> Optional[int]:
    """Pick the class with the FEWEST blocking marks as a safer filing class."""
    cb = report.get("class_breakdown") or []
    if not cb:
        return None
    scored = []
    for row in cb:
        cl = row.get("class_number") or row.get("class")
        if cl is None:
            continue
        try:
            cl_i = int(cl)
        except (TypeError, ValueError):
            continue
        blocking = int(row.get("blocking") or row.get("blocking_count") or 0)
        total = int(row.get("total") or 0)
        scored.append((blocking, total, cl_i))
    if not scored:
        return None
    scored.sort(key=lambda t: (t[0], t[1]))
    return scored[0][2]


def _device_similarity_score(report: dict) -> int:
    """0-100. Higher = stronger logo-mark similarity. Heuristic from match counts."""
    counts = report.get("summary_counts") or {}
    device_hits = 0
    for r in (report.get("all_results") or []):
        mt = (r.get("mark_type") or "").lower()
        if any(k in mt for k in ("device", "logo", "label", "composite")):
            device_hits += 1
    if device_hits == 0:
        return 0
    exact = int(counts.get("exact_matches") or 0)
    phonetic = int(counts.get("phonetic_matches") or 0)
    score = device_hits * 8 + exact * 25 + phonetic * 10
    return max(0, min(100, score))


def _alt_name_availability(alt: str, base_query: str) -> int:
    """Cheap heuristic — longer derivative names score higher availability."""
    if not alt:
        return 0
    length_bonus = min(40, len(alt) * 3)
    distinct = 40 if alt.lower() != base_query.lower() else 0
    suffix_bonus = 20 if any(alt.lower().endswith(s) for s in ("ly", "io", "hub", "labs", "co")) else 0
    return max(0, min(100, length_bonus + distinct + suffix_bonus))


def enrich_report_with_analytics(report: dict, *, enable_monitoring: bool = False) -> dict:
    """Adds analytics fields onto a report dict in-place (and returns it)."""
    if report.get("error"):
        return report
    badge_label, badge_color = _recommendation_badge(int(report.get("risk_score") or 0))
    analytics = {
        "success_probability_pct": _success_probability(report),
        "recommendation_badge":    badge_label,
        "recommendation_color":    badge_color,
        "smart_class_suggestion":  _smart_class_suggestion(report),
        "conflict_severity":       _conflict_severity(report),
        "device_similarity_score": _device_similarity_score(report),
        "monitoring_suggested":    bool(enable_monitoring or (int(report.get("risk_score") or 0) >= 40)),
    }
    # Score every alt-name suggestion the report engine produced.
    query = report.get("query") or ""
    alts = report.get("alternative_name_suggestions") or []
    scored_alts: List[Dict[str, Any]] = []
    for a in alts:
        name = a if isinstance(a, str) else (a.get("name") or a.get("suggestion") or "")
        if not name:
            continue
        scored_alts.append({"name": name, "availability_score": _alt_name_availability(name, query)})
    if scored_alts:
        analytics["alternative_names_scored"] = scored_alts

    report["analytics"] = analytics

    # Surface the analytics inside the recommendations list so that the
    # existing per-mark PDF renderer prints them without modification.
    recs = list(report.get("recommendations") or [])
    extras = [
        f"Registration success probability: {analytics['success_probability_pct']}%",
        f"Filing recommendation: {badge_label}",
        f"Conflict severity: {analytics['conflict_severity']}",
    ]
    if analytics["smart_class_suggestion"]:
        extras.append(f"Smart class suggestion: Class {analytics['smart_class_suggestion']}")
    if analytics["device_similarity_score"]:
        extras.append(f"Device-mark similarity score: {analytics['device_similarity_score']}/100")
    if analytics["monitoring_suggested"]:
        extras.append("Trademark monitoring is recommended for this mark.")
    # Avoid duplication on re-enrichment.
    seen = {r.strip().lower() for r in recs if isinstance(r, str)}
    for line in extras:
        if line.strip().lower() not in seen:
            recs.append(line)
    report["recommendations"] = recs
    return report


# ──────────────────────────────────────────────────────────────────────────────
# Portfolio analytics
# ──────────────────────────────────────────────────────────────────────────────
def compute_portfolio_analytics(items: List[dict]) -> Dict[str, Any]:
    good = [it for it in items if not it.get("error") and it.get("report")]
    total = len(items)
    if not good:
        return {
            "total_marks": total,
            "available": 0, "caution": 0, "conflict": 0,
            "average_risk": 0, "high_risk_marks": 0,
            "average_success_probability": 0,
            "filing_recommendations": [],
        }
    statuses = [it["report"].get("overall_status") for it in good]
    risks    = [int(it["report"].get("risk_score") or 0) for it in good]
    probs    = [int(it["report"].get("analytics", {}).get("success_probability_pct") or 0) for it in good]
    available = sum(1 for s in statuses if s == "AVAILABLE")
    caution   = sum(1 for s in statuses if s == "CAUTION")
    conflict  = sum(1 for s in statuses if s == "CONFLICT")
    high_risk = sum(1 for r in risks if r >= 60)

    filing_recs: List[Dict[str, Any]] = []
    for it in good:
        rep = it["report"]
        an = rep.get("analytics", {}) or {}
        filing_recs.append({
            "name": rep.get("query"),
            "verdict": rep.get("overall_status"),
            "risk": int(rep.get("risk_score") or 0),
            "success_probability_pct": an.get("success_probability_pct", 0),
            "recommendation": an.get("recommendation_badge", "CAUTION"),
            "smart_class": an.get("smart_class_suggestion"),
            "conflict_severity": an.get("conflict_severity", "NONE"),
        })

    return {
        "total_marks": total,
        "available": available,
        "caution":   caution,
        "conflict":  conflict,
        "average_risk": int(round(sum(risks) / max(1, len(risks)))),
        "high_risk_marks": high_risk,
        "average_success_probability": int(round(sum(probs) / max(1, len(probs)))),
        "filing_recommendations": filing_recs,
    }


# ──────────────────────────────────────────────────────────────────────────────
# Executive Summary cover PDF
# ──────────────────────────────────────────────────────────────────────────────
def _branding_header(story: list, branding: Dict[str, Any]) -> None:
    logo_url = branding.get("logo_data_url")
    tagline  = branding.get("tagline") or "Bulk Trademark Availability Report"
    logo_img = None
    if logo_url:
        s = _decode_logo(logo_url)
        if s:
            try:
                from PIL import Image as PILImage
                s.seek(0)
                pil = PILImage.open(s); nw, nh = pil.size; s.seek(0)
                scaled_h = min((nh / nw) * CONTENT_W if nw else 20 * mm, 36 * mm)
                logo_img = RLImage(s, width=CONTENT_W, height=scaled_h)
            except Exception:
                try:
                    s.seek(0)
                    logo_img = RLImage(s, width=CONTENT_W, height=30 * mm, kind="proportional")
                except Exception:
                    logo_img = None
    if logo_img:
        t = Table([[logo_img]], colWidths=[CONTENT_W])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0),(-1,-1), WHITE), ("BOX",(0,0),(-1,-1), 0.4, BORDER),
            ("TOPPADDING",(0,0),(-1,-1),8),("BOTTOMPADDING",(0,0),(-1,-1),8),
            ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),0),
            ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ]))
        story.append(t); story.append(Spacer(1, 6))
    else:
        hdr = Table([[Paragraph(
            (branding.get("company_name") or "TRADEMARK SPHERE").upper(),
            ParagraphStyle("hdr_b", fontName="Helvetica-Bold", fontSize=12, textColor=WHITE, leading=15),
        )]], colWidths=[CONTENT_W])
        hdr.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1), NAVY),
            ("TOPPADDING",(0,0),(-1,-1),10),("BOTTOMPADDING",(0,0),(-1,-1),10),
            ("LEFTPADDING",(0,0),(-1,-1),14),
        ]))
        story.append(hdr); story.append(Spacer(1, 6))
    if tagline:
        t = Table([[Paragraph(tagline, ParagraphStyle(
            "tg_b", fontName="Helvetica", fontSize=8.5, textColor=BLUE, leading=12))]], colWidths=[CONTENT_W])
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1), LIGHT_BLUE),
            ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
            ("LEFTPADDING",(0,0),(-1,-1),12),("BOX",(0,0),(-1,-1), 0.4, BORDER),
        ]))
        story.append(t); story.append(Spacer(1, 8))


def build_executive_summary_pdf(items: List[dict], branding: Dict[str, Any], analytics: Dict[str, Any]) -> bytes:
    """Builds the Page-1 Executive Summary as its own PDF (will be merged on top)."""
    st = _S()
    buf = BytesIO()

    footer_text = (branding.get("footer") or "").strip() or "Trademark Sphere · Bureau of Trademark Intelligence"
    watermark   = (branding.get("watermark") or "").strip()
    if watermark.upper() == "CUSTOM":
        watermark = (branding.get("custom_watermark") or "").strip()
    created = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M") + " UTC"

    pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=L_MARGIN, rightMargin=R_MARGIN,
        topMargin=14 * mm, bottomMargin=20 * mm,
        title="Bulk Trademark Availability Report",
    )
    story: list = []
    _branding_header(story, branding)

    story.append(Paragraph("Bulk Trademark Availability Report", st["h1"]))
    story.append(Paragraph(
        f"{analytics['total_marks']} marks analysed · Generated {created}",
        st["subtitle"],
    ))
    story.append(_section_rule())

    # Client / prepared-by
    client_name   = (branding.get("client_name")   or "").strip()
    client_mobile = (branding.get("client_mobile") or "").strip()
    report_date   = (branding.get("report_date")   or "").strip()
    prepared_by   = (branding.get("prepared_by")   or branding.get("footer") or "").strip()
    rows = []
    if client_name:   rows.append([Paragraph("CLIENT", st["small_bold"]),    Paragraph(client_name, st["body_bold"])])
    if client_mobile: rows.append([Paragraph("MOBILE", st["small_bold"]),    Paragraph(client_mobile, st["body"])])
    if report_date:   rows.append([Paragraph("REPORT DATE", st["small_bold"]),Paragraph(report_date, st["body"])])
    if prepared_by:   rows.append([Paragraph("PREPARED BY", st["small_bold"]),Paragraph(prepared_by, st["body"])])
    if rows:
        ci = Table(rows, colWidths=[32 * mm, CONTENT_W - 32 * mm])
        ci.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1), SUBTLE),
            ("BACKGROUND",(0,0),(0,-1),  LIGHT_BLUE),
            ("BOX",(0,0),(-1,-1), 0.5, BORDER),
            ("INNERGRID",(0,0),(-1,-1), 0.3, BORDER),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
            ("LEFTPADDING",(0,0),(-1,-1),10),("RIGHTPADDING",(0,0),(-1,-1),10),
        ]))
        story.append(ci); story.append(Spacer(1, 10))

    # KPI grid: total / available / caution / conflict / avg risk / high-risk / success%
    def _kpi(label, value, color):
        return [
            Paragraph(label, st["small_bold"]),
            Paragraph(
                f'<font color="{color}"><b>{value}</b></font>',
                ParagraphStyle("kv", fontName="Helvetica-Bold", fontSize=22, leading=26, alignment=TA_CENTER),
            ),
        ]

    cells = [
        _kpi("TOTAL MARKS",   analytics["total_marks"],   NAVY.hexval()),
        _kpi("AVAILABLE",     analytics["available"],     colors.HexColor("#166534").hexval()),
        _kpi("CAUTION",       analytics["caution"],       colors.HexColor("#92400E").hexval()),
        _kpi("CONFLICT",      analytics["conflict"],      colors.HexColor("#7F1D1D").hexval()),
        _kpi("AVG RISK",      f'{analytics["average_risk"]}/100', BLUE.hexval()),
        _kpi("HIGH RISK",     analytics["high_risk_marks"],colors.HexColor("#B45309").hexval()),
        _kpi("AVG SUCCESS %", f'{analytics["average_success_probability"]}%', colors.HexColor("#1F6FB2").hexval()),
    ]

    # Lay them out as a 4-column grid (row 1: 4 cells, row 2: 3 cells + blank)
    cw4 = CONTENT_W / 4
    header_row = [c[0] for c in cells[:4]]
    value_row  = [c[1] for c in cells[:4]]
    grid = Table([header_row, value_row], colWidths=[cw4] * 4)
    grid.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1), 0.5, BORDER),
        ("INNERGRID",(0,0),(-1,-1), 0.4, BORDER),
        ("BACKGROUND",(0,0),(-1,0), SUBTLE),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),
    ]))
    story.append(grid); story.append(Spacer(1, 8))
    header_row2 = [c[0] for c in cells[4:]] + [Paragraph("&nbsp;", st["small_bold"])]
    value_row2  = [c[1] for c in cells[4:]] + [Paragraph("&nbsp;", st["body"])]
    grid2 = Table([header_row2, value_row2], colWidths=[cw4] * 4)
    grid2.setStyle(TableStyle([
        ("BOX",(0,0),(2,-1), 0.5, BORDER),
        ("INNERGRID",(0,0),(2,-1), 0.4, BORDER),
        ("BACKGROUND",(0,0),(2,0), SUBTLE),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),
    ]))
    story.append(grid2); story.append(Spacer(1, 14))

    # ── PLAIN-LANGUAGE PORTFOLIO SUMMARY (Layman Guide) ──────────────────────
    story.append(Paragraph("WHAT THIS REPORT MEANS — PLAIN LANGUAGE GUIDE", st["eyebrow"]))
    story.append(_section_rule())

    total_m   = analytics["total_marks"]
    avail_m   = analytics["available"]
    caution_m = analytics["caution"]
    conflict_m= analytics["conflict"]
    avg_suc   = analytics["average_success_probability"]

    if conflict_m == 0 and caution_m == 0:
        guide_bg     = colors.HexColor("#DCFCE7")
        guide_border = colors.HexColor("#86EFAC")
        guide_color  = colors.HexColor("#166534")
        guide_icon   = "✔"
        guide_head   = "All Marks Look Clear — Good Position to File"
        guide_body   = (
            f"All <b>{total_m}</b> brand name(s) searched appear to be available with no major "
            "conflicts found. This is a strong position to begin the trademark registration process. "
            "Review individual mark reports below for specific details before filing."
        )
    elif conflict_m > 0 and conflict_m >= total_m // 2:
        guide_bg     = colors.HexColor("#FEE2E2")
        guide_border = colors.HexColor("#FCA5A5")
        guide_color  = colors.HexColor("#7F1D1D")
        guide_icon   = "✘"
        guide_head   = "Most Marks Have Conflicts — Review Before Filing Anything"
        guide_body   = (
            f"Out of <b>{total_m}</b> marks searched, <b>{conflict_m}</b> have significant conflicts "
            "with existing registered trademarks. Filing these without legal advice is risky — "
            "applications are likely to be rejected or challenged. Consider modifying brand names "
            "or choosing alternative names suggested in each individual report."
        )
    else:
        guide_bg     = colors.HexColor("#FEF3C7")
        guide_border = colors.HexColor("#FCD34D")
        guide_color  = colors.HexColor("#92400E")
        guide_icon   = "◐"
        guide_head   = "Mixed Results — Some Marks Ready, Others Need Review"
        guide_body   = (
            f"Out of <b>{total_m}</b> marks: <b>{avail_m}</b> appear available to file, "
            f"<b>{caution_m}</b> need caution and deeper review, and <b>{conflict_m}</b> have "
            "strong conflicts. The average chance of successful registration across this portfolio "
            f"is <b>{avg_suc}%</b>. Focus on marks marked 'Available' first, and consult a "
            "trademark attorney for those marked 'Caution' or 'Conflict'."
        )

    guide_action_parts = []
    if avail_m:
        guide_action_parts.append(f"✅ <b>{avail_m} mark(s)</b> — Safe to proceed with filing.")
    if caution_m:
        guide_action_parts.append(f"🔍 <b>{caution_m} mark(s)</b> — Get legal opinion before filing.")
    if conflict_m:
        guide_action_parts.append(f"⚠ <b>{conflict_m} mark(s)</b> — Avoid filing without significant modification or attorney review.")
    guide_action_parts.append("📋 Always verify final results on the official IP India database (ipindia.gov.in) before filing.")
    guide_actions_text = "<br/>".join(guide_action_parts)

    guide_tbl = Table([[
        Paragraph(
            f'<font color="{guide_color.hexval()}"><b>{guide_icon}  {guide_head}</b></font><br/><br/>'
            f'{guide_body}<br/><br/>'
            f'{guide_actions_text}',
            ParagraphStyle("guide_p", fontName="Helvetica", fontSize=9, textColor=colors.HexColor("#1A1A2E"),
                           leading=14),
        )
    ]], colWidths=[CONTENT_W])
    guide_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), guide_bg),
        ("BOX",           (0, 0), (-1, -1), 1.0, guide_border),
        ("TOPPADDING",    (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ("LEFTPADDING",   (0, 0), (-1, -1), 14),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 14),
    ]))
    story.append(guide_tbl)
    story.append(Spacer(1, 14))

    # Filing recommendations table
    story.append(Paragraph("FILING RECOMMENDATIONS", st["eyebrow"]))
    story.append(_section_rule())
    rec_rows = [[
        Paragraph("MARK",        st["tbl_hdr"]),
        Paragraph("VERDICT",     st["tbl_hdr"]),
        Paragraph("RISK",        st["tbl_hdr"]),
        Paragraph("SUCCESS %",   st["tbl_hdr"]),
        Paragraph("BADGE",       st["tbl_hdr"]),
        Paragraph("SMART CLASS", st["tbl_hdr"]),
        Paragraph("SEVERITY",    st["tbl_hdr"]),
    ]]
    for r in analytics.get("filing_recommendations", []):
        verdict = r.get("verdict") or "UNKNOWN"
        vp = VERDICT_PALETTE.get(verdict, VERDICT_PALETTE["CAUTION"])
        badge_label, badge_color = _recommendation_badge(int(r.get("risk") or 0))
        rec_rows.append([
            Paragraph(f'<b>{r.get("name","")}</b>',
                      ParagraphStyle("nm", fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=NAVY)),
            Paragraph(f'<font color="{vp["fg"].hexval()}"><b>{verdict}</b></font>',
                      ParagraphStyle("vd", fontName="Helvetica-Bold", fontSize=8.5, leading=11)),
            Paragraph(str(r.get("risk", "—")),
                      ParagraphStyle("rk", fontName="Helvetica-Bold", fontSize=9, leading=11, alignment=TA_CENTER)),
            Paragraph(f'{r.get("success_probability_pct",0)}%',
                      ParagraphStyle("sp", fontName="Helvetica-Bold", fontSize=9, leading=11, alignment=TA_CENTER, textColor=BLUE)),
            Paragraph(f'<font color="{badge_color}"><b>{badge_label}</b></font>',
                      ParagraphStyle("bd", fontName="Helvetica-Bold", fontSize=8.5, leading=11, alignment=TA_CENTER)),
            Paragraph(f'CL{int(r["smart_class"]):02d}' if r.get("smart_class") else "—",
                      ParagraphStyle("sc", fontName="Helvetica", fontSize=8.5, leading=11, alignment=TA_CENTER, textColor=MUTED)),
            Paragraph(r.get("conflict_severity") or "NONE",
                      ParagraphStyle("cs", fontName="Helvetica-Bold", fontSize=8.5, leading=11, alignment=TA_CENTER)),
        ])
    cws = [42*mm, 24*mm, 14*mm, 18*mm, 22*mm, 18*mm, CONTENT_W - 158*mm]
    rec_tbl = Table(rec_rows, colWidths=cws, repeatRows=1)
    rec_tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0), NAVY),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [WHITE, SUBTLE]),
        ("BOX",(0,0),(-1,-1), 0.5, BORDER),
        ("LINEBELOW",(0,0),(-1,0), 0.8, NAVY),
        ("INNERGRID",(0,1),(-1,-1), 0.3, BORDER),
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("TOPPADDING",(0,0),(-1,-1), 5),("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING",(0,0),(-1,-1), 5),("RIGHTPADDING",(0,0),(-1,-1), 5),
    ]))
    story.append(rec_tbl); story.append(Spacer(1, 10))

    disclaimer = (branding.get("disclaimer") or "").strip()
    if disclaimer:
        story.append(Paragraph("DISCLAIMER", st["eyebrow"]))
        story.append(_section_rule())
        story.append(Paragraph(disclaimer, st["footer"]))

    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "Each subject mark below is followed by its complete individual dossier "
        "(verdict, match counts, recommendations, alternative names, class-wise "
        "breakdown, and Exhibit A — all recorded matches).",
        st["footer"],
    ))

    footer_l = footer_text
    footer_r = f"Bulk report · {analytics['total_marks']} marks · {created}"
    page_cb = _make_page_cb(footer_l, footer_r, watermark)
    pdf.build(story, onFirstPage=page_cb, onLaterPages=page_cb)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# Combined PDF: cover + per-mark dossier (identical to single-report renderer)
# ──────────────────────────────────────────────────────────────────────────────
def _per_mark_doc_record(item: dict, branding: Dict[str, Any]) -> dict:
    """Construct a doc_record dict compatible with build_report_pdf()."""
    report = dict(item.get("report") or {})
    # Inject branding so build_report_pdf renders identical header/footer/watermark.
    report.setdefault("logo_data_url",    branding.get("logo_data_url"))
    report.setdefault("footer_text",      branding.get("footer", ""))
    report.setdefault("tagline",          branding.get("tagline", "Trademark Availability Report"))
    report.setdefault("watermark",        branding.get("watermark", ""))
    report.setdefault("custom_watermark", branding.get("custom_watermark", ""))
    report.setdefault("client_name",      branding.get("client_name", ""))
    report.setdefault("client_mobile",    branding.get("client_mobile", ""))
    report.setdefault("report_date",      branding.get("report_date", ""))
    return {
        "id": item.get("id"),
        "created_at": item.get("created_at") or datetime.now(timezone.utc).isoformat(),
        "report": report,
    }


def build_bulk_dossier_pdf(items: List[dict], branding: Dict[str, Any], analytics: Dict[str, Any]) -> bytes:
    """Cover page + per-mark dossier rendered via the EXISTING single-report renderer."""
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:  # pragma: no cover
        from PyPDF2 import PdfReader, PdfWriter  # type: ignore

    writer = PdfWriter()

    # 1) Executive summary cover.
    cover_bytes = build_executive_summary_pdf(items, branding, analytics)
    for page in PdfReader(BytesIO(cover_bytes)).pages:
        writer.add_page(page)

    # 2) Full individual dossier per mark.
    for item in items:
        if item.get("error") or not item.get("report"):
            continue
        try:
            pdf_bytes = build_report_pdf(_per_mark_doc_record(item, branding))
            for page in PdfReader(BytesIO(pdf_bytes)).pages:
                writer.add_page(page)
        except Exception:
            log.exception("Failed to render per-mark PDF for %s", item.get("name"))

    out = BytesIO()
    writer.write(out)
    return out.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# DOCX export (python-docx)
# ──────────────────────────────────────────────────────────────────────────────
def build_bulk_docx(items: List[dict], branding: Dict[str, Any], analytics: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10)

    # Optional logo
    logo_url = branding.get("logo_data_url")
    if logo_url:
        s = _decode_logo(logo_url)
        if s:
            try:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                s.seek(0); p.add_run().add_picture(s, width=Inches(2.2))
            except Exception:
                pass

    title = doc.add_heading(branding.get("tagline") or "Bulk Trademark Availability Report", 0)
    for r in title.runs: r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    p = doc.add_paragraph()
    p.add_run(f'Generated: {datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")}').italic = True

    # Executive summary table
    doc.add_heading("Executive Summary", 1)
    table = doc.add_table(rows=2, cols=7)
    table.style = "Light Grid Accent 1"
    headers = ["Total", "Available", "Caution", "Conflict", "Avg Risk", "High Risk", "Avg Success %"]
    values  = [
        analytics["total_marks"], analytics["available"], analytics["caution"],
        analytics["conflict"], analytics["average_risk"], analytics["high_risk_marks"],
        f'{analytics["average_success_probability"]}%',
    ]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs: r.bold = True
    for i, v in enumerate(values):
        table.rows[1].cells[i].text = str(v)

    # Filing recommendations
    doc.add_heading("Filing Recommendations", 1)
    ft = doc.add_table(rows=1, cols=7); ft.style = "Light Grid Accent 1"
    hdr = ["Mark", "Verdict", "Risk", "Success %", "Badge", "Smart Class", "Severity"]
    for i, h in enumerate(hdr):
        ft.rows[0].cells[i].text = h
        for r in ft.rows[0].cells[i].paragraphs[0].runs: r.bold = True
    for r in analytics.get("filing_recommendations", []):
        row = ft.add_row().cells
        row[0].text = str(r.get("name") or "")
        row[1].text = str(r.get("verdict") or "")
        row[2].text = str(r.get("risk") or 0)
        row[3].text = f'{r.get("success_probability_pct", 0)}%'
        row[4].text = str(r.get("recommendation") or "")
        row[5].text = f'CL{int(r["smart_class"]):02d}' if r.get("smart_class") else "—"
        row[6].text = str(r.get("conflict_severity") or "")

    # Per-mark dossier
    for item in items:
        if item.get("error") or not item.get("report"):
            continue
        rep = item["report"]
        an = rep.get("analytics") or {}
        doc.add_page_break()
        h = doc.add_heading(rep.get("query") or item.get("name") or "Mark", 1)
        for r in h.runs: r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        doc.add_paragraph(
            f'Verdict: {rep.get("overall_status","")} · Risk: {rep.get("risk_score",0)}/100 · '
            f'Success probability: {an.get("success_probability_pct",0)}% · '
            f'Badge: {an.get("recommendation_badge","")} · '
            f'Severity: {an.get("conflict_severity","")}'
        )
        if rep.get("headline"):
            doc.add_paragraph(rep["headline"])

        counts = rep.get("summary_counts") or {}
        doc.add_heading("Match Counts", 2)
        mt = doc.add_table(rows=1, cols=4); mt.style = "Light Grid Accent 1"
        for i, h in enumerate(["Exact", "Phonetic", "Contains", "Total"]):
            mt.rows[0].cells[i].text = h
            for r in mt.rows[0].cells[i].paragraphs[0].runs: r.bold = True
        mrow = mt.add_row().cells
        mrow[0].text = str(counts.get("exact_matches", 0))
        mrow[1].text = str(counts.get("phonetic_matches", 0))
        mrow[2].text = str(counts.get("contains_matches", 0))
        mrow[3].text = str(counts.get("total_results", 0))

        cb = rep.get("class_breakdown") or []
        if cb:
            doc.add_heading("Class Breakdown", 2)
            ct = doc.add_table(rows=1, cols=4); ct.style = "Light Grid Accent 1"
            for i, h in enumerate(["Class", "Total", "Blocking", "Dead"]):
                ct.rows[0].cells[i].text = h
                for r in ct.rows[0].cells[i].paragraphs[0].runs: r.bold = True
            for row in cb:
                rr = ct.add_row().cells
                cl = row.get("class_number") or row.get("class") or ""
                rr[0].text = f"Class {cl}"
                rr[1].text = str(row.get("total", "—"))
                rr[2].text = str(row.get("blocking", row.get("blocking_count", "—")))
                rr[3].text = str(row.get("dead", row.get("dead_count", "—")))

        recs = rep.get("recommendations") or []
        if recs:
            doc.add_heading("Recommendations", 2)
            for r in recs:
                doc.add_paragraph(str(r), style="List Bullet")

        alts = rep.get("analytics", {}).get("alternative_names_scored") or [
            {"name": a if isinstance(a, str) else a.get("name", ""), "availability_score": None}
            for a in (rep.get("alternative_name_suggestions") or [])
        ]
        if alts:
            doc.add_heading("Alternative Names", 2)
            at = doc.add_table(rows=1, cols=2); at.style = "Light Grid Accent 1"
            for i, hh in enumerate(["Name", "Availability Score"]):
                at.rows[0].cells[i].text = hh
                for r in at.rows[0].cells[i].paragraphs[0].runs: r.bold = True
            for a in alts:
                rr = at.add_row().cells
                rr[0].text = str(a.get("name") or "")
                rr[1].text = "" if a.get("availability_score") is None else f'{a["availability_score"]}/100'

        all_results = rep.get("all_results") or []
        if all_results:
            doc.add_heading(f'Exhibit A — All Recorded Matches ({len(all_results)})', 2)
            et = doc.add_table(rows=1, cols=7); et.style = "Light Grid Accent 1"
            for i, h in enumerate(["App. No.", "Mark", "Applicant", "Status", "Class", "Match", "Risk"]):
                et.rows[0].cells[i].text = h
                for r in et.rows[0].cells[i].paragraphs[0].runs: r.bold = True
            for row in all_results[:200]:
                rr = et.add_row().cells
                rr[0].text = str(row.get("application_no") or row.get("application_number") or "")
                rr[1].text = str(row.get("name") or "")
                rr[2].text = str(row.get("applicant") or "")
                rr[3].text = str(row.get("status") or "")
                rr[4].text = str(row.get("class") or row.get("class_number") or "")
                rr[5].text = str((row.get("match_type") or "").upper())
                rr[6].text = str(row.get("risk") or row.get("risk_score") or "—")

    # Footer paragraph
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run((branding.get("footer") or "Trademark Sphere — Bureau of Trademark Intelligence"))
    r.italic = True

    out = BytesIO(); doc.save(out); return out.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# XLSX export (openpyxl)
# ──────────────────────────────────────────────────────────────────────────────
def build_bulk_xlsx(items: List[dict], branding: Dict[str, Any], analytics: Dict[str, Any]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    bold = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1B2A4A")

    # Sheet 1: Executive Summary
    s = wb.active; s.title = "Executive Summary"
    s["A1"] = branding.get("tagline") or "Bulk Trademark Availability Report"
    s["A1"].font = Font(bold=True, size=16, color="1B2A4A")
    s["A2"] = f'Generated: {datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")}'
    s["A2"].font = Font(italic=True)

    kpi_headers = ["Total", "Available", "Caution", "Conflict", "Avg Risk", "High Risk", "Avg Success %"]
    kpi_values  = [
        analytics["total_marks"], analytics["available"], analytics["caution"],
        analytics["conflict"], analytics["average_risk"], analytics["high_risk_marks"],
        f'{analytics["average_success_probability"]}%',
    ]
    for i, h in enumerate(kpi_headers, start=1):
        c = s.cell(row=4, column=i, value=h); c.fill = header_fill; c.font = bold
        c.alignment = Alignment(horizontal="center")
    for i, v in enumerate(kpi_values, start=1):
        s.cell(row=5, column=i, value=v).alignment = Alignment(horizontal="center")

    # Sheet 2: Filing Recommendations
    rec = wb.create_sheet("Filing Recommendations")
    headers = ["Mark", "Verdict", "Risk", "Success %", "Badge", "Smart Class", "Severity"]
    for i, h in enumerate(headers, start=1):
        c = rec.cell(row=1, column=i, value=h); c.fill = header_fill; c.font = bold
    for r_i, r in enumerate(analytics.get("filing_recommendations", []), start=2):
        rec.cell(row=r_i, column=1, value=r.get("name"))
        rec.cell(row=r_i, column=2, value=r.get("verdict"))
        rec.cell(row=r_i, column=3, value=r.get("risk"))
        rec.cell(row=r_i, column=4, value=f'{r.get("success_probability_pct",0)}%')
        rec.cell(row=r_i, column=5, value=r.get("recommendation"))
        rec.cell(row=r_i, column=6, value=f'CL{int(r["smart_class"]):02d}' if r.get("smart_class") else "—")
        rec.cell(row=r_i, column=7, value=r.get("conflict_severity"))
    for col in range(1, 8):
        rec.column_dimensions[chr(64 + col)].width = 18

    # Sheet 3: All Matches (flattened)
    am = wb.create_sheet("All Matches")
    am_headers = ["Subject Mark", "App. No.", "Conflicting Mark", "Applicant", "Status",
                  "Class", "Match Type", "Risk"]
    for i, h in enumerate(am_headers, start=1):
        c = am.cell(row=1, column=i, value=h); c.fill = header_fill; c.font = bold
    row_i = 2
    for item in items:
        if item.get("error") or not item.get("report"):
            continue
        rep = item["report"]
        subj = rep.get("query") or item.get("name") or ""
        for row in (rep.get("all_results") or []):
            am.cell(row=row_i, column=1, value=subj)
            am.cell(row=row_i, column=2, value=row.get("application_no") or row.get("application_number"))
            am.cell(row=row_i, column=3, value=row.get("name"))
            am.cell(row=row_i, column=4, value=row.get("applicant"))
            am.cell(row=row_i, column=5, value=row.get("status"))
            am.cell(row=row_i, column=6, value=row.get("class") or row.get("class_number"))
            am.cell(row=row_i, column=7, value=(row.get("match_type") or "").upper())
            am.cell(row=row_i, column=8, value=row.get("risk") or row.get("risk_score"))
            row_i += 1
    for col_letter, w in zip("ABCDEFGH", [22, 16, 26, 28, 16, 10, 14, 8]):
        am.column_dimensions[col_letter].width = w

    # Sheet 4: Per-mark Dossier summary
    pm = wb.create_sheet("Per-Mark Dossier")
    pm_headers = ["Mark", "Verdict", "Risk", "Success %", "Badge", "Smart Class",
                  "Severity", "Exact", "Phonetic", "Contains", "Total", "Headline"]
    for i, h in enumerate(pm_headers, start=1):
        c = pm.cell(row=1, column=i, value=h); c.fill = header_fill; c.font = bold
    for r_i, item in enumerate([it for it in items if not it.get("error")], start=2):
        rep = item["report"]; an = rep.get("analytics") or {}; counts = rep.get("summary_counts") or {}
        pm.cell(row=r_i, column=1,  value=rep.get("query"))
        pm.cell(row=r_i, column=2,  value=rep.get("overall_status"))
        pm.cell(row=r_i, column=3,  value=rep.get("risk_score"))
        pm.cell(row=r_i, column=4,  value=f'{an.get("success_probability_pct",0)}%')
        pm.cell(row=r_i, column=5,  value=an.get("recommendation_badge"))
        pm.cell(row=r_i, column=6,  value=f'CL{int(an["smart_class_suggestion"]):02d}' if an.get("smart_class_suggestion") else "—")
        pm.cell(row=r_i, column=7,  value=an.get("conflict_severity"))
        pm.cell(row=r_i, column=8,  value=counts.get("exact_matches", 0))
        pm.cell(row=r_i, column=9,  value=counts.get("phonetic_matches", 0))
        pm.cell(row=r_i, column=10, value=counts.get("contains_matches", 0))
        pm.cell(row=r_i, column=11, value=counts.get("total_results", 0))
        pm.cell(row=r_i, column=12, value=rep.get("headline"))
    for col_letter, w in zip("ABCDEFGHIJKL", [22, 12, 8, 12, 12, 12, 12, 8, 10, 10, 8, 60]):
        pm.column_dimensions[col_letter].width = w

    out = BytesIO(); wb.save(out); return out.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# Parallel scrape orchestration (used by the router)
# ──────────────────────────────────────────────────────────────────────────────
async def run_bulk_searches(
    names: List[str],
    *,
    class_filter: Optional[int],
    device_only: bool,
    scrape_fn,
    max_parallel: int = 5,
    enable_monitoring: bool = False,
) -> List[dict]:
    """
    Runs `scrape_fn(name, class_filter, device_only=...)` for every name in
    parallel (bounded). Returns a list of items where each item has:
        { name, report?, error?, overall_status?, risk_score?, ... }

    `scrape_fn` is injected to avoid a circular import with the router/_scraper.
    """
    sem = asyncio.Semaphore(max(1, max_parallel))

    async def _one(name: str) -> dict:
        cached = cache_get(name, class_filter, device_only)
        if cached is not None:
            report = dict(cached)
        else:
            async with sem:
                try:
                    report = await scrape_fn(name, class_filter, device_only=device_only)
                except Exception as e:  # pragma: no cover
                    log.exception("bulk scrape failed for %s", name)
                    return {"name": name, "error": str(e)}
                cache_put(name, class_filter, device_only, report)

        enrich_report_with_analytics(report, enable_monitoring=enable_monitoring)
        return {
            "name": name,
            "overall_status": report.get("overall_status"),
            "risk_score":     report.get("risk_score"),
            "headline":       report.get("headline"),
            "report":         report,
            "error":          None,
        }

    return await asyncio.gather(*[_one(n) for n in names])
