"""
MIS Report exports
==================
Turns the computed MIS numbers into files a client can actually be sent:

    build_pdf_report(meta, sections)   -> PDF  (reportlab)
    build_word_report(meta, sections)  -> DOCX (python-docx)
    build_excel_workbook(meta, sheets) -> XLSX (openpyxl)

`sections` is a plain list of dicts so the report layout lives in
mis_report.py and this module stays a dumb renderer:

    {
      "title": "Financial Dashboard",
      "kpis":  [{"label": "Total Revenue", "value": 1234.0, "money": True}, ...],
      "table": {"columns": ["Party", "Outstanding"], "rows": [["ABC", 100.0]]},
      "note":  "optional paragraph"
    }
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

BRAND = "#1E4B8F"
BRAND_LIGHT = "#E8EEF7"
GREY = "#6B7280"

MAX_TABLE_ROWS = 500          # keeps a printable document printable


def fmt_money(v: Any) -> str:
    try:
        n = float(v)
    except (TypeError, ValueError):
        return str(v or "")
    neg = n < 0
    n = abs(n)
    # Indian grouping: 12,34,567.89
    whole, dec = f"{n:.2f}".split(".")
    if len(whole) > 3:
        head, tail = whole[:-3], whole[-3:]
        parts = []
        while len(head) > 2:
            parts.insert(0, head[-2:])
            head = head[:-2]
        if head:
            parts.insert(0, head)
        whole = ",".join(parts + [tail])
    s = f"₹{whole}.{dec}"
    return f"({s})" if neg else s


def fmt_value(item: Dict[str, Any]) -> str:
    v = item.get("value")
    if v is None or v == "":
        return "—"
    if item.get("money"):
        return fmt_money(v)
    if item.get("pct"):
        try:
            return f"{float(v):.2f}%"
        except (TypeError, ValueError):
            return str(v)
    return str(v)


def _cell(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float):
        return fmt_money(v)
    return str(v)


# ══════════════════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════════════════

def build_pdf_report(meta: Dict[str, Any], sections: List[Dict[str, Any]]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, PageBreak)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=15 * mm, rightMargin=15 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
        title=f"MIS Report — {meta.get('client_name','')}",
        author=meta.get("firm_name") or "Task-O-Sphere",
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Title"], fontSize=20, textColor=colors.HexColor(BRAND))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13,
                        textColor=colors.HexColor(BRAND), spaceBefore=10, spaceAfter=6)
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=8.5, textColor=colors.HexColor(GREY))
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=9)

    story: List[Any] = []
    story.append(Paragraph("Management Information System Report", h1))
    story.append(Spacer(1, 4))
    head_rows = [
        ["Client", meta.get("client_name") or "—", "Period", meta.get("period") or "—"],
        ["Generated", meta.get("generated_at") or datetime.now().strftime("%d-%b-%Y %H:%M"),
         "Prepared by", meta.get("prepared_by") or "—"],
    ]
    t = Table(head_rows, colWidths=[25 * mm, 65 * mm, 30 * mm, 60 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(BRAND_LIGHT)),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor(BRAND_LIGHT)),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#111827")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))

    for idx, sec in enumerate(sections):
        if idx:
            story.append(Spacer(1, 8))
        story.append(Paragraph(sec.get("title", ""), h2))
        if sec.get("note"):
            story.append(Paragraph(sec["note"], small))
            story.append(Spacer(1, 4))

        kpis = sec.get("kpis") or []
        if kpis:
            per_row = 3
            data = []
            for i in range(0, len(kpis), per_row):
                chunk = kpis[i:i + per_row]
                data.append([Paragraph(f"<b>{k['label']}</b><br/><font size=11>{fmt_value(k)}</font>", body)
                             for k in chunk] + [""] * (per_row - len(chunk)))
            kt = Table(data, colWidths=[60 * mm] * per_row)
            kt.setStyle(TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E5E7EB")),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(kt)
            story.append(Spacer(1, 6))

        tbl = sec.get("table")
        if tbl and tbl.get("rows"):
            cols = tbl["columns"]
            rows = tbl["rows"][:MAX_TABLE_ROWS]
            width = (A4[0] - 30 * mm)
            col_w = [width / len(cols)] * len(cols)
            data = [[Paragraph(f"<b>{c}</b>", body) for c in cols]]
            for r in rows:
                data.append([Paragraph(_cell(c), body) for c in r])
            dt = Table(data, colWidths=col_w, repeatRows=1)
            dt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CBD5E1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(dt)
            if len(tbl["rows"]) > MAX_TABLE_ROWS:
                story.append(Spacer(1, 3))
                story.append(Paragraph(
                    f"Showing first {MAX_TABLE_ROWS} of {len(tbl['rows'])} rows — "
                    "download the Excel export for the complete data.", small))

    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "This MIS report is generated from the source documents uploaded for the selected period "
        "and the manual entries recorded against it.", small))

    def _footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(GREY))
        canvas.drawString(15 * mm, 10 * mm, f"{meta.get('client_name','')} • {meta.get('period','')}")
        canvas.drawRightString(A4[0] - 15 * mm, 10 * mm, f"Page {doc_.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════
# WORD
# ══════════════════════════════════════════════════════════════════════════

def build_word_report(meta: Dict[str, Any], sections: List[Dict[str, Any]]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Management Information System Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x4B, 0x8F)

    info = document.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    pairs = [
        ("Client", meta.get("client_name") or "—"),
        ("Period", meta.get("period") or "—"),
        ("Generated", meta.get("generated_at") or datetime.now().strftime("%d-%b-%Y %H:%M")),
        ("Prepared by", meta.get("prepared_by") or "—"),
    ]
    for i, (k, v) in enumerate(pairs):
        cell_k = info.cell(i // 2, (i % 2) * 2)
        cell_v = info.cell(i // 2, (i % 2) * 2 + 1)
        cell_k.text = ""
        cell_k.paragraphs[0].add_run(k).bold = True
        cell_v.text = str(v)

    for sec in sections:
        document.add_heading(sec.get("title", ""), level=1)
        if sec.get("note"):
            p = document.add_paragraph(sec["note"])
            p.runs[0].italic = True

        kpis = sec.get("kpis") or []
        if kpis:
            rows = (len(kpis) + 1) // 2
            kt = document.add_table(rows=rows, cols=4)
            kt.style = "Table Grid"
            for i, k in enumerate(kpis):
                kc = kt.cell(i // 2, (i % 2) * 2)
                vc = kt.cell(i // 2, (i % 2) * 2 + 1)
                kc.text = ""
                kc.paragraphs[0].add_run(k["label"]).bold = True
                vc.text = fmt_value(k)
            document.add_paragraph()

        tbl = sec.get("table")
        if tbl and tbl.get("rows"):
            cols = tbl["columns"]
            rows = tbl["rows"][:MAX_TABLE_ROWS]
            dt = document.add_table(rows=1, cols=len(cols))
            dt.style = "Light Grid Accent 1"
            for i, c in enumerate(cols):
                cell = dt.rows[0].cells[i]
                cell.text = ""
                cell.paragraphs[0].add_run(str(c)).bold = True
            for r in rows:
                cells = dt.add_row().cells
                for i, v in enumerate(r):
                    cells[i].text = _cell(v)
            if len(tbl["rows"]) > MAX_TABLE_ROWS:
                note = document.add_paragraph(
                    f"Showing first {MAX_TABLE_ROWS} of {len(tbl['rows'])} rows — "
                    "download the Excel export for the complete data.")
                note.runs[0].italic = True

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════
# EXCEL
# ══════════════════════════════════════════════════════════════════════════

def build_excel_workbook(meta: Dict[str, Any], sheets: List[Dict[str, Any]]) -> bytes:
    """`sheets` = [{"name", "columns", "rows", "money_columns": [idx,...]}]"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)

    header_fill = PatternFill("solid", start_color="1E4B8F")
    header_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    title_font = Font(bold=True, size=13, color="1E4B8F", name="Calibri")
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    used_names = set()
    for sheet in sheets:
        name = (sheet.get("name") or "Sheet")[:31]
        base, i = name, 2
        while name in used_names:
            name = f"{base[:28]}_{i}"
            i += 1
        used_names.add(name)
        ws = wb.create_sheet(name)

        ws["A1"] = f"{meta.get('client_name','')} — {sheet.get('name','')}"
        ws["A1"].font = title_font
        ws["A2"] = f"Period: {meta.get('period','')}   |   Generated: {meta.get('generated_at','')}"
        ws["A2"].font = Font(size=9, color="6B7280", name="Calibri")

        cols = sheet.get("columns") or []
        start = 4
        for c, label in enumerate(cols, start=1):
            cell = ws.cell(row=start, column=c, value=str(label))
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        money_cols = set(sheet.get("money_columns") or [])
        for r, row in enumerate(sheet.get("rows") or [], start=start + 1):
            for c, val in enumerate(row, start=1):
                cell = ws.cell(row=r, column=c, value=val)
                cell.border = border
                if (c - 1) in money_cols and isinstance(val, (int, float)):
                    cell.number_format = '#,##0.00;(#,##0.00);"-"'
                    cell.alignment = Alignment(horizontal="right")

        for c, label in enumerate(cols, start=1):
            longest = len(str(label))
            for row in (sheet.get("rows") or [])[:400]:
                if c - 1 < len(row):
                    longest = max(longest, len(str(row[c - 1] if row[c - 1] is not None else "")))
            ws.column_dimensions[get_column_letter(c)].width = min(max(12, longest + 2), 55)
        ws.freeze_panes = ws.cell(row=start + 1, column=1)
        if cols and (sheet.get("rows")):
            ws.auto_filter.ref = f"A{start}:{get_column_letter(len(cols))}{start + len(sheet['rows'])}"

    if not wb.sheetnames:
        wb.create_sheet("Empty")["A1"] = "No data available for this period."

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
