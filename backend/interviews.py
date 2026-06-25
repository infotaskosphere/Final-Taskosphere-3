"""
Employee Interviews module.

Lets HR/admin staff log every candidate interviewed, the pay scale offered,
conditions, training period, department & experience — then, once a
candidate is hired, convert them straight into a system User without
re-typing anything (all fields pre-filled from the interview record but
fully editable before the account is created).
"""

import io
import os
import re
import uuid
import logging
from datetime import datetime, timezone
from typing import List, Optional, Literal

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from pydantic import BaseModel, Field, ConfigDict, EmailStr, field_validator
from bson import ObjectId

from backend.dependencies import db, get_current_user, create_audit_log, _get_perm

router = APIRouter(prefix="/interviews", tags=["Employee Interviews"])
logger = logging.getLogger(__name__)


# ====================== PERMISSIONS ======================


def _can_manage(current_user) -> bool:
    """Admin always has access. Other roles need can_view_interviews permission."""
    if getattr(current_user, "role", None) == "admin":
        return True
    return bool(_get_perm(current_user, "can_view_interviews"))


def assert_can_manage(current_user):
    if not _can_manage(current_user):
        raise HTTPException(
            status_code=403,
            detail="You do not have permission to access Employee Interviews",
        )


# ====================== MODELS ======================


class CandidateBase(BaseModel):
    model_config = ConfigDict(extra="ignore")

    full_name: str = Field(..., description="Candidate's full name")
    email: Optional[EmailStr] = None
    phone: Optional[str] = None

    position: Optional[str] = Field(
        None, description="Role applied for / interviewed for"
    )
    department: Optional[str] = None
    experience_years: Optional[float] = None
    current_company: Optional[str] = None
    skills: List[str] = Field(default_factory=list)
    education: Optional[str] = None

    interview_date: Optional[str] = None  # ISO date (YYYY-MM-DD)
    interview_time: Optional[str] = None  # HH:MM (24h)
    interviewer: Optional[str] = None

    attendance: Literal["pending", "attended", "not_attended"] = "pending"
    interview_feedback: Optional[str] = None

    pay_scale_offered: Optional[str] = None
    conditions: Optional[str] = None
    training_period: Optional[str] = None

    status: Literal[
        "scheduled", "in_review", "selected", "on_hold", "rejected", "hired"
    ] = "scheduled"

    notes: Optional[str] = None
    resume_filename: Optional[str] = None
    resume_text: Optional[str] = None

    @field_validator(
        "email",
        "phone",
        "position",
        "department",
        "current_company",
        "education",
        "interview_date",
        "interview_time",
        "interviewer",
        "interview_feedback",
        "pay_scale_offered",
        "conditions",
        "training_period",
        "notes",
        mode="before",
    )
    @classmethod
    def empty_to_none(cls, v):
        return None if v == "" else v

    @field_validator("experience_years", mode="before")
    @classmethod
    def exp_to_float(cls, v):
        if v in ("", None):
            return None
        try:
            return float(v)
        except (TypeError, ValueError):
            return None

    @field_validator("skills", mode="before")
    @classmethod
    def skills_to_list(cls, v):
        if isinstance(v, str):
            return [s.strip() for s in v.split(",") if s.strip()]
        return v or []


class CandidateCreate(CandidateBase):
    pass


class CandidateUpdate(BaseModel):
    full_name: Optional[str] = None
    email: Optional[EmailStr] = None
    phone: Optional[str] = None
    position: Optional[str] = None
    department: Optional[str] = None
    experience_years: Optional[float] = None
    current_company: Optional[str] = None
    skills: Optional[List[str]] = None
    education: Optional[str] = None
    interview_date: Optional[str] = None
    interview_time: Optional[str] = None
    interviewer: Optional[str] = None
    attendance: Optional[Literal["pending", "attended", "not_attended"]] = None
    interview_feedback: Optional[str] = None
    pay_scale_offered: Optional[str] = None
    conditions: Optional[str] = None
    training_period: Optional[str] = None
    status: Optional[
        Literal["scheduled", "in_review", "selected", "on_hold", "rejected", "hired"]
    ] = None
    notes: Optional[str] = None


class PraiseResumeRequest(BaseModel):
    resume_text: str
    position: Optional[str] = ""


class ConvertToUserRequest(BaseModel):
    """Sent by the frontend conversion dialog — every field is editable there
    before submission, even though it's pre-filled from the candidate record."""

    full_name: str
    email: EmailStr
    password: str
    role: str = "staff"
    departments: List[str] = Field(default_factory=list)
    phone: Optional[str] = None
    punch_in_time: Optional[str] = "10:30"
    grace_time: Optional[str] = "00:10"
    punch_out_time: Optional[str] = "19:00"


# ====================== HELPERS ======================


def normalize_doc(doc: dict) -> dict:
    if not doc:
        return doc
    if "_id" in doc:
        doc["id"] = str(doc.pop("_id"))
    return doc


def validate_obj_id(id_str: str) -> ObjectId:
    if not ObjectId.is_valid(id_str):
        raise HTTPException(status_code=400, detail="Invalid candidate ID")
    return ObjectId(id_str)


# Lightweight resume field extraction (regex) used as a fallback / supplement
# to whatever the AI model returns, so basic contact details are never missed.
_EMAIL_RE = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
_PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s\-()])?(?:\(?\d{2,5}\)?[\s\-]?)?\d{3,5}[\s\-]?\d{3,5}"
)
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/([^\s,]+)", re.I)


def _regex_fields(text: str) -> dict:
    out = {}
    m = _EMAIL_RE.search(text)
    if m:
        out["email"] = m.group(0)
    # Try multiple phone patterns
    for pat in [
        r"\+?\d{1,3}[\s\-]?\d{10}",  # +91 8238441686
        r"\(\d{3}\)\s?\d{3}[\s\-]?\d{4}",  # (123) 456-7890
        r"\d{10}",  # 8238441686
    ]:
        m = re.search(pat, text)
        if m:
            out["phone"] = m.group(0).strip()
            break
    m = _LINKEDIN_RE.search(text)
    if m:
        out["linkedin"] = m.group(0)
    return out


def _heuristic_fields(text: str) -> dict:
    """Section-aware heuristic extraction that works with NO external AI call.
    Used to fill any gaps the AI model leaves (or as the sole source if the
    AI call fails / no API key is configured), so parsing stays useful for
    any resume layout — not just ones the model happens to handle well."""
    import calendar
    from datetime import date as ddate

    out = {}
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return out

    # ── Name: first short, mostly-alphabetic line near the top ──────────────
    # Handles ALL CAPS names like "PUJA ADAK" as well as mixed case
    for l in lines[:8]:
        # Skip lines with email, phone numbers, URLs, or special chars
        if "@" in l or any(ch.isdigit() for ch in l):
            continue
        if re.search(r"http|www\.|\.com|\.in|linkedin|phone|email|address", l, re.I):
            continue
        # Remove bullet points and special prefixes
        clean = re.sub(r"^[•\-\*\|]\s*", "", l).strip()
        if not clean:
            continue
        words = clean.split()
        # Accept 1-6 word names, all alphabetic (with dots/hyphens/apostrophes)
        if 1 <= len(words) <= 6 and all(re.match(r"^[A-Za-z.\-']+$", w) for w in words):
            # Title-case if ALL CAPS, otherwise keep as-is
            out["full_name"] = clean.title() if clean.isupper() else clean
            break

    # ── Split into sections by common resume headings ────────────────────────
    # Expanded to handle more section names found in real resumes
    # More flexible pattern that handles variations like "CS ARTICLESHIP EXPERIENCE"
    section_pat = re.compile(
        r"^(.*\b)?"  # Optional prefix (e.g., "CS ", "TECHNICAL AND ")
        r"(SUMMARY|OBJECTIVE|PROFILE|CAREER\s*OBJECTIVE|"
        r"WORK\s*EXPERIENCE|EXPERIENCE|EMPLOYMENT(\s*HISTORY)?|"
        r"INTERNSHIP|ARTICLESHIP|TRAINING(\s*EXPERIENCE)?|"
        r"EDUCATION|ACADEMIC(\s*(BACKGROUND|QUALIFICATION|DETAILS))?|"
        r"PROFESSIONAL\s*QUALIFICATION|QUALIFICATION|CERTIFICATION|CERTIFICATIONS|"
        r"KEY\s*SKILLS|SKILLS|TECHNICAL(\s*AND\s*\w+)?\s*SKILLS|CORE\s*COMPETENCIES|COMPETENCIES|"
        r"PROJECTS|ACADEMIC\s*PROJECTS|"
        r"PERSONAL\s*DETAILS|PERSONAL\s*PROFILE|"
        r"ACHIEVEMENTS|AWARDS|EXTRA[\-\s]?CURRICULAR|LANGUAGES)"
        r"(\b.*)?$",  # Optional suffix
        re.I,
    )
    sections, current, buf = {}, "header", []
    for l in lines:
        m = section_pat.match(l)
        if m:
            # Extract the main section keyword (group 2)
            section_name = m.group(2).lower().strip()
            sections[current] = buf
            current = re.sub(r"\s+", " ", section_name)
            buf = []
        else:
            buf.append(l)
    sections[current] = buf

    def _get_section(*names):
        for n in names:
            for key in sections:
                if n in key:
                    return sections[key]
        return []

    # ── Position / Role ──────────────────────────────────────────────────────
    # Try experience/internship sections first
    exp_lines = _get_section(
        "work experience",
        "experience",
        "employment",
        "internship",
        "articleship",
        "training",
    )
    if exp_lines:
        first = exp_lines[0]
        # Try "Position at Company" pattern
        m = re.search(r"(.+?)\bat\b\s*(.+)", first, re.I)
        if m:
            out["position"] = m.group(1).strip(" ,")
            out["current_company"] = re.split(r"\s{2,}", m.group(2))[0].strip(" .")
        # Try "Trainer: X" or "Company: X" pattern
        elif re.search(r"trainer|company|organization|employer", first, re.I):
            m2 = re.search(
                r"(?:trainer|company|organization|employer)[:\s]+(.+)", first, re.I
            )
            if m2:
                out["current_company"] = m2.group(1).strip()
            # Look for role in subsequent lines
            for el in exp_lines[1:4]:
                if re.search(r"assist|help|work|manage|handle|draft|file", el, re.I):
                    # Extract first action as position hint
                    out.setdefault("position", "Trainee")
                    break
        else:
            out["position"] = first.split(",")[0].strip()

        # Try to extract duration for experience calculation
        duration_re = re.compile(
            r"(?:duration|period)[:\s]*(\d+)\s*(month|year|week)", re.I
        )
        for el in exp_lines[:5]:
            dm = duration_re.search(el)
            if dm:
                num = int(dm.group(1))
                unit = dm.group(2).lower()
                if "year" in unit:
                    out["experience_years"] = float(num)
                elif "month" in unit:
                    out["experience_years"] = round(num / 12, 1)
                elif "week" in unit:
                    out["experience_years"] = round(num / 52, 1)
                break

        # Estimate total experience by summing every "Mon YYYY - Mon YYYY" style range
        if "experience_years" not in out:
            date_re = re.compile(
                r"([A-Za-z]{3,9})\s+(\d{4})\s*[-–to]+\s*([A-Za-z]{3,9}|present)\s*(\d{4})?",
                re.I,
            )

            def month_num(name):
                name = name[:3].title()
                try:
                    return list(calendar.month_abbr).index(name)
                except ValueError:
                    return None

            months_total = 0
            for l in exp_lines:
                m2 = date_re.search(l)
                if not m2:
                    continue
                m1n, y1 = month_num(m2.group(1)), int(m2.group(2))
                if m1n is None:
                    continue
                if m2.group(3).lower() == "present":
                    end = ddate.today()
                else:
                    m2n = month_num(m2.group(3))
                    if m2n is None:
                        continue
                    end = ddate(int(m2.group(4) or y1), m2n, 1)
                start = ddate(y1, m1n, 1)
                months = (end.year - start.year) * 12 + (end.month - start.month)
                if months > 0:
                    months_total += months
            if months_total:
                out["experience_years"] = round(months_total / 12, 1)

    # ── Position from summary/objective if not found in experience ───────────
    if "position" not in out:
        summary_lines = _get_section(
            "summary", "objective", "profile", "career objective"
        )
        if summary_lines:
            first_summary = summary_lines[0]
            # Look for "seeking to apply... in [field]" or "looking for [role]"
            m = re.search(
                r"(?:seeking|looking\s*for|aspiring|as\s+a)\s+(?:a\s+)?(.+?)(?:\s+to|\s+in|\s+at|,|\.)",
                first_summary,
                re.I,
            )
            if m:
                role = m.group(1).strip()
                # Clean up common prefixes
                role = re.sub(
                    r"^(position|role|job)\s+(?:of|as|in)\s+", "", role, flags=re.I
                )
                if len(role) < 50:
                    out["position"] = role

    # ── Education ────────────────────────────────────────────────────────────
    edu_lines = _get_section(
        "education", "academic background", "academic qualification", "academic details"
    )
    if edu_lines:
        out["education"] = "; ".join(edu_lines[:4])
    else:
        # Look for degree patterns anywhere in text
        degree_re = re.compile(
            r"(Bachelor|Master|B\.|M\.|Ph\.?D|MBA|BBA|BCA|MCA|B\.Com|M\.Com|B\.Sc|M\.Sc|B\.Tech|M\.Tech|B\.A|M\.A)",
            re.I,
        )
        for l in lines[:20]:
            if degree_re.search(l):
                out["education"] = l.strip()
                break

    skill_lines = _get_section(
        "key skills", "skills", "technical skills", "core competencies", "competencies"
    )
    if skill_lines:
        skills = []
        for l in skill_lines:
            # Split by comma, bullet, semicolon, dash, or pipe
            parts = re.split(r",|•|;|\||\-\s", l)
            for p in parts:
                clean = p.strip().lstrip("-•*").strip()
                # Skip lines that look like sentences (too long) or are empty
                if clean and len(clean) < 60:
                    skills.append(clean)
        if skills:
            out["skills"] = skills[:15]

    # ── Department mapping ───────────────────────────────────────────────────
    # Expanded to cover corporate law, compliance, company secretary, etc.
    dep_map = [
        ("trademark", "TM"),
        ("legal", "Legal"),
        ("law", "Legal"),
        ("corporate law", "Legal"),
        ("company secretary", "Legal"),
        ("cs ", "Legal"),  # CS = Company Secretary
        ("secretarial", "Legal"),
        ("compliance", "ROC"),
        ("roc", "ROC"),
        ("companies act", "ROC"),
        ("sebi", "ROC"),
        ("mca", "ROC"),
        ("account", "ACC"),
        ("finance", "ACC"),
        ("tax", "TDS"),
        ("gst", "GST"),
        ("software", "IT"),
        ("developer", "IT"),
        ("information technology", "IT"),
        ("human resource", "HR"),
        (" hr ", "HR"),
        ("marketing", "Marketing"),
        ("sales", "Sales"),
        ("operations", "Operations"),
    ]
    # Build haystack from all available context
    haystack_parts = [
        out.get("position", ""),
        " ".join(out.get("skills", [])),
        out.get("education", ""),
    ]
    # Also scan the first 15 lines of the resume for keywords
    haystack_parts.extend(lines[:15])
    haystack = f" {' '.join(haystack_parts)} ".lower()
    for kw, dep in dep_map:
        if kw in haystack:
            out["department"] = dep
            break

    return out


async def _groq_vision_structured_resume(page_images_b64: list, filename: str) -> dict:
    """Use Groq vision to directly extract structured candidate fields from
    scanned PDF page images. Returns a dict of extracted fields.

    This is more reliable than transcribing text first then parsing with
    heuristics, because the vision model can understand layout and context."""
    import httpx
    import json

    groq_key = os.environ.get("GROQ_API_KEY", "")
    if not groq_key:
        logger.warning(
            "GROQ_API_KEY not set — cannot extract structured data from scanned PDF."
        )
        return {}

    content = []
    for img_b64 in page_images_b64:
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
            }
        )
    content.append(
        {
            "type": "text",
            "text": (
                "You are an expert HR assistant. Analyze this resume image and extract candidate details.\n"
                "Respond with ONLY valid JSON (no markdown, no code fences, no explanation) using these exact keys:\n"
                '{"full_name":"","email":"","phone":"","position":"","department":"",'
                '"experience_years":0,"current_company":"","skills":[],"education":""}\n\n'
                "Rules:\n"
                "- full_name: the person's full name (usually the largest text at top)\n"
                "- email: email address\n"
                "- phone: phone number with country code if present\n"
                "- position: most recent or primary job title / role sought\n"
                "- department: one of Sales, IT, HR, Accounts, Legal, Marketing, Operations, GST, TDS, ROC, TM, Other\n"
                "- experience_years: total years as a number (0 if fresher)\n"
                "- current_company: most recent or primary company name\n"
                "- skills: array of skill keywords (max 15)\n"
                "- education: highest or most relevant degree\n"
                "- Leave blank/0/[] if genuinely not found. Never invent data.\n"
            ),
        }
    )

    payload = {
        "model": "meta-llama/llama-4-scout-17b-16e-instruct",
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 2000,
    }

    try:
        async with httpx.AsyncClient(timeout=90) as client:
            resp = await client.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {groq_key}",
                    "Content-Type": "application/json",
                },
                json=payload,
            )
        if resp.status_code == 429:
            logger.warning("Groq quota exceeded during structured resume extraction.")
            return {}
        if resp.status_code != 200:
            logger.warning(
                f"Groq API error {resp.status_code} during structured resume extraction."
            )
            return {}

        raw = resp.json()["choices"][0]["message"]["content"]
        # Clean markdown code fences if present
        raw = re.sub(r"^```(json)?|```$", "", raw.strip(), flags=re.MULTILINE).strip()
        return json.loads(raw)
    except Exception as e:
        logger.warning(f"Groq vision structured extraction failed: {e}")
        return {}


async def _extract_resume_text(contents: bytes, filename: str) -> str:
    """Extract text from PDF/DOCX/TXT resume with table support.

    Text-based PDFs → Gemini (same pattern as ai_document_reader.py).
    Scanned/image PDFs → Groq vision (up to 4 pages).
    DOCX / TXT → parsed locally.
    """
    import base64

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page in pdf.pages[:10]:
                t = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                for table in page.extract_tables():
                    rows = []
                    for row in table:
                        cells = [str(c).strip() if c else "" for c in row]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        t += "\n" + "\n".join(rows)
                if t.strip():
                    text_parts.append(t)
        text = "\n\n".join(text_parts).strip()

        if text:
            # Text-based PDF — use Gemini to clean / transcribe
            try:
                import google.generativeai as genai

                gemini_key = os.environ.get("GEMINI_API_KEY", "")
                if gemini_key:
                    genai.configure(api_key=gemini_key)
                    model = genai.GenerativeModel("gemini-2.0-flash")
                    prompt = (
                        "Below is raw text extracted from a resume PDF. "
                        "Return only the cleaned, fully transcribed resume text preserving all sections, "
                        "names, dates, and details. Do not summarise or add anything.\n\n"
                        f"{text[:30000]}"
                    )
                    resp = await model.generate_content_async(prompt)
                    cleaned = (resp.text or "").strip()
                    if cleaned:
                        return cleaned
            except Exception as e:
                logger.warning(
                    f"Gemini text-PDF cleaning failed, using raw extract: {e}"
                )
            return text  # fall back to raw pdfplumber text

        # Scanned PDF (no text layer) — render pages as images → Groq vision
        try:
            import httpx

            groq_key = os.environ.get("GROQ_API_KEY", "")
            if not groq_key:
                raise HTTPException(
                    status_code=500,
                    detail="GROQ_API_KEY is not configured on the server.",
                )

            page_images_b64 = []
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                for page in pdf.pages[:4]:  # max 4 pages for Groq
                    pil_img = page.to_image(resolution=150).original
                    if pil_img.mode not in ("RGB", "L"):
                        pil_img = pil_img.convert("RGB")
                    buf = io.BytesIO()
                    pil_img.save(buf, format="JPEG", quality=85)
                    page_images_b64.append(base64.b64encode(buf.getvalue()).decode())

            if not page_images_b64:
                raise HTTPException(
                    status_code=422, detail="No pages could be rendered from this PDF."
                )

            content = []
            for img_b64 in page_images_b64:
                content.append(
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
                    }
                )
            content.append(
                {
                    "type": "text",
                    "text": "Transcribe ALL text from these resume pages exactly as it appears, preserving structure, sections, and all details.",
                }
            )

            payload = {
                "model": "meta-llama/llama-4-scout-17b-16e-instruct",
                "messages": [{"role": "user", "content": content}],
                "max_tokens": 3000,
            }
            async with httpx.AsyncClient(timeout=90) as client:
                resp = await client.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {groq_key}",
                        "Content-Type": "application/json",
                    },
                    json=payload,
                )
            if resp.status_code == 429:
                raise HTTPException(
                    status_code=429,
                    detail="Groq quota exceeded. Please wait a moment and try again.",
                )
            if resp.status_code != 200:
                raise HTTPException(
                    status_code=422,
                    detail=f"Groq API error {resp.status_code}: {resp.text[:300]}",
                )
            return resp.json()["choices"][0]["message"]["content"]
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=422, detail=f"Could not read scanned resume: {e}"
            )

    if ext == "docx":
        try:
            import docx

            d = docx.Document(io.BytesIO(contents))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as e:
            raise HTTPException(
                status_code=422, detail=f"Could not read .docx resume: {e}"
            )

    if ext == "txt":
        return contents.decode("utf-8", errors="replace")

    raise HTTPException(
        status_code=422, detail=f"Unsupported format: .{ext}. Upload PDF, DOCX, or TXT."
    )


async def _ai_structure_resume(resume_text: str) -> dict:
    """Use Gemini to extract structured candidate fields from resume text.

    Falls back gracefully if GEMINI_API_KEY is absent.
    """
    import json

    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    if not gemini_key:
        logger.warning(
            "GEMINI_API_KEY not set — skipping AI structuring, using heuristics only."
        )
        return {}

    try:
        import google.generativeai as genai

        genai.configure(api_key=gemini_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
    except ImportError:
        logger.warning("google-generativeai not installed — skipping AI structuring.")
        return {}

    prompt = (
        "Extract candidate details from this resume and respond with ONLY valid JSON "
        "(no markdown, no code fences, no explanation) with these exact keys:\n"
        '{"full_name":"","email":"","phone":"","position":"","department":"",'
        '"experience_years":0,"current_company":"","skills":[],"education":""}\n\n'
        "Rules:\n"
        "- full_name: the person's full name (usually the largest / first text)\n"
        "- position: most recent or primary job title\n"
        "- department: one of Sales, IT, HR, Accounts, Legal, Marketing, Operations, GST, TDS, ROC, TM, Other\n"
        "- experience_years: total years as a number (sum from work history dates if needed)\n"
        "- skills: array of concise skill keywords (max 15)\n"
        "- education: most relevant degree/qualification as a short string\n"
        "- Leave blank/0/[] if genuinely not found. Never invent data.\n\n"
        f"RESUME:\n{resume_text[:12000]}"
    )

    try:
        resp = await model.generate_content_async(prompt)
        raw = (resp.text or "").strip()
        raw = re.sub(r"^```(json)?|```$", "", raw, flags=re.MULTILINE).strip()
        return json.loads(raw)
    except Exception as e:
        logger.warning(f"Gemini JSON parse failed: {e}")
        return {}


# ====================== ROUTES ======================


@router.get("")
async def list_candidates(
    status_filter: Optional[str] = None,
    department: Optional[str] = None,
    search: Optional[str] = None,
    current_user=Depends(get_current_user),
):
    assert_can_manage(current_user)
    query = {}
    if status_filter:
        query["status"] = status_filter
    if department:
        query["department"] = department
    if search:
        rx = {"$regex": re.escape(search), "$options": "i"}
        query["$or"] = [
            {"full_name": rx},
            {"email": rx},
            {"phone": rx},
            {"position": rx},
        ]

    docs = (
        await db.interview_candidates.find(query).sort("created_at", -1).to_list(1000)
    )
    return [normalize_doc(d) for d in docs]


@router.get("/{candidate_id}")
async def get_candidate(candidate_id: str, current_user=Depends(get_current_user)):
    assert_can_manage(current_user)
    doc = await db.interview_candidates.find_one({"_id": validate_obj_id(candidate_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Candidate not found")
    return normalize_doc(doc)


@router.post("")
async def create_candidate(
    payload: CandidateCreate, current_user=Depends(get_current_user)
):
    assert_can_manage(current_user)
    doc = payload.model_dump()
    doc.update(
        {
            "created_by": current_user.id,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "converted_user_id": None,
        }
    )
    result = await db.interview_candidates.insert_one(doc)
    new_doc = await db.interview_candidates.find_one({"_id": result.inserted_id})
    await create_audit_log(
        current_user,
        "create",
        "interview_candidate",
        str(result.inserted_id),
        new_data=doc,
    )
    return normalize_doc(new_doc)


@router.put("/{candidate_id}")
async def update_candidate(
    candidate_id: str, payload: CandidateUpdate, current_user=Depends(get_current_user)
):
    assert_can_manage(current_user)
    oid = validate_obj_id(candidate_id)
    existing = await db.interview_candidates.find_one({"_id": oid})
    if not existing:
        raise HTTPException(status_code=404, detail="Candidate not found")

    update_data = {k: v for k, v in payload.model_dump(exclude_unset=True).items()}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.interview_candidates.update_one({"_id": oid}, {"$set": update_data})

    updated = await db.interview_candidates.find_one({"_id": oid})
    await create_audit_log(
        current_user,
        "update",
        "interview_candidate",
        candidate_id,
        old_data=existing,
        new_data=update_data,
    )
    return normalize_doc(updated)


@router.delete("/{candidate_id}")
async def delete_candidate(candidate_id: str, current_user=Depends(get_current_user)):
    assert_can_manage(current_user)
    oid = validate_obj_id(candidate_id)
    existing = await db.interview_candidates.find_one({"_id": oid})
    if not existing:
        raise HTTPException(status_code=404, detail="Candidate not found")
    await db.interview_candidates.delete_one({"_id": oid})
    await create_audit_log(
        current_user, "delete", "interview_candidate", candidate_id, old_data=existing
    )
    return {"message": "Candidate deleted"}


@router.post("/parse-resume")
async def parse_resume(
    file: UploadFile = File(...), current_user=Depends(get_current_user)
):
    """Extracts text from an uploaded resume and returns structured candidate
    fields the frontend can drop straight into the Add Candidate form.

    For scanned PDFs (no text layer), uses Groq vision to directly extract
    structured fields from page images, bypassing unreliable text extraction."""
    assert_can_manage(current_user)
    contents = await file.read()
    filename = file.filename or "resume"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Track if this is a scanned PDF for special handling
    is_scanned_pdf = False
    scanned_fields = {}
    resume_text = ""

    # ── Special handling for scanned PDFs ────────────────────────────────────
    if ext == "pdf":
        import pdfplumber
        import base64

        # First check if PDF has any extractable text
        has_text = False
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page in pdf.pages[:3]:
                t = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                if t.strip():
                    has_text = True
                    break

        if not has_text:
            # Scanned PDF — use Groq vision to directly extract structured fields
            is_scanned_pdf = True
            logger.info(
                "Scanned PDF detected — using Groq vision for structured extraction"
            )

            try:
                groq_key = os.environ.get("GROQ_API_KEY", "")
                if groq_key:
                    page_images_b64 = []
                    with pdfplumber.open(io.BytesIO(contents)) as pdf:
                        for page in pdf.pages[:4]:  # max 4 pages
                            pil_img = page.to_image(resolution=150).original
                            if pil_img.mode not in ("RGB", "L"):
                                pil_img = pil_img.convert("RGB")
                            buf = io.BytesIO()
                            pil_img.save(buf, format="JPEG", quality=85)
                            page_images_b64.append(
                                base64.b64encode(buf.getvalue()).decode()
                            )

                    if page_images_b64:
                        # Get structured fields directly from vision
                        scanned_fields = await _groq_vision_structured_resume(
                            page_images_b64, filename
                        )
                        logger.info(
                            f"Groq vision extracted fields: {list(scanned_fields.keys())}"
                        )

                        # Also get raw text transcription for AI assessment
                        import httpx

                        content = []
                        for img_b64 in page_images_b64:
                            content.append(
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{img_b64}"
                                    },
                                }
                            )
                        content.append(
                            {
                                "type": "text",
                                "text": "Transcribe ALL text from these resume pages exactly as it appears.",
                            }
                        )

                        payload = {
                            "model": "meta-llama/llama-4-scout-17b-16e-instruct",
                            "messages": [{"role": "user", "content": content}],
                            "max_tokens": 3000,
                        }
                        async with httpx.AsyncClient(timeout=90) as client:
                            resp = await client.post(
                                "https://api.groq.com/openai/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {groq_key}",
                                    "Content-Type": "application/json",
                                },
                                json=payload,
                            )
                        if resp.status_code == 200:
                            resume_text = resp.json()["choices"][0]["message"][
                                "content"
                            ]
                        else:
                            logger.warning(
                                f"Groq text transcription failed: {resp.status_code}"
                            )
                else:
                    logger.warning("GROQ_API_KEY not set — cannot process scanned PDF")
                    raise HTTPException(
                        status_code=500,
                        detail="GROQ_API_KEY is not configured on the server.",
                    )
            except HTTPException:
                raise
            except Exception as e:
                logger.error(f"Scanned PDF processing failed: {e}")
                raise HTTPException(
                    status_code=422, detail=f"Could not read scanned resume: {e}"
                )

    # ── Normal text extraction for non-scanned files ─────────────────────────
    if not is_scanned_pdf:
        resume_text = await _extract_resume_text(contents, filename)
        if not resume_text.strip():
            raise HTTPException(
                status_code=422, detail="Could not extract any text from this resume"
            )

    # ── AI structuring (Gemini for text, skip if already have scanned fields) ─
    ai_fields = {}
    if not is_scanned_pdf or not scanned_fields:
        try:
            ai_fields = await _ai_structure_resume(resume_text)
        except HTTPException:
            raise
        except Exception as e:
            logger.warning(f"AI resume structuring failed: {e}")
            ai_fields = {}

    # ── Regex and heuristic extraction ───────────────────────────────────────
    regex_f = _regex_fields(resume_text)
    heuristic_f = _heuristic_fields(resume_text)

    def first(*vals):
        """Return first non-empty value from the list."""
        for v in vals:
            if v and str(v).strip():
                return v
        return ""

    # ── Combine all sources with priority: scanned > ai > heuristic > regex ──
    fields = {
        "full_name": first(
            scanned_fields.get("full_name"),
            ai_fields.get("full_name"),
            heuristic_f.get("full_name"),
        ),
        "email": first(
            scanned_fields.get("email"),
            ai_fields.get("email"),
            regex_f.get("email"),
            heuristic_f.get("email"),
        ),
        "phone": first(
            scanned_fields.get("phone"),
            ai_fields.get("phone"),
            regex_f.get("phone"),
            heuristic_f.get("phone"),
        ),
        "position": first(
            scanned_fields.get("position"),
            ai_fields.get("position"),
            heuristic_f.get("position"),
        ),
        "department": first(
            scanned_fields.get("department"),
            ai_fields.get("department"),
            heuristic_f.get("department"),
        ),
        "experience_years": (
            scanned_fields.get("experience_years")
            or ai_fields.get("experience_years")
            or heuristic_f.get("experience_years")
            or None
        ),
        "current_company": first(
            scanned_fields.get("current_company"),
            ai_fields.get("current_company"),
            heuristic_f.get("current_company"),
        ),
        "skills": (
            scanned_fields.get("skills")
            or ai_fields.get("skills")
            or heuristic_f.get("skills")
            or []
        ),
        "education": first(
            scanned_fields.get("education"),
            ai_fields.get("education"),
            heuristic_f.get("education"),
        ),
    }

    # Log what was extracted
    filled_count = sum(
        1
        for v in fields.values()
        if v and (isinstance(v, list) and len(v) > 0 or str(v).strip())
    )
    logger.info(
        f"Resume parsing complete: {filled_count}/9 fields filled (scanned={is_scanned_pdf})"
    )

    return {
        "filename": filename,
        "resume_text": resume_text[:20000],
        "fields": fields,
    }


async def _ai_assess_resume(resume_text: str, position: str) -> dict:
    """Use Gemini to produce a structured candidate assessment (verdict, score,
    strengths, concerns, standout skills, experience quality, and suggested
    interview questions).  Falls back to a sensible default dict if the AI
    call fails or no API key is configured."""
    import json

    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    if not gemini_key:
        logger.warning("GEMINI_API_KEY not set — skipping AI resume assessment.")
        return _default_assessment()

    try:
        import google.generativeai as genai

        genai.configure(api_key=gemini_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
    except ImportError:
        logger.warning(
            "google-generativeai not installed — skipping AI resume assessment."
        )
        return _default_assessment()

    position_hint = (
        f"The candidate applied for: {position}."
        if position
        else "No specific position was mentioned."
    )

    prompt = (
        "You are an expert HR recruiter evaluating a candidate's resume.\n"
        f"{position_hint}\n\n"
        "Analyse the resume below and respond with ONLY valid JSON "
        "(no markdown, no code fences, no explanation) using these exact keys:\n"
        "{\n"
        '  "verdict": "Strong Hire" | "Hire" | "Maybe" | "Pass",\n'
        '  "fit_score": <integer 0-100>,\n'
        '  "summary": "<one concise sentence summarising the candidate>",\n'
        '  "standout_skills": ["<skill>", ...],\n'
        '  "strengths": ["<strength>", ...],\n'
        '  "concerns": ["<concern>", ...],\n'
        '  "experience_quality": "Excellent" | "Good" | "Average" | "Poor",\n'
        '  "recommended_questions": ["<question>", ...]\n'
        "}\n\n"
        "Rules:\n"
        "- verdict: Strong Hire (80+), Hire (60-79), Maybe (40-59), Pass (<40) based on fit_score.\n"
        "- fit_score: 0-100 integer reflecting overall suitability.\n"
        "- standout_skills: max 5 key skills that stand out.\n"
        "- strengths: max 5 bullet points.\n"
        "- concerns: max 5 bullet points (can be empty array if none).\n"
        "- experience_quality: rate the depth and relevance of work experience.\n"
        "- recommended_questions: 3-5 interview questions to ask this candidate.\n"
        "- Never invent data not supported by the resume.\n\n"
        f"RESUME:\n{resume_text[:12000]}"
    )

    try:
        resp = await model.generate_content_async(prompt)
        raw = (resp.text or "").strip()
        raw = re.sub(r"^```(json)?|```$", "", raw, flags=re.MULTILINE).strip()
        data = json.loads(raw)

        # Validate / normalise
        valid_verdicts = {"Strong Hire", "Hire", "Maybe", "Pass"}
        if data.get("verdict") not in valid_verdicts:
            score = int(data.get("fit_score", 50))
            if score >= 80:
                data["verdict"] = "Strong Hire"
            elif score >= 60:
                data["verdict"] = "Hire"
            elif score >= 40:
                data["verdict"] = "Maybe"
            else:
                data["verdict"] = "Pass"

        data.setdefault("fit_score", 50)
        data.setdefault("summary", "Candidate profile assessed.")
        data.setdefault("standout_skills", [])
        data.setdefault("strengths", [])
        data.setdefault("concerns", [])
        data.setdefault("recommended_questions", [])

        valid_exp = {"Excellent", "Good", "Average", "Poor"}
        if data.get("experience_quality") not in valid_exp:
            data["experience_quality"] = "Average"

        return data
    except Exception as e:
        logger.warning(f"Gemini resume assessment failed: {e}")
        return _default_assessment()


def _default_assessment() -> dict:
    """Fallback assessment when AI is unavailable."""
    return {
        "verdict": "Maybe",
        "fit_score": 50,
        "summary": "AI assessment unavailable — please review manually.",
        "standout_skills": [],
        "strengths": [],
        "concerns": [],
        "experience_quality": "Average",
        "recommended_questions": [],
    }


@router.post("/praise-resume-json")
async def praise_resume_json(
    payload: PraiseResumeRequest, current_user=Depends(get_current_user)
):
    """AI-powered resume assessment: returns verdict, fit score, strengths,
    concerns, standout skills, experience quality, and suggested interview
    questions for the uploaded resume."""
    assert_can_manage(current_user)

    if not payload.resume_text or not payload.resume_text.strip():
        raise HTTPException(status_code=422, detail="No resume text provided")

    assessment = await _ai_assess_resume(payload.resume_text, payload.position or "")
    return assessment


@router.post("/{candidate_id}/convert-to-user")
async def convert_to_user(
    candidate_id: str,
    payload: ConvertToUserRequest,
    current_user=Depends(get_current_user),
):
    """Creates a real system User account from a hired candidate.
    All fields arrive from the (editable) conversion form on the frontend —
    nothing here is silently reused from the interview record."""
    assert_can_manage(current_user)
    oid = validate_obj_id(candidate_id)
    candidate = await db.interview_candidates.find_one({"_id": oid})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    if candidate.get("converted_user_id"):
        raise HTTPException(
            status_code=400,
            detail="This candidate has already been converted to a user",
        )

    existing_user = await db.users.find_one({"email": payload.email}, {"_id": 0})
    if existing_user:
        raise HTTPException(
            status_code=400, detail="A user with this email already exists"
        )

    # Late imports avoid a circular import with backend.server at module load time
    from backend.server import get_password_hash, DEFAULT_ROLE_PERMISSIONS

    if (
        payload.role in ("admin", "manager", "superadmin")
        and current_user.role != "admin"
    ):
        raise HTTPException(
            status_code=403, detail="Only an admin can assign that role"
        )

    user_id = str(uuid.uuid4())
    default_permissions = DEFAULT_ROLE_PERMISSIONS.get(payload.role, {})

    new_user = {
        "id": user_id,
        "email": payload.email,
        "full_name": payload.full_name,
        "role": payload.role,
        "password": get_password_hash(payload.password),
        "departments": payload.departments or [],
        "phone": payload.phone,
        "punch_in_time": payload.punch_in_time or "10:30",
        "grace_time": payload.grace_time or "00:10",
        "punch_out_time": payload.punch_out_time or "19:00",
        "is_active": False,
        "status": "pending_approval",
        "approved_by": None,
        "approved_at": None,
        "permissions": default_permissions,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source_candidate_id": str(oid),
    }
    await db.users.insert_one(new_user)

    await db.interview_candidates.update_one(
        {"_id": oid},
        {
            "$set": {
                "status": "hired",
                "converted_user_id": user_id,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }
        },
    )

    await create_audit_log(
        current_user,
        "convert_to_user",
        "interview_candidate",
        candidate_id,
        new_data={"user_id": user_id},
    )

    new_user.pop("password", None)
    new_user.pop("_id", None)
    return {
        "message": "Candidate converted to user — pending admin approval",
        "user": new_user,
    }
