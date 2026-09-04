"""
ROC Sphere — Companies Act (India) compliance & document automation module.

Self-contained backend router. Nothing here is imported by any other
existing module, so dropping this file into backend/ does not touch any
existing behaviour. It is wired into the app with exactly two lines in
server.py (see INTEGRATION.md):

    from backend.roc_sphere import router as roc_sphere_router
    api_router.include_router(roc_sphere_router)

Covers:
  - Company master (linked to an existing Client, or standalone)
  - Directors / shareholders register
  - Upload & best-effort parse of AOC-4 / MGT-7 / MGT-7A PDFs plus the
    separate macro-enabled MGT-7A shareholder workbook
  - Share-transfer register, SH-4 instrument and share-certificate drafts
  - Companies Act 2013 compliance checklist engine (heuristic, based on
    company category/size — see COMPLIANCE_RULES below)
  - Word (.docx) generation for: Board Resolution, Notice of Meeting
    (Board/EGM/AGM), Minutes of Meeting (Board/General), Register of
    Members / List of Shareholders, and a printable Compliance Checklist

IMPORTANT — legal disclaimer baked into the product, not just this
comment: MCA thresholds and formats change (e.g. the "small company"
paid-up capital/turnover limits were revised effective 1 Dec 2025). The
checklist engine is a drafting aid, not a substitute for a professional's
judgement, and COMPLIANCE_RULES should be reviewed periodically against
the current Companies Act / MCA rules.
"""

import io
import logging
import re
import uuid
from datetime import datetime, timezone, date
from typing import Optional, List, Any, Dict

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import Response
from pydantic import BaseModel, Field, ConfigDict

from backend.dependencies import db, get_current_user, check_module_permission
from backend.models import User

# Permission flags used here (see backend/models.py DEFAULT_ROLE_PERMISSIONS
# and backend/dependencies.py MODULE_ACTION_MAP):
#   can_view_roc_sphere   — open the page, view company masters/checklist,
#                            generate & download documents
#   can_manage_roc_sphere — create/edit/delete company masters, upload
#                            AOC-4/MGT-7 to prefill them
# Admin-granted-only by default (same pattern as GST Reconciliation,
# Trademark Sphere and the Salary Slip Generator) — toggled per-user from
# Settings → Permission Governance → Compliance → "ROC Sphere".
VIEW = check_module_permission("roc_sphere", "view")
CREATE = check_module_permission("roc_sphere", "create")
EDIT = check_module_permission("roc_sphere", "edit")
DELETE = check_module_permission("roc_sphere", "delete")

logger = logging.getLogger("roc_sphere")
router = APIRouter(prefix="/roc-sphere", tags=["roc-sphere"])

COMPANIES = db.roc_companies
DOCS_LOG = db.roc_documents
CLIENTS = db.clients

COMPANY_CLIENT_TYPES = {"pvt_ltd", "PVT_LTD", "public_ltd", "section_8", "llp", "LLP", "opc"}


def _uid() -> str:
    return str(uuid.uuid4())


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _who(user: User) -> str:
    return getattr(user, "full_name", None) or getattr(user, "username", None) or "—"


def _fmt_date(v: Any) -> str:
    if not v:
        return "—"
    if isinstance(v, (datetime, date)):
        return v.strftime("%d-%m-%Y")
    s = str(v)
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(s[:10], fmt).strftime("%d-%m-%Y")
        except ValueError:
            continue
    return s


def _num(v: Any) -> float:
    try:
        return float(v)
    except (TypeError, ValueError):
        return 0.0


# ─────────────────────────────────────────────────────────────────────────
# MODELS
# ─────────────────────────────────────────────────────────────────────────

class Director(BaseModel):
    model_config = ConfigDict(extra="ignore")
    name: str
    din: Optional[str] = None
    designation: Optional[str] = "Director"  # Director / Managing Director / Whole-time Director / Additional Director
    date_of_appointment: Optional[str] = None
    date_of_cessation: Optional[str] = None
    pan: Optional[str] = None
    address: Optional[str] = None


class Shareholder(BaseModel):
    model_config = ConfigDict(extra="ignore")
    name: str
    folio_no: Optional[str] = None
    pan: Optional[str] = None
    holder_type: Optional[str] = None
    category: Optional[str] = None
    details: Optional[str] = None
    class_of_shares: Optional[str] = "Equity"
    security_type: Optional[str] = None
    nationality: Optional[str] = None
    gender: Optional[str] = None
    identifier_type: Optional[str] = None
    occupation: Optional[str] = None
    shares_held: float = 0
    face_value: Optional[float] = 10
    total_value: Optional[float] = None
    percentage: Optional[float] = None
    address: Optional[str] = None


class ShareTransferRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    transfer_date: Optional[str] = None
    transferor_name: str
    transferee_name: str
    transferor_folio_no: Optional[str] = None
    transferee_folio_no: Optional[str] = None
    share_certificate_no: Optional[str] = None
    distinctive_from: Optional[str] = None
    distinctive_to: Optional[str] = None
    number_of_shares: float = 0
    class_of_shares: str = "Equity"
    nominal_value_per_share: float = 10
    consideration: float = 0
    stamp_duty: float = 0
    board_resolution_date: Optional[str] = None
    instrument_date: Optional[str] = None
    instrument_received_date: Optional[str] = None
    sh4_status: str = "Pending review"
    remarks: Optional[str] = None
    update_register: bool = True


class ShareCertificateRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    certificate_no: str
    issue_date: Optional[str] = None
    holder_name: str
    holder_address: Optional[str] = None
    folio_no: Optional[str] = None
    class_of_shares: str = "Equity"
    number_of_shares: float = 0
    distinctive_from: Optional[str] = None
    distinctive_to: Optional[str] = None
    nominal_value_per_share: float = 10
    amount_paid_per_share: float = 0
    joint_holders: List[str] = Field(default_factory=list)
    remarks: Optional[str] = None


class Auditor(BaseModel):
    model_config = ConfigDict(extra="ignore")
    name: Optional[str] = None
    firm_reg_no: Optional[str] = None
    membership_no: Optional[str] = None
    appointed_from: Optional[str] = None
    appointed_till: Optional[str] = None


class RocCompanyIn(BaseModel):
    model_config = ConfigDict(extra="ignore")
    client_id: Optional[str] = None
    company_name: str
    cin: Optional[str] = None
    category: str = "private"          # private | public | opc | section_8 | llp
    is_small_company: Optional[bool] = None   # None = auto-compute from capital/turnover
    listed: bool = False
    roc_office: Optional[str] = None
    pan: Optional[str] = None
    date_of_incorporation: Optional[str] = None
    registered_office_address: Optional[str] = None
    authorized_capital: Optional[float] = 0
    paid_up_capital: Optional[float] = 0
    last_year_turnover: Optional[float] = 0
    financial_year_end: Optional[str] = "31-03"
    last_agm_date: Optional[str] = None
    last_board_meeting_date: Optional[str] = None
    directors: List[Director] = Field(default_factory=list)
    designated_partners: List[Director] = Field(default_factory=list)
    partners: List[Director] = Field(default_factory=list)
    shareholders: List[Shareholder] = Field(default_factory=list)
    master_data: Dict[str, Any] = Field(default_factory=dict)
    mgt_shareholder_data: Dict[str, Any] = Field(default_factory=dict)
    financial_data: Dict[str, Any] = Field(default_factory=dict)  # key figures pulled from AOC-4 (see FINANCIAL_DATA_FIELDS)
    annual_return_data: Dict[str, Any] = Field(default_factory=dict)
    audit_report_data: Dict[str, Any] = Field(default_factory=dict)
    board_report_data: Dict[str, Any] = Field(default_factory=dict)
    share_transfers: List[Dict[str, Any]] = Field(default_factory=list)
    share_certificates: List[Dict[str, Any]] = Field(default_factory=list)
    roc_form_uploads: List[Dict[str, Any]] = Field(default_factory=list)
    auditor: Optional[Auditor] = None
    notes: Optional[str] = None


class RocCompanyOut(RocCompanyIn):
    id: str
    created_at: datetime
    updated_at: datetime
    created_by: Optional[str] = None


class ResolutionItem(BaseModel):
    model_config = ConfigDict(extra="ignore")
    particulars: str                 # short heading, e.g. "Opening of Bank Account"
    resolution_text: str             # the "RESOLVED THAT ..." body
    proposed_by: Optional[str] = None
    seconded_by: Optional[str] = None


class BoardResolutionRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    meeting_date: str
    meeting_time: Optional[str] = "11:00 AM"
    venue: Optional[str] = "Registered Office of the Company"
    directors_present: List[str] = Field(default_factory=list)
    chairman: Optional[str] = None
    resolutions: List[ResolutionItem]


class MeetingNoticeRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    meeting_type: str = "board"      # board | agm | egm
    meeting_date: str
    meeting_time: Optional[str] = "11:00 AM"
    venue: Optional[str] = "Registered Office of the Company"
    notice_date: Optional[str] = None
    agenda_items: List[str] = Field(default_factory=list)
    special_business: List[ResolutionItem] = Field(default_factory=list)


class MinutesRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    meeting_type: str = "board"      # board | agm | egm
    meeting_date: str
    meeting_time: Optional[str] = "11:00 AM"
    venue: Optional[str] = "Registered Office of the Company"
    chairman: Optional[str] = None
    directors_present: List[str] = Field(default_factory=list)
    directors_absent: List[str] = Field(default_factory=list)
    attendees_other: List[str] = Field(default_factory=list)
    quorum_present: bool = True
    resolutions: List[ResolutionItem] = Field(default_factory=list)
    discussion_notes: Optional[str] = None


# ─────────────────────────────────────────────────────────────────────────
# COMPANY MASTER — CRUD
# ─────────────────────────────────────────────────────────────────────────

@router.get("/clients-eligible")
async def list_eligible_clients(current_user: User = Depends(VIEW)):
    """Clients that are registered entities (not proprietors) and don't yet
    have a ROC Sphere company master — used to populate the 'create from
    client' picker."""
    cursor = CLIENTS.find({"client_type": {"$in": list(COMPANY_CLIENT_TYPES)}})
    clients = [c async for c in cursor]
    existing = {c["client_id"] async for c in COMPANIES.find({"client_id": {"$ne": None}}, {"client_id": 1}) if c.get("client_id")}
    out = []
    for c in clients:
        if c.get("id") in existing:
            continue
        out.append({
            "client_id": c.get("id"),
            "company_name": c.get("company_name"),
            "client_type": c.get("client_type"),
            "pan": c.get("pan"),
            "date_of_incorporation": c.get("date_of_incorporation"),
            "address": c.get("address"),
        })
    return out


CLIENT_CATEGORY_MAP = {
    "pvt_ltd": "private", "PVT_LTD": "private",
    "private_limited": "private", "private_limited_company": "private",
    "public_ltd": "public", "public_limited": "public", "public_limited_company": "public",
    "section_8": "section_8", "section8": "section_8", "section_8_company": "section_8",
    "llp": "llp", "LLP": "llp", "limited_liability_partnership": "llp",
    "opc": "opc", "one_person_company": "opc",
}


def _normalize_roc_category(company: dict, client: Optional[dict] = None) -> str:
    """Return the reliable ROC entity category for legacy and imported records.

    Older ROC records can have category='private' even when the underlying
    client was imported as LLP.  LLPIN/master-data/name are stronger signals
    than that stale category, so they are checked first.
    """
    name = str(company.get("company_name") or "").strip().lower()
    llpin = str(company.get("llpin") or "").strip()
    master = company.get("master_data") or {}
    if not llpin:
        llpin = str(master.get("llpin") or master.get("llpin_number") or "").strip()

    raw_client_type = str((client or {}).get("client_type") or "").strip().lower().replace("-", "_").replace(" ", "_")
    mapped_client_type = CLIENT_CATEGORY_MAP.get(raw_client_type)

    # An LLP suffix / LLPIN is definitive for this filter. This deliberately
    # overrides a stale 'private' category on old records.
    if llpin or name.endswith(" llp") or name.endswith("llp") or " limited liability partnership" in name:
        return "llp"
    if mapped_client_type:
        return mapped_client_type

    raw = str(company.get("category") or "").strip().lower().replace("-", "_").replace(" ", "_")
    return CLIENT_CATEGORY_MAP.get(raw, raw or "private")


async def _sync_companies_from_clients() -> int:
    """Auto-provision a ROC Sphere company master for every Client record
    that is a registered entity (Pvt/Public Ltd, LLP, OPC, Section 8) and
    doesn't have one yet, prefilled from that client's CIN/PAN/address/
    incorporation date. Existing ROC Sphere records (and anything a user
    has since edited on them) are never touched -- this only fills the gap
    for clients that have no linked record at all. Returns count created.
    """
    cursor = CLIENTS.find({"client_type": {"$in": list(COMPANY_CLIENT_TYPES)}})
    clients = [c async for c in cursor]
    if not clients:
        return 0
    existing_client_ids = {
        c["client_id"]
        async for c in COMPANIES.find({"client_id": {"$ne": None}}, {"client_id": 1})
        if c.get("client_id")
    }
    now = _now()
    new_docs = []
    for c in clients:
        cid = c.get("id")
        if not cid or cid in existing_client_ids:
            continue
        contact_people = c.get("contact_persons") or []
        is_llp = c.get("client_type") in ("llp", "LLP")
        new_docs.append({
            "id": _uid(),
            "client_id": cid,
            "company_name": c.get("company_name") or "Unnamed Company",
            "cin": c.get("cin"),
            "category": CLIENT_CATEGORY_MAP.get(c.get("client_type"), "private"),
            "is_small_company": None,
            "listed": False,
            "roc_office": None,
            "pan": c.get("pan"),
            "date_of_incorporation": (str(c.get("date_of_incorporation"))[:10] if c.get("date_of_incorporation") else None),
            "registered_office_address": c.get("address"),
            "authorized_capital": 0,
            "paid_up_capital": 0,
            "last_year_turnover": 0,
            "financial_year_end": "31-03",
            "last_agm_date": None,
            "last_board_meeting_date": None,
            "directors": [] if is_llp else contact_people,
            "designated_partners": contact_people if is_llp else [],
            "partners": contact_people if is_llp else [],
            "shareholders": [],
            "master_data": {},
            "mgt_shareholder_data": {},
            "annual_return_data": {},
            "audit_report_data": {},
            "board_report_data": {},
            "share_transfers": [],
            "share_certificates": [],
            "roc_form_uploads": [],
            "auditor": None,
            "notes": None,
            "created_at": now,
            "updated_at": now,
            "created_by": "Auto-synced from Clients",
        })
    if new_docs:
        await COMPANIES.insert_many(new_docs)
    return len(new_docs)


@router.get("/companies")
async def list_companies(
    q: Optional[str] = Query(None),
    current_user: User = Depends(VIEW),
):
    await _sync_companies_from_clients()
    query: Dict[str, Any] = {}
    if q:
        query["company_name"] = {"$regex": re.escape(q), "$options": "i"}
    cursor = COMPANIES.find(query).sort("company_name", 1)
    items = [c async for c in cursor]

    # Repair legacy ROC masters in-place. In particular, old records may have
    # category='private' although their linked Client/name is an LLP. Doing
    # this here makes the API response and the stored value agree, so the UI
    # filter/count cannot show LLP (0) while LLP records are visible.
    client_ids = [c.get("client_id") for c in items if c.get("client_id")]
    clients_by_id = {}
    if client_ids:
        client_cursor = CLIENTS.find({"id": {"$in": client_ids}}, {"_id": 0})
        clients_by_id = {c.get("id"): c async for c in client_cursor if c.get("id")}

    for c in items:
        client = clients_by_id.get(c.get("client_id"))
        normalized = _normalize_roc_category(c, client)
        if c.get("category") != normalized:
            await COMPANIES.update_one(
                {"id": c.get("id")},
                {"$set": {"category": normalized, "updated_at": _now()}},
            )
            c["category"] = normalized
        c.pop("_id", None)
    return items


@router.post("/companies/sync-from-clients")
async def sync_from_clients_endpoint(current_user: User = Depends(CREATE)):
    """Manual re-sync trigger (e.g. a 'Sync from Clients' button) -- same
    logic as the automatic sync on GET /companies, exposed separately so
    the UI can show how many were newly added after a bulk client import."""
    created = await _sync_companies_from_clients()
    return {"created": created}


@router.get("/companies/{company_id}")
async def get_company(company_id: str, current_user: User = Depends(VIEW)):
    c = await COMPANIES.find_one({"id": company_id})
    if not c:
        raise HTTPException(404, "Company not found")
    # Client contact persons are the source for the initial ROC people list.
    # Preserve any richer ROC edits, but backfill empty registers from Client.
    if c.get("client_id"):
        client = await CLIENTS.find_one({"id": c["client_id"]}, {"_id": 0, "contact_persons": 1})
        contacts = (client or {}).get("contact_persons") or []
        if contacts:
            if c.get("category") == "llp":
                c.setdefault("designated_partners", contacts)
                c.setdefault("partners", contacts)
                if not c.get("designated_partners"):
                    c["designated_partners"] = contacts
                if not c.get("partners"):
                    c["partners"] = contacts
            elif not c.get("directors"):
                c["directors"] = contacts
    c.pop("_id", None)
    return c


@router.post("/companies")
async def create_company(payload: RocCompanyIn, current_user: User = Depends(CREATE)):
    now = _now()
    doc = payload.model_dump()
    doc["id"] = _uid()
    doc["created_at"] = now
    doc["updated_at"] = now
    doc["created_by"] = _who(current_user)
    await COMPANIES.insert_one(doc)
    await _sync_company_to_client(doc)
    doc.pop("_id", None)
    return doc


@router.put("/companies/{company_id}")
async def update_company(company_id: str, payload: RocCompanyIn, current_user: User = Depends(EDIT)):
    existing = await COMPANIES.find_one({"id": company_id})
    if not existing:
        raise HTTPException(404, "Company not found")
    doc = payload.model_dump()
    doc["updated_at"] = _now()
    await COMPANIES.update_one({"id": company_id}, {"$set": doc})
    await _sync_company_to_client({**existing, **doc})
    merged = {**existing, **doc}
    merged.pop("_id", None)
    return merged


async def _sync_company_to_client(company: Dict[str, Any]) -> None:
    """Keep the linked Client as the searchable source of truth as well.

    ROC has richer role-specific records than Clients, so the compact client
    contact_persons list is populated from directors or LLP partners while
    the complete source data is retained in dedicated fields.
    """
    client_id = company.get("client_id")
    if not client_id:
        return
    is_llp = company.get("category") == "llp"
    people = (company.get("designated_partners") or []) + (company.get("partners") or []) if is_llp else (company.get("directors") or [])
    contacts = []
    for person in people:
        if not person or not person.get("name"):
            continue
        contacts.append({
            "name": person.get("name"),
            "designation": person.get("designation") or ("Designated Partner" if is_llp else "Director"),
            "din": person.get("din"),
            "pan": person.get("pan"),
            "email": person.get("email"),
            "phone": person.get("phone"),
        })
    update = {
        "company_name": company.get("company_name"),
        "pan": company.get("pan"),
        "date_of_incorporation": company.get("date_of_incorporation"),
        "address": company.get("registered_office_address"),
        "contact_persons": contacts,
        "roc_directors": company.get("directors") or [],
        "roc_designated_partners": company.get("designated_partners") or [],
        "roc_partners": company.get("partners") or [],
        "roc_shareholders": company.get("shareholders") or [],
        "roc_master_data": company.get("master_data") or {},
        "roc_mgt_shareholder_data": company.get("mgt_shareholder_data") or {},
        "roc_form_uploads": company.get("roc_form_uploads") or [],
    }
    if is_llp:
        update["llpin"] = company.get("cin")
    else:
        update["cin"] = company.get("cin")
    await CLIENTS.update_one({"id": client_id}, {"$set": update})


@router.delete("/companies/{company_id}")
async def delete_company(company_id: str, current_user: User = Depends(DELETE)):
    res = await COMPANIES.delete_one({"id": company_id})
    if not res.deleted_count:
        raise HTTPException(404, "Company not found")
    await DOCS_LOG.delete_many({"company_id": company_id})
    return {"deleted": True}


# ─────────────────────────────────────────────────────────────────────────
# AOC-4 / MGT-7 / MGT-7A / AOC-2 / Board & Auditor report UPLOAD → EXTRACTION
# ─────────────────────────────────────────────────────────────────────────
# Every MCA acknowledgement PDF is a different form with a different
# purpose, and each field extracted here is only ever pulled from the form
# that actually carries that data on the MCA record:
#
#   Company Master fields  (CIN, name, address, authorised capital,
#                            AGM date, board meeting date)
#                             — read from whichever recognised ROC form
#                               states them (present on most of AOC-4,
#                               AOC-2, MGT-7/7A, Board's/Auditor's Report)
#   Directors / KMP register — ONLY from MGT-7 / MGT-7A (the Annual Return
#                               is the filing that carries the statutory
#                               Director/Signatory register; AOC-4, AOC-2,
#                               and the Board's/Auditor's Report extracts
#                               do not, and must never populate it)
#   Shareholders register    — ONLY from MGT-7 / MGT-7A, same reasoning
#   Financial data            — ONLY from AOC-4 (Balance Sheet / P&L /
#                               Net Worth figures; AOC-2 and the Board's/
#                               Auditor's Report extracts carry no
#                               standardised financial figures worth
#                               trusting for this)
#   Statutory Auditor details — ONLY from AOC-4 (Auditor Details block)
#
# This routing is enforced by _identify_roc_form_type() below and the
# per-form-type dispatch in upload_master_data(): a field is never applied
# from a form that isn't its statutory source, so one wrong/unusual PDF in
# a batch can no longer pollute Directors & Shareholders (this replaced an
# earlier version that scanned the full text of every uploaded form for
# anything DIN/PAN-shaped, which produced garbage director rows out of
# AOC-4/AOC-2/Board & Auditor report boilerplate — see CHANGELOG below).
#
# This is a heuristic text-scrape (same approach as backend/compliance.py's
# parse_compliance_dates), not an XBRL parser — every applied field is
# still shown to the user in the "Extracted fields" preview and can be
# corrected on the Company Master / Directors & Shareholders tabs.

ROC_FORM_ALLOWED_EXT = (".pdf", ".xlsx", ".xlsm", ".csv")
ROC_FORM_RECOGNIZED = (
    # order matters: more specific labels (mgt-7a) must be checked before
    # their substrings (mgt-7)
    ("mgt-7a", r"mgt[- ]?7a"),
    ("mgt-7", r"mgt[- ]?7\b"),
    ("mgt-7a-attachment", r"details? of (?:share|debenture)|shareholder.*(?:xls|xlsx|xlsm)"),
    ("aoc-4", r"\baoc[- ]?4\b"),
    ("aoc-2", r"\baoc[- ]?2\b"),
    ("board-report", r"extract of board.?s report|board.?s report"),
    ("auditor-report", r"extract of auditor.?s report|auditor.?s report"),
    ("dir-12", r"dir[- ]?12"),
    ("adt-1", r"adt[- ]?1"),
    ("inc-22", r"inc[- ]?22"),
    ("pas-3", r"pas[- ]?3"),
    ("mgt-14", r"mgt[- ]?14"),
    ("dpt-3", r"dpt[- ]?3"),
)

# Forms whose MCA-prescribed content includes the statutory Director/
# Signatory register and the shareholder/member register.
DIRECTOR_SHAREHOLDER_SOURCE_TYPES = {"mgt-7", "mgt-7a", "mgt-7a-attachment"}
# Form whose MCA-prescribed content includes the audited Balance Sheet,
# Statement of Profit & Loss and Auditor Details block.
FINANCIAL_SOURCE_TYPE = "aoc-4"

# Keys populated on company.financial_data by an AOC-4 upload — kept as a
# named set so the frontend/UI and this parser stay in sync.
FINANCIAL_DATA_FIELDS = (
    "period_from", "period_to", "total_income", "total_expenses",
    "profit_before_tax", "profit_after_tax", "net_worth", "share_capital",
    "reserves_and_surplus", "balance_sheet_total", "turnover",
)


def _identify_roc_form_type(filename: str, text: str) -> str:
    """Best-effort filing-type label — checked against both the filename
    and the extracted text, since MCA acknowledgement PDFs are sometimes
    downloaded/renamed generically. This label is load-bearing: it decides
    which fields (if any) a given upload is allowed to touch, not just an
    audit-trail cosmetic."""
    hay = f"{filename}\n{text[:2000]}".lower()
    if (filename or "").lower().endswith((".xlsx", ".xlsm", ".csv")) and re.search(
        r"shareholder|debenture holder|security held|mgt[- ]?7", hay, re.I
    ):
        return "mgt-7a-attachment"
    for label, pattern in ROC_FORM_RECOGNIZED:
        if re.search(pattern, hay):
            return label
    return "roc-form"


def _extract_text_from_upload(filename: str, raw: bytes) -> str:
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for page in pdf.pages[:20]:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        if name.endswith((".xlsx", ".xlsm", ".xls")):
            import openpyxl
            wb = openpyxl.load_workbook(
                io.BytesIO(raw),
                data_only=True,
                read_only=True,
                keep_vba=name.endswith(".xlsm"),
            )
            lines = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    lines.append(" ".join(str(c) for c in row if c is not None))
            wb.close()
            return "\n".join(lines)
        if name.endswith(".csv"):
            return raw.decode("utf-8", errors="ignore")
        return raw.decode("utf-8", errors="ignore")
    except Exception as e:  # pragma: no cover
        logger.warning("roc_sphere: text extraction failed for %s: %s", filename, e)
        return ""


def _parse_mgt_shareholder_workbook(raw: bytes) -> List[Dict[str, Any]]:
    """Read the separate MCA MGT-7/MGT-7A shareholder attachment.

    MCA supplies this attachment as a macro-enabled workbook.  The
    shareholder table can move between sheets and the sheet may contain
    instructions above it, so locate the header row by its labels instead of
    relying on a fixed sheet/cell range.  VBA is never executed.
    """
    rows: List[Dict[str, Any]] = []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(
            io.BytesIO(raw), read_only=True, data_only=True, keep_vba=True
        )
        for ws in wb.worksheets:
            header_row = None
            headers: Dict[str, int] = {}
            for row in ws.iter_rows(values_only=True):
                values = [str(v).strip() if v is not None else "" for v in row]
                lowered = [v.lower() for v in values]
                if any("name of shareholder" in v for v in lowered):
                    header_row = values
                    headers = {v.lower(): i for i, v in enumerate(values) if v}
                    break
            if not header_row:
                continue

            def cell(values: List[Any], label: str) -> Any:
                idx = next((i for h, i in headers.items() if label in h), None)
                return values[idx] if idx is not None and idx < len(values) else None

            for row in ws.iter_rows(values_only=True):
                values = list(row)
                name = cell(values, "name of shareholder")
                if not name or not str(name).strip():
                    continue
                name = re.sub(r"\s+", " ", str(name).strip())
                if name.lower().startswith("name of shareholder"):
                    continue
                share_count = _num(cell(values, "number of security"))
                face_value = _num(cell(values, "nominal value per security"))
                total_value = _num(cell(values, "total amount of securities"))
                rows.append({
                    "name": name,
                    "holder_type": cell(values, "type of shareholder"),
                    "category": cell(values, "category of shareholder"),
                    "details": cell(values, "details of shareholder"),
                    "class_of_shares": cell(values, "class of security") or cell(values, "type of security") or "Equity",
                    "folio_no": str(cell(values, "folio number") or "").strip() or None,
                    "nationality": cell(values, "nationality"),
                    "gender": cell(values, "gender"),
                    "identifier_type": cell(values, "type of identifier"),
                    "pan": cell(values, "identification no"),
                    "occupation": cell(values, "occupation"),
                    "shares_held": share_count,
                    "face_value": face_value or 10,
                    "total_value": total_value or share_count * (face_value or 10),
                    "percentage": None,
                })
        wb.close()
    except Exception as e:  # pragma: no cover
        logger.warning("roc_sphere: shareholder workbook parse failed: %s", e)
        return []

    total = sum(_num(row.get("shares_held")) for row in rows)
    if total:
        for row in rows:
            row["percentage"] = round(_num(row.get("shares_held")) / total * 100, 2)
    return rows


# ── director / shareholder register (MGT-7 / MGT-7A only) ─────────────────

def _parse_people(text: str) -> Dict[str, List[Dict[str, Any]]]:
    """Read the tabular director/signatory block in MCA MGT-7/MGT-7A text.

    A director "id" line must be a real DIN (8 digits) or a PAN
    (5 letters + 4 digits + 1 letter) — the previous version accepted any
    6-20 character alphanumeric line, which matched pincodes, page
    footers, dropdown option text and other boilerplate found on AOC-4/
    AOC-2/Board & Auditor report pages and produced garbage director rows.
    Callers must only invoke this for form_type in
    DIRECTOR_SHAREHOLDER_SOURCE_TYPES.
    """
    people: List[Dict[str, Any]] = []
    block_match = re.search(
        r"(?:Directors?/Signatory Details|Director/SignatoryDetails)(.*?)(?:Charges|No Records found|$)",
        text, re.I | re.S,
    )
    block = block_match.group(1) if block_match else text
    lines = [re.sub(r"\s+", " ", x).strip(" -:\t") for x in block.splitlines()]
    din_or_pan_re = re.compile(r"^\d{8}$|^[A-Z]{5}\d{4}[A-Z]$", re.I)
    for i, line in enumerate(lines):
        if not din_or_pan_re.fullmatch(line.replace(" ", "")):
            continue
        candidates = [x for x in lines[i + 1:i + 5] if x]
        name = next((x for x in candidates if re.search(r"[A-Za-z]", x) and len(x) >= 5
                     and not re.fullmatch(r"(Director|Promoter|Signatory|Active|Yes|No)", x, re.I)), None)
        if not name or any(p.get("din") == line for p in people):
            continue
        designation = next((x for x in candidates if re.search(
            r"director|partner|manager|secretary", x, re.I)), "Director")
        people.append({
            "name": name,
            "din": line,
            "designation": designation.title(),
            "date_of_appointment": next((x for x in candidates if re.fullmatch(
                r"\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}", x)), None),
        })
    return {"people": people}


def _parse_mgt_shareholders(text: str) -> List[Dict[str, Any]]:
    """Parse shareholder rows when the attached MGT-7/MGT-7A includes them.

    The filed MGT-7A often references a separate XLSM attachment and
    therefore contains only the shareholder count. In that case returning
    [] is correct and preserves the existing shareholder register instead
    of fabricating names. Callers must only invoke this for form_type in
    DIRECTOR_SHAREHOLDER_SOURCE_TYPES.
    """
    rows = []
    for line in text.splitlines():
        line = re.sub(r"\s+", " ", line).strip()
        m = re.match(r"^\d+\s+([A-Za-z][A-Za-z .,&'-]{3,})\s+([A-Z]{5}\d{4}[A-Z])?\s*(\d[\d,]*)\s*$", line)
        if m:
            rows.append({"name": m.group(1).strip(), "pan": m.group(2), "shares_held": _num(m.group(3).replace(",", "")), "class_of_shares": "Equity"})
    return rows


def _first_amount(text: str, patterns: List[str]) -> Optional[float]:
    for pattern in patterns:
        match = re.search(pattern, text, re.I | re.M)
        if match:
            return _num(match.group(1).replace(",", ""))
    return None


def parse_mgt_annual_return(text: str) -> Dict[str, Any]:
    """Extract structured annual-return facts from MGT-7/MGT-7A.

    The separate XLSM attachment is the source for member rows; this parser
    captures the form-level facts that are useful for pre-filling filings and
    compliance decisions.
    """
    out: Dict[str, Any] = {}
    turnover = _first_amount(text, [
        r"\*?\s*Turnover\s+(-?[\d,]+(?:\.\d+)?)",
        r"Turnover\s*\(in Rs\.\)\s+(-?[\d,]+(?:\.\d+)?)",
    ])
    net_worth = _first_amount(text, [
        r"\*?\s*Net worth of the Company\s+(-?[\d,]+(?:\.\d+)?)",
        r"Net worth of the company\s+(-?[\d,]+(?:\.\d+)?)",
    ])
    paid_up = _first_amount(text, [
        r"Paid Up capital\s+(-?[\d,]+(?:\.\d+)?)",
        r"Paid-up capital\s+(-?[\d,]+(?:\.\d+)?)",
    ])
    if turnover is not None:
        out["turnover"] = turnover
    if net_worth is not None:
        out["net_worth"] = net_worth
    if paid_up is not None:
        out["paid_up_capital"] = paid_up

    count = _first_amount(text, [
        r"Number of shareholder/ debenture holder\s+([\d,]+)",
        r"Total number of shareholders \(Promoters \+ Other than promoters\)\s+([\d,]+(?:\.\d+)?)",
    ])
    if count is not None:
        out["shareholder_count"] = int(count)

    meeting_matches = re.findall(r"\*?Number of meetings held\s+([\d,]+)", text, re.I)
    meeting_count = max((_num(v) for v in meeting_matches), default=None)
    if meeting_count is not None:
        out["board_meetings_held"] = int(meeting_count)

    activity = re.search(
        r"\d+\s+([A-Z])\s+\d+\s+(.+?)\s+(\d+(?:\.\d+)?)\s*$",
        text,
        re.I | re.M,
    )
    if activity:
        out["principal_business_activity"] = {
            "main_activity_group_code": activity.group(1),
            "description": re.sub(r"\s+", " ", activity.group(2)).strip(),
            "turnover_percentage": _num(activity.group(3)),
        }

    out["filing_source"] = "MGT-7A / MGT-7"
    return out


def parse_auditor_report(text: str) -> Dict[str, Any]:
    """Capture audit-report facts for review and Board's Report drafting."""
    out: Dict[str, Any] = {}
    qualified = _first_amount(text, [
        r"Number of qualifications, reservation or adverse remark or disclaimer\s+([\d,]+)",
    ])
    if qualified is not None:
        out["qualifications_count"] = int(qualified)
    out["caro_applicable"] = bool(re.search(
        r"whether companies auditors report order.*?applicable.*?\bYes\b",
        text,
        re.I | re.S,
    ))
    opinion = _extract_section(text, "Opinion of the auditor", "Basis of Opinion")
    if opinion:
        out["opinion"] = opinion
    basis = _extract_section(text, "Basis of Opinion", "Emphasis of matter")
    if basis:
        out["basis_of_opinion"] = basis
    other = _extract_section(text, "State other matters as per Rule 11", "State any other matters")
    if other:
        out["rule_11_other_matters"] = other
    controls = _extract_section(text, "Reporting on the Internal Financial Controls", "Attachments")
    if controls:
        out["internal_financial_controls"] = controls
    return out


def _extract_section(text: str, start_label: str, end_label: str) -> Optional[str]:
    match = re.search(
        rf"{re.escape(start_label)}(.*?){re.escape(end_label)}",
        text,
        re.I | re.S,
    )
    if not match:
        return None
    value = re.sub(r"\s+", " ", match.group(1)).strip(" :-")
    return value[:4000] if value else None


def parse_board_report(text: str) -> Dict[str, Any]:
    """Capture Board's Report disclosures as structured filing context."""
    out: Dict[str, Any] = {}
    meetings = _first_amount(text, [r"Number of meetings held\s+([\d,]+)"])
    if meetings is not None:
        out["board_meetings_held"] = int(meetings)
    for key, label, end in (
        ("state_of_affairs", "Description of state of company’s affairs", "Disclosure relating to amounts"),
        ("reserves_recommendation", "Disclosure relating to amounts if any which is proposed to carry to any reserves", "Disclosures relating to amount recommended"),
        ("dividend_recommendation", "Disclosures relating to amount recommended to be paid as dividend", "Details of material changes"),
        ("material_changes", "Details of material changes and commitment occurred during period", "Disclosure of statement on development"),
        ("risk_management", "Disclosure of statement on development and implementation of risk management policy", "CSR details"),
        ("financial_summary", "Disclosure of financial summary or highlights", "Disclosure of change in nature of business"),
        ("business_change", "Disclosure of change in nature of business", "Details of directors or key managerial personnel"),
    ):
        value = _extract_section(text, label, end)
        if value:
            out[key] = value
    out["csr_applicable"] = bool(re.search(
        r"whether CSR is applicable as per section 135\s+Yes",
        text,
        re.I,
    ))
    return out


def _parse_mgt_annual_return_directors(text: str) -> List[Dict[str, Any]]:
    """Parse the DIN/name/attendance table embedded in MGT-7/MGT-7A."""
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    people: List[Dict[str, Any]] = []
    din_re = re.compile(r"^\d{8}$")
    for i, line in enumerate(lines):
        if not din_re.fullmatch(line):
            continue
        window = [x for x in lines[i + 1:i + 12] if x]
        name_parts: List[str] = []
        for candidate in window:
            if re.fullmatch(r"\d+(?:\.\d+)?", candidate):
                break
            if re.search(
                r"number of|meeting|attendance|whether|director|board|yes|no|name of|date of|designation|net worth|turnover|share holding|capital|pattern",
                candidate,
                re.I,
            ):
                break
            if re.search(r"[A-Za-z]", candidate):
                name_parts.append(candidate)
        name = " ".join(name_parts[:3]) if name_parts else None
        if not name or any(p.get("din") == line for p in people):
            continue
        people.append({
            "name": name,
            "din": line,
            "designation": "Director",
            "date_of_appointment": None,
            "attendance": {
                "meetings_entitled": _first_amount(
                    " ".join(window),
                    [r"(\d+)\s+\d+\s+\d+(?:\.\d+)?"],
                ),
            },
        })
    return people


def _parse_directors_for_form(form_type: str, filename: str, raw: bytes, text: str) -> List[Dict[str, Any]]:
    """Directors register — gated to MGT-7/MGT-7A. Prefers the bordered
    table-grid reader (far more reliable than a text-line scan) and only
    falls back to the text-regex scan above when no bordered table is
    found, e.g. an MGT-7A that was flattened/scanned oddly."""
    if form_type not in DIRECTOR_SHAREHOLDER_SOURCE_TYPES:
        return []
    if (filename or "").lower().endswith((".xlsx", ".xlsm", ".csv")):
        return []
    people = _parse_mgt_annual_return_directors(text) if form_type in {"mgt-7", "mgt-7a"} else []
    if not people:
        people = _parse_master_director_tables(raw) if (filename or "").lower().endswith(".pdf") else []
    if not people:
        people = _parse_people(text)["people"]
    return people


def _parse_shareholders_for_form(form_type: str, filename: str, raw: bytes, text: str) -> List[Dict[str, Any]]:
    """Shareholder register — gated to MGT-7/MGT-7A."""
    if form_type not in DIRECTOR_SHAREHOLDER_SOURCE_TYPES:
        return []
    if (filename or "").lower().endswith((".xlsx", ".xlsm")):
        return _parse_mgt_shareholder_workbook(raw)
    return _parse_mgt_shareholders(text)


# ── general Company Master fields (any recognised ROC form) ───────────────
# MCA e-form PDFs render each field as a label followed by its value —
# sometimes on the same line ("*Name of the Company PRODIGIST … LIMITED"
# on AOC-2/Board's/Auditor's Report), sometimes with the value in the box
# below the label ("*Corporate identity number (CIN)" / value on the next
# line, as on AOC-4). _extract_labeled_value() below tries the same-line
# remainder first and only falls back to scanning forward when that
# remainder doesn't validate — this is what keeps e.g. a bare label like
# "Address of the registered office of the company" (nothing left over
# after stripping the label text) from being read as its own value.

STOP_MARKER_RE = re.compile(r"^\(?[a-hj-vx-z]\)\s|^\(?[ivx]{1,4}\)\s", re.I)
_TRUNC_ENDING_WORDS = ("private", "pvt", "the", "of", "and", "&", "-")

_DATE_VALUE_RE = re.compile(r"^[0-3]?\d[/\-][01]?\d[/\-]\d{2,4}$")
_NUM_VALUE_RE = re.compile(r"^[\d,]+(\.\d+)?$")
_CIN_VALUE_RE = re.compile(r"^[A-Z0-9]{21}$")
_NAME_VALUE_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9 &.,'\-]{1,90}$")
_ANY_TEXT_VALUE_RE = re.compile(r"^.{4,150}$")
_ALNUM_ID_VALUE_RE = re.compile(r"^[A-Za-z0-9]{4,20}$")


def _find_label_line(lines: List[str], phrases: List[str], start: int = 0):
    """First line at/after `start` containing any phrase (longer phrases
    should be listed first so a more specific label wins over a shorter
    one that happens to be a substring of it)."""
    low = [p.lower() for p in phrases]
    for i in range(start, len(lines)):
        ll = lines[i].lower()
        for p in low:
            if p in ll:
                return i, p
    return -1, None


def _extract_labeled_value(lines: List[str], label_idx: int, phrase: str, validator, join_multiline: bool = False, max_lookahead: int = 6) -> Optional[str]:
    line = lines[label_idx]
    ll = line.lower()
    p_idx = ll.find(phrase)
    remainder = line[p_idx + len(phrase):].strip(" :()-\t*") if p_idx != -1 else ""
    if remainder and validator(remainder):
        collected = [remainder]
        # a same-line value that trails off mid-word ("… PRIVATE") is
        # completed from the immediate next line ("LIMITED"); a complete
        # value is left alone so the next label's text is never pulled in.
        if join_multiline and remainder.strip().lower().split(" ")[-1] in _TRUNC_ENDING_WORDS:
            for j in range(label_idx + 1, min(label_idx + 3, len(lines))):
                cand = lines[j].strip()
                if not cand or STOP_MARKER_RE.match(cand):
                    break
                if validator(cand):
                    collected.append(cand)
                break
        return " ".join(collected)
    collected: List[str] = []
    for j in range(label_idx + 1, min(label_idx + 1 + max_lookahead, len(lines))):
        cand = lines[j].strip()
        if not cand:
            continue
        is_stop = bool(STOP_MARKER_RE.match(cand))
        ok = bool(validator(cand))
        if is_stop and (collected or not ok):
            break
        if ok:
            collected.append(cand)
            if not join_multiline:
                break
        elif collected:
            break
    return " ".join(collected) if collected else None


# (field, label phrases — longest/most specific first, validator, join multi-line?, lookahead)
ROC_GENERAL_FIELD_SPECS = [
    ("cin", ["corporate identity number", "llpin"], lambda s: bool(_CIN_VALUE_RE.match(s)), False, 3),
    ("company_name", ["name of the company", "name of company", "llp name"], lambda s: bool(_NAME_VALUE_RE.match(s)), True, 3),
    ("registered_office_address",
     ["address of the registered office of the company", "address of the registered office", "registered office address"],
     lambda s: bool(_ANY_TEXT_VALUE_RE.match(s)), True, 6),
    ("authorized_capital", ["authorised capital of the company", "authorized capital of the company"], lambda s: bool(_NUM_VALUE_RE.match(s)), False, 3),
    ("last_agm_date", ["date of agm"], lambda s: bool(_DATE_VALUE_RE.match(s)), False, 3),
    ("last_board_meeting_date", ["date of board of directors", "date of board meeting"], lambda s: bool(_DATE_VALUE_RE.match(s)), False, 3),
]


def parse_roc_general_fields(text: str) -> Dict[str, Any]:
    """Company Master fields common to (most) ROC forms — safe to apply
    regardless of which recognised form type the upload turned out to be,
    since none of these are director/shareholder/financial data."""
    lines = [re.sub(r"\s+", " ", x).strip() for x in text.splitlines()]
    extracted: Dict[str, Any] = {}
    for field, phrases, validator, join_multi, lookahead in ROC_GENERAL_FIELD_SPECS:
        idx, phrase = _find_label_line(lines, phrases)
        if idx == -1:
            continue
        val = _extract_labeled_value(lines, idx, phrase, validator, join_multi, lookahead)
        if val:
            extracted[field] = _num(val) if field == "authorized_capital" else val
    return extracted


# ── financial data + auditor (AOC-4 only) ──────────────────────────────────
# AOC-4's Balance Sheet / P&L segments print each line item and its
# current/previous-period figures on one line ("(IX) Profit before tax
# (VII-VIII) -94343.00 -132689.00"), so these are simple same-line regexes
# rather than the label/value scan used for the general fields above.

AOC4_FINANCIAL_PATTERNS = {
    "total_income": r"\(III\)\s*Total Income.*?(-?[\d,]+\.\d{2})",
    "total_expenses": r"Total expenses\s+(-?[\d,]+\.\d{2})",
    "profit_before_tax": r"\(IX\)\s*Profit before tax.*?(-?[\d,]+\.\d{2})",
    "profit_after_tax": r"\(XV\)\s*Profit\s*/\(Loss\).*?(-?[\d,]+\.\d{2})",
    "net_worth": r"Net Worth of the company\s+(-?[\d,]+(?:\.\d+)?)",
    "share_capital": r"\(a\)\s*Share capital\s+(-?[\d,]+)",
    "reserves_and_surplus": r"\(b\)\s*Reserves and surplus\s+(-?[\d,]+)",
    "turnover": r"Domestic turnover\s+(?:\n\s*)?(?:\(i\)\s*Sale of goods manufactured\s+)?(-?[\d,]+(?:\.\d+)?)",
}


def parse_aoc4_financials(text: str) -> Dict[str, Any]:
    """Balance Sheet / P&L key figures — only meaningful for form_type ==
    'aoc-4'; callers must gate on that before calling this."""
    out: Dict[str, Any] = {}
    for field, pattern in AOC4_FINANCIAL_PATTERNS.items():
        m = re.search(pattern, text)
        if m:
            out[field] = _num(m.group(1))
    lines = [re.sub(r"\s+", " ", x).strip() for x in text.splitlines()]

    def nearby_amount(labels: List[str], lookahead: int = 12, prefer_last: bool = False) -> Optional[float]:
        for i, line in enumerate(lines):
            if not any(label.lower() in line.lower() for label in labels):
                continue
            values = []
            for candidate in lines[i + 1:i + 1 + lookahead]:
                if re.fullmatch(r"-?[\d,]+(?:\.\d+)?", candidate):
                    values.append(_num(candidate))
            if values:
                if prefer_last:
                    # Printed forms often put a two-digit row number before
                    # the actual amount and the next row's zero after it.
                    substantive = [value for value in values if abs(value) > 100]
                    return (substantive[-1] if substantive else values[-1])
                return values[0]
        return None

    # MCA's printable AOC-4 places some labels, row numbers and values on
    # separate lines. Prefer the label-aware value over a regex hit that can
    # accidentally capture a row index (for example, Net Worth's "42").
    net_worth = nearby_amount(["Net Worth of the company"], prefer_last=True)
    if net_worth is not None:
        out["net_worth"] = net_worth
    pbt = nearby_amount(["Profit before exceptional", "Profit before tax"])
    if pbt is not None:
        out["profit_before_tax"] = pbt
    pat = nearby_amount(["Profit/(Loss) for the period from continuing operations", "Profit /(Loss) (XI+XIV)"])
    if pat is not None:
        out["profit_after_tax"] = pat
    if "total_income" not in out and out.get("total_expenses") is not None and out.get("profit_before_tax") is not None:
        out["total_income"] = out["total_expenses"] + out["profit_before_tax"]
    if "turnover" not in out:
        # Fallback for PDFs where the "Domestic turnover" label is on its
        # own line and the first operating-revenue row follows later.
        m = re.search(
            r"Domestic turnover.*?\(i\)\s*Sale of goods manufactured\s+(-?[\d,]+(?:\.\d+)?)",
            text,
            re.I | re.S,
        )
        if m:
            out["turnover"] = _num(m.group(1))
    # the Balance Sheet's grand total is printed twice (Equity & Liabilities
    # total, then Assets total) with identical figures — take the first.
    m = re.search(r"^Total\s+(-?[\d,]+\.\d{2})\s+(-?[\d,]+\.\d{2})\s*$", text, re.M)
    if m:
        out["balance_sheet_total"] = _num(m.group(1))
    idx, _ = _find_label_line(lines, ["financial year to which financial statements relates"])
    if idx != -1:
        idx_from, p1 = _find_label_line(lines, ["from (dd/mm/yyyy)"], idx)
        if idx_from != -1:
            v = _extract_labeled_value(lines, idx_from, p1, lambda s: bool(_DATE_VALUE_RE.match(s)), False, 2)
            if v:
                out["period_from"] = v
            idx_to, p2 = _find_label_line(lines, ["to (dd/mm/yyyy)"], idx_from + 1)
            if idx_to != -1:
                v = _extract_labeled_value(lines, idx_to, p2, lambda s: bool(_DATE_VALUE_RE.match(s)), False, 2)
                if v:
                    out["period_to"] = v
    return out


def parse_aoc4_auditor(text: str) -> Dict[str, Any]:
    """Statutory Auditor name + firm registration/membership number —
    only meaningful for form_type == 'aoc-4'; callers must gate on that."""
    lines = [re.sub(r"\s+", " ", x).strip() for x in text.splitlines()]
    out: Dict[str, Any] = {}
    idx, p = _find_label_line(lines, ["name of the auditor or auditor's firm", "name of the auditor"])
    if idx != -1:
        v = _extract_labeled_value(lines, idx, p, lambda s: bool(_NAME_VALUE_RE.match(s)) and "address" not in s.lower(), True, 2)
        if v:
            out["name"] = v
    idx, p = _find_label_line(lines, ["auditor's firm's registration number", "membership number of auditor"])
    if idx != -1:
        v = _extract_labeled_value(lines, idx, p, lambda s: bool(_ALNUM_ID_VALUE_RE.match(s)), False, 2)
        if v:
            out["firm_reg_no"] = v
    return out


@router.post("/companies/{company_id}/upload-master-data")
async def upload_master_data(
    company_id: str,
    files: List[UploadFile] = File(...),
    apply: bool = Form(False),
    source_type: str = Form("roc"),
    current_user: User = Depends(CREATE),
):
    """Upload ROC Forms — AOC-4, AOC-2, MGT-7, MGT-7A, Board's/Auditor's
    Report extracts, DIR-12, ADT-1, INC-22, PAS-3, MGT-14, DPT-3
    acknowledgement PDFs — and best-effort extract filing data to prefill
    the company master. Each field is only ever taken from its statutory
    source form (see the block comment above this section): Company
    Master fields from any recognised form, the Director/Signatory
    register and Shareholder register only from MGT-7/MGT-7A, and
    financial data + Auditor details only from AOC-4.

    This is the *filing extraction* path and is intentionally separate
    from the Master Data importer below (`/companies/{id}/master-data/
    fetch`), which reads the MCA "Company/LLP Master Data" export
    instead. Robust to partial batch failures: one unreadable file no
    longer aborts the whole upload.
    """
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    if not files:
        raise HTTPException(400, "Choose at least one ROC form or MGT-7/MGT-7A attachment")

    results = []
    errors: List[str] = []
    conflicts: List[Dict[str, Any]] = []
    roc_extracted: Dict[str, Any] = {}
    total_size = 0

    for uploaded in files:
        filename = uploaded.filename or "uploaded-file"
        if not filename.lower().endswith(ROC_FORM_ALLOWED_EXT):
            errors.append(f"{filename}: skipped — use PDF, XLSX, XLSM or CSV for ROC filing data")
            continue
        raw = await uploaded.read()
        total_size += len(raw)
        if total_size > 50 * 1024 * 1024:
            errors.append(f"{filename}: skipped — combined upload exceeds the 50MB limit")
            continue
        if not raw:
            errors.append(f"{filename}: skipped — empty file")
            continue
        try:
            text = _extract_text_from_upload(filename, raw)
        except Exception as e:  # a single corrupt/scanned PDF must not kill the batch
            logger.warning("roc_sphere upload-roc-form: failed to parse %s: %s", filename, e)
            errors.append(f"{filename}: could not be read — it may be scanned/image-only or password-protected")
            continue
        if not text.strip():
            errors.append(f"{filename}: could not be read — it may be scanned/image-only or password-protected")
            continue

        form_type = _identify_roc_form_type(filename, text)
        extracted: Dict[str, Any] = parse_roc_general_fields(text)

        directors = _parse_directors_for_form(form_type, filename, raw, text)
        if directors:
            extracted["_directors"] = directors
        shareholders = _parse_shareholders_for_form(form_type, filename, raw, text)
        if shareholders:
            extracted["_shareholders"] = shareholders

        if form_type == FINANCIAL_SOURCE_TYPE:
            financials = parse_aoc4_financials(text)
            if financials:
                extracted["_financials"] = financials
            auditor = parse_aoc4_auditor(text)
            if auditor:
                extracted["_auditor"] = auditor
        if form_type in {"mgt-7", "mgt-7a"} and filename.lower().endswith(".pdf"):
            annual_return = parse_mgt_annual_return(text)
            if annual_return:
                extracted["_annual_return"] = annual_return
        if form_type == "auditor-report":
            audit_report = parse_auditor_report(text)
            if audit_report:
                extracted["_audit_report"] = audit_report
        if form_type == "board-report":
            board_report = parse_board_report(text)
            if board_report:
                extracted["_board_report"] = board_report

        extracted["_source_type"] = form_type
        fields_found = {k: v for k, v in extracted.items() if not k.startswith("_")}
        results.append({"filename": filename, "source_type": form_type, "extracted": fields_found,
                         "fields_found": len(fields_found),
                         "director_rows": len(extracted.get("_directors") or []),
                         "shareholder_rows": len(extracted.get("_shareholders") or [])})
        for key, value in extracted.items():
            if not key.startswith("_") and value:
                if key in roc_extracted and roc_extracted[key] != value:
                    conflicts.append({
                        "field": key,
                        "kept": value,
                        "previous": roc_extracted[key],
                        "source": filename,
                        "message": "Later upload value is shown as the candidate; review before Apply.",
                    })
                roc_extracted[key] = value
        if extracted.get("_directors"):
            roc_extracted["_directors"] = (roc_extracted.get("_directors") or []) + extracted["_directors"]
        if extracted.get("_shareholders"):
            roc_extracted["_shareholders"] = (roc_extracted.get("_shareholders") or []) + extracted["_shareholders"]
        if extracted.get("_financials"):
            roc_extracted["_financials"] = {**(roc_extracted.get("_financials") or {}), **extracted["_financials"]}
        if extracted.get("_auditor"):
            roc_extracted["_auditor"] = {**(roc_extracted.get("_auditor") or {}), **extracted["_auditor"]}
        if extracted.get("_annual_return"):
            roc_extracted["_annual_return"] = {**(roc_extracted.get("_annual_return") or {}), **extracted["_annual_return"]}
        if extracted.get("_audit_report"):
            roc_extracted["_audit_report"] = {**(roc_extracted.get("_audit_report") or {}), **extracted["_audit_report"]}
        if extracted.get("_board_report"):
            roc_extracted["_board_report"] = {**(roc_extracted.get("_board_report") or {}), **extracted["_board_report"]}

    if not any(k for k in roc_extracted if not k.startswith("_")) and not any(
        roc_extracted.get(k) for k in ("_directors", "_shareholders", "_financials", "_auditor", "_annual_return", "_audit_report", "_board_report")
    ):
        return {
            "extracted": {},
            "results": results,
            "applied": False,
            "errors": errors,
            "conflicts": conflicts,
            "message": errors[0] if errors and len(errors) == len(files) else
                       "Could not confidently extract fields from these forms — please enter details manually.",
        }

    if apply:
        clean = {k: v for k, v in roc_extracted.items() if not k.startswith("_") and v}
        if roc_extracted.get("_directors"):
            if company.get("category") == "llp":
                clean["designated_partners"] = roc_extracted["_directors"]
                clean["partners"] = roc_extracted["_directors"]
            else:
                clean["directors"] = roc_extracted["_directors"]
        if roc_extracted.get("_shareholders"):
            clean["shareholders"] = roc_extracted["_shareholders"]
        if roc_extracted.get("_financials"):
            clean["financial_data"] = {**(company.get("financial_data") or {}), **roc_extracted["_financials"]}
            if roc_extracted["_financials"].get("turnover") is not None:
                clean["last_year_turnover"] = roc_extracted["_financials"]["turnover"]
        if roc_extracted.get("_auditor"):
            existing_auditor = dict(company.get("auditor") or {})
            existing_auditor.update({k: v for k, v in roc_extracted["_auditor"].items() if v})
            clean["auditor"] = existing_auditor
        if roc_extracted.get("_annual_return"):
            annual_return = {**(company.get("annual_return_data") or {}), **roc_extracted["_annual_return"]}
            clean["annual_return_data"] = annual_return
            if annual_return.get("turnover") is not None:
                clean["last_year_turnover"] = annual_return["turnover"]
            if annual_return.get("paid_up_capital") is not None and not _num(company.get("paid_up_capital")):
                clean["paid_up_capital"] = annual_return["paid_up_capital"]
        if roc_extracted.get("_audit_report"):
            clean["audit_report_data"] = {**(company.get("audit_report_data") or {}), **roc_extracted["_audit_report"]}
        if roc_extracted.get("_board_report"):
            clean["board_report_data"] = {**(company.get("board_report_data") or {}), **roc_extracted["_board_report"]}
        clean["mgt_shareholder_data"] = {k: v for k, v in roc_extracted.items() if not k.startswith("_") and k not in ("directors", "shareholders", "financial_data", "auditor")}
        clean["roc_form_uploads"] = (company.get("roc_form_uploads") or []) + [
            {"filename": r["filename"], "form_type": r["source_type"], "uploaded_at": _now().isoformat()}
            for r in results
        ]
        clean["updated_at"] = _now()
        await COMPANIES.update_one({"id": company_id}, {"$set": clean})
        await _sync_company_to_client({**company, **clean})

    visible = {k: v for k, v in roc_extracted.items() if not k.startswith("_")}
    if roc_extracted.get("_directors"):
        visible["directors"] = roc_extracted["_directors"]
    if roc_extracted.get("_shareholders"):
        visible["shareholders"] = roc_extracted["_shareholders"]
    if roc_extracted.get("_financials"):
        visible["financial_data"] = roc_extracted["_financials"]
    if roc_extracted.get("_auditor"):
        visible["auditor"] = roc_extracted["_auditor"]
    if roc_extracted.get("_annual_return"):
        visible["annual_return_data"] = roc_extracted["_annual_return"]
    if roc_extracted.get("_audit_report"):
        visible["audit_report_data"] = roc_extracted["_audit_report"]
    if roc_extracted.get("_board_report"):
        visible["board_report_data"] = roc_extracted["_board_report"]
    return {"extracted": visible, "results": results, "applied": bool(apply), "errors": errors, "conflicts": conflicts}


# ─────────────────────────────────────────────────────────────────────────
# MASTER DATA IMPORTER — separate extraction path from Upload ROC Forms
# ─────────────────────────────────────────────────────────────────────────
# Reads the MCA "View Company/LLP Master Data" export (the printable PDF
# from the MCA portal's Master Data service, or an MCA master-data
# XLSX/CSV). Own text-reading strategy (PDF text-flow order, since the
# Master Data PDF is a two-column label/value layout that default PDF text
# extraction jumbles), own label dictionary, own field mapper. Deliberately
# does not share the ROC-form parsers/upload_master_data() above — a change
# to ROC filing parsing must never silently change Master Data parsing.
#
# Behaviour mirrors "Smart Import" on the Clients page: upload → fields are
# fetched and applied to the Company Master automatically, no manual
# per-field Apply step.

MASTER_DATA_ALLOWED_EXT = (".pdf", ".xlsx", ".xls", ".csv")

# (internal field name, label phrases as they appear on the MCA export)
MASTER_DATA_LABELS: List[tuple] = [
    ("cin", ("cin", "llpin")),
    ("company_name", ("company name", "name of the company", "llp name")),
    ("roc_name", ("roc name",)),
    ("registration_number", ("registration number",)),
    ("date_of_incorporation", ("date of incorporation",)),
    ("email", ("email id", "email")),
    ("registered_office_address", ("registered address", "registered office address")),
    ("listed", ("listed in stock exchange",)),
    ("company_category_raw", ("category of company",)),
    ("company_subcategory_raw", ("subcategory of the company",)),
    ("class_of_company_raw", ("class of company",)),
    ("active_compliance_raw", ("active compliance",)),
    ("authorized_capital", ("authorised capital", "authorized capital")),
    ("paid_up_capital", ("paid up capital", "paid-up capital")),
    ("last_agm_date", ("date of last agm",)),
    ("date_of_balance_sheet", ("date of balance sheet",)),
    ("company_status", ("company status",)),
    ("roc_office", ("roc (name and office)", "roc name and office")),
    ("rd_region", ("rd (name and region)", "rd name and region")),
    ("pan", ("pan", "permanent account number")),
    # stop-only markers — never stored, just tell the collector where a
    # multi-line value (e.g. the wrapped registered address) ends
    ("_stop_books_address", ("address at which the books of account",)),
    ("_stop_index_of_charges", ("index of charges",)),
    ("_stop_jurisdiction", ("jurisdiction",)),
    ("_stop_director_block", ("director/signatory details", "director / signatory details")),
]


def _match_master_label(line_lower: str):
    for field, phrases in MASTER_DATA_LABELS:
        for p in phrases:
            if line_lower == p or line_lower.startswith(p + " ") or line_lower.startswith(p + ":"):
                return field, p
    return None, None


def _read_master_data_text(filename: str, raw: bytes) -> str:
    """Own file→text reader for Master Data uploads. PDFs use text-flow
    ordering (reading order by position rather than raw stream order) so
    the label/value pairs on the MCA export don't get scrambled — the
    default extraction used for ROC forms gets this layout wrong."""
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            import pdfplumber
            parts = []
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for page in pdf.pages[:8]:
                    t = None
                    try:
                        t = page.extract_text(use_text_flow=True)
                    except Exception:
                        t = None
                    parts.append(t or page.extract_text() or "")
            return "\n".join(parts)
        if name.endswith((".xlsx", ".xls")):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True)
            lines = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        lines.append(" ".join(cells))
            return "\n".join(lines)
        if name.endswith(".csv"):
            return raw.decode("utf-8-sig", errors="ignore")
        return raw.decode("utf-8", errors="ignore")
    except Exception as e:  # pragma: no cover
        logger.warning("roc_sphere master-data: text extraction failed for %s: %s", filename, e)
        return ""


def _parse_master_director_tables(raw: bytes) -> List[Dict[str, Any]]:
    """Read the Director/Signatory table straight from the PDF's table
    grid (bordered on the MCA Master Data export) rather than scanning
    text lines — far more reliable than regex here since the table's
    column count varies (older exports omit the 'Category' column) but
    the last three columns are always Date of Appointment / Cessation
    Date / Signatory, letting position alone locate every field."""
    directors: List[Dict[str, Any]] = []
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                for table in (page.extract_tables() or []):
                    if not table or len(table) < 2:
                        continue
                    header_line = " ".join(str(c or "") for c in table[0]).lower()
                    if "din" not in header_line:
                        continue
                    for row in table[1:]:
                        if not row or len(row) < 4 or not any(row):
                            continue
                        name = str(row[2] or "").replace("\n", " ").strip()
                        din = str(row[1] or "").replace("\n", " ").strip()
                        if not name or not din or not re.search(r"[A-Za-z]", name) or not re.search(r"\d", din):
                            continue
                        appt_i, cess_i = len(row) - 3, len(row) - 2
                        directors.append({
                            "name": re.sub(r"\s+", " ", name),
                            "din": din,
                            "designation": (str(row[3] or "").replace("\n", " ").strip() or "Director"),
                            "date_of_appointment": str(row[appt_i] or "").strip() or None if appt_i > 0 else None,
                            "date_of_cessation": (str(row[cess_i] or "").strip() or None) if 0 <= cess_i < len(row) and str(row[cess_i] or "").strip() not in ("", "-") else None,
                        })
    except Exception as e:  # pragma: no cover
        logger.warning("roc_sphere master-data: director table extraction failed: %s", e)
    return directors


def _map_master_category(class_of_company_raw: Optional[str], category_raw: Optional[str], is_llp_hint: bool) -> Optional[str]:
    if is_llp_hint:
        return "llp"
    c = (class_of_company_raw or "").strip().lower()
    cat = (category_raw or "").strip().lower()
    if "one person" in cat:
        return "opc"
    if "section 8" in cat or "section-8" in cat or "u/s 8" in cat:
        return "section_8"
    if "public" in c:
        return "public"
    if "private" in c:
        return "private"
    return None


def extract_mca_master_data(filename: str, raw: bytes) -> Dict[str, Any]:
    """Parse an MCA Master Data export (PDF/XLSX/CSV) into Company Master
    fields. Separate parser/algorithm from the ROC-form parsers above —
    walks the document as an ordered label→value sequence (handling both
    same-line values like 'CIN U80900GJ...' and MCA's multi-line wrapped
    values like the registered address) rather than regex-scanning the
    whole blob."""
    text = _read_master_data_text(filename, raw)
    lines = [re.sub(r"\s+", " ", x).strip() for x in text.splitlines()]
    lines = [x for x in lines if x]

    extracted: Dict[str, Any] = {}
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        low = line.lower().rstrip(":")
        field, phrase = _match_master_label(low)
        if not field:
            i += 1
            continue
        if field.startswith("_stop_"):
            i += 1
            continue
        remainder = line[len(phrase):].strip(" :\t-")
        if remainder:
            value_parts, j = [remainder], i + 1
        else:
            value_parts, j = [], i + 1
            while j < n:
                nxt_field, _ = _match_master_label(lines[j].lower().rstrip(":"))
                if nxt_field:
                    break
                value_parts.append(lines[j])
                j += 1
        value = " ".join(p for p in value_parts if p).strip(" ,")
        if value and value not in ("-", "—", "NA", "N/A") and field not in extracted:
            extracted[field] = value
        i = j if j > i else i + 1

    for f in ("authorized_capital", "paid_up_capital"):
        if f in extracted:
            digits = re.sub(r"[^\d.]", "", str(extracted[f]))
            extracted[f] = _num(digits) if digits else 0.0

    if "listed" in extracted:
        extracted["listed"] = str(extracted["listed"]).strip().lower().startswith("y")

    is_llp_hint = bool(re.search(r"\bLLPIN\b", text, re.I))
    mapped_category = _map_master_category(
        extracted.pop("class_of_company_raw", None), extracted.get("company_category_raw"), is_llp_hint)
    if mapped_category:
        extracted["category"] = mapped_category

    # Table-grid extraction (PDF only) is far more reliable than the
    # regex line-scan, so prefer it and only fall back for XLSX/CSV or a
    # PDF whose director table has no visible borders.
    people = _parse_master_director_tables(raw) if (filename or "").lower().endswith(".pdf") else []
    if not people:
        people = _parse_people(text)["people"]
    if people:
        extracted["_directors"] = people

    extracted["_source_type"] = "master-data"
    extracted["_source_file"] = filename
    extracted["_chars_scanned"] = len(text)
    return extracted


@router.post("/companies/{company_id}/master-data/fetch")
async def fetch_master_data(
    company_id: str,
    files: List[UploadFile] = File(...),
    current_user: User = Depends(CREATE),
):
    """Master Data tab — upload the MCA 'View Company/LLP Master Data' PDF
    (or an MCA master-data XLSX/CSV) and the Company Master + Director/
    Signatory register are fetched and applied automatically, the same way
    Smart Import on the Clients page auto-fills a client from an uploaded
    document. Always applies (no manual per-field Apply step); returns
    exactly which fields changed so the UI can summarise it."""
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    if not files:
        raise HTTPException(400, "Choose at least one Master Data file")

    merged: Dict[str, Any] = {}
    directors: List[Dict[str, Any]] = []
    results, errors = [], []
    total_size = 0

    for uploaded in files:
        filename = uploaded.filename or "master-data-file"
        if not filename.lower().endswith(MASTER_DATA_ALLOWED_EXT):
            errors.append(f"{filename}: unsupported file type — use PDF, XLSX or CSV")
            continue
        raw = await uploaded.read()
        total_size += len(raw)
        if total_size > 25 * 1024 * 1024:
            errors.append(f"{filename}: skipped — combined upload exceeds the 25MB limit")
            continue
        if not raw:
            errors.append(f"{filename}: skipped — empty file")
            continue
        try:
            extracted = extract_mca_master_data(filename, raw)
        except Exception as e:
            logger.warning("roc_sphere master-data: failed to parse %s: %s", filename, e)
            errors.append(f"{filename}: could not be read — file may be corrupted or password-protected")
            continue

        fields = {k: v for k, v in extracted.items() if not k.startswith("_")}
        results.append({"filename": filename, "fields_found": len(fields)})
        for k, v in fields.items():
            if v not in (None, "", 0):
                merged[k] = v
        for d in extracted.get("_directors") or []:
            if not any((existing.get("din") or "").upper() == (d.get("din") or "").upper() and d.get("din") for existing in directors):
                directors.append(d)

    if not merged and not directors:
        return {
            "applied": False,
            "fields_applied": [],
            "results": results,
            "errors": errors or ["Could not confidently extract Master Data fields from this file — please check it's the MCA Master Data export, or enter details manually."],
        }

    clean: Dict[str, Any] = {}
    for f in ("cin", "company_name", "registered_office_address", "date_of_incorporation",
              "authorized_capital", "paid_up_capital", "last_agm_date", "pan", "category", "listed"):
        if merged.get(f) not in (None, ""):
            clean[f] = merged[f]

    effective_category = clean.get("category") or company.get("category")
    if directors:
        if effective_category == "llp":
            clean["designated_partners"] = directors
            clean["partners"] = directors
        else:
            clean["directors"] = directors

    existing_master = dict(company.get("master_data") or {})
    existing_master.update({k: v for k, v in {
        "roc_name": merged.get("roc_name"),
        "registration_number": merged.get("registration_number"),
        "email": merged.get("email"),
        "company_status": merged.get("company_status"),
        "roc_office": merged.get("roc_office"),
        "rd_region": merged.get("rd_region"),
        "date_of_balance_sheet": merged.get("date_of_balance_sheet"),
        "active_compliance": merged.get("active_compliance_raw"),
        "company_subcategory": merged.get("company_subcategory_raw"),
    }.items() if v not in (None, "")})
    existing_master["last_fetched_at"] = _now().isoformat()
    existing_master["last_fetched_by"] = _who(current_user)
    existing_master["source_files"] = [r["filename"] for r in results]
    clean["master_data"] = existing_master

    clean["roc_form_uploads"] = (company.get("roc_form_uploads") or []) + [
        {"filename": r["filename"], "form_type": "master-data", "uploaded_at": _now().isoformat()}
        for r in results
    ]
    clean["updated_at"] = _now()

    await COMPANIES.update_one({"id": company_id}, {"$set": clean})
    await _sync_company_to_client({**company, **clean})
    updated = await COMPANIES.find_one({"id": company_id})
    updated.pop("_id", None)

    return {
        "applied": True,
        "fields_applied": sorted(k for k in clean.keys() if k not in ("master_data", "roc_form_uploads", "updated_at")),
        "company": updated,
        "results": results,
        "errors": errors,
    }


# ─────────────────────────────────────────────────────────────────────────
# STATUTORY REGISTERS — share transfers, certificates and SH-4
# ─────────────────────────────────────────────────────────────────────────

def _same_holder(left: Dict[str, Any], name: str, folio: Optional[str]) -> bool:
    if folio and str(left.get("folio_no") or "").strip().lower() == str(folio).strip().lower():
        return True
    return bool(name) and str(left.get("name") or "").strip().lower() == name.strip().lower()


def _apply_transfer_to_register(
    shareholders: List[Dict[str, Any]], transfer: Dict[str, Any]
) -> Dict[str, Any]:
    """Apply a reviewed transfer to the in-app register of members.

    This intentionally reports what could not be matched instead of silently
    inventing a transferor holding. A CS can correct the register and rerun
    the record with `update_register` disabled when the legal instrument is
    still under review.
    """
    quantity = _num(transfer.get("number_of_shares"))
    if quantity <= 0:
        return {"updated": False, "reason": "Number of shares must be greater than zero"}
    transferor = next(
        (s for s in shareholders if _same_holder(s, transfer.get("transferor_name", ""), transfer.get("transferor_folio_no"))),
        None,
    )
    transferee = next(
        (s for s in shareholders if _same_holder(s, transfer.get("transferee_name", ""), transfer.get("transferee_folio_no"))),
        None,
    )
    if not transferor:
        return {"updated": False, "reason": "Transferor was not found in the current shareholder register"}
    if _num(transferor.get("shares_held")) < quantity:
        return {"updated": False, "reason": "Transferor holding is lower than the transfer quantity"}

    transferor["shares_held"] = _num(transferor.get("shares_held")) - quantity
    if transferee:
        transferee["shares_held"] = _num(transferee.get("shares_held")) + quantity
    else:
        shareholders.append({
            "name": transfer.get("transferee_name"),
            "folio_no": transfer.get("transferee_folio_no"),
            "class_of_shares": transfer.get("class_of_shares") or "Equity",
            "shares_held": quantity,
            "face_value": _num(transfer.get("nominal_value_per_share")) or 10,
            "percentage": None,
        })

    total = sum(_num(s.get("shares_held")) for s in shareholders)
    if total:
        for shareholder in shareholders:
            shareholder["percentage"] = round(_num(shareholder.get("shares_held")) / total * 100, 2)
    return {
        "updated": True,
        "transferor_remaining": transferor["shares_held"],
        "transferee_new_holding": next(
            (_num(s.get("shares_held")) for s in shareholders if _same_holder(
                s, transfer.get("transferee_name", ""), transfer.get("transferee_folio_no")
            )),
            quantity,
        ),
    }


@router.get("/companies/{company_id}/statutory-records")
async def get_statutory_records(company_id: str, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    shareholders = company.get("shareholders") or []
    return {
        "company_name": company.get("company_name"),
        "cin": company.get("cin"),
        "shareholders": shareholders,
        "financial_data": company.get("financial_data") or {},
        "annual_return_data": company.get("annual_return_data") or {},
        "audit_report_data": company.get("audit_report_data") or {},
        "board_report_data": company.get("board_report_data") or {},
        "share_transfers": company.get("share_transfers") or [],
        "share_certificates": company.get("share_certificates") or [],
        "register_status": {
            "members": bool(shareholders),
            "share_transfer_register": bool(company.get("share_transfers")),
            "share_certificates": bool(company.get("share_certificates")),
        },
        "disclaimer": "Generated records are working drafts. Verify the executed instrument, stamp duty, board approvals, folio balances and applicable MCA requirements before signing or filing.",
    }


@router.get("/companies/{company_id}/filing-preparation")
async def get_filing_preparation(company_id: str, current_user: User = Depends(VIEW)):
    """Return source-separated working data for the next AOC-4/MGT-7A cycle."""
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    aoc4 = company.get("financial_data") or {}
    mgt7a = company.get("annual_return_data") or {}
    required = {
        "company_name": company.get("company_name"),
        "cin": company.get("cin"),
        "registered_office_address": company.get("registered_office_address"),
        "period_to": aoc4.get("period_to"),
        "turnover": mgt7a.get("turnover") or aoc4.get("turnover") or company.get("last_year_turnover"),
        "net_worth": mgt7a.get("net_worth") or aoc4.get("net_worth"),
        "auditor": (company.get("auditor") or {}).get("name"),
        "directors": company.get("directors") or [],
        "shareholders": company.get("shareholders") or [],
    }
    missing = [key for key, value in required.items() if value in (None, "", [])]
    return {
        "company_id": company_id,
        "aoc4": aoc4,
        "mgt7a": mgt7a,
        "audit_report": company.get("audit_report_data") or {},
        "board_report": company.get("board_report_data") or {},
        "required_working_fields": required,
        "missing_working_fields": missing,
        "source_rules": {
            "financials_and_auditor": "AOC-4",
            "directors_and_shareholders": "MGT-7 / MGT-7A and its shareholder attachment",
            "disclosure_context": "Auditor's Report and Board's Report",
        },
    }


@router.post("/companies/{company_id}/share-transfers")
async def create_share_transfer(
    company_id: str,
    payload: ShareTransferRequest,
    current_user: User = Depends(CREATE),
):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    if _num(payload.number_of_shares) <= 0:
        raise HTTPException(422, "Number of shares must be greater than zero")
    transfer = payload.model_dump()
    transfer.update({"id": _uid(), "created_at": _now().isoformat(), "created_by": _who(current_user)})
    shareholders = [dict(s) for s in (company.get("shareholders") or [])]
    register_update = {"updated": False, "reason": "Register update was not requested"}
    if payload.update_register:
        register_update = _apply_transfer_to_register(shareholders, transfer)
        if not register_update["updated"]:
            raise HTTPException(422, register_update["reason"])
    transfers = list(company.get("share_transfers") or [])
    transfers.append(transfer)
    clean = {"share_transfers": transfers, "updated_at": _now()}
    if payload.update_register:
        clean["shareholders"] = shareholders
    await COMPANIES.update_one({"id": company_id}, {"$set": clean})
    await _sync_company_to_client({**company, **clean})
    return {"transfer": transfer, "register_update": register_update, "shareholders": shareholders}


@router.post("/companies/{company_id}/share-certificates")
async def create_share_certificate(
    company_id: str,
    payload: ShareCertificateRequest,
    current_user: User = Depends(CREATE),
):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    certificate = payload.model_dump()
    certificate.update({"id": _uid(), "created_at": _now().isoformat(), "created_by": _who(current_user)})
    certificates = list(company.get("share_certificates") or [])
    certificates.append(certificate)
    await COMPANIES.update_one(
        {"id": company_id},
        {"$set": {"share_certificates": certificates, "updated_at": _now()}},
    )
    return {"certificate": certificate}


# ─────────────────────────────────────────────────────────────────────────
# COMPLIANCE CHECKLIST ENGINE
# ─────────────────────────────────────────────────────────────────────────
# Heuristic, config-driven so it's easy to keep current. Review
# thresholds against the latest MCA notifications periodically —
# "small company" limits were last revised (paid-up ≤ Rs 10 cr, turnover
# ≤ Rs 100 cr) with effect from 1 Dec 2025.

SMALL_CO_PAID_UP_LIMIT = 10_00_00_000     # Rs 10 crore
SMALL_CO_TURNOVER_LIMIT = 100_00_00_000   # Rs 100 crore


def _is_llp(company: Dict[str, Any]) -> bool:
    return company.get("category") == "llp"


def _is_small_company(company: Dict[str, Any]) -> bool:
    # "Small company" is a Companies Act, 2013 concept only — never applies to an LLP.
    if _is_llp(company):
        return False
    if company.get("is_small_company") is not None:
        return bool(company["is_small_company"])
    if company.get("category") in ("public", "section_8"):
        return False
    paid_up = _num(company.get("paid_up_capital"))
    turnover = _num(company.get("last_year_turnover"))
    return paid_up <= SMALL_CO_PAID_UP_LIMIT and turnover <= SMALL_CO_TURNOVER_LIMIT


# LLP Act, 2008 / LLP Rules audit-applicability thresholds
LLP_AUDIT_TURNOVER_LIMIT = 40_00_000       # Rs. 40 lakh turnover
LLP_AUDIT_CONTRIBUTION_LIMIT = 25_00_000   # Rs. 25 lakh partners' contribution


def build_llp_compliance_checklist(company: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Compliance checklist under the LLP Act, 2008 / LLP Rules — entirely
    separate from the Companies Act, 2013 checklist below, since an LLP has
    no share capital, no Board/AGM in the company-law sense, no AOC-4/MGT-7/
    ADT-1/CSR/MGT-14/PAS-6, etc. Keyed off Designated Partners, not Directors."""
    items: List[Dict[str, Any]] = []
    num_dp = len(company.get("directors") or [])  # "directors" list doubles as Designated Partners for LLP records
    turnover = _num(company.get("last_year_turnover"))
    contribution = _num(company.get("paid_up_capital"))  # "paid_up_capital" field doubles as total partners' contribution
    audit_applicable = turnover > LLP_AUDIT_TURNOVER_LIMIT or contribution > LLP_AUDIT_CONTRIBUTION_LIMIT

    def add(form, particulars, due, frequency, applicable=True, notes=None):
        items.append({
            "form": form, "particulars": particulars, "due_date_rule": due,
            "frequency": frequency, "applicable": applicable, "notes": notes,
        })

    add("Form 11", "Annual Return of the LLP", "Within 60 days of FY close (by 30th May)", "Annual")
    add("Form 8", "Statement of Account & Solvency (incl. Statement of Solvency by Designated Partners)",
        "Within 30 days from end of 6 months of FY close (by 30th October)", "Annual")
    add("Statutory Audit", "Audit of accounts by a Chartered Accountant",
        "Before filing Form 8", "Annual",
        applicable=audit_applicable,
        notes=f"Mandatory only if turnover > ₹{LLP_AUDIT_TURNOVER_LIMIT:,} or partners' contribution > ₹{LLP_AUDIT_CONTRIBUTION_LIMIT:,}; otherwise a self-declaration of solvency suffices.")
    add("DIR-3 KYC", "KYC of every Designated Partner holding a DIN/DPIN", "By 30th September every year", "Annual", applicable=num_dp > 0)
    add("Form 3", "Filing of LLP Agreement / any changes to the LLP Agreement", "Within 30 days of execution/change", "Event-based")
    add("Form 4", "Notice of appointment/cessation/change of a partner or Designated Partner, or change of name/address of a partner",
        "Within 30 days of the event", "Event-based")
    add("Form 15", "Notice of change of registered office of the LLP", "Within 30 days of the change", "Event-based")
    add("Income Tax Return", "Filing of the LLP's income tax return",
        "31st July (no audit) / 31st October (audit applicable)", "Annual",
        notes="Due date shifts to 31st Oct if the LLP is subject to tax audit / transfer-pricing audit.")
    add("GST Returns", "GSTR-1 / GSTR-3B (and annual return GSTR-9) if GST-registered", "Monthly/Quarterly + Annual", "Periodic",
        applicable=bool(company.get("gst_registered", True)),
        notes="Applicable only if the LLP holds a GST registration — verify against the client's actual GSTIN status.")
    add("Meeting of Designated Partners", "Periodic meeting of Designated Partners as required by the LLP Agreement",
        "As per the LLP Agreement (no statutory AGM/Board Meeting requirement under the LLP Act)", "As per LLP Agreement")
    add("Register of Partners / Designated Partners", "Internal registers to be maintained & kept updated", "Ongoing", "Ongoing")

    return items


def build_compliance_checklist(company: Dict[str, Any]) -> List[Dict[str, Any]]:
    category = company.get("category", "private")
    if category == "llp":
        return build_llp_compliance_checklist(company)
    is_opc = category == "opc"
    is_small = _is_small_company(company)
    is_section8 = category == "section_8"
    num_directors = len(company.get("directors") or [])

    items: List[Dict[str, Any]] = []

    def add(form, particulars, due, frequency, applicable=True, notes=None):
        items.append({
            "form": form,
            "particulars": particulars,
            "due_date_rule": due,
            "frequency": frequency,
            "applicable": applicable,
            "notes": notes,
        })

    # Annual filings
    add("AOC-4" + (" (XBRL)" if company.get("listed") else ""), "Filing of Financial Statements", "Within 30 days of AGM", "Annual")
    add("MGT-7A" if is_small or is_opc else "MGT-7", "Annual Return", "Within 60 days of AGM", "Annual")
    add("ADT-1", "Appointment/Ratification of Statutory Auditor", "Within 15 days of AGM (on appointment)", "As applicable")
    add("DIR-3 KYC", "KYC of every Director holding a DIN", "By 30th September every year", "Annual", applicable=num_directors > 0)
    add("DPT-3", "Return of Deposits / particulars of transactions not treated as deposits", "By 30th June every year", "Annual")
    add("MSME-1", "Half-yearly return of outstanding dues to Micro & Small Enterprises", "29th April & 31st October", "Half-Yearly")

    # Meetings
    if not is_opc:
        add("AGM", "Annual General Meeting", "Within 6 months of FY end (9 months for first AGM); gap between two AGMs ≤ 15 months", "Annual", applicable=not is_opc)
    board_meetings_required = 2 if (is_small or is_opc) else 4
    add("Board Meeting", f"Minimum {board_meetings_required} Board Meetings in a calendar year, gap ≤ 120 days between two meetings",
        "Ongoing", "Quarterly" if board_meetings_required == 4 else "Half-Yearly")

    # Registers / compliance
    add("MBP-1", "Director's disclosure of interest in other entities", "First Board Meeting of the FY / on change", "Annual + event-based")
    add("DIR-8", "Director's declaration of non-disqualification", "First Board Meeting of the FY", "Annual")
    add("Register of Members / Charges / Directors", "Statutory registers to be maintained & kept updated", "Ongoing", "Ongoing")
    add("CSR-2", "Report on CSR", "As notified (with AOC-4 or separately)", "Annual",
        applicable=_num(company.get("paid_up_capital")) >= 5_00_00_000 or _num(company.get("last_year_turnover")) >= 100_00_00_000,
        notes="Applicable only if CSR provisions (Sec 135) trigger — verify against latest profit criterion too.")

    if is_section8:
        add("CSR-1", "Registration for undertaking CSR activities (if applicable)", "Before undertaking CSR work", "One-time", notes="Section 8 company specific")

    if company.get("category") == "public" or company.get("listed"):
        add("MGT-14", "Filing of certain Board/Special resolutions with ROC", "Within 30 days of passing", "Event-based")

    add("PAS-6", "Reconciliation of Share Capital Audit Report (unlisted companies with dematerialised shares)", "Within 60 days of half-year end", "Half-Yearly", applicable=not is_opc and not is_section8)

    return items


@router.get("/companies/{company_id}/compliance-checklist")
async def get_compliance_checklist(company_id: str, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    checklist = build_compliance_checklist(company)
    is_llp = _is_llp(company)
    return {
        "company_name": company.get("company_name"),
        "category": company.get("category"),
        "is_small_company": None if is_llp else _is_small_company(company),
        "checklist": checklist,
        "generated_at": _now().isoformat(),
        "disclaimer": (
            "Generated from the LLP's category/turnover/contribution using current "
            "LLP Act, 2008 / LLP Rules thresholds as configured in the app. Verify "
            "against the latest MCA notifications before relying on it for filing."
            if is_llp else
            "Generated from company category/capital/turnover using current "
            "Companies Act 2013 thresholds as configured in the app. Verify "
            "against the latest MCA notifications before relying on it for filing."
        ),
    }


FREQUENCY_GROUP_ORDER = [
    "Annual", "Half-Yearly", "Quarterly", "Periodic", "Event-based",
    "As applicable", "Ongoing", "As per LLP Agreement",
]


@router.get("/companies/{company_id}/applicable-compliances")
async def get_applicable_compliances(company_id: str, current_user: User = Depends(VIEW)):
    """Full list of ROC compliances currently applicable to this company —
    a dashboard view driven by the *same* checklist engine as the
    Compliance Checklist tab, but only the items that apply, grouped by
    frequency. Because Upload ROC Forms and the Master Data importer both
    write straight into the Company Master fields this reads (capital,
    turnover, director/partner count, category), applicability here always
    reflects the latest uploaded data with no separate wiring needed."""
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    checklist = build_compliance_checklist(company)
    applicable = [item for item in checklist if item.get("applicable")]

    groups: Dict[str, List[Dict[str, Any]]] = {}
    for item in applicable:
        groups.setdefault(item.get("frequency") or "Other", []).append(item)
    ordered_groups = [{"frequency": g, "items": groups[g]} for g in FREQUENCY_GROUP_ORDER if g in groups]
    ordered_groups += [{"frequency": g, "items": items} for g, items in groups.items() if g not in FREQUENCY_GROUP_ORDER]

    is_llp = _is_llp(company)
    master_data = company.get("master_data") or {}
    return {
        "company_name": company.get("company_name"),
        "category": company.get("category"),
        "cin": company.get("cin"),
        "is_small_company": None if is_llp else _is_small_company(company),
        "total_forms_tracked": len(checklist),
        "total_applicable": len(applicable),
        "not_applicable": len(checklist) - len(applicable),
        "groups": ordered_groups,
        "master_data_last_fetched": master_data.get("last_fetched_at"),
        "roc_forms_uploaded": len(company.get("roc_form_uploads") or []),
        "generated_at": _now().isoformat(),
        "disclaimer": (
            "Reflects the LLP Act, 2008 / LLP Rules obligations currently applicable to this LLP, "
            "computed live from its category/turnover/contribution. Verify against the latest MCA "
            "notifications before relying on it for filing."
            if is_llp else
            "Reflects the Companies Act, 2013 obligations currently applicable to this company, "
            "computed live from its category/capital/turnover/director count — including any values "
            "fetched from uploaded Master Data or ROC Forms. Verify against the latest MCA "
            "notifications before relying on it for filing."
        ),
    }


# ─────────────────────────────────────────────────────────────────────────
# DOCX GENERATION
# ─────────────────────────────────────────────────────────────────────────

def _base_doc():
    from docx import Document
    from docx.shared import Pt
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    sections = document.sections
    for s in sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = document.sections[0].left_margin
    return document


def _heading(document, text, size=14, center=True, bold=True, underline=False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = document.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size)
    return p


def _para(document, text, center=False, bold=False, italic=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = document.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def build_board_resolution_doc(company: Dict[str, Any], req: BoardResolutionRequest, prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    cin = company.get("cin") or "—"
    is_llp = _is_llp(company)
    body_label = "Designated Partners of the LLP" if is_llp else "Board of Directors of the Company"
    present_label = "Designated Partners Present" if is_llp else "Directors Present"
    sign_label = "Designated Partner" if is_llp else "Director / Company Secretary"
    din_label = "DIN/DPIN" if is_llp else "DIN/Membership No."
    _heading(d, name, size=16)
    _para(d, f"CIN: {cin}" if not is_llp else f"LLPIN: {cin}", center=True)
    _para(d, f"Registered Office: {company.get('registered_office_address') or '—'}", center=True)
    d.add_paragraph()
    _heading(d, "EXTRACT OF MINUTES / CERTIFIED TRUE COPY OF RESOLUTION(S)", size=13)
    _para(
        d,
        f"Passed at the meeting of the {body_label} held on "
        f"{_fmt_date(req.meeting_date)} at {req.meeting_time} at {req.venue}."
    )
    if req.directors_present:
        _para(d, f"{present_label}: " + ", ".join(req.directors_present))
    if req.chairman:
        _para(d, f"Chairman of the Meeting: {req.chairman}")
    d.add_paragraph()

    for i, r in enumerate(req.resolutions, 1):
        _para(d, f"{i}. {r.particulars}", bold=True)
        _para(d, f'"RESOLVED THAT {r.resolution_text.strip().rstrip(".")}."')
        if r.proposed_by or r.seconded_by:
            bits = []
            if r.proposed_by:
                bits.append(f"Proposed by: {r.proposed_by}")
            if r.seconded_by:
                bits.append(f"Seconded by: {r.seconded_by}")
            _para(d, "   " + " | ".join(bits), italic=True)
        d.add_paragraph()

    _para(d, "\nCertified True Copy")
    d.add_paragraph()
    _para(d, "For " + name, bold=True)
    d.add_paragraph()
    d.add_paragraph()
    _para(d, sign_label)
    _para(d, f"{din_label}: __________________")
    _para(d, f"Date: {_fmt_date(datetime.now())}", )
    _para(d, f"Prepared by: {prepared_by} (Taskosphere ROC Sphere)", italic=True)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_notice_doc(company: Dict[str, Any], req: MeetingNoticeRequest, prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    is_llp = _is_llp(company)
    if is_llp:
        label = {
            "board": "NOTICE OF MEETING OF DESIGNATED PARTNERS",
            "agm": "NOTICE OF MEETING OF PARTNERS",
            "egm": "NOTICE OF MEETING OF PARTNERS (EXTRA-ORDINARY)",
        }.get(req.meeting_type, "NOTICE OF MEETING")
    else:
        label = {
            "board": "NOTICE OF BOARD MEETING",
            "agm": "NOTICE OF ANNUAL GENERAL MEETING",
            "egm": "NOTICE OF EXTRA-ORDINARY GENERAL MEETING",
        }.get(req.meeting_type, "NOTICE OF MEETING")
    sign_label = "Designated Partner" if is_llp else "Director / Company Secretary"
    order_label = "By Order of the Designated Partners" if is_llp else "By Order of the Board"
    _heading(d, name, size=16)
    _para(d, f"LLPIN: {company.get('cin') or '—'}" if is_llp else f"CIN: {company.get('cin') or '—'}", center=True)
    _para(d, f"Registered Office: {company.get('registered_office_address') or '—'}", center=True)
    d.add_paragraph()
    _heading(d, label, size=13, underline=True)
    _para(d, f"Notice dated: {_fmt_date(req.notice_date or datetime.now())}")
    d.add_paragraph()
    if req.meeting_type == "board":
        body = "Designated Partners of the LLP" if is_llp else "Board of Directors of the Company"
        _para(d, f"NOTICE is hereby given that a meeting of the {body} will be held on "
                 f"{_fmt_date(req.meeting_date)} at {req.meeting_time} at {req.venue}, to transact the following business:")
    else:
        body = "Partners of the LLP" if is_llp else "members of the Company"
        _para(d, f"NOTICE is hereby given that the {label.split('OF ')[-1].title()} of the {body} will be held on "
                 f"{_fmt_date(req.meeting_date)} at {req.meeting_time} at {req.venue}, to transact the following business:")
    d.add_paragraph()
    if req.agenda_items:
        _para(d, "ORDINARY BUSINESS / AGENDA:", bold=True)
        for i, item in enumerate(req.agenda_items, 1):
            _para(d, f"{i}. {item}")
        d.add_paragraph()
    if req.special_business:
        _para(d, "SPECIAL BUSINESS:", bold=True)
        for i, r in enumerate(req.special_business, 1):
            _para(d, f"{i}. {r.particulars}", bold=True)
            _para(d, f'"RESOLVED THAT {r.resolution_text.strip().rstrip(".")}."')
        d.add_paragraph()

    if req.meeting_type != "board":
        _para(d, "NOTES:", bold=True)
        if is_llp:
            _para(d, "1. Attendance, quorum and voting shall be governed by the provisions of the LLP Agreement.")
            _para(d, "2. A Partner may be represented by an authorised representative only if expressly permitted by the LLP Agreement.")
        else:
            _para(d, "1. A member entitled to attend and vote is entitled to appoint a proxy to attend and vote instead of "
                     "himself/herself, and such proxy need not be a member of the Company.")
            _para(d, "2. Proxies, in order to be effective, must be received at the Registered Office not less than 48 hours "
                     "before the commencement of the meeting.")
    d.add_paragraph()
    _para(d, order_label)
    _para(d, "For " + name, bold=True)
    d.add_paragraph()
    _para(d, sign_label)
    _para(d, f"Prepared by: {prepared_by} (Taskosphere ROC Sphere)", italic=True)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_minutes_doc(company: Dict[str, Any], req: MinutesRequest, prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    is_llp = _is_llp(company)
    if is_llp:
        label = {
            "board": "MINUTES OF THE MEETING OF THE DESIGNATED PARTNERS",
            "agm": "MINUTES OF THE MEETING OF THE PARTNERS",
            "egm": "MINUTES OF THE MEETING OF THE PARTNERS (EXTRA-ORDINARY)",
        }.get(req.meeting_type, "MINUTES OF MEETING")
    else:
        label = {
            "board": "MINUTES OF THE MEETING OF THE BOARD OF DIRECTORS",
            "agm": "MINUTES OF THE ANNUAL GENERAL MEETING",
            "egm": "MINUTES OF THE EXTRA-ORDINARY GENERAL MEETING",
        }.get(req.meeting_type, "MINUTES OF MEETING")
    present_label = "Designated Partners Present" if is_llp else "Directors Present"
    absent_label = "Designated Partners Absent (Leave of Absence granted)" if is_llp else "Directors Absent (Leave of Absence granted)"
    other_label = "Other Partners / Attendees Present" if is_llp else "Members / Attendees Present"
    _heading(d, name, size=16)
    _para(d, f"LLPIN: {company.get('cin') or '—'}" if is_llp else f"CIN: {company.get('cin') or '—'}", center=True)
    d.add_paragraph()
    _heading(d, label, size=13, underline=True)
    _para(d, f"Held on {_fmt_date(req.meeting_date)} at {req.meeting_time} at {req.venue}.")
    d.add_paragraph()
    if req.meeting_type == "board":
        _para(d, f"{present_label}: " + (", ".join(req.directors_present) or "—"))
        if req.directors_absent:
            _para(d, f"{absent_label}: " + ", ".join(req.directors_absent))
    else:
        _para(d, f"{present_label}: " + (", ".join(req.directors_present) or "—"))
        if req.attendees_other:
            _para(d, f"{other_label}: " + ", ".join(req.attendees_other))
    if req.chairman:
        _para(d, f"{req.chairman} chaired the meeting.")
    _para(
        d,
        "Quorum was confirmed to be present."
        if req.quorum_present else
        ("NOTE: Quorum was NOT present — meeting stands adjourned as per the LLP Agreement." if is_llp
         else "NOTE: Quorum was NOT present — meeting stands adjourned as per Companies Act / AoA provisions.")
    )
    d.add_paragraph()

    if req.discussion_notes:
        _para(d, "DISCUSSION:", bold=True)
        _para(d, req.discussion_notes)
        d.add_paragraph()

    if req.resolutions:
        _para(d, "RESOLUTIONS PASSED:", bold=True)
        for i, r in enumerate(req.resolutions, 1):
            _para(d, f"{i}. {r.particulars}", bold=True)
            _para(d, f'"RESOLVED THAT {r.resolution_text.strip().rstrip(".")}."')
            d.add_paragraph()

    _para(d, "There being no other business, the meeting concluded with a vote of thanks to the Chair.")
    d.add_paragraph()
    d.add_paragraph()
    _para(d, "Designated Partner" if is_llp else "Chairman", bold=True)
    _para(d, f"Prepared by: {prepared_by} (Taskosphere ROC Sphere)", italic=True)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_shareholders_doc(company: Dict[str, Any], prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    _heading(d, name, size=16)
    _para(d, f"CIN: {company.get('cin') or '—'}", center=True)
    d.add_paragraph()
    _heading(d, "REGISTER OF MEMBERS / LIST OF SHAREHOLDERS", size=13, underline=True)
    _para(d, f"As on: {_fmt_date(datetime.now())}")
    d.add_paragraph()

    shareholders = company.get("shareholders") or []
    total_shares = sum(_num(s.get("shares_held")) for s in shareholders) or 1

    table = d.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Sl. No.", "Name of Member", "Folio No. / PAN", "Class of Shares", "No. of Shares Held", "% Holding"]):
        hdr[i].text = ""
        hdr[i].paragraphs[0].add_run(h).bold = True

    for idx, s in enumerate(shareholders, 1):
        row = table.add_row().cells
        pct = s.get("percentage")
        if pct is None:
            pct = round(_num(s.get("shares_held")) / total_shares * 100, 2)
        row[0].text = str(idx)
        row[1].text = s.get("name", "")
        row[2].text = " / ".join(filter(None, [s.get("folio_no"), s.get("pan")])) or "—"
        row[3].text = s.get("class_of_shares") or "Equity"
        row[4].text = f"{_num(s.get('shares_held')):,.0f}"
        row[5].text = f"{pct}%"

    d.add_paragraph()
    _para(d, f"Total Paid-up Share Capital: Rs. {_num(company.get('paid_up_capital')):,.0f}", bold=True)
    _para(d, f"Authorized Share Capital: Rs. {_num(company.get('authorized_capital')):,.0f}", bold=True)
    d.add_paragraph()
    _para(d, f"Prepared by: {prepared_by} (Taskosphere ROC Sphere)", italic=True)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_share_transfer_register_doc(company: Dict[str, Any], prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    _heading(d, name, size=16)
    _para(d, f"CIN: {company.get('cin') or '—'}", center=True)
    d.add_paragraph()
    _heading(d, "REGISTER OF SHARE TRANSFERS", size=13, underline=True)
    _para(d, "Working register — update after receipt, stamping, approval and registration of the instrument.")
    d.add_paragraph()
    table = d.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = [
        "Sl. No.", "Date", "Transferor", "Transferee", "Folio / Certificate",
        "Distinctive Nos.", "Shares", "Consideration (Rs.)", "SH-4 Status",
    ]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for index, transfer in enumerate(company.get("share_transfers") or [], 1):
        row = table.add_row().cells
        distinctive = " - ".join(filter(None, [
            str(transfer.get("distinctive_from") or ""),
            str(transfer.get("distinctive_to") or ""),
        ])) or "—"
        row[0].text = str(index)
        row[1].text = _fmt_date(transfer.get("transfer_date"))
        row[2].text = transfer.get("transferor_name") or "—"
        row[3].text = transfer.get("transferee_name") or "—"
        row[4].text = " / ".join(filter(None, [
            transfer.get("transferor_folio_no"),
            transfer.get("share_certificate_no"),
        ])) or "—"
        row[5].text = distinctive
        row[6].text = f"{_num(transfer.get('number_of_shares')):,.0f}"
        row[7].text = f"{_num(transfer.get('consideration')):,.2f}"
        row[8].text = transfer.get("sh4_status") or "Pending review"
    d.add_paragraph()
    _para(d, "Prepared by: " + prepared_by + " (ROC Sphere)", italic=True)
    _para(d, "This register is a controlled working draft and must be reconciled with the Register of Members and executed SH-4 instruments.", italic=True)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_sh4_doc(company: Dict[str, Any], req: ShareTransferRequest, prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    _heading(d, "FORM NO. SH-4", size=16)
    _heading(d, "SECURITIES TRANSFER FORM", size=13, underline=True)
    _para(d, "(Pursuant to Section 56 of the Companies Act, 2013 and applicable rules)", center=True, italic=True)
    d.add_paragraph()
    details = [
        ("Name of company", name),
        ("CIN", company.get("cin") or "—"),
        ("Registered office", company.get("registered_office_address") or "—"),
        ("Date of execution", _fmt_date(req.instrument_date or req.transfer_date)),
        ("Class of securities", req.class_of_shares),
        ("Number of securities transferred", f"{_num(req.number_of_shares):,.0f}"),
        ("Nominal value per security", f"Rs. {_num(req.nominal_value_per_share):,.2f}"),
        ("Consideration", f"Rs. {_num(req.consideration):,.2f}"),
        ("Distinctive numbers", " - ".join(filter(None, [req.distinctive_from, req.distinctive_to])) or "—"),
        ("Existing share certificate no.", req.share_certificate_no or "—"),
        ("Transferor / registered holder", req.transferor_name),
        ("Transferor folio no.", req.transferor_folio_no or "—"),
        ("Transferee", req.transferee_name),
        ("Transferee folio no.", req.transferee_folio_no or "To be allotted"),
        ("Stamp duty", f"Rs. {_num(req.stamp_duty):,.2f}"),
        ("Date instrument received by company", _fmt_date(req.instrument_received_date)),
    ]
    table = d.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in details:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
    d.add_paragraph()
    _para(d, "Declaration by transferor", bold=True)
    _para(d, "I / We hereby transfer the above securities to the transferee named above, subject to the terms and conditions applicable to the company and the Companies Act, 2013.")
    d.add_paragraph()
    _para(d, "Transferor signature: ______________________________")
    _para(d, "Transferee signature: ______________________________")
    _para(d, "Witness name, address and signature: ______________________________")
    d.add_paragraph()
    _para(d, "For office use", bold=True)
    _para(d, "Board approval date: " + _fmt_date(req.board_resolution_date))
    _para(d, "Registration / SH-4 status: " + (req.sh4_status or "Pending review"))
    _para(d, "Company authorised signatory: ______________________________")
    _para(d, "Prepared by: " + prepared_by + " (ROC Sphere)", italic=True)
    _para(d, "Draft only — verify stamp duty, execution, witness details, board approval and applicable exemptions before use.", italic=True)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_share_certificate_doc(company: Dict[str, Any], req: ShareCertificateRequest, prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    _heading(d, name, size=16)
    _para(d, f"CIN: {company.get('cin') or '—'}", center=True)
    _heading(d, "SHARE CERTIFICATE", size=14, underline=True)
    _para(d, "Certificate No. " + req.certificate_no, center=True, bold=True)
    d.add_paragraph()
    _para(d, f"This is to certify that {req.holder_name} is/are the registered holder(s) of the following {req.class_of_shares} share(s) in the Company, subject to its Memorandum and Articles of Association.")
    table = d.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    values = [
        ("Registered holder", req.holder_name),
        ("Address", req.holder_address or "—"),
        ("Folio number", req.folio_no or "—"),
        ("Class of shares", req.class_of_shares),
        ("Number of shares", f"{_num(req.number_of_shares):,.0f}"),
        ("Nominal value per share", f"Rs. {_num(req.nominal_value_per_share):,.2f}"),
        ("Amount paid per share", f"Rs. {_num(req.amount_paid_per_share):,.2f}"),
        ("Distinctive numbers", " - ".join(filter(None, [req.distinctive_from, req.distinctive_to])) or "—"),
        ("Date of issue", _fmt_date(req.issue_date)),
    ]
    for label, value in values:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
    d.add_paragraph()
    _para(d, "Authorised signatory: ______________________________")
    _para(d, "Authorised signatory: ______________________________")
    _para(d, "Prepared by: " + prepared_by + " (ROC Sphere)", italic=True)
    _para(d, "Draft only — verify certificate numbering, Register of Members, share allotment/transfer records and applicable statutory requirements before issue.", italic=True)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_checklist_doc(company: Dict[str, Any], checklist: List[Dict[str, Any]], prepared_by: str) -> bytes:
    d = _base_doc()
    name = company.get("company_name", "").upper()
    _heading(d, name, size=16)
    _para(d, f"CIN: {company.get('cin') or '—'}", center=True)
    d.add_paragraph()
    _heading(d, "ROC / COMPANIES ACT, 2013 — COMPLIANCE CHECKLIST", size=13, underline=True)
    _para(d, f"Generated on: {_fmt_date(datetime.now())}")
    d.add_paragraph()

    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Form / Item", "Particulars", "Due Date Rule", "Frequency", "Applicable"]):
        hdr[i].text = ""
        hdr[i].paragraphs[0].add_run(h).bold = True
    for item in checklist:
        row = table.add_row().cells
        row[0].text = item["form"]
        row[1].text = item["particulars"] + (f" ({item['notes']})" if item.get("notes") else "")
        row[2].text = item["due_date_rule"]
        row[3].text = item["frequency"]
        row[4].text = "Yes" if item["applicable"] else "No"

    d.add_paragraph()
    _para(d, "This checklist is a drafting aid generated from the company's category, capital and turnover as "
             "recorded in Taskosphere. Please verify against the latest MCA notifications before filing.", italic=True)
    _para(d, f"Prepared by: {prepared_by} (Taskosphere ROC Sphere)", italic=True)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _docx_response(content: bytes, filename: str) -> Response:
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


def _safe(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "_", text or "").strip("_") or "Document"


async def _log_doc(company_id: str, doc_type: str, filename: str, user: User):
    await DOCS_LOG.insert_one({
        "id": _uid(),
        "company_id": company_id,
        "doc_type": doc_type,
        "filename": filename,
        "generated_at": _now(),
        "generated_by": _who(user),
    })


@router.post("/companies/{company_id}/generate/board-resolution")
async def generate_board_resolution(company_id: str, req: BoardResolutionRequest, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    try:
        content = build_board_resolution_doc(company, req, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"Board_Resolution_{_safe(company.get('company_name'))}_{_safe(req.meeting_date)}.docx"
    await _log_doc(company_id, "board_resolution", fname, current_user)
    return _docx_response(content, fname)


@router.post("/companies/{company_id}/generate/notice")
async def generate_notice(company_id: str, req: MeetingNoticeRequest, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    try:
        content = build_notice_doc(company, req, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"Notice_{req.meeting_type.upper()}_{_safe(company.get('company_name'))}_{_safe(req.meeting_date)}.docx"
    await _log_doc(company_id, f"notice_{req.meeting_type}", fname, current_user)
    return _docx_response(content, fname)


@router.post("/companies/{company_id}/generate/minutes")
async def generate_minutes(company_id: str, req: MinutesRequest, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    try:
        content = build_minutes_doc(company, req, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"Minutes_{req.meeting_type.upper()}_{_safe(company.get('company_name'))}_{_safe(req.meeting_date)}.docx"
    await _log_doc(company_id, f"minutes_{req.meeting_type}", fname, current_user)
    return _docx_response(content, fname)


@router.get("/companies/{company_id}/generate/shareholders")
async def generate_shareholders(company_id: str, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    try:
        content = build_shareholders_doc(company, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"Shareholders_{_safe(company.get('company_name'))}.docx"
    await _log_doc(company_id, "shareholders", fname, current_user)
    return _docx_response(content, fname)


@router.get("/companies/{company_id}/generate/share-transfer-register")
async def generate_share_transfer_register(company_id: str, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    try:
        content = build_share_transfer_register_doc(company, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"Share_Transfer_Register_{_safe(company.get('company_name'))}.docx"
    await _log_doc(company_id, "share_transfer_register", fname, current_user)
    return _docx_response(content, fname)


@router.post("/companies/{company_id}/generate/sh-4")
async def generate_sh4(company_id: str, req: ShareTransferRequest, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    try:
        content = build_sh4_doc(company, req, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"SH-4_{_safe(req.transferor_name)}_to_{_safe(req.transferee_name)}_{_safe(req.transfer_date or datetime.now().date().isoformat())}.docx"
    await _log_doc(company_id, "sh4", fname, current_user)
    return _docx_response(content, fname)


@router.post("/companies/{company_id}/generate/share-certificate")
async def generate_share_certificate(company_id: str, req: ShareCertificateRequest, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    try:
        content = build_share_certificate_doc(company, req, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"Share_Certificate_{_safe(req.certificate_no)}_{_safe(req.holder_name)}.docx"
    await _log_doc(company_id, "share_certificate", fname, current_user)
    return _docx_response(content, fname)


@router.get("/companies/{company_id}/generate/checklist")
async def generate_checklist_doc(company_id: str, current_user: User = Depends(VIEW)):
    company = await COMPANIES.find_one({"id": company_id})
    if not company:
        raise HTTPException(404, "Company not found")
    checklist = build_compliance_checklist(company)
    try:
        content = build_checklist_doc(company, checklist, _who(current_user))
    except ImportError as e:
        raise HTTPException(500, f"Document generator not installed on the server: {e}")
    fname = f"Compliance_Checklist_{_safe(company.get('company_name'))}.docx"
    await _log_doc(company_id, "checklist", fname, current_user)
    return _docx_response(content, fname)


@router.get("/companies/{company_id}/documents")
async def list_generated_documents(company_id: str, current_user: User = Depends(VIEW)):
    cursor = DOCS_LOG.find({"company_id": company_id}).sort("generated_at", -1)
    items = [x async for x in cursor]
    for x in items:
        x.pop("_id", None)
    return items
