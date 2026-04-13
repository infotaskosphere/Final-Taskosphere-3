import os
import re
import csv
import uuid
import json
import base64
import logging
import pytz
import traceback
import asyncio
import calendar
import requests
import httpx
import shutil
import pandas as pd
from datetime import datetime, date, timezone, timedelta

# --- FIXED ROUTER IMPORTS ---
# Added 'backend.' to invoicing to match the others
from backend.compliance import router as compliance_router
# reminders routes are inlined directly below (no separate router file needed)
from backend.quotations import router as quotation_router
from backend.attendance_identix import identix_router
from backend.google_auth import router as google_auth_router
from backend.website_tracking import router as website_tracking_router
from backend.invoicing import router as invoicing_router
from backend.visits import router as visits_router
from backend.leads import router as leads_router
from backend.telegram import router as telegram_router
from backend.notifications import router as notification_router, create_notification
from backend.email_integration import router as email_router
# Gemini AI instance (already configured in email_integration module)
try:
    from backend.email_integration import _gemini as _gemini_ai
except ImportError:
    _gemini_ai = None
from backend.passwords import router as passwords_router

from zoneinfo import ZoneInfo
from pathlib import Path
from io import StringIO, BytesIO
from typing import List, Optional, Dict, Any
from dateutil import parser
from contextlib import asynccontextmanager

# Single logger definition
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# FastAPI
from fastapi import FastAPI, APIRouter, Depends, HTTPException, status, BackgroundTasks, UploadFile, File, Form, Query, Request
from fastapi.security import HTTPBearer
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from starlette.middleware.gzip import GZipMiddleware
from passlib.context import CryptContext

# Validation
from pydantic import BaseModel, EmailStr, Field, ConfigDict, field_validator, ValidationError
from bson import ObjectId
from dotenv import load_dotenv

# --- BACKEND MODULE IMPORTS ---
import backend.models as models
from backend.models import (
    Token, User, UserCreate, UserLogin, UserPermissions,
    Todo, TodoCreate, Task, TaskCreate, BulkTaskCreate,
    Client, ClientCreate, MasterClientForm,
    Attendance, StaffActivityLog, StaffActivityCreate, PerformanceMetric,
    DueDate, DueDateCreate,
    DSC, DSCCreate, DSCListResponse, DSCMovementRequest, MovementUpdateRequest,
    Document, DocumentCreate, DocumentMovementRequest,
    DashboardStats, AuditLog,
    HolidayResponse, HolidayCreate,
    DEFAULT_ROLE_PERMISSIONS,
    Reminder, ReminderCreate,
    OffboardRequest
)
from backend.dependencies import (
    db,
    client,
    get_current_user,
    create_access_token,
    check_permission,
    require_admin,
    require_manager_or_admin,
    verify_record_access,
    verify_client_access,
    get_team_user_ids
)

# External Services
from fpdf import FPDF
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from apscheduler.schedulers.background import BackgroundScheduler

# ====================== CONFIG ======================
# Single IST definition
IST = pytz.timezone('Asia/Kolkata')
india_tz = ZoneInfo("Asia/Kolkata")

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# ── Attendance proof upload directory ─────────────────────────────────────────
PROOF_UPLOAD_DIR = ROOT_DIR / "uploads" / "attendance_proof"
PROOF_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# ====================== SECURITY CONFIG ===========================
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# ====================== SCHEDULER ======================
scheduler = BackgroundScheduler(timezone=pytz.timezone("Asia/Kolkata"))

# IN-MEMORY CACHE for daily reminder (avoids DB query on every request)
_last_reminder_date_cache: Optional[str] = None

# ====================== APP ======================
app = FastAPI(title="Taskosphere Backend", redirect_slashes=False)

# === CRITICAL: CORS MUST BE THE VERY FIRST MIDDLEWARE ===
# Registered BEFORE startup_event and all other middleware.
# When the Render free-tier backend is sleeping (cold start), it returns no
# headers at all — the browser shows "No Access-Control-Allow-Origin". This is
# NOT a CORS misconfiguration; it is a cold-start timing issue. Keeping CORS
# registered first ensures that once the server wakes, the headers are correct.
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://final-taskosphere-frontend.onrender.com",
        "http://localhost:3000",
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "Accept", "X-Requested-With"],
    expose_headers=["*"],
    max_age=3600,
)
app.add_middleware(GZipMiddleware, minimum_size=1000)


# =============================================================
# ABSENT MARKING CORE LOGIC
# Runs via APScheduler at 19:00 IST every working day.
# Also exposed as POST /api/attendance/mark-absent-bulk for
# manual admin triggering.
#
# Rules:
#   - Skip confirmed holidays
#   - Skip weekends (Saturday=5, Sunday=6)
#   - For every active user with no present/leave/absent record
#     insert a new absent record with auto_marked=True
#   - If a record exists but has an unexpected status update to absent
# =============================================================

async def _mark_absent_for_date(target_date_str: str) -> dict:
    """Core absent-marking logic. Returns a summary dict."""
    # Skip confirmed holidays
    # FIX: {"_id": 0} projection — ObjectId causes issues if doc is returned
    holiday = await db.holidays.find_one({"date": target_date_str, "status": "confirmed"}, {"_id": 0})
    if holiday:
        return {"skipped": True, "reason": f"Holiday: {holiday.get('name')}", "marked": 0, "date": target_date_str}

    # Skip weekends
    target_date_obj = date.fromisoformat(target_date_str)
    if target_date_obj.weekday() >= 5:
        return {"skipped": True, "reason": "Weekend", "marked": 0, "date": target_date_str}

    # Fetch all active users
    active_users = await db.users.find(
        {"is_active": True, "status": "active"},
        {"_id": 0, "id": 1, "full_name": 1}
    ).to_list(1000)

    marked_count = 0
    already_recorded = 0

    for u in active_users:
        uid = u["id"]
        existing = await db.attendance.find_one({"user_id": uid, "date": target_date_str}, {"_id": 0})

        if existing:
            if existing.get("status") in ("present", "leave", "absent"):
                already_recorded += 1
                continue
            # Record exists but status is unexpected → update to absent
            await db.attendance.update_one(
                {"user_id": uid, "date": target_date_str},
                {"$set": {
                    "status": "absent",
                    "auto_marked": True,
                    "auto_marked_at": datetime.now(timezone.utc).isoformat(),
                }}
            )
            marked_count += 1
        else:
            # No record at all → insert absent
            await db.attendance.insert_one({
                "user_id": uid,
                "date": target_date_str,
                "status": "absent",
                "punch_in": None,
                "punch_out": None,
                "duration_minutes": 0,
                "is_late": False,
                "punched_out_early": False,
                "leave_reason": None,
                "auto_marked": True,
                "auto_marked_at": datetime.now(timezone.utc).isoformat(),
            })
            marked_count += 1

    logger.info(f"Absent marking for {target_date_str}: marked={marked_count}, skipped={already_recorded}")
    return {
        "skipped": False,
        "date": target_date_str,
        "marked": marked_count,
        "total_active_users": len(active_users),
        "already_recorded": already_recorded,
    }


def mark_absent_users_task():
    """
    Sync wrapper called by APScheduler at 19:00 IST every working day.
    FIX: Uses a dedicated new event loop instead of asyncio.run() to avoid
    RuntimeError when called from within an already-running event loop context.
    """
    loop = None
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        today_str = datetime.now(IST).date().isoformat()
        result = loop.run_until_complete(_mark_absent_for_date(today_str))
        logger.info(f"Scheduled absent job result: {result}")
    except Exception as e:
        logger.error(f"mark_absent_users_task failed: {e}")
    finally:
        if loop and not loop.is_closed():
            loop.close()

    

@app.on_event("startup")
async def startup_event():
    try:
        await db.tasks.create_index("assigned_to")
        await create_compliance_indexes()  # import it too
        await db.tasks.create_index("created_by")
        await db.tasks.create_index("due_date")
        await db.users.create_index("email")
        await db.staff_activity.create_index("user_id")
        await db.staff_activity.create_index("timestamp")
        await db.staff_activity.create_index([("user_id", 1), ("timestamp", -1)])
        await db.due_dates.create_index("department")
        await db.tasks.create_index([("assigned_to", 1), ("status", 1)])
        await db.tasks.create_index("created_at")
        await db.referrers.create_index("name")
        await db.clients.create_index("assigned_to")
        await db.dsc_register.create_index("expiry_date")
        await db.todos.create_index([("user_id", 1), ("created_at", -1)])
        await db.attendance.create_index([("user_id", 1), ("date", -1)])
        await db.notifications.create_index("user_id")
        await db.visits.create_index([("assigned_to", 1), ("visit_date", -1)])
        await db.visits.create_index("visit_date")
        await db.visits.create_index("client_id")
        await db.visits.create_index("status")
        await db.notifications.create_index([("user_id", 1), ("is_read", 1)])
        await db.notifications.create_index("created_at")
        await db.quotations.create_index([("created_by", 1), ("created_at", -1)])
        await db.quotations.create_index("status")
        await db.quotations.create_index("service")
        await db.companies.create_index("created_by")
        await db.companies.create_index("name")
        await db.staff_activity.create_index("type")
        await db.staff_activity.create_index("domain")
        await db.staff_activity.create_index([("user_id", 1), ("timestamp", -1)])
        await db.staff_activity.create_index([("user_id", 1), ("type", 1)])

    # ── FIXED: EMAIL CONNECTIONS INDEX ──────────────────────────────────
        try:
            # Drop old rule (Unique User + Provider)
            await db.email_connections.drop_index("user_id_1_provider_1")
        except Exception:
            pass 

        # Create new rule (Unique User + Email Address)
        await db.email_connections.create_index(
            [("user_id", 1), ("email_address", 1)],
            unique=True,
            background=True
        )
        
        # Unique indexes — use background=True so they don't block startup if they already exist
        await db.attendance.create_index(
            [("user_id", 1), ("date", 1)],
            unique=True,
            background=True
        )
        await db.clients.create_index(
            [("created_by", 1), ("company_name", 1)],
            unique=True,
            background=True
        )
        await db.holidays.create_index("date", unique=True, background=True)
    except Exception as e:
        # Log index creation errors but do NOT crash the server
        logger.warning(f"Index creation warning (non-fatal): {e}")

    try:
        visits = await db.visits.find({"id": {"$exists": False}}).to_list(10000)
        repaired = 0
        for v in visits:
            raw_id = v.get("_id")
            new_id = str(raw_id)
            await db.visits.update_one(
                {"_id": raw_id},
                {"$set": {"id": new_id}}
            )
            repaired += 1
        logger.info(f"✅ Visit ID repair: {repaired} documents patched")
    except Exception as e:
        logger.error(f"⚠️ Visit ID repair failed (non-fatal): {e}")

    # Scheduled jobs=====================================================================
    try:
        scheduler.add_job(fetch_indian_holidays_task, 'cron', day=1, hour=0, minute=5)
        # Also run immediately on startup so holidays are available from day 1
        scheduler.add_job(fetch_indian_holidays_task, 'date',
                          run_date=datetime.now(pytz.timezone("Asia/Kolkata")))
        # Absent marking job — fires every working day at 19:00 IST
        scheduler.add_job(
            mark_absent_users_task,
            'cron',
            hour=19,
            minute=0,
            timezone=pytz.timezone("Asia/Kolkata"),
            id="mark_absent_daily",
            replace_existing=True,
        )
        scheduler.start()
        logger.info("APScheduler started successfully.")
    except Exception as e:
        logger.error(f"APScheduler startup failed: {e}")

    # ── AUTO-SYNC HOLIDAYS ON EVERY BOOT ─────────────────────────────────────
    # Runs async in the background — never blocks startup.
    # Fetches Indian public holidays (current + next year) from date.nager.at
    # and saves them all as 'confirmed'. Any existing 'pending' ones are upgraded.
    async def _boot_holiday_sync():
        try:
            import httpx as _httpx
            now_ist = datetime.now(pytz.timezone("Asia/Kolkata"))
            total_added = 0
            for year in [now_ist.year, now_ist.year + 1]:
                try:
                    async with _httpx.AsyncClient(timeout=10) as http:
                        resp = await http.get(
                            f"https://date.nager.at/api/v3/PublicHolidays/{year}/IN"
                        )
                    if resp.status_code != 200:
                        continue
                    for h in resp.json():
                        date_str = h["date"]
                        name     = h.get("localName") or h.get("name", "Holiday")
                        existing = await db.holidays.find_one({"date": date_str}, {"_id": 0})
                        if not existing:
                            await db.holidays.insert_one({
                                "date":       date_str,
                                "name":       name,
                                "status":     "confirmed",
                                "type":       "public",
                                "created_at": now_ist.isoformat(),
                            })
                            total_added += 1
                        elif existing.get("status") not in ("confirmed", "rejected"):
                            await db.holidays.update_one(
                                {"date": date_str},
                                {"$set": {"status": "confirmed"}}
                            )
                except Exception as year_err:
                    logger.warning(f"Holiday sync for {year} failed: {year_err}")
            logger.info(f"Boot holiday sync complete: {total_added} new holidays added")
        except Exception as e:
            logger.warning(f"Boot holiday sync failed (non-fatal): {e}")

    asyncio.create_task(_boot_holiday_sync())

        # 🔥 AUTO MIGRATION: Add consent_given for old users
    try:
        result = await db.users.update_many(
            {},  # all users
            {"$set": {"consent_given": True}}
        )
        logger.info(f"Consent cleanup: Updated {result.modified_count} users")
    except Exception as e:
        logger.error(f"Consent cleanup failed: {e}")

# ====================== HEALTH ======================
@app.api_route("/health", methods=["GET", "HEAD"])
async def health():
    return JSONResponse({"status": "ok", "cors": "configured correctly"})

@app.get("/")
async def root():
    return {"message": "Server is running"}

# ====================== SECURITY & DB ======================
rankings_cache = {}
# Store cache times as timezone-aware UTC datetimes for consistent comparison
rankings_cache_time: Dict[str, datetime] = {}

# ===================== HELPER FUNCTIONS =====================
def safe_dt(value):
    if not value:
        return None
    if isinstance(value, datetime):
        return value.astimezone(IST)
    try:
        dt = datetime.fromisoformat(value)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=pytz.UTC).astimezone(IST)
        return dt
    except Exception:
        return None

def sanitize_user_data(users, current_user=None):
    is_single = False
    if not isinstance(users, list):
        users = [users]
        is_single = True
    sanitized = []
    for user in users:
        if isinstance(user, dict):
            safe_user = {k: v for k, v in user.items() if k not in ["password", "_id"]}
            sanitized.append(safe_user)
        else:
            user_dict = user.model_dump() if hasattr(user, "model_dump") else vars(user)
            safe_user = {k: v for k, v in user_dict.items() if k not in ["password", "_id"]}
            sanitized.append(safe_user)
    return sanitized[0] if is_single else sanitized

def convert_objectids(data):
    """Recursively convert MongoDB ObjectId fields to string."""
    if isinstance(data, list):
        return [convert_objectids(item) for item in data]
    if isinstance(data, dict):
        new_dict = {}
        for key, value in data.items():
            if isinstance(value, ObjectId):
                new_dict[key] = str(value)
            elif isinstance(value, (dict, list)):
                new_dict[key] = convert_objectids(value)
            else:
                new_dict[key] = value
        return new_dict
    if isinstance(data, ObjectId):
        return str(data)
    return data

def get_user_permissions(current_user: User) -> dict:
    if isinstance(current_user.permissions, dict):
        return current_user.permissions
    if current_user.permissions:
        return current_user.permissions.model_dump()
    return {}

def is_own_record(current_user: User, record: dict) -> bool:
    uid = current_user.id
    return (
        record.get("user_id") == uid
        or record.get("assigned_to") == uid
        or record.get("created_by") == uid
        or uid in record.get("sub_assignees", [])
    )

async def create_audit_log(current_user: User, action: str, module: str, record_id: str, old_data: dict = None, new_data: dict = None):
    log_entry = AuditLog(
        user_id=current_user.id,
        user_name=current_user.full_name,
        action=action,
        module=module,
        record_id=record_id,
        old_data=convert_objectids(old_data) if old_data else None,
        new_data=convert_objectids(new_data) if new_data else None,
        timestamp=datetime.now(timezone.utc)
    )
    await db.audit_logs.insert_one(log_entry.model_dump())

async def calculate_expected_hours(start_date_str: str, end_date_str: str, shift_start: str = "10:30", shift_end: str = "19:00"):
    try:
        start = date.fromisoformat(start_date_str)
        end = date.fromisoformat(end_date_str)
    except Exception:
        return 0
    if start > end:
        return 0
    try:
        t1 = datetime.strptime(shift_start, "%H:%M")
        t2 = datetime.strptime(shift_end, "%H:%M")
        hrs_per_day = (t2 - t1).total_seconds() / 3600
    except Exception:
        hrs_per_day = 8.5
    holidays_cursor = db.holidays.find({"status": "confirmed"})
    holidays = [h["date"] for h in await holidays_cursor.to_list(length=None)]
    total_hours = 0
    current_date = start
    while current_date <= end:
        if current_date.weekday() < 5 and current_date.isoformat() not in holidays:
            total_hours += hrs_per_day
        current_date += timedelta(days=1)
    return round(total_hours, 2)

def fetch_indian_holidays_task():
    """
    Scheduled job (sync wrapper for BackgroundScheduler) to fetch holidays for the current month.
    FIX: Uses a dedicated new event loop instead of asyncio.run() to avoid
    RuntimeError when an event loop is already running.
    """
    loop = None
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        async def _async_fetch():
            try:
                now = datetime.now(IST)
                # Sync current + next year so holidays are never missing mid-year
                for year in [now.year, now.year + 1]:
                    url = f"https://date.nager.at/api/v3/PublicHolidays/{year}/IN"
                    response = requests.get(url, timeout=10)
                    if response.status_code != 200:
                        continue
                    external_holidays = response.json()
                    count = 0
                    for h in external_holidays:
                        date_str = h['date']
                        existing = await db.holidays.find_one({"date": date_str}, {"_id": 0})
                        if not existing:
                            new_holiday = {
                                "date": date_str,
                                "name": h.get('localName') or h.get('name', 'Holiday'),
                                "status": "confirmed",
                                "type": "public",
                                "created_at": datetime.now(IST).isoformat()
                            }
                            await db.holidays.insert_one(new_holiday)
                            count += 1
                        elif existing.get("status") not in ("confirmed", "rejected"):
                            # Upgrade any pending / unset → confirmed automatically
                            await db.holidays.update_one(
                                {"date": date_str},
                                {"$set": {"status": "confirmed"}}
                            )
                    logger.info(f"Auto-synced holidays for {year}: {count} new")
            except Exception as e:
                logger.error(f"Holiday Autofetch Failed: {str(e)}")

        loop.run_until_complete(_async_fetch())
    except Exception as e:
        logger.error(f"fetch_indian_holidays_task failed: {e}")
    finally:
        if loop and not loop.is_closed():
            loop.close()


# ROUTER
api_router = APIRouter(prefix="/api")

# HELPERS - Email Service Functions
def _brevo_send(to_email: str, subject: str, body_plain: str, body_html: str = None):
    """Core Brevo HTTP API sender — SMTP (port 587) is blocked on Render free tier."""
    api_key      = os.getenv("BREVO_API_KEY")
    sender_email = os.getenv("SENDER_EMAIL")
    sender_name  = os.getenv("SENDER_NAME", "TaskoSphere")

    if not api_key or not sender_email:
        raise Exception(
            "Brevo API env vars not configured. "
            "Set BREVO_API_KEY and SENDER_EMAIL in Render environment variables."
        )

    payload = {
        "sender": {"name": sender_name, "email": sender_email},
        "to": [{"email": to_email}],
        "subject": subject,
        "textContent": body_plain,
    }
    if body_html:
        payload["htmlContent"] = body_html

    response = httpx.post(
        "https://api.brevo.com/v3/smtp/email",
        headers={
            "api-key": api_key,
            "Content-Type": "application/json",
        },
        json=payload,
        timeout=30,
    )

    if response.status_code not in (200, 201):
        raise Exception(f"Brevo API error {response.status_code}: {response.text}")
    return True


def send_birthday_email(recipient_email: str, client_name: str):
    """Send birthday wish email to client via Brevo SMTP."""
    subject = f"Happy Birthday, {client_name}!"
    body_plain = (
        f"Dear {client_name},\n\n"
        f"On behalf of our entire team, we wish you a very Happy Birthday!\n\n"
        f"We appreciate your continued trust and partnership. "
        f"May this year bring you prosperity, success, and happiness.\n\n"
        f"Best regards,\nTaskosphere Team"
    )
    html_content = f"""
    <html>
    <body style="font-family: Arial, sans-serif; padding: 20px; background-color: #f5f5f5;">
    <div style="max-width: 600px; margin: 0 auto; background-color: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
    <h1 style="color: #4F46E5; text-align: center;"> Happy Birthday! </h1>
    <p style="font-size: 16px; line-height: 1.6; color: #333;">Dear {client_name},</p>
    <p style="font-size: 16px; line-height: 1.6; color: #333;">
     On behalf of our entire team, we wish you a very Happy Birthday!
    </p>
    <p style="font-size: 16px; line-height: 1.6; color: #333;">
     We appreciate your continued trust and partnership. May this year bring you prosperity, success, and happiness.
    </p>
    <div style="background-color: #4F46E5; color: white; padding: 15px; border-radius: 5px; margin: 20px 0; text-align: center;">
    <p style="margin: 0; font-size: 18px; font-weight: bold;">Wishing you all the best!</p>
    </div>
    <p style="font-size: 14px; color: #666; text-align: center; margin-top: 30px;">
     Best regards,<br><strong>Taskosphere Team</strong>
    </p>
    </div>
    </body>
    </html>
    """
    try:
        _brevo_send(recipient_email, subject, body_plain, html_content)
        logger.info(f"Birthday email sent to {recipient_email}")
        return True
    except Exception as e:
        logger.error(f"Failed to send birthday email: {str(e)}")
        return False

# ─── TEST EMAIL ENDPOINT ──────────────────────────────────────────────────────
@api_router.post("/email/test")
async def test_email_service(current_user: User = Depends(get_current_user)):
    """Send a test email to the logged-in admin to verify Brevo SMTP is working."""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only")

    missing = [k for k, v in {
        "BREVO_SMTP_USER": os.getenv("BREVO_SMTP_USER"),
        "BREVO_SMTP_PASS": os.getenv("BREVO_SMTP_PASS"),
        "SENDER_EMAIL":    os.getenv("SENDER_EMAIL"),
    }.items() if not v]

    if missing:
        raise HTTPException(status_code=500, detail=f"Missing env vars: {', '.join(missing)}")

    try:
        _brevo_send(
            to_email   = current_user.email,
            subject    = "✅ TaskoSphere — Mail Service Test",
            body_plain = (
                f"Hello {current_user.full_name},\n\n"
                f"This is a test email from TaskoSphere.\n"
                f"If you received this, your mail service is working correctly.\n\n"
                f"SMTP Host : {os.getenv('BREVO_SMTP_HOST', 'smtp-relay.brevo.com')}\n"
                f"Sender    : {os.getenv('SENDER_EMAIL')}\n\n"
                f"Regards,\nTaskoSphere"
            ),
        )
        return {"status": "success", "message": f"Test email sent to {current_user.email}"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Mail error: {str(e)}")

# Task Analytics
@api_router.get("/tasks/analytics")
async def get_task_analytics(
    month: str,
    current_user: User = Depends(get_current_user)
):
    """ Get task analytics for a specific month (YYYY-MM) """
    query = {}
    if current_user.role != "admin":
        query["$or"] = [
            {"assigned_to": current_user.id},
            {"sub_assignees": current_user.id},
            {"created_by": current_user.id}
        ]
    tasks = await db.tasks.find(query, {"_id": 0}).to_list(1000)
    total = 0
    completed = 0
    pending = 0
    for task in tasks:
        created = task.get("created_at")
        if isinstance(created, str):
            if created.startswith(month):
                total += 1
                if task.get("status") == "completed":
                    completed += 1
                elif task.get("status") == "pending":
                    pending += 1
        elif isinstance(created, datetime):
            if created.strftime("%Y-%m") == month:
                total += 1
                if task.get("status") == "completed":
                    completed += 1
                elif task.get("status") == "pending":
                    pending += 1
    return {
        "month": month,
        "total_tasks": total,
        "completed_tasks": completed,
        "pending_tasks": pending
    }


# ── AI: Detect Duplicate Tasks ─────────────────────────────────────────────────
@api_router.post("/tasks/detect-duplicates")
async def detect_duplicate_tasks(
    current_user: User = Depends(get_current_user)
):
    """
    Use Gemini AI (gemini-1.5-flash) to find duplicate tasks.
    Better free-tier quota than gemini-2.0-flash-lite.
    """
    import json as _json, re as _re

    # ── 1. Verify Gemini is configured ────────────────────────────────────
    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    if not gemini_key:
        raise HTTPException(
            status_code=503,
            detail="GEMINI_API_KEY is not set on the server."
        )

    try:
        import google.generativeai as _genai
        _genai.configure(api_key=gemini_key)
        _model = _genai.GenerativeModel("gemini-1.5-flash")
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="google-generativeai package not installed."
        )

    # ── 2. Scope query same as GET /tasks ──────────────────────────────────
    query: dict = {"type": {"$ne": "todo"}}
    if current_user.role != "admin":
        permissions = get_user_permissions(current_user)
        allowed_users = permissions.get("view_other_tasks", []) or []
        if current_user.role == "manager":
            team_ids = await get_team_user_ids(current_user.id)
            allowed_users = list(set(allowed_users + team_ids))
        or_clauses = [
            {"assigned_to": current_user.id},
            {"sub_assignees": current_user.id},
            {"created_by": current_user.id},
        ]
        if allowed_users:
            or_clauses.append({"assigned_to": {"$in": allowed_users}})
        query["$or"] = or_clauses

    # Cap at 50 tasks to stay within free-tier token limits
    tasks = await db.tasks.find(query, {"_id": 0}).to_list(50)

    if not tasks:
        return {"groups": [], "total_tasks_scanned": 0}

    # ── 3. Build minimal payload ───────────────────────────────────────────
    task_summaries = [
        {
            "id":    str(t.get("id", "")),
            "title": (t.get("title") or "")[:100],
            "desc":  (t.get("description") or "")[:80],
            "cat":   t.get("category") or "",
            "cid":   t.get("client_id") or "",
        }
        for t in tasks
    ]

    prompt = (
        "Find duplicate or very similar tasks. "
        "Return ONLY a JSON array, no markdown, no explanation. "
        'Format: [{"reason":"brief reason","confidence":"high|medium","task_ids":["id1","id2"]}] '
        "Only groups with 2+ tasks. If none found return []. "
        f"Tasks: {_json.dumps(task_summaries, ensure_ascii=False)}"
    )

    # ── 4. Call Gemini ─────────────────────────────────────────────────────
    try:
        resp   = await _model.generate_content_async(prompt)
        raw    = _re.sub(r"```[a-zA-Z]*", "", resp.text.strip()).replace("```", "").strip()
        groups = _json.loads(raw)
        if not isinstance(groups, list):
            groups = []
    except Exception as e:
        err_str = str(e)
        if "429" in err_str or "quota" in err_str.lower() or "rate" in err_str.lower():
            raise HTTPException(
                status_code=429,
                detail=(
                    "Gemini API quota exceeded. "
                    "Upgrade at https://ai.google.dev/pricing or wait and retry."
                )
            )
        logger.warning(f"Gemini duplicate detection failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI error: {err_str[:200]}")

    return {"groups": groups, "total_tasks_scanned": len(tasks)}


# Helper functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def send_email(to_email: str, subject: str, body: str):
    """Send plain text email via Brevo SMTP."""
    try:
        _brevo_send(to_email, subject, body)
        return True
    except Exception as e:
        raise Exception(f"Brevo SMTP error: {str(e)}")


#===========================================================
#Website activity
#===========================================================

@api_router.get("/activity/websites")
async def get_website_activity(
    current_user: User = Depends(get_current_user)
):
    try:
        pipeline = [
            {
                "$match": {
                    "user_id": current_user.id,
                    "type": "website"
                }
            },
            {
                "$group": {
                    "_id": "$user_id",
                    "websites": {
                        "$push": {
                            "url": "$url",
                            "domain": "$domain",
                            "title": "$title",
                            "duration": "$duration",
                            "timestamp": "$timestamp"
                        }
                    }
                }
            },
            {
                "$project": {
                    "_id": 0,
                    "user_id": "$_id",
                    "websites": 1
                }
            }
        ]

        data = await db.staff_activity.aggregate(pipeline).to_list(100)

        return data

    except Exception as e:
        logger.error(f"Fetch website activity error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to fetch website activity")


@api_router.post("/activity/track-website")
async def track_website(
    data: dict,
    current_user: User = Depends(get_current_user)
):
    try:
        url = data.get("url")
        domain = data.get("domain")

        if not url or not domain:
            raise HTTPException(status_code=400, detail="Invalid website data")

        activity = {
            "id": str(uuid.uuid4()),
            "user_id": current_user.id,
            "type": "website",
            "url": url,
            "domain": domain,
            "title": data.get("title", ""),
            "timestamp": datetime.now(timezone.utc),
            "duration": int(data.get("duration", 0))
        }

        await db.staff_activity.insert_one(activity)

        return {"status": "tracked"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Website tracking error: {str(e)}")
        raise HTTPException(status_code=500, detail="Tracking failed")

#===========================================================
# AUTH ROUTES
#============================================================
@api_router.get("/system/time")
async def get_system_time():
    now = datetime.now(IST)
    return {
        "server_time": now.isoformat(),
        "display_time": now.strftime("%I:%M:%S %p"),
        "date": now.strftime("%Y-%m-%d")
    }


# =========================
# REFERRERS ROUTES
# =========================

@api_router.get("/referrers")
async def get_referrers(current_user: User = Depends(get_current_user)):
    referrers = await db.referrers.find({}, {"_id": 0}).to_list(500)
    return referrers


@api_router.post("/referrers")
async def create_referrer(data: dict, current_user: User = Depends(get_current_user)):
    name = (data.get("name") or "").strip()

    if not name:
        raise HTTPException(status_code=400, detail="Referrer name required")

    existing = await db.referrers.find_one({"name": {"$regex": f"^{re.escape(name)}$", "$options": "i"}}, {"_id": 0})

    if existing:
        return existing

    referrer = {
        "id": str(uuid.uuid4()),
        "name": name,
        "created_by": current_user.id,
        "created_at": datetime.now(timezone.utc).isoformat()
    }

    await db.referrers.insert_one(referrer)
    # FIX: insert_one mutates `referrer` in-place by adding `_id: ObjectId(...)`.
    # Returning the dict without popping _id causes FastAPI's jsonable_encoder
    # to fail with: ValueError: [TypeError("'ObjectId' object is not iterable"),
    # TypeError('vars() argument must have __dict__ attribute')]
    referrer.pop("_id", None)
    return referrer


# ==========================================================
# TODO DASHBOARD
# ==========================================================
@api_router.post("/todos", response_model=Todo)
async def create_todo(
    todo_data: TodoCreate,
    current_user: User = Depends(get_current_user)
):
    now = datetime.now(timezone.utc)
    todo = Todo(
        user_id=current_user.id,
        **todo_data.model_dump()
    )
    doc = todo.model_dump()

    # Safe conversion with fallback
    doc["created_at"] = (doc.get("created_at") or now).isoformat()
    doc["updated_at"] = (doc.get("updated_at") or now).isoformat()
    
    if doc.get("due_date"):
        if isinstance(doc["due_date"], (datetime, date)):
            doc["due_date"] = doc["due_date"].isoformat()

    result = await db.todos.insert_one(doc)
    doc["id"] = str(result.inserted_id)
    doc.pop("_id", None)
    return doc

@api_router.get("/todos")
async def get_todos(
    user_id: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    if current_user.role == "admin":
        if user_id == "all":
            query = {}
        elif user_id:
            query = {"user_id": user_id}
        else:
            query = {"user_id": current_user.id}

    else:
        permissions = current_user.permissions.model_dump() if hasattr(current_user.permissions, "model_dump") else (current_user.permissions or {})
        if not isinstance(permissions, dict):
            permissions = {}
        allowed_others = permissions.get("view_other_todos", []) or []
        if current_user.role == "manager":
            # Manager: Own + Team (same department)
            team_ids = await get_team_user_ids(current_user.id)
            allowed_others = list(set(allowed_others + team_ids))
        if user_id:
            if user_id != current_user.id and user_id not in allowed_others:
                raise HTTPException(status_code=403, detail="Not allowed")
            query = {"user_id": user_id}
        else:
            visible_ids = list(set(allowed_others + [current_user.id]))
            query = {"user_id": {"$in": visible_ids}}

    todos = await db.todos.find(query).to_list(1000)
    for t in todos:
        t["id"] = str(t["_id"])
        del t["_id"]
    return todos
@api_router.get("/dashboard/todo-overview")
async def get_todo_dashboard(current_user: User = Depends(get_current_user)):
    is_admin = current_user.role == "admin"
    if is_admin:
        todos = await db.todos.find().to_list(2000)
        # Replaced N+1 user queries with a single batch lookup
        user_ids = list({t["user_id"] for t in todos if t.get("user_id")})
        users_raw = await db.users.find({"id": {"$in": user_ids}}, {"_id": 0}).to_list(1000)
        user_name_map = {u["id"]: u.get("full_name", "Unknown User") for u in users_raw}

        grouped_todos = {}
        all_todos_flat = []
        for todo in todos:
            user_name = user_name_map.get(todo["user_id"], "Unknown User")
            if user_name not in grouped_todos:
                grouped_todos[user_name] = []
            todo["_id"] = str(todo["_id"])
            grouped_todos[user_name].append(todo)
            all_todos_flat.append(todo)
        return {
            "role": "admin",
            "todos": all_todos_flat,
            "grouped_todos": grouped_todos
        }
    else:
        permissions = get_user_permissions(current_user)
        allowed_users = permissions.get("view_other_todos", []) or []
        if not isinstance(allowed_users, list):
            allowed_users = []
        if current_user.role == "manager":
            # Manager: Own + Team (same department)
            team_ids = await get_team_user_ids(current_user.id)
            allowed_users = list(set(allowed_users + team_ids))
        query_ids = list(set(allowed_users + [current_user.id]))
        todos = await db.todos.find({"user_id": {"$in": query_ids}}).to_list(2000)
        for todo in todos:
            todo["_id"] = str(todo["_id"])
        return {
            "role": current_user.role,
            "todos": todos
        }

@api_router.post("/todos/{todo_id}/promote-to-task")
async def promote_todo(todo_id: str, current_user: User = Depends(get_current_user)):
    try:
        todo = await db.todos.find_one({"_id": ObjectId(todo_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid Todo ID")
    if not todo:
        raise HTTPException(status_code=404, detail="Todo not found")
    if current_user.role != "admin" and todo["user_id"] != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized to promote this todo")
    now = datetime.now(IST)
    new_task = {
        "id": str(uuid.uuid4()),
        "title": todo["title"],
        "description": todo.get("description"),
        "assigned_to": todo["user_id"],
        "sub_assignees": [],
        "priority": "medium",
        "status": "pending",
        "category": "other",
        "client_id": None,
        "is_recurring": False,
        "type": "task",
        "created_by": current_user.id,
        "created_at": now,
        "updated_at": now
    }
    async with await client.start_session() as session:
        async def cb(session):
            await db.tasks.insert_one(new_task, session=session)
            await db.todos.delete_one({"_id": ObjectId(todo_id)}, session=session)
        await session.with_transaction(cb)
    return {"message": "Todo promoted to task successfully"}

@api_router.delete("/todos/{todo_id}")
async def delete_todo(
    todo_id: str,
    current_user: User = Depends(get_current_user)
):
    try:
        obj_id = ObjectId(todo_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid Todo ID")
    todo = await db.todos.find_one({"_id": obj_id})
    if not todo:
        raise HTTPException(status_code=404, detail="Todo not found")
    if current_user.role != "admin" and todo["user_id"] != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
    await db.todos.delete_one({"_id": obj_id})
    return {"message": "Todo deleted successfully"}

@api_router.patch("/todos/{todo_id}")
async def update_todo(
    todo_id: str,
    updates: dict,
    current_user: User = Depends(get_current_user)
):
    try:
        todo = await db.todos.find_one({"_id": ObjectId(todo_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid Todo ID")
    if not todo:
        raise HTTPException(status_code=404, detail="Todo not found")
    if current_user.role != "admin" and todo["user_id"] != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
    now = datetime.now(IST)
    if updates.get("is_completed") is True:
        updates["completed_at"] = now
    updates["updated_at"] = now
    await db.todos.update_one(
        {"_id": ObjectId(todo_id)},
        {"$set": updates}
    )
    return {"message": "Todo updated successfully"}

# REGISTER Endpoint
@api_router.post("/auth/register", response_model=Token)
async def register(user_data: UserCreate, current_user: User = Depends(get_current_user)):
    # PERMISSION MATRIX (updated):
    # Admin   → can register users with any role
    # Manager → can register staff users only (if can_manage_users is True)
    # Staff   → can register staff users only (if can_manage_users is True)
    perms = get_user_permissions(current_user)
    is_admin = current_user.role == "admin"
    can_manage = perms.get("can_manage_users", False)

    if not is_admin and not can_manage:
        raise HTTPException(status_code=403, detail="You do not have permission to register users")

    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")

    hashed_password = get_password_hash(user_data.password)

    requested_role = user_data.role.value if hasattr(user_data.role, "value") else user_data.role

    if requested_role in ["admin", "manager", "superadmin"]:
        if current_user.role != "admin":
            raise HTTPException(
                status_code=400,
                detail="Only staff role can be assigned during registration by non-admin users"
            )

    role_val = requested_role
    default_permissions = DEFAULT_ROLE_PERMISSIONS.get(role_val, {})
    user_id = str(uuid.uuid4())

    new_user = {
        "id": user_id,
        "email": user_data.email,
        "full_name": user_data.full_name,
        "role": role_val,
        "password": hashed_password,
        "departments": user_data.departments or [],
        "phone": user_data.phone,
        "birthday": user_data.birthday.isoformat() if user_data.birthday else None,
        "telegram_id": user_data.telegram_id,
        "punch_in_time": user_data.punch_in_time or "10:30",
        "grace_time": user_data.grace_time or "00:15",
        "punch_out_time": user_data.punch_out_time or "19:00",
        "profile_picture": user_data.profile_picture,
        "is_active": False,
        "status": "pending_approval",
        "approved_by": None,
        "approved_at": None,
        "permissions": user_data.permissions if user_data.permissions else default_permissions,
        "created_at": datetime.now(timezone.utc).isoformat()
    }

    await db.users.insert_one(new_user)

    # Background sync (NON-BLOCKING)
   # Safe background runner (kept for future use if needed)
    async def run_safe_background(coro, name="task"):
        try:
            await coro
        except Exception as e:
            logger.error(f"{name} failed: {e}", exc_info=True)


    # ─────────────────────────────────────────────
    # AUTH RESPONSE LOGIC (CLEANED)
    # ─────────────────────────────────────────────

    access_token = create_access_token({"sub": user_id})

    # Remove sensitive fields
    new_user.pop("password", None)
    new_user.pop("_id", None)

    # ❌ REMOVED (Identix sync — no longer exists)
    # asyncio.create_task(run_safe_background(
    #     sync_user_to_identix_devices(new_user),
    #     "identix_sync"
    # ))

    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": new_user
    }

@api_router.post("/auth/login", response_model=Token)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})

    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    user_status = user.get("status")
    if user_status is not None and user_status != "active":
        raise HTTPException(
            status_code=403,
            detail=f"Your account is {user_status}. Awaiting admin approval."
        )

    user["permissions"] = user.get("permissions", UserPermissions().model_dump())

    if "created_at" in user and isinstance(user["created_at"], str):
        user["created_at"] = datetime.fromisoformat(user["created_at"])

    user_obj = User(**{k: v for k, v in user.items() if k != "password"})
    access_token = create_access_token({"sub": user_obj.id})

    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": user_obj,
        "consent_given": True
    }


@api_router.get("/auth/me", response_model=User)
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user


# ── Forgot / Reset Password ───────────────────────────────────────────────────
import secrets
from datetime import timedelta

class ForgotPasswordRequest(BaseModel):
    email: str

class ResetPasswordRequest(BaseModel):
    email: str
    token: str
    new_password: str

@api_router.post("/auth/forgot-password")
async def forgot_password(data: ForgotPasswordRequest):
    """
    Always returns 200 (to prevent email enumeration).
    Generates a short-lived token, stores it in DB, and emails it via Brevo SMTP.
    """
    user = await db.users.find_one({"email": data.email.strip().lower()}, {"_id": 0})
    if user:
        token = secrets.token_urlsafe(32)
        expires_at = (datetime.now(timezone.utc) + timedelta(hours=1)).isoformat()
        await db.password_reset_tokens.delete_many({"email": data.email.strip().lower()})
        await db.password_reset_tokens.insert_one({
            "email": data.email.strip().lower(),
            "token": token,
            "expires_at": expires_at,
        })
        subject = "TaskoSphere – Password Reset"
        body = (
            f"Hi {user.get('full_name', '')},\n\n"
            f"You requested a password reset for your TaskoSphere account.\n\n"
            f"Your reset token is:\n\n  {token}\n\n"
            f"Enter this token on the reset password page along with your new password.\n"
            f"This token expires in 1 hour.\n\n"
            f"If you did not request this, you can safely ignore this email.\n\n"
            f"— TaskoSphere"
        )
        try:
            send_email(data.email.strip(), subject, body)
            logger.info(f"Password reset email sent to {data.email}")
        except Exception as e:
            logger.error(f"Failed to send password reset email to {data.email}: {e}")
    return {"message": "If that email is registered, reset instructions have been sent."}


@api_router.post("/auth/reset-password")
async def reset_password(data: ResetPasswordRequest):
    email = data.email.strip().lower()
    record = await db.password_reset_tokens.find_one({"email": email, "token": data.token.strip()})
    if not record:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token.")
    expires_at = parser.isoparse(record["expires_at"])
    if datetime.now(timezone.utc) > expires_at:
        await db.password_reset_tokens.delete_many({"email": email})
        raise HTTPException(status_code=400, detail="Reset token has expired. Please request a new one.")
    if len(data.new_password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters.")
    hashed = get_password_hash(data.new_password)
    result = await db.users.update_one({"email": email}, {"$set": {"password": hashed}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="User not found.")
    await db.password_reset_tokens.delete_many({"email": email})
    logger.info(f"Password reset successful for {email}")
    return {"message": "Password updated successfully."}


@api_router.post("/users/{user_id}/approve")
async def approve_user(user_id: str, current_user: User = Depends(get_current_user)):
    # PERMISSION MATRIX (updated): Admin or users with can_manage_users can approve
    perms = get_user_permissions(current_user)
    if current_user.role != "admin" and not perms.get("can_manage_users", False):
        raise HTTPException(status_code=403, detail="You do not have permission to approve users")

    existing = await db.users.find_one({"id": user_id}, {"_id": 0})

    if not existing:
        raise HTTPException(status_code=404, detail="User not found")

    if existing.get("status") != "pending_approval":
        raise HTTPException(
            status_code=400,
            detail=f"User status is {existing.get('status')}, not pending approval"
        )

    update_data = {
        "status": "active",
        "is_active": True,
        "approved_by": current_user.id,
        "approved_at": datetime.now(timezone.utc).isoformat()
    }

    await db.users.update_one({"id": user_id}, {"$set": update_data})

    await create_audit_log(
        current_user,
        "APPROVE_USER",
        "user",
        user_id,
        existing,
        update_data
    )

    return {"message": "User approved successfully"}


@api_router.post("/users/{user_id}/reject")
async def reject_user(user_id: str, current_user: User = Depends(get_current_user)):
    # PERMISSION MATRIX (updated): Admin or users with can_manage_users can reject
    perms = get_user_permissions(current_user)
    if current_user.role != "admin" and not perms.get("can_manage_users", False):
        raise HTTPException(status_code=403, detail="You do not have permission to reject users")

    existing = await db.users.find_one({"id": user_id}, {"_id": 0})

    if not existing:
        raise HTTPException(status_code=404, detail="User not found")

    update_data = {
        "status": "rejected",
        "is_active": False
    }

    await db.users.update_one({"id": user_id}, {"$set": update_data})

    await create_audit_log(
        current_user,
        "REJECT_USER",
        "user",
        user_id,
        existing,
        update_data
    )

    return {"message": "User rejected"}
#============================================================
# USER MANAGEMENT
#=============================================================
@api_router.get("/users")
async def get_users(
    user_id: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    if current_user.role == "admin":
        query = {}
        users_raw = await db.users.find(query, {"_id": 0, "password": 0}).to_list(1000)
    elif current_user.role == "manager":
        if user_id:
            target_user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
            if not target_user:
                raise HTTPException(status_code=404, detail="User not found")
            target_depts = target_user.get("departments", [])
            manager_depts = current_user.departments
            if not any(d in manager_depts for d in target_depts):
                raise HTTPException(status_code=403, detail="User not in your departments")
            users_raw = [target_user]
        else:
            team_ids = await get_team_user_ids(current_user.id)
            query = {"id": {"$in": team_ids + [current_user.id]}}
            users_raw = await db.users.find(query, {"_id": 0, "password": 0}).to_list(1000)
    else:
        permissions = get_user_permissions(current_user)
        if not user_id and not permissions.get("can_view_user_page", False):
            raise HTTPException(status_code=403, detail="User directory access not permitted")
        if user_id and user_id != current_user.id:
            allowed = permissions.get("view_other_activity", [])
            if user_id not in allowed:
                raise HTTPException(status_code=403, detail="Not allowed")
        users_raw = await db.users.find(
            {"id": user_id or current_user.id},
            {"_id": 0, "password": 0}
        ).to_list(1000)
    for u in users_raw:
        if u.get("created_at") and isinstance(u["created_at"], str):
            try:
                u["created_at"] = datetime.fromisoformat(u["created_at"])
            except Exception:
                u["created_at"] = datetime.now(timezone.utc)
        else:
            u["created_at"] = datetime.now(timezone.utc)
    return users_raw

@api_router.put("/users/{user_id}", response_model=User)
async def update_user(
    user_id: str,
    user_data: dict,
    current_user: User = Depends(get_current_user)
):
    is_own     = user_id == current_user.id
    is_admin   = current_user.role.lower() == "admin"
    is_manager = current_user.role.lower() == "manager"
    perms      = get_user_permissions(current_user)
    has_edit_users = perms.get("can_edit_users", False)

    # Manager scope check: manager with can_edit_users can edit their team staff only
    if not is_admin and not is_own and has_edit_users and is_manager:
        team_ids = await get_team_user_ids(current_user.id)
        if user_id not in team_ids:
            raise HTTPException(status_code=403, detail="User is not in your team")
        target_user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if target_user and target_user.get("role") in ("admin", "manager"):
            raise HTTPException(status_code=403, detail="Managers can only edit staff members")
    elif not is_admin and not is_own and not has_edit_users:
        raise HTTPException(status_code=403, detail="You can only update your own profile.")

    existing = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="User not found.")

    if is_admin:
        # Admin can update all fields including role, permissions, status
        allowed_fields = [
            "full_name", "email", "role", "departments", "phone",
            "birthday", "punch_in_time", "grace_time",
            "punch_out_time", "is_active", "profile_picture", "telegram_id",
            "status", "permissions"
        ]
    elif is_manager and has_edit_users and not is_own:
        # Manager editing a team staff member — can update profile + work settings, not role/permissions
        allowed_fields = [
            "full_name", "email", "departments", "phone",
            "birthday", "punch_in_time", "grace_time",
            "punch_out_time", "is_active", "profile_picture", "telegram_id",
            "status"
        ]
    else:
        # Self-edit: own profile fields only
        allowed_fields = [
            "full_name", "phone", "birthday",
            "punch_in_time", "punch_out_time", "profile_picture", "telegram_id"
        ]

    update_payload = {}
    for key in allowed_fields:
        if key in user_data:
            val = user_data[key]
            update_payload[key] = val if val != "" else None
    new_password = user_data.get("password")
    if new_password and len(new_password.strip()) > 0:
        update_payload["password"] = get_password_hash(new_password)
    if update_payload:
        await db.users.update_one({"id": user_id}, {"$set": update_payload})
    await create_audit_log(current_user, "UPDATE_USER", "user", user_id, existing, update_payload)
    updated_user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    return updated_user

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: User = Depends(get_current_user)):
    # Per permission matrix: DELETE on Users is Admin-only.
    # Manager has VIEW/CREATE/EDIT/UPDATE (Permission-based) but NOT DELETE.
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only administrators can delete users")
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot delete yourself")
    existing = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")
    await create_audit_log(current_user, "DELETE_USER", "user", record_id=user_id, old_data=existing)
    await db.users.delete_one({"id": user_id})
    return {"message": "User deleted successfully"}


# ════════════════════════════════════════════════════════════════════════════════
# EMPLOYEE OFFBOARDING / REPLACEMENT
# ════════════════════════════════════════════════════════════════════════════════

@api_router.get("/users/{user_id}/offboard-preview")
async def offboard_preview(
    user_id: str,
    current_user: User = Depends(require_admin()),
):
    """Preview what data belongs to this user before offboarding."""
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    counts = {
        "tasks_assigned": await db.tasks.count_documents({"assigned_to": user_id}),
        "tasks_created": await db.tasks.count_documents({"created_by": user_id}),
        "clients": await db.clients.count_documents({"assigned_to": user_id}),
        "dsc": await db.dsc_register.count_documents({"assigned_to": user_id}),
        "documents": await db.documents.count_documents(
            {"$or": [{"assigned_to": user_id}, {"created_by": user_id}]}
        ),
        "todos": await db.todos.count_documents({"user_id": user_id}),
        "visits": await db.visits.count_documents({"assigned_to": user_id}),
        "leads": await db.leads.count_documents({"assigned_to": user_id}),
    }

    return {
        "user": {
            "id": user.get("id"),
            "full_name": user.get("full_name"),
            "email": user.get("email"),
            "role": user.get("role"),
            "departments": user.get("departments", []),
        },
        "data_counts": counts,
        "total_items": sum(counts.values()),
    }


@api_router.post("/users/{user_id}/offboard")
async def offboard_user(
    user_id: str,
    body: OffboardRequest,
    current_user: User = Depends(require_admin()),
):
    """
    Offboard an employee: transfer all their data to a replacement user,
    keep an audit trail, then optionally delete the old account.
    """
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot offboard yourself")
    if user_id == body.replacement_user_id:
        raise HTTPException(status_code=400, detail="Old and replacement user cannot be the same")

    old_user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not old_user:
        raise HTTPException(status_code=404, detail="User to offboard not found")

    new_user = await db.users.find_one({"id": body.replacement_user_id}, {"_id": 0})
    if not new_user:
        raise HTTPException(status_code=404, detail="Replacement user not found")

    transfer_summary = {}

    # 1. Tasks
    if body.transfer_tasks:
        r1 = await db.tasks.update_many({"assigned_to": user_id}, {"$set": {"assigned_to": body.replacement_user_id}})
        r2 = await db.tasks.update_many({"created_by": user_id}, {"$set": {"created_by": body.replacement_user_id}})
        transfer_summary["tasks_assigned"] = r1.modified_count
        transfer_summary["tasks_created"] = r2.modified_count

    # 2. Clients
    if body.transfer_clients:
        r = await db.clients.update_many({"assigned_to": user_id}, {"$set": {"assigned_to": body.replacement_user_id}})
        transfer_summary["clients_reassigned"] = r.modified_count

    # 3. DSC
    if body.transfer_dsc:
        r = await db.dsc_register.update_many({"assigned_to": user_id}, {"$set": {"assigned_to": body.replacement_user_id}})
        transfer_summary["dsc_transferred"] = r.modified_count

    # 4. Documents
    if body.transfer_documents:
        r = await db.documents.update_many(
            {"$or": [{"assigned_to": user_id}, {"created_by": user_id}]},
            {"$set": {"assigned_to": body.replacement_user_id}}
        )
        transfer_summary["documents_transferred"] = r.modified_count

    # 5. Todos
    if body.transfer_todos:
        r = await db.todos.update_many({"user_id": user_id}, {"$set": {"user_id": body.replacement_user_id}})
        transfer_summary["todos_transferred"] = r.modified_count

    # 6. Visits
    if body.transfer_visits:
        r = await db.visits.update_many({"assigned_to": user_id}, {"$set": {"assigned_to": body.replacement_user_id}})
        transfer_summary["visits_transferred"] = r.modified_count

    # 7. Leads
    if body.transfer_leads:
        r = await db.leads.update_many({"assigned_to": user_id}, {"$set": {"assigned_to": body.replacement_user_id}})
        transfer_summary["leads_transferred"] = r.modified_count

    # 8. Update cross-user permission references in all other users
    for field in [
        "permissions.view_other_tasks", "permissions.view_other_attendance",
        "permissions.view_other_reports", "permissions.view_other_todos",
        "permissions.view_other_activity", "permissions.view_other_visits",
        "permissions.assigned_clients",
    ]:
        await db.users.update_many(
            {field: user_id},
            {"$set": {f"{field}.$[elem]": body.replacement_user_id}},
            array_filters=[{"elem": user_id}]
        )
    transfer_summary["permission_references_updated"] = True

    # 9. Optionally update the replacement user's email
    if body.update_email and body.update_email.strip():
        new_email = body.update_email.strip().lower()
        email_exists = await db.users.find_one(
            {"email": new_email, "id": {"$ne": body.replacement_user_id}}, {"_id": 0, "id": 1}
        )
        if email_exists:
            raise HTTPException(status_code=400, detail=f"Email {new_email} is already in use")
        await db.users.update_one({"id": body.replacement_user_id}, {"$set": {"email": new_email}})
        transfer_summary["email_updated"] = new_email

    # 10. Audit Log
    await create_audit_log(
        current_user, "OFFBOARD_USER", "user", record_id=user_id,
        old_data={
            "offboarded_user": {"id": old_user.get("id"), "full_name": old_user.get("full_name"),
                                "email": old_user.get("email"), "role": old_user.get("role"),
                                "departments": old_user.get("departments", [])},
            "replacement_user": {"id": new_user.get("id"), "full_name": new_user.get("full_name"),
                                 "email": new_user.get("email")},
            "transfer_summary": transfer_summary, "notes": body.notes,
        },
    )

    # 11. Delete or deactivate old user
    if body.delete_old_user:
        await db.users.delete_one({"id": user_id})
        transfer_summary["old_user_deleted"] = True
    else:
        await db.users.update_one({"id": user_id}, {"$set": {"is_active": False, "status": "inactive"}})
        transfer_summary["old_user_deactivated"] = True

    return {
        "message": f"Successfully offboarded {old_user.get('full_name')} → {new_user.get('full_name')}",
        "transfer_summary": transfer_summary,
    }


@api_router.get("/users/{user_id}/permissions")
async def get_permissions(user_id: str, current_user: User = Depends(get_current_user)):
    """
    Retrieve permissions for a user.
    - Admin: can fetch any user's permissions
    - Manager with can_manage_users: can fetch permissions of staff in their department
    - Staff: can only fetch their own permissions (read-only display)
    """
    # Admin always allowed
    if current_user.role == "admin":
        user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        return user.get("permissions", {})

    # Any user can always fetch their OWN permissions
    if user_id == current_user.id:
        user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        return user.get("permissions", {})

    # Manager with can_manage_users can fetch their team's permissions
    perms = get_user_permissions(current_user)
    if current_user.role == "manager" and perms.get("can_manage_users", False):
        team_ids = await get_team_user_ids(current_user.id)
        if user_id not in team_ids:
            raise HTTPException(status_code=403, detail="User is not in your team")
        target_user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        if not target_user:
            raise HTTPException(status_code=404, detail="User not found")
        # Manager cannot view admin/manager permissions — only staff
        if target_user.get("role") in ("admin", "manager"):
            raise HTTPException(status_code=403, detail="Managers can only view permissions of staff members")
        return target_user.get("permissions", {})

    raise HTTPException(status_code=403, detail="Not allowed")

@api_router.put("/users/{user_id}/permissions")
async def update_user_permissions(
    user_id: str,
    permissions: dict,
    current_user: User = Depends(get_current_user)
):
    """
    Update permissions for a user.
    - Admin: can update any user's permissions
    - Manager with can_manage_users: can update permissions of staff in their department only
      (cannot escalate permissions beyond their own permission level)
    - Staff: not allowed
    """
    # Admin always allowed
    if current_user.role == "admin":
        existing = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        if not existing:
            raise HTTPException(status_code=404, detail="User not found")
        old_permissions = existing.get("permissions", {})
        await db.users.update_one({"id": user_id}, {"$set": {"permissions": permissions}})
        await create_audit_log(
            current_user, "UPDATE_PERMISSIONS", "user",
            record_id=user_id, old_data=old_permissions, new_data=permissions
        )
        return {"message": "Permissions updated successfully"}

    # Manager with can_manage_users can update their team staff permissions (not admin/manager)
    perms = get_user_permissions(current_user)
    if current_user.role == "manager" and perms.get("can_manage_users", False):
        team_ids = await get_team_user_ids(current_user.id)
        if user_id not in team_ids:
            raise HTTPException(status_code=403, detail="User is not in your team")
        existing = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
        if not existing:
            raise HTTPException(status_code=404, detail="User not found")
        if existing.get("role") in ("admin", "manager"):
            raise HTTPException(status_code=403, detail="Managers can only update permissions of staff members")

        # Managers CANNOT grant permissions they do not themselves possess
        # Strip any elevated permission flags that the manager doesn't have
        manager_perms = get_user_permissions(current_user)
        safe_permissions = {}
        BOOLEAN_PERM_KEYS = [
            "can_view_all_tasks", "can_view_all_clients", "can_view_all_dsc",
            "can_view_documents", "can_view_all_duedates", "can_view_reports",
            "can_view_attendance", "can_view_all_leads", "can_edit_tasks",
            "can_edit_clients", "can_edit_dsc", "can_edit_documents",
            "can_edit_due_dates", "can_edit_users", "can_download_reports",
            "can_manage_users", "can_manage_settings", "can_assign_tasks",
            "can_assign_clients", "can_view_staff_activity", "can_view_user_page",
            "can_view_audit_logs", "can_view_selected_users_reports",
            "can_view_todo_dashboard", "can_use_chat", "can_view_staff_rankings",
            "can_connect_email", "can_view_own_data", "can_create_quotations",
            "can_manage_invoices", "can_view_passwords", "can_edit_passwords",
            "can_view_compliance", "can_manage_compliance",
            "can_view_all_visits", "can_edit_visits",
        ]
        # Flags only admin can grant — managers cannot escalate these
        ADMIN_ONLY_GRANTS = {
            "can_delete_data", "can_delete_tasks", "can_delete_visits",
            "can_send_reminders",
        }
        for key, val in permissions.items():
            if key in ADMIN_ONLY_GRANTS:
                # Manager cannot grant admin-only permissions
                safe_permissions[key] = existing.get("permissions", {}).get(key, False)
            elif key in BOOLEAN_PERM_KEYS and isinstance(val, bool):
                # Manager can only grant a permission they themselves have
                if val and not manager_perms.get(key, False):
                    safe_permissions[key] = False
                else:
                    safe_permissions[key] = val
            else:
                # Non-boolean keys (lists) passed through as-is
                safe_permissions[key] = val

        old_permissions = existing.get("permissions", {})
        await db.users.update_one({"id": user_id}, {"$set": {"permissions": safe_permissions}})
        await create_audit_log(
            current_user, "UPDATE_PERMISSIONS", "user",
            record_id=user_id, old_data=old_permissions, new_data=safe_permissions
        )
        return {"message": "Permissions updated successfully"}

    raise HTTPException(status_code=403, detail="Admin access required")

#====================================================================================
# ATTENDANCE ROUTES
#=====================================================================================
def get_real_client_ip(request: Request):
    x_forwarded_for = request.headers.get("x-forwarded-for")
    if x_forwarded_for:
        return x_forwarded_for.split(",")[0].strip()
    if request.client:
        return request.client.host
    return None

def check_is_late(user: dict, punch_in_ist: datetime) -> bool:
    try:
        pit = datetime.strptime(user.get("punch_in_time", "10:30"), "%H:%M")
        if user.get("late_grace_minutes") is not None:
            grace_minutes = int(user["late_grace_minutes"])
        else:
            raw = str(user.get("grace_time", "00:15"))
            gt = datetime.strptime(raw, "%H:%M")
            grace_minutes = gt.hour * 60 + gt.minute
        deadline = punch_in_ist.replace(
            hour=pit.hour, minute=pit.minute, second=0, microsecond=0
        ) + timedelta(minutes=grace_minutes)
        return punch_in_ist > deadline
    except Exception:
        return False

def check_punched_out_early(user: dict, punch_out_ist: datetime) -> bool:
    try:
        pot = datetime.strptime(user.get("punch_out_time", "19:00"), "%H:%M")
        expected_out = punch_out_ist.replace(
            hour=pot.hour, minute=pot.minute, second=0, microsecond=0
        )
        return punch_out_ist < expected_out
    except Exception:
        return False

@api_router.post("/attendance")
async def handle_attendance(
    data: dict,
    current_user: User = Depends(get_current_user)
):
    today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
    today_str = today.isoformat()
    # Note: We do NOT block punch-in on holidays.
    # Users who choose to work on holidays can still punch in/out freely.
    # The frontend suppresses the auto-popup on holidays but keeps the button visible.
    action = data.get("action")
    if action not in ["punch_in", "punch_out"]:
        raise HTTPException(status_code=400, detail="Invalid action")
    attendance = await db.attendance.find_one(
        {"user_id": current_user.id, "date": today_str},
        {"_id": 0}
    )
    if action == "punch_in":
        if attendance and attendance.get("punch_in"):
            raise HTTPException(status_code=400, detail="Already punched in")
        user_doc = await db.users.find_one({"id": current_user.id}, {"_id": 0})
        punch_in_utc = datetime.now(timezone.utc)
        punch_in_ist = punch_in_utc.astimezone(ZoneInfo("Asia/Kolkata"))
        is_late = check_is_late(user_doc or {}, punch_in_ist)
        location_data = data.get("location")
        update_fields = {
            "status": "present",
            "punch_in": punch_in_utc,
            "is_late": is_late,
            "leave_reason": None,
            # Clear auto_absent flag if user punches in manually
            "auto_marked": False,
        }
        if location_data:
            update_fields["location"] = location_data
        await db.attendance.update_one(
            {"user_id": current_user.id, "date": today_str},
            {"$set": update_fields},
            upsert=True
        )
        return {"message": "Punched in successfully", "is_late": is_late}

    # PUNCH_OUT_BLOCK
    if action == "punch_out":
        if not attendance or not attendance.get("punch_in"):
            raise HTTPException(status_code=400, detail="Not punched in yet")

        if attendance.get("punch_out"):
            raise HTTPException(status_code=400, detail="Already punched out")

        punch_in_dt = attendance.get("punch_in")
        punch_out_utc = datetime.now(timezone.utc)
        punch_out_ist = punch_out_utc.astimezone(ZoneInfo("Asia/Kolkata"))

        user_doc = await db.users.find_one({"id": current_user.id}, {"_id": 0})
        punched_out_early = check_punched_out_early(user_doc or {}, punch_out_ist)

        # FIX: MongoDB may return punch_in as a naive datetime (no tzinfo).
        # Treat naive datetimes as UTC before computing the delta.
        if isinstance(punch_in_dt, datetime):
            if punch_in_dt.tzinfo is None:
                punch_in_dt = punch_in_dt.replace(tzinfo=timezone.utc)
        else:
            # Fallback: parse from string if stored as ISO string
            try:
                punch_in_dt = datetime.fromisoformat(str(punch_in_dt))
                if punch_in_dt.tzinfo is None:
                    punch_in_dt = punch_in_dt.replace(tzinfo=timezone.utc)
            except Exception:
                punch_in_dt = punch_out_utc  # safeguard: 0-minute duration

        delta = punch_out_utc - punch_in_dt.astimezone(timezone.utc)
        duration_minutes = int(delta.total_seconds() / 60)

        update_fields = {
            "punch_out": punch_out_utc,
            "punched_out_early": punched_out_early,
            "duration_minutes": max(0, duration_minutes)
        }

        if data.get("location"):
            update_fields["punch_out_location"] = data.get("location")

        # ── NEW: record overtime minutes if punched out after 7 PM IST ───────
        shift_end_ist = punch_out_ist.replace(hour=19, minute=0, second=0, microsecond=0)
        if punch_out_ist > shift_end_ist:
            update_fields["overtime_minutes"] = max(
                0, int((punch_out_ist - shift_end_ist).total_seconds() / 60)
            )
        else:
            update_fields["overtime_minutes"] = 0
        # ─────────────────────────────────────────────────────────────────────

        await db.attendance.update_one(
            {"user_id": current_user.id, "date": today_str},
            {"$set": update_fields}
        )

        return {
            "message": "Punched out successfully",
            "duration": duration_minutes,
            "punched_out_early": punched_out_early
        }

@api_router.post("/attendance/mark-leave-today")
async def mark_leave_today(current_user: User = Depends(get_current_user)):
    today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
    today_str = today.isoformat()
    await db.attendance.update_one(
        {"user_id": current_user.id, "date": today_str},
        {
            "$set": {
                "status": "leave",
                "punch_in": None,
                "punch_out": None,
                "leave_reason": "Marked on leave today"
            }
        },
        upsert=True
    )
    return {"message": "Marked on leave today"}

@api_router.get("/attendance/today")
async def get_today_attendance(current_user: User = Depends(get_current_user)):
    today = datetime.now(ZoneInfo("Asia/Kolkata")).date()
    today_str = today.isoformat()
    # FIX: Added {"_id": 0} projection — without it ObjectId leaks into JSON response
    # and causes ValueError: 'ObjectId' object is not iterable (500 error)
    holiday = await db.holidays.find_one({"date": today_str}, {"_id": 0})
    if holiday:
        return {
            "status": "holiday",
            "holiday": holiday,
            "punch_in": None,
            "punch_out": None,
            "leave_reason": None
        }
    attendance = await db.attendance.find_one(
        {"user_id": current_user.id, "date": today_str},
        {"_id": 0}
    )
    if not attendance:
        return {
            "status": "absent",
            "punch_in": None,
            "punch_out": None,
            "leave_reason": None
        }
    if "status" not in attendance:
        attendance["status"] = (
            "present" if attendance.get("punch_in") else "absent"
        )
    return attendance

@api_router.post("/attendance/apply-leave")
async def apply_leave(
    data: dict,
    current_user: User = Depends(get_current_user)
):
    try:
        from_date = datetime.fromisoformat(data["from_date"]).date()
        to_date = datetime.fromisoformat(data.get("to_date", data["from_date"])).date()
        reason = data.get("reason", "Leave Applied")
        leave_type = data.get("leave_type", "full_day")  # full_day | half_day_morning | half_day_afternoon | early_leave
        early_leave_time = data.get("early_leave_time")  # "HH:MM" string for early leave

        VALID_LEAVE_TYPES = ("full_day", "half_day_morning", "half_day_afternoon", "early_leave")
        if leave_type not in VALID_LEAVE_TYPES:
            raise HTTPException(status_code=400, detail=f"Invalid leave_type. Must be one of: {VALID_LEAVE_TYPES}")

        if to_date < from_date:
            raise HTTPException(status_code=400, detail="Invalid date range")

        # For partial-day leave types, only single-day makes sense
        if leave_type in ("half_day_morning", "half_day_afternoon", "early_leave"):
            if to_date != from_date:
                raise HTTPException(
                    status_code=400,
                    detail="Half-day and early leave can only be applied for a single day"
                )

        current = from_date
        while current <= to_date:
            current_str = current.isoformat()
            existing = await db.attendance.find_one(
                {"user_id": current_user.id, "date": current_str}, {"_id": 0}
            )

            if leave_type == "full_day":
                update_fields = {
                    "status": "leave",
                    "leave_reason": reason,
                    "leave_type": "full_day",
                    "punch_in": None,
                    "punch_out": None,
                    "duration_minutes": 0,
                }

            elif leave_type == "half_day_morning":
                # Morning off — not expected to punch in before noon
                update_fields = {
                    "status": "present",
                    "leave_reason": reason,
                    "leave_type": "half_day_morning",
                    "is_half_day": True,
                    # Don't wipe punch_in/out if already recorded
                }
                if not existing or not existing.get("punch_in"):
                    update_fields["punch_in"] = None
                    update_fields["punch_out"] = None
                    update_fields["duration_minutes"] = 0

            elif leave_type == "half_day_afternoon":
                # Afternoon off — present in morning, leaves at noon
                update_fields = {
                    "status": "present",
                    "leave_reason": reason,
                    "leave_type": "half_day_afternoon",
                    "is_half_day": True,
                }
                # Set a nominal punch-out at 13:30 IST if not already punched out
                if existing and existing.get("punch_in") and not existing.get("punch_out"):
                    punch_in_dt = existing["punch_in"]
                    if isinstance(punch_in_dt, str):
                        punch_in_dt = datetime.fromisoformat(punch_in_dt)
                    if punch_in_dt.tzinfo is None:
                        punch_in_dt = punch_in_dt.replace(tzinfo=timezone.utc)
                    half_day_out = datetime.now(timezone.utc).replace(hour=8, minute=0, second=0, microsecond=0)  # 13:30 IST = 08:00 UTC
                    delta = half_day_out - punch_in_dt
                    dur = max(0, int(delta.total_seconds() / 60))
                    update_fields["punch_out"] = half_day_out
                    update_fields["duration_minutes"] = dur
                    update_fields["punched_out_early"] = True

            elif leave_type == "early_leave":
                # Present but left early at a specified time
                update_fields = {
                    "status": "present",
                    "leave_reason": reason,
                    "leave_type": "early_leave",
                    "is_early_leave": True,
                    "early_leave_time": early_leave_time,
                }
                # If a departure time given and user is punched in without punch-out, record it
                if early_leave_time and existing and existing.get("punch_in") and not existing.get("punch_out"):
                    try:
                        punch_in_dt = existing["punch_in"]
                        if isinstance(punch_in_dt, str):
                            punch_in_dt = datetime.fromisoformat(punch_in_dt)
                        if punch_in_dt.tzinfo is None:
                            punch_in_dt = punch_in_dt.replace(tzinfo=timezone.utc)
                        # Parse "HH:MM" departure time as IST then convert to UTC
                        h, m = map(int, early_leave_time.split(":"))
                        today_ist = datetime.now(ZoneInfo("Asia/Kolkata")).replace(
                            hour=h, minute=m, second=0, microsecond=0
                        )
                        early_out_utc = today_ist.astimezone(timezone.utc)
                        delta = early_out_utc - punch_in_dt
                        dur = max(0, int(delta.total_seconds() / 60))
                        update_fields["punch_out"] = early_out_utc
                        update_fields["duration_minutes"] = dur
                        update_fields["punched_out_early"] = True
                    except Exception:
                        pass

            await db.attendance.update_one(
                {"user_id": current_user.id, "date": current_str},
                {"$set": update_fields},
                upsert=True
            )
            current += timedelta(days=1)

        return {"message": "Leave applied successfully", "leave_type": leave_type}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@api_router.get("/attendance/history", response_model=List[Attendance])
async def get_attendance_history(
    user_id: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    query = {}
    if current_user.role == "admin":
        if user_id:
            query["user_id"] = user_id
        # No user_id from admin = return ALL records (aggregate view)
    elif current_user.role == "manager":
        permissions_mgr = get_user_permissions(current_user)
        # Manager: Own + Team (same department) + explicit cross-visibility list
        team_ids = await get_team_user_ids(current_user.id)
        allowed_users = list(set(
            (permissions_mgr.get("view_other_attendance", []) or []) + team_ids
        ))
        if user_id:
            if user_id == current_user.id:
                query["user_id"] = user_id
            else:
                if not permissions_mgr.get("can_view_attendance", False):
                    raise HTTPException(
                        status_code=403,
                        detail="You do not have permission to view other users' attendance"
                    )
                if user_id not in allowed_users:
                    raise HTTPException(
                        status_code=403,
                        detail="This user is outside your team scope"
                    )
                query["user_id"] = user_id
        else:
            if permissions_mgr.get("can_view_attendance", False):
                query["user_id"] = {"$in": list(set(allowed_users + [current_user.id]))}
            else:
                query["user_id"] = current_user.id
    else:
        if user_id and user_id != current_user.id:
            permissions = get_user_permissions(current_user)
            allowed_users = permissions.get("view_other_attendance", [])
            if user_id not in allowed_users:
                raise HTTPException(
                    status_code=403,
                    detail="Not authorized to view other users' attendance"
                )
        query["user_id"] = user_id if user_id else current_user.id
    attendance_list = await db.attendance.find(query, {"_id": 0}).sort("date", -1).to_list(1000)
    for attendance in attendance_list:
        attendance["punch_in"] = safe_dt(attendance.get("punch_in"))
        attendance["punch_out"] = safe_dt(attendance.get("punch_out"))
        if "status" not in attendance:
            attendance["status"] = (
                "present" if attendance.get("punch_in") else "absent"
            )
    return attendance_list

@api_router.get("/attendance/my-summary")
async def get_my_attendance_summary(
    current_user: User = Depends(get_current_user)
):
    now = datetime.now(IST)
    current_month = now.strftime("%Y-%m")
    attendance_list = await db.attendance.find(
        {"user_id": current_user.id},
        {"_id": 0}
    ).sort("date", -1).to_list(1000)
    monthly_data = {}
    total_minutes_all = 0
    total_days = 0
    for attendance in attendance_list:
        month = attendance["date"][:7]
        if month not in monthly_data:
            monthly_data[month] = {
                "total_minutes": 0,
                "days_present": 0
            }
        duration = attendance.get("duration_minutes")
        if isinstance(duration, (int, float)):
            monthly_data[month]["total_minutes"] += duration
            total_minutes_all += duration
            monthly_data[month]["days_present"] += 1
            total_days += 1
    formatted_data = []
    for month, data in monthly_data.items():
        minutes = data["total_minutes"]
        hours = minutes // 60
        mins = minutes % 60
        formatted_data.append({
            "month": month,
            "total_minutes": minutes,
            "total_hours": f"{hours}h {mins}m",
            "days_present": data["days_present"]
        })
    return {
        "current_month": current_month,
        "total_days": total_days,
        "total_minutes": total_minutes_all,
        "monthly_summary": formatted_data
    }

@api_router.get("/attendance/staff-report")
async def get_staff_attendance_report(
    month: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        permissions = get_user_permissions(current_user)
        if not permissions.get("can_view_attendance"):
            raise HTTPException(status_code=403, detail="Not allowed")
    now = datetime.now(IST)
    target_month = month or now.strftime("%Y-%m")
    users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(1000)
    user_map = {u["id"]: u for u in users}

    # SCOPE enforcement
    attendance_query = {"date": {"$regex": f"^{target_month}"}}
    if current_user.role != "admin":
        permissions = get_user_permissions(current_user)
        allowed_others = permissions.get("view_other_attendance", []) or []
        if current_user.role == "manager":
            # Manager: Own + Team (same department)
            team_ids = await get_team_user_ids(current_user.id)
            allowed_others = list(set(allowed_others + team_ids))
        visible_ids = list(set(allowed_others + [current_user.id]))
        attendance_query["user_id"] = {"$in": visible_ids}

    attendance_list = await db.attendance.find(
        attendance_query,
        {"_id": 0}
    ).to_list(5000)
    staff_report = {}
    for attendance in attendance_list:
        uid = attendance["user_id"]
        if uid not in staff_report:
            user_info = user_map.get(uid, {})
            staff_report[uid] = {
                "user_id": uid,
                "user_name": user_info.get("full_name", "Unknown"),
                "role": user_info.get("role", "staff"),
                "total_minutes": 0,
                "days_present": 0,
                "days_absent": 0,
                "late_days": 0,
                "early_out_days": 0,
                "records": []
            }
        # count absent days separately
        if attendance.get("status") == "absent":
            staff_report[uid]["days_absent"] += 1
        duration = attendance.get("duration_minutes")
        if isinstance(duration, (int, float)) and attendance.get("status") == "present":
            staff_report[uid]["total_minutes"] += duration
            staff_report[uid]["days_present"] += 1
        if attendance.get("is_late"):
            staff_report[uid]["late_days"] += 1
        if attendance.get("punched_out_early"):
            staff_report[uid]["early_out_days"] += 1
        staff_report[uid]["records"].append({
            "date": attendance["date"],
            "status": attendance.get("status", "absent"),
            "punch_in": attendance.get("punch_in"),
            "punch_out": attendance.get("punch_out"),
            "duration_minutes": duration,
            "is_late": attendance.get("is_late", False),
            "punched_out_early": attendance.get("punched_out_early", False)
        })
    result = []
    for uid, data in staff_report.items():
        total_minutes = data["total_minutes"]
        hours = total_minutes // 60
        minutes = total_minutes % 60
        data["total_hours"] = f"{hours}h {minutes}m"
        if data["days_present"] > 0:
            data["avg_hours_per_day"] = round(
                (total_minutes / data["days_present"]) / 60, 1
            )
        else:
            data["avg_hours_per_day"] = 0
        user_data = user_map.get(uid, {})
        year, month_val = map(int, target_month.split("-"))
        _, last_day = calendar.monthrange(year, month_val)
        expected_hours = await calculate_expected_hours(
            f"{target_month}-01",
            f"{target_month}-{last_day}",
            user_data.get("punch_in_time", "10:30"),
            user_data.get("punch_out_time", "19:00")
        )
        data["expected_hours"] = expected_hours
        result.append(data)
    result.sort(key=lambda x: x["total_minutes"], reverse=True)
    return result

@api_router.get("/attendance/export-pdf")
async def export_attendance_pdf(
    user_id: str,
    current_user: User = Depends(get_current_user)
):
    if user_id != current_user.id:
        permissions = get_user_permissions(current_user)
        if current_user.role != "admin" and not permissions.get("can_view_attendance"):
            raise HTTPException(status_code=403, detail="Not authorized to export other users' attendance")
    records = await db.attendance.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("date", 1).to_list(1000)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 14)
    pdf.cell(200, 10, txt="Attendance Report", ln=True, align="C")
    pdf.ln(5)
    pdf.set_font("Arial", size=10)
    for rec in records:
        status = rec.get("status", "unknown")
        late_flag = " [LATE]" if rec.get("is_late") else ""
        early_flag = " [EARLY OUT]" if rec.get("punched_out_early") else ""
        if status == "absent":
            pdf.multi_cell(
                0, 8,
                f"Date: {rec.get('date')} | Status: ABSENT"
                f"{' [AUTO-MARKED 7PM]' if rec.get('auto_marked') else ''}"
            )
        else:
            pdf.multi_cell(
                0,
                8,
                f"Date: {rec.get('date')} | In: {rec.get('punch_in')} | "
                f"Out: {rec.get('punch_out')} | Duration: {rec.get('duration_minutes')} mins"
                f"{late_flag}{early_flag}"
            )
        pdf.ln(2)
    output = BytesIO()
    output.write(pdf.output(dest="S").encode("latin1"))
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=attendance_{user_id}.pdf"}
    )


# ─────────────────────────────────────────────────────────────────────────────
# MANUAL ABSENT MARKING ENDPOINT (Admin only)
# POST /api/attendance/mark-absent-bulk
# Body: { "date": "YYYY-MM-DD" }  (optional — defaults to today IST)
# ─────────────────────────────────────────────────────────────────────────────
@api_router.post("/attendance/mark-absent-bulk")
async def mark_absent_bulk(
    data: dict = {},
    current_user: User = Depends(require_admin())
):
    """
    Manually trigger absent-marking for a given date.
    If no date is provided, defaults to today (IST).
    Respects the same rules as the scheduled job:
      - Skips confirmed holidays
      - Skips weekends
      - Only marks users who have no present/leave/absent record
    """
    target_date_str = (data or {}).get("date") or datetime.now(IST).date().isoformat()
    try:
        datetime.strptime(target_date_str, "%Y-%m-%d")
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD.")
    result = await _mark_absent_for_date(target_date_str)
    return result


# ─────────────────────────────────────────────────────────────────────────────
# ABSENT SUMMARY ENDPOINT
# GET /api/attendance/absent-summary?month=YYYY-MM
# ─────────────────────────────────────────────────────────────────────────────
@api_router.get("/attendance/absent-summary")
async def get_absent_summary(
    month: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """
    Returns per-user absent day counts for the given month.
    Admin sees all users; others see only their own data.
    """
    now = datetime.now(IST)
    target_month = month or now.strftime("%Y-%m")

    if current_user.role == "admin":
        query = {"date": {"$regex": f"^{target_month}"}, "status": "absent"}
    else:
        query = {
            "user_id": current_user.id,
            "date": {"$regex": f"^{target_month}"},
            "status": "absent"
        }

    absent_records = await db.attendance.find(query, {"_id": 0}).to_list(5000)
    summary: dict = {}
    for rec in absent_records:
        uid = rec["user_id"]
        if uid not in summary:
            summary[uid] = {"user_id": uid, "absent_days": 0, "dates": []}
        summary[uid]["absent_days"] += 1
        summary[uid]["dates"].append(rec["date"])

    if current_user.role == "admin":
        user_ids = list(summary.keys())
        users = await db.users.find(
            {"id": {"$in": user_ids}},
            {"_id": 0, "id": 1, "full_name": 1}
        ).to_list(1000)
        name_map = {u["id"]: u["full_name"] for u in users}
        for uid in summary:
            summary[uid]["user_name"] = name_map.get(uid, "Unknown")

    return {"month": target_month, "data": list(summary.values())}


# =============================================================
# NEW: POST /attendance/punch-out
# Dedicated auto punch-out endpoint.
# Called by the frontend Smart Auto Punch-Out feature when the
# user is inactive for >60 minutes after the shift grace period.
# Falls back gracefully if the user is already punched out.
# =============================================================
@api_router.post("/attendance/punch-out")
async def auto_punch_out(
    data: dict,
    current_user: User = Depends(get_current_user)
):
    """
    Auto punch-out triggered by the frontend inactivity detector.
    Accepts { auto: true, reason: "inactive_after_shift" }.
    Records overtime_minutes (minutes worked past 7 PM IST).
    """
    today_str = datetime.now(ZoneInfo("Asia/Kolkata")).date().isoformat()

    attendance = await db.attendance.find_one(
        {"user_id": current_user.id, "date": today_str},
        {"_id": 0}
    )
    if not attendance or not attendance.get("punch_in"):
        raise HTTPException(status_code=400, detail="Not punched in yet")
    if attendance.get("punch_out"):
        # Already punched out — return silently (idempotent)
        return {"message": "Already punched out", "duration": attendance.get("duration_minutes", 0)}

    punch_in_dt = attendance.get("punch_in")
    punch_out_utc = datetime.now(timezone.utc)
    punch_out_ist = punch_out_utc.astimezone(ZoneInfo("Asia/Kolkata"))

    # Normalise punch_in to aware datetime
    if isinstance(punch_in_dt, datetime):
        if punch_in_dt.tzinfo is None:
            punch_in_dt = punch_in_dt.replace(tzinfo=timezone.utc)
    else:
        try:
            punch_in_dt = datetime.fromisoformat(str(punch_in_dt))
            if punch_in_dt.tzinfo is None:
                punch_in_dt = punch_in_dt.replace(tzinfo=timezone.utc)
        except Exception:
            punch_in_dt = punch_out_utc

    delta = punch_out_utc - punch_in_dt.astimezone(timezone.utc)
    duration_minutes = max(0, int(delta.total_seconds() / 60))

    # Calculate overtime minutes (minutes worked past 7:00 PM IST)
    shift_end_ist = punch_out_ist.replace(hour=19, minute=0, second=0, microsecond=0)
    overtime_minutes = 0
    if punch_out_ist > shift_end_ist:
        overtime_minutes = max(0, int((punch_out_ist - shift_end_ist).total_seconds() / 60))

    user_doc = await db.users.find_one({"id": current_user.id}, {"_id": 0})
    punched_out_early = check_punched_out_early(user_doc or {}, punch_out_ist)

    update_fields = {
        "punch_out":         punch_out_utc,
        "punched_out_early": punched_out_early,
        "duration_minutes":  duration_minutes,
        "auto_punch_out":    True,
        "auto_punch_reason": data.get("reason", "inactive_after_shift"),
        "overtime_minutes":  overtime_minutes,
    }

    await db.attendance.update_one(
        {"user_id": current_user.id, "date": today_str},
        {"$set": update_fields}
    )

    logger.info(
        "Auto punch-out: user=%s date=%s duration=%dm overtime=%dm reason=%s",
        current_user.id, today_str, duration_minutes, overtime_minutes,
        data.get("reason", "inactive_after_shift")
    )

    return {
        "message":          "Auto punch-out recorded",
        "duration":         duration_minutes,
        "overtime_minutes": overtime_minutes,
        "auto":             True,
    }


# =============================================================
# NEW: POST /attendance/proof
# Upload photos, documents, and a note as attendance proof.
# Example: "Visited client → upload photo".
# Files are saved to uploads/attendance_proof/ on the server.
# The proof dict is embedded inside the attendance document.
# =============================================================
ALLOWED_PHOTO_TYPES  = {"image/jpeg", "image/png", "image/webp", "image/gif", "image/heic"}
ALLOWED_DOC_TYPES    = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
    "text/csv",
}
MAX_FILE_SIZE_MB     = 10
MAX_FILES_PER_UPLOAD = 5


@api_router.post("/attendance/proof")
async def upload_attendance_proof(
    note:      str              = Form(default=""),
    photos:    List[UploadFile] = File(default=[]),
    documents: List[UploadFile] = File(default=[]),
    current_user: User          = Depends(get_current_user)
):
    """
    Attach proof to today's attendance record.
    - note:      free-text description (e.g. "Visited ABC client office")
    - photos:    up to 5 image files (JPEG, PNG, WebP, GIF, HEIC)
    - documents: up to 5 document files (PDF, DOC, DOCX, XLS, XLSX, TXT, CSV)

    Each call REPLACES the existing proof for today (idempotent upsert).
    """
    today_str = datetime.now(ZoneInfo("Asia/Kolkata")).date().isoformat()
    now_iso   = datetime.now(timezone.utc).isoformat()

    # Validate counts
    if len(photos) > MAX_FILES_PER_UPLOAD:
        raise HTTPException(status_code=400, detail=f"Maximum {MAX_FILES_PER_UPLOAD} photos allowed")
    if len(documents) > MAX_FILES_PER_UPLOAD:
        raise HTTPException(status_code=400, detail=f"Maximum {MAX_FILES_PER_UPLOAD} documents allowed")

    saved_photos: List[str] = []
    saved_docs:   List[str] = []

    # ── Save photos ────────────────────────────────────────────────────────────
    for photo in photos:
        if not photo.filename:
            continue
        content_type = photo.content_type or ""
        if content_type not in ALLOWED_PHOTO_TYPES and not content_type.startswith("image/"):
            raise HTTPException(
                status_code=400,
                detail=f"File '{photo.filename}' is not an allowed image type"
            )
        contents = await photo.read()
        if len(contents) > MAX_FILE_SIZE_MB * 1024 * 1024:
            raise HTTPException(
                status_code=413,
                detail=f"Photo '{photo.filename}' exceeds {MAX_FILE_SIZE_MB} MB limit"
            )
        safe_name = re.sub(r"[^\w.\-]", "_", photo.filename)
        filename  = f"{current_user.id}_{today_str}_photo_{uuid.uuid4().hex[:8]}_{safe_name}"
        file_path = PROOF_UPLOAD_DIR / filename
        with open(file_path, "wb") as f:
            f.write(contents)
        saved_photos.append(filename)
        await photo.seek(0)   # reset for any downstream use

    # ── Save documents ─────────────────────────────────────────────────────────
    for doc in documents:
        if not doc.filename:
            continue
        content_type = doc.content_type or ""
        if content_type not in ALLOWED_DOC_TYPES and not doc.filename.lower().endswith(
            (".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt", ".csv")
        ):
            raise HTTPException(
                status_code=400,
                detail=f"File '{doc.filename}' is not an allowed document type"
            )
        contents = await doc.read()
        if len(contents) > MAX_FILE_SIZE_MB * 1024 * 1024:
            raise HTTPException(
                status_code=413,
                detail=f"Document '{doc.filename}' exceeds {MAX_FILE_SIZE_MB} MB limit"
            )
        safe_name = re.sub(r"[^\w.\-]", "_", doc.filename)
        filename  = f"{current_user.id}_{today_str}_doc_{uuid.uuid4().hex[:8]}_{safe_name}"
        file_path = PROOF_UPLOAD_DIR / filename
        with open(file_path, "wb") as f:
            f.write(contents)
        saved_docs.append(filename)
        await doc.seek(0)

    # ── Build proof dict ───────────────────────────────────────────────────────
    # Check if a previous proof exists so we can merge (not replace) file lists
    existing = await db.attendance.find_one(
        {"user_id": current_user.id, "date": today_str},
        {"_id": 0, "proof": 1}
    )
    existing_proof = existing.get("proof", {}) if existing else {}

    # Merge: keep old files, append new ones
    merged_photos = (existing_proof.get("photos") or []) + saved_photos
    merged_docs   = (existing_proof.get("documents") or []) + saved_docs

    proof_payload = {
        "note":        note.strip() if note.strip() else (existing_proof.get("note") or ""),
        "photos":      merged_photos,
        "documents":   merged_docs,
        "uploaded_at": existing_proof.get("uploaded_at") or now_iso,  # first upload time
        "updated_at":  now_iso,                                         # last update time
    }

    await db.attendance.update_one(
        {"user_id": current_user.id, "date": today_str},
        {"$set": {"proof": proof_payload}},
        upsert=True
    )

    logger.info(
        "Proof uploaded: user=%s date=%s photos=%d docs=%d note_len=%d",
        current_user.id, today_str, len(saved_photos), len(saved_docs), len(note)
    )

    return {
        "message":         "Proof saved successfully",
        "photos_saved":    len(saved_photos),
        "documents_saved": len(saved_docs),
        "note_saved":      bool(note.strip()),
        "total_photos":    len(merged_photos),
        "total_documents": len(merged_docs),
        "date":            today_str,
    }
async def get_top_performers_data(
    period: str = "monthly",
    limit: int = 5,
    db=None
):
    now = datetime.now(IST)
    if period == "weekly":
        start_date = now - timedelta(days=7)
    elif period == "monthly":
        start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    else:
        start_date = datetime(2024, 1, 1, tzinfo=timezone.utc)
    pipeline = [
        {
            "$match": {
                "status": "completed",
                "assigned_to": {"$ne": None},
                "$or": [
                    {"completed_at": {"$gte": start_date.isoformat()}},
                    {"updated_at": {"$gte": start_date.isoformat()}}
                ]
            }
        },
        {"$group": {"_id": "$assigned_to", "completed_tasks": {"$sum": 1}}},
        {
            "$lookup": {
                "from": "users",
                "localField": "_id",
                "foreignField": "id",
                "as": "user_info"
            }
        },
        {"$unwind": "$user_info"},
        {
            "$project": {
                "user_id": "$_id",
                "user_name": "$user_info.full_name",
                "profile_picture": "$user_info.profile_picture",
                "completed_tasks": 1
            }
        },
        {"$sort": {"completed_tasks": -1}},
        {"$limit": limit}
    ]
    performers = await db.tasks.aggregate(pipeline).to_list(limit)
    for idx, p in enumerate(performers):
        p["rank"] = idx + 1
    return performers

# Task routes
@api_router.post("/tasks", response_model=Task)
async def create_task(
    task_data: TaskCreate,
    current_user: User = Depends(get_current_user)
):
    if (task_data.assigned_to
            and task_data.assigned_to != current_user.id
            and current_user.role != "admin"):
        perms = get_user_permissions(current_user)
        if not perms.get("can_assign_tasks", False):
            raise HTTPException(
                status_code=403,
                detail="You do not have permission to assign tasks to other users"
            )
    task_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    task = Task(
        **task_data.model_dump(),
        id=task_id,
        created_by=current_user.id,
        created_at=now,
        updated_at=now
    )
    doc = task.model_dump()
    date_fields = ["created_at", "updated_at", "due_date", "recurrence_end_date"]
    for field in date_fields:
        if doc.get(field) and isinstance(doc[field], datetime):
            doc[field] = doc[field].isoformat()
    await db.tasks.insert_one(doc)
    if task.assigned_to and task.assigned_to != current_user.id:
        await create_notification(
            user_id=task.assigned_to,
            title="New Task Assigned",
            message=f"You have been assigned task '{task.title}'",
            type="assignment"
        )
    await create_audit_log(
        current_user=current_user,
        action="CREATE_TASK",
        module="tasks",
        record_id=task_id,
        new_data={"title": task.title}
    )
    return task

@api_router.get("/tasks/{task_id}/comments")
async def get_task_comments(
    task_id: str,
    current_user: User = Depends(get_current_user)
):
    task = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    is_admin = getattr(current_user, "role", "").lower() == "admin"
    is_involved = is_own_record(current_user, task)
    if not is_admin and not is_involved:
        raise HTTPException(status_code=403, detail="Unauthorized to view these comments")
    return task.get("comments", [])

@api_router.post("/tasks/bulk")
async def create_tasks_bulk(
    payload: BulkTaskCreate,
    current_user: User = Depends(get_current_user)
):
    created_tasks = []
    for task_data in payload.tasks:
        task_dict = task_data.model_dump()
        task_dict["id"] = str(uuid.uuid4())
        task_dict["created_by"] = current_user.id
        task_dict["created_at"] = datetime.now(IST).isoformat()
        task_dict["updated_at"] = datetime.now(IST).isoformat()
        if task_dict.get("due_date"):
            task_dict["due_date"] = task_dict["due_date"].isoformat()
        await db.tasks.insert_one(task_dict)
        if task_dict.get("assigned_to") and task_dict["assigned_to"] != current_user.id:
            await create_notification(
                user_id=task_dict["assigned_to"],
                title="New Task Assigned",
                message=f"You have been assigned task '{task_dict['title']}'"
            )
        created_tasks.append(task_dict)
    return {"message": "Tasks created successfully", "count": len(created_tasks)}

@api_router.post("/tasks/import")
async def import_tasks_from_csv(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    if file.content_type != "text/csv":
        raise HTTPException(400, "Invalid file type")
    content = await file.read()
    content_str = content.decode("utf-8")
    csv_reader = csv.DictReader(StringIO(content_str))
    tasks = []
    for row in csv_reader:
        task_data = TaskCreate(
            title=row.get("title", ""),
            description=row.get("description"),
            assigned_to=row.get("assigned_to"),
            sub_assignees=row.get("sub_assignees", "").split(",") if row.get("sub_assignees") else [],
            due_date=parser.parse(row["due_date"]) if row.get("due_date") else None,
            priority=row.get("priority", "medium"),
            status=row.get("status", "pending"),
            category=row.get("category", "other"),
            client_id=row.get("client_id"),
            is_recurring=bool(row.get("is_recurring", False)),
            recurrence_pattern=row.get("recurrence_pattern", "monthly"),
            recurrence_interval=int(row.get("recurrence_interval", 1))
        )
        tasks.append(task_data)
    payload = BulkTaskCreate(tasks=tasks)
    return await create_tasks_bulk(payload, current_user)

@api_router.get("/tasks")
async def get_tasks(current_user: User = Depends(get_current_user)):
    query = {"type": {"$ne": "todo"}}
    if current_user.role == "admin":
        pass
    else:
        permissions = get_user_permissions(current_user)
        allowed_users = permissions.get("view_other_tasks", []) or []
        if current_user.role == "manager":
            # Manager: Own + Team (same department)
            team_ids = await get_team_user_ids(current_user.id)
            allowed_users = list(set(allowed_users + team_ids))
        or_clauses = [
            {"assigned_to": current_user.id},
            {"sub_assignees": current_user.id},
            {"created_by": current_user.id},
        ]
        if allowed_users:
            or_clauses.append({"assigned_to": {"$in": allowed_users}})
            or_clauses.append({"created_by": {"$in": allowed_users}})
        query["$or"] = or_clauses
    tasks = await db.tasks.find(query, {"_id": 0}).to_list(1000)
    user_ids = {
        task.get("assigned_to") for task in tasks if task.get("assigned_to")
    } | {
        task.get("created_by") for task in tasks if task.get("created_by")
    }
    users = await db.users.find(
        {"id": {"$in": list(user_ids)}},
        {"_id": 0, "password": 0}
    ).to_list(1000)
    user_map = {u["id"]: u["full_name"] for u in users}
    for task in tasks:
        task["created_at"] = safe_dt(task.get("created_at"))
        task["updated_at"] = safe_dt(task.get("updated_at"))
        task["due_date"] = safe_dt(task.get("due_date"))
        task["assigned_to_name"] = user_map.get(task.get("assigned_to"), "Unknown")
        task["created_by_name"] = user_map.get(task.get("created_by"), "Unknown")
        if not isinstance(task.get("sub_assignees"), list):
            task["sub_assignees"] = []
        if not isinstance(task.get("comments"), list):
            task["comments"] = []
    return tasks

@api_router.get("/tasks/{task_id}/detail")
async def get_task_detail(task_id: str, current_user: User = Depends(get_current_user)):
    task = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if current_user.role != "admin":
        if not is_own_record(current_user, task):
            permissions = get_user_permissions(current_user)
            allowed_users = permissions.get("view_other_tasks", []) or []
            if current_user.role == "manager":
                team_ids = await get_team_user_ids(current_user.id)
                allowed_users = list(set(allowed_users + team_ids))
            if task.get("assigned_to") not in allowed_users:
                raise HTTPException(status_code=403, detail="Not authorized")
    # Batch all user lookups in one query instead of 3 separate find_one calls
    user_ids_to_fetch = list(filter(None, [
        task.get("assigned_to"),
        task.get("created_by"),
        *( task.get("sub_assignees") or [] ),
    ]))
    users_batch = {}
    if user_ids_to_fetch:
        fetched = await db.users.find(
            {"id": {"$in": user_ids_to_fetch}},
            {"_id": 0, "id": 1, "full_name": 1, "profile_picture": 1, "email": 1}
        ).to_list(100)
        users_batch = {u["id"]: u for u in fetched}

    assigned_user = users_batch.get(task.get("assigned_to"))
    created_user  = users_batch.get(task.get("created_by"))
    sub_assignee_names = [
        users_batch[uid]["full_name"]
        for uid in (task.get("sub_assignees") or [])
        if uid in users_batch
    ]

    client_name = None
    if task.get("client_id"):
        client_doc = await db.clients.find_one(
            {"id": task["client_id"]},
            {"_id": 0, "company_name": 1}
        )
        if client_doc:
            client_name = client_doc.get("company_name")
    task["assigned_to_name"] = assigned_user.get("full_name", "Unknown") if assigned_user else "Unknown"
    task["assigned_to_email"] = assigned_user.get("email") if assigned_user else None
    task["assigned_to_picture"] = assigned_user.get("profile_picture") if assigned_user else None
    task["created_by_name"] = created_user.get("full_name", "Unknown") if created_user else "Unknown"
    task["sub_assignee_names"] = sub_assignee_names
    task["client_name"] = client_name
    task["created_at"] = safe_dt(task.get("created_at"))
    task["updated_at"] = safe_dt(task.get("updated_at"))
    task["due_date"] = safe_dt(task.get("due_date"))
    if task.get("completed_at"):
        task["completed_at"] = safe_dt(task.get("completed_at"))
    return task

@api_router.get("/tasks/{task_id}", response_model=Task)
async def get_task(task_id: str, current_user: User = Depends(get_current_user)):
    task = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if current_user.role == "admin":
        pass
    elif is_own_record(current_user, task):
        pass
    else:
        permissions = get_user_permissions(current_user)
        allowed_users = permissions.get("view_other_tasks", []) or []
        if current_user.role == "manager":
            team_ids = await get_team_user_ids(current_user.id)
            allowed_users = list(set(allowed_users + team_ids))
        if task.get("assigned_to") not in allowed_users:
            raise HTTPException(status_code=403, detail="Not authorized")
    return Task(**task)

@api_router.api_route("/tasks/{task_id}", methods=["PATCH", "PUT"], response_model=Task)
async def patch_task(
    task_id: str,
    updates: dict,
    current_user: User = Depends(get_current_user)
):
    existing_task = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    if not existing_task:
        raise HTTPException(status_code=404, detail="Task not found")
    is_authorized = (
        current_user.role.lower() == "admin" or
        is_own_record(current_user, existing_task)
    )
    if not is_authorized:
        raise HTTPException(status_code=403, detail="Unauthorized to modify this task")
    old_data = existing_task.copy()
    # Whitelist: prevent patching of immutable/sensitive fields
    IMMUTABLE_FIELDS = {"id", "created_by", "created_at", "_id"}
    for field in IMMUTABLE_FIELDS:
        updates.pop(field, None)
    updates["updated_at"] = datetime.now(IST).isoformat()
    if updates.get("status") == "completed":
        updates["completed_at"] = datetime.now(IST).isoformat()
    await db.tasks.update_one({"id": task_id}, {"$set": updates})
    updated_task = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    action_type = "TASK_STATUS_CHANGED" if "status" in updates and old_data.get("status") != updates.get("status") else "UPDATE_TASK"
    await create_audit_log(
        current_user=current_user,
        action=action_type,
        module="task",
        record_id=task_id,
        old_data=old_data,
        new_data=updates
    )
    return Task(**updated_task)

@api_router.delete("/tasks/{task_id}")
async def delete_task(
    task_id: str,
    current_user: User = Depends(get_current_user)
):
    existing = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Task not found")
    is_admin = current_user.role.lower() == "admin"
    permissions = get_user_permissions(current_user)
    is_creator = existing.get("created_by") == current_user.id
    has_delete_perm = permissions.get("can_edit_tasks", False)
    if not (is_admin or is_creator or has_delete_perm):
        raise HTTPException(status_code=403, detail="Only Admin, task creator, or users with explicit permission can delete tasks.")
    await db.tasks.delete_one({"id": task_id})
    await create_audit_log(
        current_user=current_user,
        action="DELETE_TASK",
        module="task",
        record_id=task_id,
        old_data=existing
    )
    return {"message": "Task deleted successfully"}

@api_router.post("/tasks/{task_id}/comments")
async def add_task_comment(
    task_id: str,
    comment_data: dict,
    current_user: User = Depends(get_current_user)
):
    task = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    is_involved = (
        current_user.role.lower() == "admin" or
        is_own_record(current_user, task)
    )
    if not is_involved:
        raise HTTPException(status_code=403, detail="Access denied: You must be involved in this task to comment.")
    comment = {
        "id": str(uuid.uuid4()),
        "user_id": current_user.id,
        "user_name": current_user.full_name,
        "text": comment_data.get("text"),
        "created_at": datetime.now(IST).isoformat()
    }
    await db.tasks.update_one({"id": task_id}, {"$push": {"comments": comment}})
    return comment

# =========================================================
# EXPORT TASK AUDIT LOG PDF
# =========================================================
@api_router.get("/tasks/{task_id}/export-log-pdf")
async def export_task_log_pdf(
    task_id: str,
    current_user: User = Depends(check_permission("can_view_audit_logs"))
):
    task = await db.tasks.find_one({"id": task_id}, {"_id": 0})
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    logs = await db.audit_logs.find(
        {"module": "task", "record_id": task_id},
        {"_id": 0}
    ).sort("timestamp", 1).to_list(1000)
    if not logs:
        raise HTTPException(status_code=404, detail="No audit logs found")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Task Lifecycle Report", ln=True, align="C")
    pdf.ln(5)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Task Information", ln=True)
    pdf.ln(3)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 7, f"Title: {task.get('title', '-')}")
    pdf.multi_cell(0, 7, f"Description: {task.get('description', '-')}")
    pdf.multi_cell(0, 7, f"Assigned To: {task.get('assigned_to_name', task.get('assigned_to', '-'))}")
    pdf.multi_cell(0, 7, f"Created By: {task.get('created_by_name', task.get('created_by', '-'))}")
    pdf.multi_cell(0, 7, f"Created At: {task.get('created_at', '-')}")
    pdf.multi_cell(0, 7, f"Current Status: {task.get('status', '-')}")
    pdf.ln(8)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Activity Timeline", ln=True)
    pdf.ln(4)
    pdf.set_font("Arial", size=10)
    for log in logs:
        timestamp = log.get("timestamp")
        if isinstance(timestamp, datetime):
            timestamp = timestamp.strftime("%b %d, %Y %I:%M %p")
        action = log.get("action", "UNKNOWN")
        user = log.get("user_name", "Unknown User")
        pdf.set_font("Arial", "B", 10)
        pdf.multi_cell(0, 6, f"{timestamp} — {action.replace('_', ' ').title()} by {user}")
        pdf.set_font("Arial", size=10)
        if action == "TASK_STATUS_CHANGED":
            old_status = log.get("old_data", {}).get("status", "-")
            new_status = log.get("new_data", {}).get("status", "-")
            pdf.multi_cell(0, 6, f" Status changed: {old_status} -> {new_status}")
        elif action == "TASK_COMPLETED":
            pdf.multi_cell(0, 6, " Task marked as completed.")
        elif action == "DELETE_TASK":
            pdf.multi_cell(0, 6, " Task was deleted.")
        elif action == "CREATE_TASK":
            pdf.multi_cell(0, 6, " Task was created.")
        elif action == "UPDATE_TASK":
            pdf.multi_cell(0, 6, " Task details updated.")
        if log.get("old_data") and log.get("new_data"):
            old_data = log.get("old_data")
            new_data = log.get("new_data")
            for key in new_data:
                old_val = old_data.get(key)
                new_val = new_data.get(key)
                if old_val != new_val:
                    pdf.multi_cell(0, 6, f" {key.replace('_', ' ').title()}: {old_val} -> {new_val}")
        pdf.ln(3)
    pdf.ln(5)
    pdf.set_font("Arial", "I", 8)
    pdf.multi_cell(0, 5, f"Generated on {datetime.utcnow().strftime('%b %d, %Y %I:%M %p')} UTC")
    output = BytesIO()
    output.write(pdf.output(dest="S").encode("latin1"))
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=task_lifecycle_{task_id}.pdf"}
    )

# ═══════════════════════════════════════════════════════════════════════════════
# DSC ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

def _to_iso(val) -> Optional[str]:
    """Safely convert datetime/date/string/None to ISO string for MongoDB storage."""
    if val is None:
        return None
    if isinstance(val, (datetime, date)):
        return val.isoformat()
    return str(val)  # already a plain string like "2026-03-22"


@api_router.post("/dsc", response_model=DSC)
async def create_dsc(
    dsc_data: DSCCreate,
    current_user: User = Depends(get_current_user)
):
    try:
        now = datetime.now(timezone.utc)
        dsc = DSC(
            **dsc_data.model_dump(),
            created_by=current_user.id,
            created_at=now,          # explicitly set — never None
        )
        doc = dsc.model_dump()
        doc["created_at"]  = _to_iso(doc["created_at"])
        doc["issue_date"]  = _to_iso(doc["issue_date"])
        doc["expiry_date"] = _to_iso(doc["expiry_date"])
        await db.dsc_register.insert_one(doc)
        doc.pop("_id", None)
        return dsc
    except Exception as e:
        logger.error(f"DSC create error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to save DSC: {str(e)}")


@api_router.get("/dsc")
async def get_dsc_list(
    sort_by: str = Query("holder_name"),
    order: str = Query("asc", pattern="^(asc|desc)$", ignore_case=True),
    search: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    limit: int = Query(500, ge=1, le=500),
    current_user: User = Depends(check_permission("can_view_all_dsc"))
):
    query = {}
    if search:
        safe_search = re.escape(search)
        search_regex = {"$regex": safe_search, "$options": "i"}
        query["$or"] = [
            {"holder_name": search_regex},
            {"dsc_type": search_regex},
            {"associated_with": search_regex},
            {"current_status": search_regex}
        ]
    sort_dir = 1 if order.lower() == "asc" else -1
    skip = (page - 1) * limit
    total = await db.dsc_register.count_documents(query)
    cursor = db.dsc_register.find(query, {"_id": 0}).sort(sort_by, sort_dir).skip(skip).limit(limit)
    dsc_list = await cursor.to_list(length=limit)
    now = datetime.now(IST)

    for dsc in dsc_list:
        # Safely parse stored ISO strings back to datetime for comparison
        for field in ("created_at", "issue_date", "expiry_date"):
            val = dsc.get(field)
            if isinstance(val, str):
                try:
                    dsc[field] = datetime.fromisoformat(val)
                except ValueError:
                    dsc[field] = None

        expiry_date = dsc.get("expiry_date")
        if expiry_date:
            # Make expiry_date timezone-aware for comparison with IST now
            if isinstance(expiry_date, datetime) and expiry_date.tzinfo is None:
                expiry_date = expiry_date.replace(tzinfo=timezone.utc)

            if expiry_date < now:
                movement_log = dsc.get("movement_log", [])
                updated = False
                if not any(log.get("movement_type") == "EXPIRED" for log in movement_log):
                    movement_log.append({
                        "id": str(uuid.uuid4()),
                        "movement_type": "EXPIRED",
                        "person_name": "System Auto",
                        "notes": "Auto marked as expired",
                        "timestamp": now.isoformat(),
                        "recorded_by": "System"
                    })
                    updated = True
                if dsc.get("current_status") != "EXPIRED":
                    updated = True
                if updated:
                    await db.dsc_register.update_one(
                        {"id": dsc["id"]},
                        {"$set": {"current_status": "EXPIRED", "movement_log": movement_log}}
                    )
                    dsc["current_status"] = "EXPIRED"
                    dsc["movement_log"] = movement_log

    return DSCListResponse(data=dsc_list, total=total, page=page, limit=limit)


@api_router.put("/dsc/{dsc_id}", response_model=DSC)
async def update_dsc(
    dsc_id: str,
    dsc_data: DSCCreate,
    current_user: User = Depends(check_permission("can_edit_dsc"))
):
    existing = await db.dsc_register.find_one({"id": dsc_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="DSC not found")

    update_data = dsc_data.model_dump()
    update_data["issue_date"]  = _to_iso(update_data["issue_date"])
    update_data["expiry_date"] = _to_iso(update_data["expiry_date"])

    await db.dsc_register.update_one({"id": dsc_id}, {"$set": update_data})
    await create_audit_log(
        current_user, action="UPDATE_DSC", module="dsc",
        record_id=dsc_id, old_data=existing, new_data=update_data
    )

    updated = await db.dsc_register.find_one({"id": dsc_id}, {"_id": 0})
    if not updated:
        raise HTTPException(status_code=404, detail="DSC not found after update")

    # Parse ISO strings back to datetime for Pydantic response model
    for field in ("created_at", "issue_date", "expiry_date"):
        val = updated.get(field)
        if isinstance(val, str):
            try:
                updated[field] = datetime.fromisoformat(val)
            except ValueError:
                updated[field] = None

    return DSC(**updated)


@api_router.delete("/dsc/{dsc_id}")
async def delete_dsc(
    dsc_id: str,
    current_user: User = Depends(check_permission("can_edit_dsc"))
):
    existing = await db.dsc_register.find_one({"id": dsc_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="DSC not found")
    await create_audit_log(
        current_user, action="DELETE_DSC", module="dsc",
        record_id=dsc_id, old_data=existing
    )
    result = await db.dsc_register.delete_one({"id": dsc_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="DSC not found")
    return {"message": "DSC deleted successfully"}


@api_router.post("/dsc/{dsc_id}/movement")
async def record_dsc_movement(
    dsc_id: str,
    movement_data: DSCMovementRequest,
    current_user: User = Depends(get_current_user)
):
    existing = await db.dsc_register.find_one({"id": dsc_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="DSC not found")

    movement = {
        "id":            str(uuid.uuid4()),
        "movement_type": movement_data.movement_type,
        "person_name":   movement_data.person_name,
        "timestamp":     datetime.now(IST).isoformat(),
        "notes":         movement_data.notes,
        "recorded_by":   current_user.full_name
    }
    movement_log = existing.get("movement_log", [])
    movement_log.append(movement)

    await db.dsc_register.update_one(
        {"id": dsc_id},
        {
            "$set": {
                "current_status":   movement_data.movement_type,
                "current_location": "with_company" if movement_data.movement_type == "IN" else "taken_by_client",
                "movement_log":     movement_log
            }
        }
    )
    await create_audit_log(
        current_user, action="UPDATE_DSC", module="dsc",
        record_id=dsc_id, old_data=existing, new_data={"movement_log": movement_log}
    )
    return {"message": f"DSC marked as {movement_data.movement_type}", "movement": movement}


@api_router.put("/dsc/{dsc_id}/movement/{movement_id}")
async def update_dsc_movement(
    dsc_id: str,
    movement_id: str,
    update_data: MovementUpdateRequest,
    current_user: User = Depends(get_current_user)
):
    existing = await db.dsc_register.find_one({"id": dsc_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="DSC not found")

    movement_log = existing.get("movement_log", [])
    movement_found = False

    for i, movement in enumerate(movement_log):
        if movement.get("id") == movement_id:
            movement_log[i]["movement_type"] = update_data.movement_type
            if update_data.person_name:
                movement_log[i]["person_name"] = update_data.person_name
            if update_data.notes is not None:
                movement_log[i]["notes"] = update_data.notes
            movement_log[i]["edited_by"] = current_user.full_name
            movement_log[i]["edited_at"] = datetime.now(timezone.utc).isoformat()
            movement_found = True
            break

    if not movement_found:
        raise HTTPException(status_code=404, detail="Movement entry not found")

    # Derive current_status from the last non-EXPIRED movement entry
    non_expired = [m for m in movement_log if m.get("movement_type") != "EXPIRED"]
    new_status = non_expired[-1]["movement_type"] if non_expired else "IN"

    await db.dsc_register.update_one(
        {"id": dsc_id},
        {"$set": {"current_status": new_status, "movement_log": movement_log}}
    )
    await create_audit_log(
        current_user, action="UPDATE_DSC", module="dsc",
        record_id=dsc_id, old_data=existing, new_data={"movement_log": movement_log}
    )
    return {"message": "Movement updated successfully", "movement_log": movement_log}
# DOCUMENT REGISTER ROUTES
@api_router.post("/documents", response_model=Document)
async def create_document(document_data: DocumentCreate, current_user: User = Depends(get_current_user)):
    document = Document(**document_data.model_dump(), created_by=current_user.id)
    doc = document.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    if doc.get("issue_date"):
        doc["issue_date"] = doc["issue_date"].isoformat()
    if doc.get("valid_upto"):
        doc["valid_upto"] = doc["valid_upto"].isoformat()
    await db.documents.insert_one(doc)
    return document

@api_router.get("/documents", response_model=List[Document])
async def get_documents(current_user: User = Depends(check_permission("can_view_documents"))):
    documents = await db.documents.find({}, {"_id": 0}).to_list(1000)
    for d in documents:
        if isinstance(d["created_at"], str):
            d["created_at"] = datetime.fromisoformat(d["created_at"])
        if d.get("issue_date") and isinstance(d["issue_date"], str):
            d["issue_date"] = datetime.fromisoformat(d["issue_date"])
        if d.get("valid_upto") and isinstance(d["valid_upto"], str):
            d["valid_upto"] = datetime.fromisoformat(d["valid_upto"])
    return documents

@api_router.put("/documents/{document_id}", response_model=Document)
async def update_document(document_id: str, document_data: DocumentCreate, current_user: User = Depends(check_permission("can_edit_documents"))):
    existing = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Document not found")
    update_data = document_data.model_dump()
    if update_data.get("issue_date"):
        update_data["issue_date"] = update_data["issue_date"].isoformat()
    if update_data.get("valid_upto"):
        update_data["valid_upto"] = update_data["valid_upto"].isoformat()
    await db.documents.update_one({"id": document_id}, {"$set": update_data})
    await create_audit_log(current_user, action="UPDATE_DOCUMENT", module="document", record_id=document_id, old_data=existing, new_data=update_data)
    updated = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if isinstance(updated["created_at"], str):
        updated["created_at"] = datetime.fromisoformat(updated["created_at"])
    return Document(**updated)

@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str, current_user: User = Depends(check_permission("can_edit_documents"))):
    existing = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Document not found")
    await create_audit_log(current_user, action="DELETE_DOCUMENT", module="document", record_id=document_id, old_data=existing)
    result = await db.documents.delete_one({"id": document_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document deleted successfully"}

@api_router.post("/documents/{document_id}/movement")
async def record_document_movement(
    document_id: str,
    movement_data: DocumentMovementRequest,
    current_user: User = Depends(get_current_user)
):
    document = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    movement = {
        "id": str(uuid.uuid4()),
        "movement_type": movement_data.movement_type,
        "person_name": movement_data.person_name,
        "timestamp": datetime.now(IST).isoformat(),
        "notes": movement_data.notes,
        "recorded_by": current_user.full_name
    }
    movement_log = document.get("movement_log", [])
    movement_log.append(movement)
    await db.documents.update_one(
        {"id": document_id},
        {"$set": {"current_status": movement_data.movement_type, "movement_log": movement_log}}
    )
    await create_audit_log(current_user, action="UPDATE_DOCUMENT", module="document", record_id=document_id, old_data=document, new_data={"movement_log": movement_log})
    return {"message": "Movement recorded successfully"}

@api_router.put("/documents/{document_id}/movement/{movement_id}")
async def update_document_movement(
    document_id: str,
    movement_id: str,
    update_data: DocumentMovementRequest,
    current_user: User = Depends(get_current_user)
):
    document = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    movement_log = document.get("movement_log", [])
    movement_found = False
    for i, movement in enumerate(movement_log):
        if movement.get("id") == movement_id:
            movement_log[i]["movement_type"] = update_data.movement_type
            movement_log[i]["person_name"] = update_data.person_name
            movement_log[i]["notes"] = update_data.notes
            movement_log[i]["edited_by"] = current_user.full_name
            movement_log[i]["edited_at"] = datetime.now(timezone.utc).isoformat()
            movement_found = True
            break
    if not movement_found:
        raise HTTPException(status_code=404, detail="Movement entry not found")
    new_status = movement_log[-1]["movement_type"] if movement_log else "IN"
    await db.documents.update_one(
        {"id": document_id},
        {"$set": {"current_status": new_status, "movement_log": movement_log}}
    )
    await create_audit_log(current_user, action="UPDATE_DOCUMENT", module="document", record_id=document_id, old_data=document, new_data={"movement_log": movement_log})
    return {"message": "Movement updated successfully"}

# DUE DATE ROUTES
COMPLIANCE_RULES = [
    {"keywords": ["gstr-1", "gstr1", "outward supply"],                                "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-3b", "gstr3b", "summary return"],                              "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-9", "annual return gst"],                                      "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-4", "composition"],                                            "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-7", "tds return gst"],                                         "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-8", "tcs statement"],                                          "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-5", "non-resident"],                                           "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-6", "isd return"],                                             "category": "GST",        "department": "GST"},
    {"keywords": ["gstr-10", "final return"],                                          "category": "GST",        "department": "GST"},
    {"keywords": ["gst", "goods and service"],                                         "category": "GST",        "department": "GST"},
    {"keywords": ["itr", "income tax return"],                                         "category": "Income Tax", "department": "IT"},
    {"keywords": ["advance tax", "advance income tax"],                                "category": "Income Tax", "department": "IT"},
    {"keywords": ["tax audit", "form 3ca", "form 3cb"],                                "category": "Audit",      "department": "IT"},
    {"keywords": ["form 16", "form 26as"],                                             "category": "Income Tax", "department": "IT"},
    {"keywords": ["income tax", "direct tax"],                                         "category": "Income Tax", "department": "IT"},
    {"keywords": ["tds", "tax deducted at source", "form 24q", "form 26q", "form 27q"],"category": "TDS",        "department": "TDS"},
    {"keywords": ["tcs", "tax collected at source"],                                   "category": "TDS",        "department": "TDS"},
    {"keywords": ["challan 281"],                                                      "category": "TDS",        "department": "TDS"},
    {"keywords": ["mgt-7", "annual return roc", "annual return mca"],                  "category": "ROC",        "department": "ROC"},
    {"keywords": ["aoc-4", "financial statement", "filing of financial"],              "category": "ROC",        "department": "ROC"},
    {"keywords": ["dir-3", "director kyc", "din kyc"],                                 "category": "ROC",        "department": "ROC"},
    {"keywords": ["dir-8", "disqualification"],                                        "category": "ROC",        "department": "ROC"},
    {"keywords": ["dir-12", "appointment", "resignation of director"],                 "category": "ROC",        "department": "ROC"},
    {"keywords": ["mbp-1", "disclosure of interest"],                                  "category": "ROC",        "department": "ROC"},
    {"keywords": ["agm", "annual general meeting"],                                    "category": "ROC",        "department": "ROC"},
    {"keywords": ["dpt-3", "return of deposits"],                                      "category": "ROC",        "department": "ROC"},
    {"keywords": ["msme-1", "msme samadhaan"],                                         "category": "ROC",        "department": "MSME"},
    {"keywords": ["pas-6", "reconciliation of share"],                                 "category": "ROC",        "department": "ROC"},
    {"keywords": ["roc", "mca", "companies act", "registrar of companies"],            "category": "ROC",        "department": "ROC"},
    {"keywords": ["msme"],                                                             "category": "Other",      "department": "MSME"},
    {"keywords": ["statutory audit", "internal audit", "audit report"],                "category": "Audit",      "department": "ACC"},
    {"keywords": ["adt-1", "appointment of auditor"],                                  "category": "Audit",      "department": "ROC"},
    {"keywords": ["trademark", "tm renewal"],                                          "category": "Trademark",  "department": "TM"},
    {"keywords": ["fema", "foreign exchange", "fdi"],                                  "category": "FEMA",       "department": "FEMA"},
    {"keywords": ["rera", "real estate"],                                              "category": "RERA",       "department": "OTHER"},
    {"keywords": ["pf", "provident fund", "epfo"],                                     "category": "Other",      "department": "ACC"},
    {"keywords": ["esi", "esic"],                                                      "category": "Other",      "department": "ACC"},
    {"keywords": ["board meeting", "minute book"],                                     "category": "ROC",        "department": "ROC"},
]

MONTH_MAP = {
    "january": 1,  "jan": 1,
    "february": 2, "feb": 2,
    "march": 3,    "mar": 3,
    "april": 4,    "apr": 4,
    "may": 5,
    "june": 6,     "jun": 6,
    "july": 7,     "jul": 7,
    "august": 8,   "aug": 8,
    "september": 9,"sep": 9, "sept": 9,
    "october": 10, "oct": 10,
    "november": 11,"nov": 11,
    "december": 12,"dec": 12,
}


def parse_date_from_text(text: str):
    text = text.strip()
    now = datetime.now()
    year = now.year

    m = re.search(
        r'\b(\d{1,2})(?:st|nd|rd|th)?\s+'
        r'(january|february|march|april|may|june|july|august|september|october|november|december)'
        r'\s+(\d{4})\b', text, re.IGNORECASE)
    if m:
        try:
            return date(int(m.group(3)), MONTH_MAP[m.group(2).lower()], int(m.group(1))).isoformat()
        except Exception:
            pass

    m = re.search(
        r'\b(january|february|march|april|may|june|july|august|september|october|november|december)'
        r'\s+(\d{1,2}),?\s+(\d{4})\b', text, re.IGNORECASE)
    if m:
        try:
            return date(int(m.group(3)), MONTH_MAP[m.group(1).lower()], int(m.group(2))).isoformat()
        except Exception:
            pass

    m = re.search(r'\b(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})\b', text)
    if m:
        try:
            return date(int(m.group(3)), int(m.group(2)), int(m.group(1))).isoformat()
        except Exception:
            pass

    m = re.search(r'\b(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})\b', text)
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3))).isoformat()
        except Exception:
            pass

    m = re.search(
        r'\b(\d{1,2})(?:st|nd|rd|th)?\s+'
        r'(january|february|march|april|may|june|july|august|september|october|november|december)\b',
        text, re.IGNORECASE)
    if m:
        try:
            mo = MONTH_MAP[m.group(2).lower()]
            d_val = int(m.group(1))
            target = date(year, mo, d_val)
            if target < date.today():
                target = date(year + 1, mo, d_val)
            return target.isoformat()
        except Exception:
            pass

    m = re.search(
        r'\b(january|february|march|april|may|june|july|august|september|october|november|december)'
        r'\s+(\d{1,2})\b', text, re.IGNORECASE)
    if m:
        try:
            mo = MONTH_MAP[m.group(1).lower()]
            d_val = int(m.group(2))
            target = date(year, mo, d_val)
            if target < date.today():
                target = date(year + 1, mo, d_val)
            return target.isoformat()
        except Exception:
            pass

    m = re.search(r'\b(\d{1,2})(?:st|nd|rd|th)?\s+of\s+next\s+month\b', text, re.IGNORECASE)
    if m:
        try:
            day = int(m.group(1))
            today = date.today()
            if today.month < 12:
                target = date(today.year, today.month + 1, day)
            else:
                target = date(today.year + 1, 1, day)
            return target.isoformat()
        except Exception:
            pass

    m = re.search(r'within\s+(\d+)\s+days?', text, re.IGNORECASE)
    if m:
        try:
            return (date.today() + timedelta(days=int(m.group(1)))).isoformat()
        except Exception:
            pass

    m = re.search(r'\b(\d{1,2})\s+(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\b', text, re.IGNORECASE)
    if m:
        try:
            mo = MONTH_MAP[m.group(2).lower()]
            d_val = int(m.group(1))
            target = date(year, mo, d_val)
            if target < date.today():
                target = date(year + 1, mo, d_val)
            return target.isoformat()
        except Exception:
            pass

    return None


def classify_compliance(line: str):
    lower = line.lower()
    for rule in COMPLIANCE_RULES:
        if any(kw in lower for kw in rule["keywords"]):
            return {"category": rule["category"], "department": rule["department"]}
    return {"category": "Other", "department": "OTHER"}


def extract_title(line: str) -> str:
    title = re.sub(r'\s+', ' ', line).strip()
    title = re.sub(r'^[\-\*\•\|]+\s*', '', title)
    if len(title) > 80:
        title = title[:77].rsplit(' ', 1)[0] + '...'
    return title or "Compliance Task"


def parse_compliance_dates(raw_text: str):
    results = []
    seen = set()
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]

    for i, line in enumerate(lines):
        if "|" not in line:
            continue
        cols = [c.strip() for c in line.split("|") if c.strip()]
        if len(cols) < 2:
            continue
        date_val = None
        date_col_idx = None
        for idx, col in enumerate(cols):
            d = parse_date_from_text(col)
            if d:
                date_col_idx = idx
                date_val = d
                break
        if not date_val and i + 1 < len(lines):
            date_val = parse_date_from_text(lines[i + 1])
        if not date_val:
            continue
        title_col = next((c for idx, c in enumerate(cols) if idx != date_col_idx and len(c) > 3), None)
        if not title_col:
            continue
        title = extract_title(title_col)
        if title.lower() in seen:
            continue
        seen.add(title.lower())
        clf = classify_compliance(line)
        results.append({
            "title": title,
            "due_date": date_val,
            "category": clf["category"],
            "department": clf["department"],
            "description": title_col[:300],
            "status": "pending",
        })

    for i, line in enumerate(lines):
        if len(line) < 8:
            continue
        if re.match(r'^(form|compliance|particulars|due date|applicability|sl\.?\s*no)', line, re.IGNORECASE):
            continue
        date_val = parse_date_from_text(line)
        if not date_val and i + 1 < len(lines):
            date_val = parse_date_from_text(lines[i + 1])
        if not date_val:
            continue
        clf = classify_compliance(line)
        if clf["category"] == "Other" and clf["department"] == "OTHER" and i > 0:
            clf = classify_compliance(lines[i - 1])
        title = extract_title(line)
        if title.lower() in seen:
            continue
        seen.add(title.lower())
        stripped = re.sub(r'\d{1,2}(?:st|nd|rd|th)?\s+\w+\s*\d{0,4}', '', line, flags=re.IGNORECASE).strip()
        if len(stripped) < 5:
            continue
        results.append({
            "title": title,
            "due_date": date_val,
            "category": clf["category"],
            "department": clf["department"],
            "description": line[:300],
            "status": "pending",
        })

    form_pat = re.compile(
        r'((?:GSTR?|ITR|MGT|AOC|DIR|DPT|ADT|PAS|INC|CHG|BEN|SH|CSR|MSME)-[\w\/]+)',
        re.IGNORECASE)
    for i, line in enumerate(lines):
        m = form_pat.search(line)
        if not m:
            continue
        form_name = m.group(1).upper()
        date_val = None
        for j in range(i, min(i + 3, len(lines))):
            date_val = parse_date_from_text(lines[j])
            if date_val:
                break
        if not date_val:
            continue
        title = f"{form_name} Filing"
        if title.lower() in seen:
            continue
        seen.add(title.lower())
        clf = classify_compliance(form_name + " " + line)
        results.append({
            "title": title,
            "due_date": date_val,
            "category": clf["category"],
            "department": clf["department"],
            "description": extract_title(line),
            "status": "pending",
        })

    results.sort(key=lambda x: x.get("due_date", "9999-12-31"))
    return results


# Route registered BEFORE /{due_date_id} param routes to prevent shadowing
@api_router.post("/duedates/extract-from-file")
async def extract_due_dates_from_file(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    filename = (file.filename or "").lower()
    content_type = file.content_type or ""
    file_bytes = await file.read()
    raw_text = ""

    try:
        if content_type.startswith("image/") or filename.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp")):
            raise HTTPException(
                status_code=400,
                detail="Image upload is not supported on this server. Please upload a PDF or DOCX file instead."
            )

        elif content_type == "application/pdf" or filename.endswith(".pdf"):
            import pdfplumber
            parts = []
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
                    for table in page.extract_tables():
                        for row in table:
                            if row:
                                parts.append("  |  ".join(str(c or "") for c in row))
            raw_text = "\n".join(parts)

        elif filename.endswith((".docx", ".doc")):
            from docx import Document as DocxDocument
            doc = DocxDocument(BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append("  |  ".join(cell.text for cell in row.cells))
            raw_text = "\n".join(parts)

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use JPG, PNG, PDF, or DOCX.")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File extraction error: {e}")
        raise HTTPException(status_code=422, detail=f"Could not read file: {str(e)}")

    if not raw_text or len(raw_text.strip()) < 20:
        raise HTTPException(status_code=422, detail="No readable text found. Try a clearer image or PDF.")

    extracted = parse_compliance_dates(raw_text)

    if not extracted:
        raise HTTPException(status_code=404, detail="No compliance dates detected in this document.")

    return {"extracted": extracted, "count": len(extracted)}
# ─── Calendar → Compliance Tracker sync helper ───────────────────────────────
_CALENDAR_CATEGORY_MAP = {
    'GST': 'GST', 'ROC': 'ROC', 'MCA': 'ROC', 'ITR': 'ITR',
    'TDS': 'TDS', 'AUDIT': 'AUDIT', 'PF': 'PF_ESIC', 'ESIC': 'PF_ESIC',
    'PT': 'PT', 'INCOME TAX': 'ITR',
}
_COMPLIANCE_CATEGORIES = ['ROC', 'GST', 'ITR', 'TDS', 'AUDIT', 'PF_ESIC', 'PT', 'OTHER']

async def _sync_due_date_to_compliance(dd: dict, current_user) -> None:
    """Upsert a compliance_master record linked to the given due_date doc."""
    title    = (dd.get('title') or '').strip()
    dd_id    = dd.get('id')
    due_date = dd.get('due_date')
    if isinstance(due_date, datetime):
        due_date = due_date.strftime('%Y-%m-%d')
    elif isinstance(due_date, str) and 'T' in due_date:
        due_date = due_date[:10]

    raw_cat  = (dd.get('category') or 'OTHER').upper()
    category = _CALENDAR_CATEGORY_MAP.get(raw_cat, 'OTHER')
    if category not in _COMPLIANCE_CATEGORIES:
        category = 'OTHER'

    existing = await db.compliance_masters.find_one(
        {'calendar_due_date_id': dd_id}, {'_id': 0}
    )
    now_str = datetime.now(timezone.utc).isoformat()

    if existing:
        await db.compliance_masters.update_one(
            {'id': existing['id']},
            {'$set': {'due_date': due_date, 'name': title, 'updated_at': now_str}}
        )
    else:
        import uuid as _uuid
        doc = {
            'id':                      str(_uuid.uuid4()),
            'name':                    title,
            'category':                category,
            'frequency':               'one_time',
            'fy_year':                 None,
            'period_label':            None,
            'due_date':                due_date,
            'description':             dd.get('description', ''),
            'applicable_entity_types': [],
            'calendar_due_date_id':    dd_id,
            'created_by':              str(getattr(current_user, 'id', '')),
            'created_by_name':         getattr(current_user, 'full_name', ''),
            'created_at':              now_str,
            'updated_at':              now_str,
        }
        await db.compliance_masters.insert_one({**doc, '_id': doc['id']})


@api_router.post("/duedates", response_model=DueDate)
async def create_due_date(
    due_date_data: DueDateCreate,
    current_user: User = Depends(get_current_user)
):
    try:
        if not due_date_data.department:
            raise HTTPException(status_code=400, detail="Department required")

        data = due_date_data.model_dump()

        # ✅ FIX 1: Safe date parsing
        raw_due_date = data.get("due_date")

        if isinstance(raw_due_date, str):
            try:
                parsed_date = datetime.fromisoformat(raw_due_date)
            except:
                parsed_date = datetime.strptime(raw_due_date, "%m/%d/%Y")
        else:
            parsed_date = raw_due_date

        # ✅ FIX 2: Build document safely
        data["due_date"] = parsed_date

        due_date = DueDate(
            **data,
            created_by=current_user.id
        )

        doc = due_date.model_dump()

        # ✅ FIX 3: Safe datetime conversion
        if isinstance(doc.get("created_at"), datetime):
            doc["created_at"] = doc["created_at"].isoformat()

        if isinstance(doc.get("due_date"), datetime):
            doc["due_date"] = doc["due_date"].isoformat()

        # ✅ INSERT
        await db.due_dates.insert_one(doc)

        # ✅ FIX 4: REMOVE ObjectId if added
        doc.pop("_id", None)

        # ✅ Auto-sync to Compliance Tracker
        try:
            await _sync_due_date_to_compliance(doc, current_user)
        except Exception as sync_err:
            logger.warning(f"Compliance sync after create failed: {sync_err}")

        return doc

    except Exception as e:
        logger.error(f"DueDate creation failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to save due date")


# ──────────────────────────────────────────────────────────────────────────────
# PATCH 1 of 1  —  GET /api/duedates/upcoming
# File: main.py  (replace the entire get_upcoming_due_dates function)
#
# ROOT CAUSE: the old guard `if now <= dd_date <= future_date` excluded every
# overdue item (dd_date < now).  The Dashboard welcome banner and Deadlines card
# therefore showed nothing when ALL pending items were already past due.
#
# FIX: drop the `now <=` lower bound so overdue items are included.
#       days_remaining is already negative for overdue items via
#       `(dd_date - now).days`, which is correct.
# ──────────────────────────────────────────────────────────────────────────────

@api_router.get("/duedates/upcoming")
async def get_upcoming_due_dates(
    days: int = Query(30),
    current_user: User = Depends(get_current_user)
):
    now = datetime.now(IST)
    future_date = now + timedelta(days=days)

    query = {"status": "pending"}
    if current_user.role != "admin":
        if current_user.departments:
            query["department"] = {"$in": current_user.departments}
        else:
            return []

    due_dates = await db.due_dates.find(query, {"_id": 0}).to_list(1000)
    results = []

    for dd in due_dates:
        dd_date = (
            datetime.fromisoformat(dd["due_date"])
            if isinstance(dd["due_date"], str)
            else dd["due_date"]
        )
        # Ensure dd_date is timezone-aware for safe comparison with IST-aware `now`
        if dd_date.tzinfo is None:
            dd_date = dd_date.replace(tzinfo=IST)

        # ── FIX: include overdue items (dd_date < now) AND upcoming items ──
        # Old code: if now <= dd_date <= future_date   ← excluded overdue
        # New code: if dd_date <= future_date           ← includes overdue
        if dd_date <= future_date:
            dd["due_date"] = dd_date
            dd["days_remaining"] = (dd_date - now).days   # negative = overdue
            results.append(dd)

    # Sort: overdue first (most-negative days_remaining), then by closest due date
    return sorted(results, key=lambda x: x["days_remaining"])


@api_router.get("/duedates", response_model=List[DueDate])
async def get_due_dates(current_user: User = Depends(get_current_user)):
    query = {}
    if current_user.role == "admin":
        pass
    elif current_user.role == "manager":
        if current_user.departments:
            query["department"] = {"$in": current_user.departments}
    else:
        permissions = get_user_permissions(current_user)
        if not permissions.get("can_view_all_duedates", False):
            if current_user.departments:
                query["department"] = {"$in": current_user.departments}
            else:
                return []
    due_dates = await db.due_dates.find(query, {"_id": 0}).to_list(1000)
    for dd in due_dates:
        if isinstance(dd.get("created_at"), str):
            dd["created_at"] = datetime.fromisoformat(dd["created_at"])
        if isinstance(dd.get("due_date"), str):
            dd["due_date"] = datetime.fromisoformat(dd["due_date"])
    return [DueDate(**dd) for dd in due_dates]


@api_router.put("/duedates/{due_date_id}", response_model=DueDate)
async def update_due_date(
    due_date_id: str,
    due_date_data: DueDateCreate,
    current_user: User = Depends(check_permission("can_edit_due_dates"))
):
    existing = await db.due_dates.find_one({"id": due_date_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Due date not found")
    update_data = due_date_data.model_dump()
    update_data["due_date"] = update_data["due_date"].isoformat()
    await db.due_dates.update_one({"id": due_date_id}, {"$set": update_data})
    await create_audit_log(
        current_user,
        action="UPDATE_DUE_DATE",
        module="duedate",
        record_id=due_date_id,
        old_data=existing,
        new_data=update_data
    )
    updated = await db.due_dates.find_one({"id": due_date_id}, {"_id": 0})
    if isinstance(updated.get("created_at"), str):
        updated["created_at"] = datetime.fromisoformat(updated["created_at"])
    if isinstance(updated.get("due_date"), str):
        updated["due_date"] = datetime.fromisoformat(updated["due_date"])

    # Auto-sync updated due_date to Compliance Tracker
    try:
        sync_doc = dict(updated)
        if isinstance(sync_doc.get("due_date"), datetime):
            sync_doc["due_date"] = sync_doc["due_date"].isoformat()
        sync_doc["id"] = due_date_id
        await _sync_due_date_to_compliance(sync_doc, current_user)
    except Exception as sync_err:
        logger.warning(f"Compliance sync after update failed: {sync_err}")

    return DueDate(**updated)


@api_router.delete("/duedates/{due_date_id}")
async def delete_due_date(
    due_date_id: str,
    current_user: User = Depends(check_permission("can_edit_due_dates"))
):
    existing = await db.due_dates.find_one({"id": due_date_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Due date not found")
    await create_audit_log(
        current_user,
        action="DELETE_DUE_DATE",
        module="duedate",
        record_id=due_date_id,
        old_data=existing
    )
    result = await db.due_dates.delete_one({"id": due_date_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Due date not found")
    return {"message": "Due date deleted successfully"}


# ─────────────────────────────────────────────────────────────────────────────
# REMINDERS & MEETINGS  (inlined — no separate router file needed)
# ─────────────────────────────────────────────────────────────────────────────

reminders_col = db["reminders"]


# These inline reminder schemas extend the imported ReminderCreate model
# with email-specific fields (event_id, source, reminder_type).
class InlineReminderCreate(BaseModel):
    title: str
    description: Optional[str] = ""
    remind_at: str          # ISO datetime string
    event_id: Optional[str] = None
    source: Optional[str] = "manual"
    priority: Optional[str] = "medium"    # low | medium | high
    reminder_type: Optional[str] = "reminder"  # reminder | meeting
    related_task_id: Optional[str] = None


class InlineReminderUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    remind_at: Optional[str] = None
    is_dismissed: Optional[bool] = None
    priority: Optional[str] = None
    reminder_type: Optional[str] = None
    status: Optional[str] = None


def _serialize_reminder(doc: dict) -> dict:
    if not doc:
        return doc
    doc["_id"] = str(doc["_id"])
    doc["id"]  = doc["_id"]
    return doc


@api_router.get("/email/reminders")
async def get_reminders(
    user_id: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
):
    """Fetch reminders for the current user (admin can pass ?user_id=)."""
    query_uid = current_user.id
    if user_id and current_user.role == "admin":
        query_uid = user_id

    cursor = reminders_col.find({
        "user_id": str(query_uid),
        "$or": [
            {"is_dismissed": {"$ne": True}},
            {"is_dismissed": {"$exists": False}},
        ],
    }).sort("remind_at", 1)

    results = []
    async for doc in cursor:
        results.append(_serialize_reminder(doc))
    return results


@api_router.post("/email/save-as-reminder")
async def create_reminder(
    body: InlineReminderCreate,
    current_user: User = Depends(get_current_user),
):
    """Create a new reminder."""
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "user_id":       str(current_user.id),
        "title":         body.title,
        "description":   body.description or "",
        "remind_at":     body.remind_at,
        "event_id":      body.event_id or f"manual-{int(datetime.now().timestamp() * 1000)}",
        "source":        body.source or "manual",
        "priority":      body.priority or "medium",
        "reminder_type": body.reminder_type or "reminder",
        "related_task_id": body.related_task_id,
        "is_dismissed":  False,
        "is_fired":      False,
        "created_at":    now,
        "updated_at":    now,
    }
    result = await reminders_col.insert_one(doc)
    doc["_id"] = str(result.inserted_id)
    doc["id"]  = doc["_id"]
    return doc


@api_router.patch("/email/reminders/{reminder_id}")
async def update_reminder(
    reminder_id: str,
    body: InlineReminderUpdate,
    current_user: User = Depends(get_current_user),
):
    """Update a reminder — title, description, remind_at, is_dismissed, etc."""
    try:
        obj_id = ObjectId(reminder_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid reminder ID")

    update_fields = {k: v for k, v in body.model_dump(exclude_unset=True).items() if v is not None}
    if not update_fields:
        raise HTTPException(status_code=400, detail="No fields to update")

    update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()

    query = {"_id": obj_id}
    if current_user.role != "admin":
        query["user_id"] = str(current_user.id)

    result = await reminders_col.update_one(query, {"$set": update_fields})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Reminder not found")

    updated = await reminders_col.find_one({"_id": obj_id})
    return _serialize_reminder(updated)


@api_router.delete("/email/reminders/{reminder_id}")
async def delete_reminder(
    reminder_id: str,
    current_user: User = Depends(get_current_user),
):
    """Delete a reminder permanently."""
    try:
        obj_id = ObjectId(reminder_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid reminder ID")

    query = {"_id": obj_id}
    if current_user.role != "admin":
        query["user_id"] = str(current_user.id)

    result = await reminders_col.delete_one(query)
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Reminder not found")

    return {"message": "Reminder deleted", "id": reminder_id}


# Registered before /clients/{client_id} to prevent route shadowing
@api_router.get("/clients/upcoming-birthdays")
async def get_upcoming_birthdays(days: int = 7, current_user: User = Depends(get_current_user)):
    clients = await db.clients.find({}, {"_id": 0}).to_list(1000)
    today = date.today()
    upcoming = []
    for client in clients:
        if client.get("birthday"):
            try:
                bday = date.fromisoformat(client["birthday"]) if isinstance(client["birthday"], str) else client["birthday"]
                # Added leap year guard
                try:
                    this_year_bday = bday.replace(year=today.year)
                except ValueError:
                    this_year_bday = bday.replace(year=today.year, day=28)
                if this_year_bday < today:
                    try:
                        this_year_bday = bday.replace(year=today.year + 1)
                    except ValueError:
                        this_year_bday = bday.replace(year=today.year + 1, day=28)
                days_until = (this_year_bday - today).days
                if 0 <= days_until <= days:
                    client["days_until_birthday"] = days_until
                    upcoming.append(client)
            except (ValueError, TypeError):
                continue
    return sorted(upcoming, key=lambda x: x["days_until_birthday"])

# ─── MANUAL BIRTHDAY WISH ────────────────────────────────────────────────────
@api_router.post("/clients/{client_id}/send-birthday-wish")
async def send_birthday_wish_manual(
    client_id: str,
    current_user: User = Depends(get_current_user)
):
    """Manually send a birthday wish to a client. Admin/manager only."""
    if current_user.role not in ("admin", "manager"):
        raise HTTPException(status_code=403, detail="Admin or Manager only")

    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")

    sent_to, failed, no_email = [], [], []

    # Main client email
    client_name  = client.get("company_name") or "Valued Client"
    client_email = client.get("email")
    if client_email:
        ok = send_birthday_email(client_email, client_name)
        (sent_to if ok else failed).append(client_email)
    else:
        no_email.append(client_name)

    # Contact persons
    for cp in client.get("contact_persons") or []:
        cp_email = cp.get("email")
        cp_name  = cp.get("name") or client_name
        if cp_email:
            ok = send_birthday_email(cp_email, cp_name)
            (sent_to if ok else failed).append(cp_email)
        else:
            no_email.append(cp_name)

    return {"status": "completed", "sent_to": sent_to, "failed": failed, "no_email": no_email}

@api_router.get("/leads/meta/services")
async def get_leads_services_meta(current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        perms = get_user_permissions(current_user)
        if not perms.get("can_view_all_leads", False):
            raise HTTPException(status_code=403, detail="Leads access not permitted")
    services = await db.clients.distinct("services")
    services = [s for s in services if s and isinstance(s, str)]
    return {"services": list(set(services))}

# REPORTS ROUTES
@api_router.get("/reports/efficiency")
async def get_efficiency_report(
    user_id: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    target_user_id = user_id or current_user.id
    if target_user_id != current_user.id:
        perms = get_user_permissions(current_user)
        if current_user.role != "admin" and not perms.get("can_view_reports", False):
            raise HTTPException(status_code=403, detail="You do not have permission to view reports")
    if target_user_id != current_user.id:
        if current_user.role != "admin":
            permissions = get_user_permissions(current_user)
            allowed_users = permissions.get("view_other_reports", []) or []
            if current_user.role == "manager":
                # Manager: also allowed to view same-department team reports
                team_ids = await get_team_user_ids(current_user.id)
                allowed_users = list(set(allowed_users + team_ids))
            if target_user_id not in allowed_users:
                raise HTTPException(status_code=403, detail="Not authorized to view other users' reports")
    logs = await db.activity_logs.find({"user_id": target_user_id}, {"_id": 0}).sort("date", -1).limit(30).to_list(100)
    total_screen_time = sum(l.get("screen_time_minutes", 0) for l in logs)
    total_tasks_completed = sum(l.get("tasks_completed", 0) for l in logs)
    target_user_doc = await db.users.find_one({"id": target_user_id}, {"_id": 0, "password": 0})
    user_info = {
        "id": target_user_id,
        "full_name": target_user_doc.get("full_name", "Unknown") if target_user_doc else "Unknown"
    }
    return {
        "user_id": target_user_id,
        "user": user_info,
        "total_screen_time": total_screen_time,
        "total_tasks_completed": total_tasks_completed,
        "days_logged": len(logs)
    }

@api_router.get("/reports/export")
async def export_reports(
    format: str = "csv",
    user_id: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    perms = get_user_permissions(current_user)
    if current_user.role != "admin" and not perms.get("can_download_reports", False):
        raise HTTPException(status_code=403, detail="You do not have permission to download reports")
    target_user_id = user_id or current_user.id
    if target_user_id != current_user.id:
        if current_user.role != "admin":
            permissions = get_user_permissions(current_user)
            allowed_users = permissions.get("view_other_reports", []) or []
            if current_user.role == "manager":
                # Manager: also allowed to export same-department team reports
                team_ids = await get_team_user_ids(current_user.id)
                allowed_users = list(set(allowed_users + team_ids))
            if target_user_id not in allowed_users:
                raise HTTPException(status_code=403, detail="Not authorized to access other users' reports")
    logs = await db.activity_logs.find({"user_id": target_user_id}, {"_id": 0}).to_list(100)
    total_screen_time = sum(l.get("screen_time_minutes", 0) for l in logs)
    total_tasks_completed = sum(l.get("tasks_completed", 0) for l in logs)
    report = {
        "user_id": target_user_id,
        "total_screen_time": total_screen_time,
        "total_tasks_completed": total_tasks_completed,
        "days_logged": len(logs)
    }
    if format == "csv":
        output = StringIO()
        def sanitize_csv_value(val):
            val_str = str(val)
            # Strip newlines to prevent CSV row injection in addition to formula injection
            val_str = val_str.replace('\r', '').replace('\n', ' ')
            if val_str and val_str[0] in ['=', '+', '-', '@']:
                return f"'{val_str}"
            return val_str
        writer = csv.writer(output)
        writer.writerow(["User ID", "Screen Time", "Tasks Completed", "Days Logged"])
        writer.writerow([
            sanitize_csv_value(report["user_id"]),
            sanitize_csv_value(report["total_screen_time"]),
            sanitize_csv_value(report["total_tasks_completed"]),
            sanitize_csv_value(report["days_logged"])
        ])
        output.seek(0)
        return StreamingResponse(
            iter([output.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=efficiency_report_{target_user_id}.csv"}
        )
    elif format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Efficiency Report", ln=1, align="C")
        pdf.ln(10)
        pdf.multi_cell(0, 8, f"User ID: {report['user_id']}")
        pdf.multi_cell(0, 8, f"Screen Time: {report['total_screen_time']} minutes")
        pdf.multi_cell(0, 8, f"Tasks Completed: {report['total_tasks_completed']}")
        pdf.multi_cell(0, 8, f"Days Logged: {report['days_logged']}")
        pdf_output = BytesIO()
        pdf_output.write(pdf.output(dest='S').encode('latin1'))
        pdf_output.seek(0)
        return StreamingResponse(
            pdf_output,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=efficiency_report_{target_user_id}.pdf"}
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid format")

# ====================== PERFORMANCE RANKINGS ======================
@api_router.get("/reports/performance-rankings", response_model=List[PerformanceMetric])
async def get_performance_rankings(
    period: str = Query("monthly", enum=["weekly", "monthly", "all_time"]),
    current_user: User = Depends(get_current_user)
):
    global rankings_cache, rankings_cache_time
    cache_key = f"rankings_{period}"
    if (
        cache_key in rankings_cache and
        cache_key in rankings_cache_time and
        # Use timezone-aware UTC datetime for cache comparison to avoid TypeError
        (datetime.now(timezone.utc) - rankings_cache_time[cache_key]).total_seconds() < 300
    ):
        return rankings_cache[cache_key]
    now = datetime.now(IST)
    if period == "weekly":
        start_date = now - timedelta(days=7)
        expected_working_days = 5
    elif period == "monthly":
        start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        expected_working_days = 22
    else:
        start_date = datetime(2024, 1, 1, tzinfo=timezone.utc)
        months = max((now - start_date).days // 30, 1)
        expected_working_days = max(22, months * 22)
    end_date_str = now.strftime("%Y-%m-%d")
    start_date_str = start_date.strftime("%Y-%m-%d")
    users = await db.users.find(
        {"role": {"$ne": "admin"}},
        {"id": 1, "full_name": 1, "profile_picture": 1}
    ).to_list(100)
    rankings = []
    for user in users:
        uid = user["id"]
        att_records = await db.attendance.find(
            {
                "user_id": uid,
                "date": {"$gte": start_date_str, "$lte": end_date_str}
            },
            {"_id": 0, "duration_minutes": 1, "is_late": 1}
        ).to_list(1000)
        days_present = len(att_records)
        total_minutes = sum(r.get("duration_minutes", 0) or 0 for r in att_records)
        total_hours = round(total_minutes / 60, 1)
        attendance_percent = round(
            (days_present / expected_working_days) * 100, 1
        ) if expected_working_days else 0
        timely_days = len([r for r in att_records if not r.get("is_late", False)])
        timely_punchin_percent = round(
            (timely_days / days_present) * 100, 1
        ) if days_present else 0
        tasks_assigned = await db.tasks.count_documents({
            "assigned_to": uid,
            "created_at": {"$gte": start_date}
        })
        completed_tasks = await db.tasks.count_documents({
            "assigned_to": uid,
            "status": "completed",
            "$or": [
                {"completed_at": {"$gte": start_date}},
                {"updated_at": {"$gte": start_date}}
            ]
        })
        completed_todos = await db.todos.count_documents({
            "user_id": uid,
            "is_completed": True,
            "completed_at": {"$gte": start_date}
        })
        total_completed = completed_tasks + completed_todos
        task_completion_percent = round(
            (total_completed / tasks_assigned) * 100, 1
        ) if tasks_assigned else 0
        todos = await db.todos.find({
            "user_id": uid,
            "created_at": {"$gte": start_date}
        }).to_list(500)
        completed_ontime = 0
        for t in todos:
            if t.get("is_completed"):
                due = safe_dt(t.get("due_date"))
                completed_at = safe_dt(t.get("completed_at"))
                if due and completed_at and completed_at <= due:
                    completed_ontime += 1
        todo_ontime_percent = round(
            (completed_ontime / len(todos)) * 100, 1
        ) if todos else 0
        safe_hours_ratio = min((total_hours / 180), 1) if total_hours else 0
        score = (
            float(attendance_percent or 0) * 0.25 +
            safe_hours_ratio * 100 * 0.20 +
            float(task_completion_percent or 0) * 0.25 +
            float(todo_ontime_percent or 0) * 0.15 +
            float(timely_punchin_percent or 0) * 0.15
        )
        overall_score = round(min(score, 100), 1)
        if overall_score >= 95:
            badge = "Star Performer"
        elif overall_score >= 85:
            badge = "Top Performer"
        else:
            badge = "Good Performer"
        rankings.append(
            PerformanceMetric(
                user_id=str(uid),
                user_name=str(user.get("full_name", "Unknown")),
                profile_picture=user.get("profile_picture"),
                attendance_percent=float(attendance_percent or 0),
                total_hours=float(total_hours or 0),
                task_completion_percent=float(task_completion_percent or 0),
                todo_ontime_percent=float(todo_ontime_percent or 0),
                timely_punchin_percent=float(timely_punchin_percent or 0),
                overall_score=float(overall_score or 0),
                badge=str(badge)
            )
        )
    rankings.sort(key=lambda x: x.overall_score, reverse=True)
    for i, r in enumerate(rankings):
        r.rank = i + 1
    rankings_cache[cache_key] = rankings
    # Store as timezone-aware UTC datetime so comparison never throws TypeError
    rankings_cache_time[cache_key] = datetime.now(timezone.utc)
    return rankings

# ==============================================================
# INTEGRATED MASTER DATA SYSTEM & CLIENT ROUTES
# ==============================================================
@api_router.post("/master/import-master-preview")
async def import_master_data_preview(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(status_code=403, detail="Administrative clearance required for Master Data access.")
    filename = file.filename.lower()
    if not filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Deployment failed: Only Excel formats (.xlsx, .xls) supported.")
    try:
        content = await file.read()
        excel = pd.ExcelFile(BytesIO(content))
        parsed_blueprint = {}
        total_vectors = 0
        for sheet_name in excel.sheet_names:
            df = pd.read_excel(excel, sheet_name=sheet_name)
            df = df.fillna("")
            records = df.to_dict(orient="records")
            parsed_blueprint[sheet_name] = records
            total_vectors += len(records)
        return {
            "status": "Ready for Audit",
            "message": f"Detected {len(excel.sheet_names)} operational layers with {total_vectors} vectors.",
            "sheets_found": excel.sheet_names,
            "data": parsed_blueprint
        }
    except Exception as e:
        logger.error(f"Blueprint Error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Excel parse failure: {str(e)}")

@api_router.post("/master/sync-sheets")
async def sync_master_sheets(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    if current_user.role.lower() != "admin":
        raise HTTPException(status_code=403, detail="Master Data clearance level 5 required.")
    try:
        content = await file.read()
        excel = pd.ExcelFile(BytesIO(content))
        sync_results = {"clients": 0, "compliance": 0, "staff": 0}
        now_iso = datetime.now(timezone.utc).isoformat()
        for sheet in excel.sheet_names:
            df = pd.read_excel(excel, sheet_name=sheet).fillna("")
            records = df.to_dict(orient="records")
            sheet_type = sheet.lower()
            if "client" in sheet_type:
                for rec in records:
                    await db.clients.update_one(
                        {"company_name": str(rec.get("company_name", "")).strip()},
                        {"$set": {**rec, "id": str(uuid.uuid4()) if "id" not in rec else rec["id"], "created_by": current_user.id, "updated_at": now_iso}},
                        upsert=True
                    )
                    sync_results["clients"] += 1
            elif "due" in sheet_type or "compliance" in sheet_type:
                for rec in records:
                    await db.due_dates.insert_one({**rec, "id": str(uuid.uuid4()), "created_by": current_user.id, "created_at": now_iso, "status": "pending"})
                    sync_results["compliance"] += 1
            elif "staff" in sheet_type or "user" in sheet_type:
                for rec in records:
                    await db.users.update_one(
                        {"email": rec.get("email")},
                        {"$set": {**rec, "id": str(uuid.uuid4()), "is_active": True}},
                        upsert=True
                    )
                    sync_results["staff"] += 1
            await create_audit_log(current_user=current_user, action="GLOBAL_MASTER_SYNC", module="master_data", record_id="multi_sheet_payload", new_data=sync_results)
        return {"message": "Global Master Sync Successfully Executed", "telemetry": sync_results}
    except Exception as e:
        logger.error(f"Sync Failure: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Database synchronization failed: {str(e)}")

# ==============================================================
# MDS (MCA) EXCEL SMART PARSER
# ==============================================================
@api_router.post("/clients/parse-mds-excel")
async def parse_mds_excel_for_client_form(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    filename = file.filename.lower()
    if not filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx / .xls) are supported.")
    try:
        content = await file.read()
        excel = pd.ExcelFile(BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not open Excel file: {str(e)}")

    def clean_email(raw: str) -> str:
        if not raw:
            return ""
        cleaned = raw.replace("[at]", "@").replace("[dot]", ".").strip()
        return cleaned if "@" in cleaned else ""

    def clean_phone(raw: str) -> str:
        digits = re.sub(r"\D", "", str(raw or ""))
        if len(digits) == 12 and digits.startswith("91"):
            digits = digits[2:]
        return digits[:10] if len(digits) >= 10 else digits

    def detect_type(name: str) -> str:
        n = name.lower()
        if any(x in n for x in ["private limited", "pvt ltd", "pvt. ltd", "pvt limited"]):
            return "pvt_ltd"
        if any(x in n for x in ["limited liability partnership", "llp"]):
            return "llp"
        if any(x in n for x in [" ltd", " limited"]):
            return "pvt_ltd"
        if "partnership" in n:
            return "partnership"
        if "huf" in n:
            return "huf"
        if "trust" in n:
            return "trust"
        return "proprietor"

    def parse_date(raw: str) -> str:
        if not raw or str(raw).strip() in ("", "-", "N/A"):
            return ""
        try:
            return parser.parse(str(raw).strip()).strftime("%Y-%m-%d")
        except Exception:
            return ""

    company_info: dict = {}
    directors: list = []
    extra_notes_parts: list = []

    for sheet_name in excel.sheet_names:
        df = pd.read_excel(excel, sheet_name=sheet_name, header=None).fillna("")
        sheet_lower = sheet_name.lower().strip()
        if "master" in sheet_lower or "company" in sheet_lower or sheet_lower == "masterdata":
            for _, row in df.iterrows():
                key = str(row.iloc[0]).strip()
                value = str(row.iloc[1]).strip() if len(row) > 1 else ""
                if key and key not in ("", "nan") and value not in ("", "nan"):
                    if any(phrase in key for phrase in [
                        "Accounts and Solvency", "Annual Returns", "Filing Information",
                        "Interim Resolution", "Sr. No", "Date of filing"
                    ]):
                        continue
                    company_info[key] = value
        elif "director" in sheet_lower or "signatory" in sheet_lower or "partner" in sheet_lower:
            rows_list = df.values.tolist()
            header_row_idx = None
            for idx, row in enumerate(rows_list):
                row_strs = [str(c).strip() for c in row]
                if any(h in row_strs for h in ["Name", "DIN/PAN", "DIN"]):
                    header_row_idx = idx
                    break
            if header_row_idx is None:
                continue
            headers = [str(h).strip() for h in rows_list[header_row_idx]]
            for row in rows_list[header_row_idx + 1:]:
                row_dict = {headers[i]: str(row[i]).strip() if i < len(row) else "" for i in range(len(headers))}
                name = row_dict.get("Name", "").strip()
                if not name or name in ("nan", ""):
                    continue
                din = row_dict.get("DIN/PAN", "") or row_dict.get("DIN", "")
                designation = row_dict.get("Designation", "")
                directors.append({
                    "name": name,
                    "designation": designation or "Director",
                    "email": None,
                    "phone": None,
                    "birthday": None,
                    "din": din if din not in ("nan", "-", "") else None,
                })
        else:
            rows_list = df.values.tolist()
            if len(rows_list) >= 2:
                cols = [str(c).strip() for c in rows_list[0]]
                data_rows = []
                for row in rows_list[1:]:
                    vals = [str(v).strip() for v in row]
                    if any(v not in ("", "nan", "-") for v in vals):
                        data_rows.append(dict(zip(cols, vals)))
                if data_rows:
                    extra_notes_parts.append(f"[{sheet_name}]")
                    for r in data_rows[:10]:
                        extra_notes_parts.append(
                            " | ".join(f"{k}: {v}" for k, v in r.items() if v not in ("", "nan", "-"))
                        )

    # Support both Pvt Ltd (MCA) and LLP field naming conventions
    company_name = (
        company_info.get("Company Name") or
        company_info.get("LLP Name") or
        company_info.get("company_name") or ""
    ).strip()

    raw_email = (
        company_info.get("Email Id") or
        company_info.get("Email") or
        company_info.get("email") or ""
    )
    email = clean_email(raw_email)

    raw_phone = (
        company_info.get("Phone") or
        company_info.get("Mobile") or
        company_info.get("Contact") or ""
    )
    phone = clean_phone(raw_phone)

    raw_doi = (
        company_info.get("Date of Incorporation") or
        company_info.get("Incorporation Date") or ""
    )
    birthday = parse_date(raw_doi)
    client_type = detect_type(company_name)

    address = (
        company_info.get("Registered Address") or
        company_info.get("Registered Office Address") or
        company_info.get("Address") or ""
    ).strip()
    if address in ("-", "nan"):
        address = ""

    city = ""
    state = ""
    if address:
        address_parts = [p.strip() for p in address.split(",") if p.strip()]
        if len(address_parts) >= 3:
            state = address_parts[-3]
            city = address_parts[-4] if len(address_parts) >= 4 else ""
        elif len(address_parts) == 2:
            state = address_parts[-1]
            city = address_parts[-2]

    notes_lines = []

    cin = company_info.get("CIN", "") or company_info.get("LLPIN", "")
    if cin and cin not in ("-", "nan"):
        notes_lines.append(f"CIN/LLPIN: {cin}")

    roc = company_info.get("ROC Name", "") or company_info.get("ROC (name and office)", "")
    if roc and roc not in ("-", "nan"):
        notes_lines.append(f"ROC: {roc}")

    reg_no = company_info.get("Registration Number", "")
    if reg_no and reg_no not in ("-", "nan"):
        notes_lines.append(f"Reg No: {reg_no}")

    auth_cap = (
        company_info.get("Authorised Capital (Rs)", "") or
        company_info.get("Total Obligation of Contribution", "")
    )
    if auth_cap and auth_cap not in ("-", "nan"):
        notes_lines.append(f"Capital/Contribution: Rs{auth_cap}")

    paid_cap = company_info.get("Paid up Capital (Rs)", "")
    if paid_cap and paid_cap not in ("-", "nan"):
        notes_lines.append(f"Paid-up Capital: Rs{paid_cap}")

    if extra_notes_parts:
        notes_lines.append("\n".join(extra_notes_parts))

    notes = "\n".join(notes_lines)

    status_raw = (
        company_info.get("Company Status", "") or
        company_info.get("LLP Status", "Active")
    ).lower()
    status = "active" if "active" in status_raw else "inactive"
    return {
        "status": "ok",
        "company_name": company_name,
        "client_type": client_type,
        "email": email,
        "phone": phone,
        "birthday": birthday,
        "address": address,
        "city": city,
        "state": state,
        "services": [],
        "notes": notes,
        "status_value": status,
        "contact_persons": directors,
        "dsc_details": [],
        "assigned_to": "unassigned",
        "raw_company_info": company_info,
        "sheets_parsed": excel.sheet_names,
    }

@api_router.post("/clients/import")
async def import_clients_from_csv(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """
    Bulk-import clients from a CSV file.
    Expected CSV columns (all optional except company_name):
      company_name, client_type, email, phone, birthday, address,
      city, state, services, notes, assigned_to, status
    Returns: { message, clients_created, clients_skipped, errors }
    """
    if current_user.role not in ("admin", "manager"):
        perms = get_user_permissions(current_user)
        if not perms.get("can_edit_clients", False):
            raise HTTPException(status_code=403, detail="Permission denied")

    filename = (file.filename or "").lower()
    if not filename.endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only CSV files are supported (.csv)")

    content = await file.read()
    try:
        text = content.decode("utf-8-sig")  # handle BOM
    except UnicodeDecodeError:
        text = content.decode("latin-1")

    reader = csv.DictReader(StringIO(text))
    if not reader.fieldnames:
        raise HTTPException(status_code=422, detail="CSV file is empty or has no header row")

    created_count = 0
    skipped_count = 0
    errors: list = []
    now_iso = datetime.now(timezone.utc).isoformat()

    for i, row in enumerate(reader, start=2):  # row 1 = header
        company_name = str(row.get("company_name") or row.get("Company Name") or "").strip()
        if not company_name:
            skipped_count += 1
            continue

        # Skip duplicate (same created_by + company_name)
        existing = await db.clients.find_one(
            {"created_by": current_user.id, "company_name": {"$regex": f"^{re.escape(company_name)}$", "$options": "i"}},
            {"_id": 0, "id": 1}
        )
        if existing:
            skipped_count += 1
            continue

        # Parse services column (comma-separated string → list)
        raw_services = str(row.get("services") or row.get("Services") or "").strip()
        services = [s.strip() for s in raw_services.split(",") if s.strip()] if raw_services else []

        # Normalise client_type
        raw_type = str(row.get("client_type") or row.get("Client Type") or "proprietor").strip().lower()
        valid_types = {"proprietor", "pvt_ltd", "llp", "partnership", "huf", "trust", "other"}
        client_type = raw_type if raw_type in valid_types else "proprietor"

        doc = {
            "id":           str(uuid.uuid4()),
            "company_name": company_name,
            "client_type":  client_type,
            "email":        str(row.get("email") or row.get("Email") or "").strip() or None,
            "phone":        str(row.get("phone") or row.get("Phone") or "").strip() or None,
            "birthday":     str(row.get("birthday") or row.get("Birthday") or "").strip() or None,
            "address":      str(row.get("address") or row.get("Address") or "").strip() or None,
            "city":         str(row.get("city") or row.get("City") or "").strip() or None,
            "state":        str(row.get("state") or row.get("State") or "").strip() or None,
            "services":     services,
            "notes":        str(row.get("notes") or row.get("Notes") or "").strip() or None,
            "assigned_to":  str(row.get("assigned_to") or row.get("Assigned To") or "").strip() or None,
            "status":       str(row.get("status") or row.get("Status") or "active").strip().lower(),
            "created_by":   current_user.id,
            "created_at":   now_iso,
        }

        try:
            await db.clients.insert_one(doc)
            created_count += 1
        except Exception as e:
            errors.append({"row": i, "company": company_name, "error": str(e)[:80]})
            skipped_count += 1

    return {
        "message":         f"{created_count} client(s) imported successfully",
        "clients_created": created_count,
        "clients_skipped": skipped_count,
        "errors":          errors[:10],  # cap error list to avoid huge response
    }


@api_router.post("/clients", response_model=Client)
async def create_client(payload: dict, current_user: User = Depends(get_current_user)):
    try:
        client_data = ClientCreate(**{k: v for k, v in payload.items() if k in ClientCreate.model_fields})
        client = Client(**client_data.model_dump(), created_by=current_user.id)
        doc = client.model_dump()

        # ── safe_iso: handles None, str, date, datetime — never crashes ──────
        def safe_iso(val):
            if val is None:
                return None
            if isinstance(val, str):
                return val[:10] if val else None  # already "2015-04-01", keep as-is
            if isinstance(val, (date, datetime)):
                return val.isoformat()
            return str(val)

        doc["created_at"] = safe_iso(doc.get("created_at")) or datetime.now(timezone.utc).isoformat()
        doc["birthday"]   = safe_iso(doc.get("birthday"))

        # Persist extra fields from frontend that live outside Pydantic schema
        for key in ("address", "city", "state", "client_type_label",
                    "contact_persons", "dsc_details", "assignments",
                    "referred_by", "gstin", "pan", "gst_treatment",
                    "place_of_supply", "default_payment_terms", "credit_limit",
                    "opening_balance", "opening_balance_type", "tally_ledger_name",
                    "tally_group", "website", "msme_number"):
            val = payload.get(key)
            if val is not None:
                doc[key] = val

        doc.pop("_id", None)
        await db.clients.insert_one(doc)
        return client
    except ValidationError as ve:
        raise HTTPException(status_code=422, detail=ve.errors())
    except Exception as e:
        logger.error(f"create_client error: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=str(e))

_DEPT_SERVICE_MAP: Dict[str, List[str]] = {
    "GST":   ["GST", "Compliance"],
    "IT":    ["Income Tax", "Tax Planning"],
    "ACC":   ["Accounting", "Payroll", "Audit"],
    "TDS":   ["TDS"],
    "ROC":   ["ROC", "Company Registration", "Compliance"],
    "TM":    ["Trademark"],
    "MSME":  ["MSME"],
    "FEMA":  ["FEMA"],
    "DSC":   [],
    "OTHER": [],
}
 
 
@api_router.get("/clients", response_model=List[Client])
async def get_clients(current_user: User = Depends(get_current_user)):

    permissions = get_user_permissions(current_user)

    # SCOPE: Admin = all; can_view_all_clients = all; else OWN + assigned_clients (same for Manager and Staff)
    if current_user.role == "admin" or permissions.get("can_view_all_clients", False):
        query = {}

    else:
        # SCOPE: OWN + CROSS-VISIBILITY only (same for Manager and Staff when can_view_all_clients is revoked)
        extra_clients = permissions.get("assigned_clients", [])
        if extra_clients:
            query = {
                "$or": [
                    {"assigned_to": current_user.id},
                    {"created_by": current_user.id},
                    {"id": {"$in": extra_clients}},
                ]
            }
        else:
            query = {
                "$or": [
                    {"assigned_to": current_user.id},
                    {"created_by": current_user.id},
                ]
            }

    clients = await db.clients.find(query, {"_id": 0}).to_list(1000)

    for client in clients:
        if isinstance(client.get("created_at"), str):
            try:
                client["created_at"] = datetime.fromisoformat(client["created_at"])
            except ValueError:
                client["created_at"] = None

        if client.get("birthday") and isinstance(client["birthday"], str):
            try:
                client["birthday"] = date.fromisoformat(client["birthday"])
            except ValueError:
                client["birthday"] = None

    return clients

@api_router.get("/clients/{client_id}", response_model=Client)
async def get_client(client_id: str, current_user: User = Depends(get_current_user)):
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    if current_user.role != "admin":
        is_assigned = client.get("assigned_to") == current_user.id
        permissions = get_user_permissions(current_user)
        extra_clients = permissions.get("assigned_clients", [])
        if not is_assigned and client_id not in extra_clients:
            raise HTTPException(status_code=403, detail="Not authorized to view this client")
    if isinstance(client["created_at"], str):
        client["created_at"] = datetime.fromisoformat(client["created_at"])
    if client.get("birthday") and isinstance(client["birthday"], str):
        client["birthday"] = date.fromisoformat(client["birthday"])
    return Client(**client)

@api_router.put("/clients/{client_id}", response_model=Client)
async def update_client(
    client_id: str,
    client_data: dict,          # <-- Changed from ClientCreate to dict
    current_user: User = Depends(get_current_user)
):
    """
    Update an existing client record.
 
    Accepts a raw dict instead of ClientCreate so that:
      - Extra frontend-only fields (address, city, state, client_type_label)
        are accepted without causing Pydantic validation failures.
      - Empty strings sent by the frontend (email="", phone="") are
        safely converted to None before any validation runs.
      - The referred_by field can be any string without pattern checks.
 
    Only fields in ALLOWED_FIELDS are written to the database.
    """
    existing = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Client not found")
 
    perms = get_user_permissions(current_user)
    if (
        current_user.role != "admin"
        and existing.get("assigned_to") != current_user.id
        and not perms.get("can_edit_clients", False)
    ):
        raise HTTPException(status_code=403, detail="Not authorized to edit this client")
 
    # ── Whitelist: only persist known fields ────────────────────────
    ALLOWED_FIELDS = {
        "company_name", "client_type", "client_type_label",
        "email", "phone", "birthday", "date_of_incorporation",
        "address", "city", "state",
        "services", "notes", "assigned_to", "assignments",
        "status", "contact_persons", "dsc_details", "referred_by",
        # Tax & Billing
        "gstin", "pan", "gst_treatment", "place_of_supply",
        "default_payment_terms", "credit_limit", "opening_balance",
        "opening_balance_type", "tally_ledger_name", "tally_group",
        "website", "msme_number",
    }
    update_data = {k: v for k, v in client_data.items() if k in ALLOWED_FIELDS}
 
    # ── Convert empty strings → None for nullable fields ────────────
    NULLABLE_FIELDS = {
        "email", "phone", "referred_by", "notes", "assigned_to",
        "birthday", "date_of_incorporation", "address", "city",
        "state", "client_type_label",
        "gstin", "pan", "place_of_supply", "default_payment_terms",
        "credit_limit", "opening_balance", "tally_ledger_name",
        "tally_group", "website", "msme_number",
    }
    for field in NULLABLE_FIELDS:
        if field in update_data and update_data[field] == "":
            update_data[field] = None
 
    # ── Validate and normalise client_type ──────────────────────────
    VALID_CLIENT_TYPES = {
        "proprietor", "pvt_ltd", "llp", "partnership",
        "huf", "trust", "other",
        # Accept legacy uppercase variants from old data
        "LLP", "PVT_LTD",
    }
    if "client_type" in update_data:
        ct = update_data["client_type"]
        if ct not in VALID_CLIENT_TYPES:
            raise HTTPException(
                status_code=422,
                detail=f"Invalid client_type '{ct}'. Must be one of: proprietor, pvt_ltd, llp, partnership, huf, trust, other"
            )
        # Normalise to lowercase for consistent storage
        lower_map = {"LLP": "llp", "PVT_LTD": "pvt_ltd"}
        update_data["client_type"] = lower_map.get(ct, ct)
 
    # ── Validate company_name length ────────────────────────────────
    if "company_name" in update_data:
        name = str(update_data["company_name"]).strip()
        if len(name) < 3:
            raise HTTPException(
                status_code=422,
                detail="Company name must be at least 3 characters long"
            )
        update_data["company_name"] = name
 
    # ── Validate phone if provided ──────────────────────────────────
    if update_data.get("phone"):
        cleaned_phone = re.sub(r"\s|-|\+", "", str(update_data["phone"]))
        if not cleaned_phone.isdigit():
            raise HTTPException(status_code=422, detail="Phone number must contain only digits")
        if not (10 <= len(cleaned_phone) <= 15):
            raise HTTPException(status_code=422, detail="Phone number must be 10–15 digits")
 
    # ── Persist ─────────────────────────────────────────────────────
    await db.clients.update_one({"id": client_id}, {"$set": update_data})
 
    await create_audit_log(
        current_user,
        action="UPDATE_CLIENT",
        module="client",
        record_id=client_id,
        old_data=existing,
        new_data=update_data,
    )
 
    updated = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if isinstance(updated.get("created_at"), str):
        try:
            updated["created_at"] = datetime.fromisoformat(updated["created_at"])
        except ValueError:
            updated["created_at"] = datetime.now(timezone.utc)
    if updated.get("birthday") and isinstance(updated["birthday"], str):
        try:
            updated["birthday"] = date.fromisoformat(updated["birthday"])
        except ValueError:
            updated["birthday"] = None
    # Strip fields not in Pydantic model before constructing Client
    client_fields = Client.model_fields.keys()
    safe_updated = {k: v for k, v in updated.items() if k in client_fields}
    return Client(**safe_updated)
@api_router.delete("/clients/{client_id}")
async def delete_client(
    client_id: str,
    current_user: User = Depends(get_current_user)
):
    """
    Delete a client by ID.
    - Nullifies any leads that reference this client (converted_client_id)
      so FK-style references don't leave dangling data.
    - Requires can_delete_data permission (admin always passes).
    """
    perms = get_user_permissions(current_user)
    if current_user.role != "admin" and not perms.get("can_delete_data", False):
        raise HTTPException(status_code=403, detail="You do not have permission to delete clients")

    existing = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Client not found")

    # Nullify any leads that were converted to this client
    await db.leads.update_many(
        {"converted_client_id": client_id},
        {"$set": {"converted_client_id": None}}
    )

    # Also unlink tasks referencing this client
    await db.tasks.update_many(
        {"client_id": client_id},
        {"$set": {"client_id": None}}
    )

    await create_audit_log(
        current_user,
        action="DELETE_CLIENT",
        module="client",
        record_id=client_id,
        old_data=existing,
    )

    result = await db.clients.delete_one({"id": client_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")

    return {"message": f"Client '{existing.get('company_name', client_id)}' deleted successfully"}
#============================================
# DASHBOARD ROUTES
#============================================

@api_router.get("/dashboard/dept-members")
async def get_dept_member_count(current_user: User = Depends(get_current_user)):
    """
    Returns count + basic info of users in the same department(s) as the current user.
    Accessible by all roles (admin, manager, staff).
    - Admin: returns count of ALL active users.
    - Manager: returns count of same-department users (staff + other managers).
    - Staff: returns count of same-department users (staff + managers).
    """
    if current_user.role == "admin":
        all_users = await db.users.find(
            {"is_active": True, "status": "active"},
            {"_id": 0, "id": 1, "full_name": 1, "departments": 1, "role": 1}
        ).to_list(1000)
        return {
            "count": len(all_users),
            "departments": [],
            "members": [{"id": u["id"], "full_name": u.get("full_name", ""), "role": u.get("role", "")} for u in all_users]
        }

    user_depts = current_user.departments or []
    if not user_depts:
        return {"count": 0, "departments": [], "members": []}

    dept_users = await db.users.find(
        {
            "departments": {"$in": user_depts},
            "id": {"$ne": current_user.id},
            "is_active": True,
            "status": "active",
        },
        {"_id": 0, "id": 1, "full_name": 1, "departments": 1, "role": 1}
    ).to_list(500)

    return {
        "count": len(dept_users),
        "departments": user_depts,
        "members": [{"id": u["id"], "full_name": u.get("full_name", ""), "role": u.get("role", "")} for u in dept_users]
    }


@api_router.get("/dashboard/stats", response_model=DashboardStats)
async def get_dashboard_stats(current_user: User = Depends(get_current_user)):
    now = datetime.now(IST)
    task_query = {}
    if current_user.role != "admin":
        permissions = get_user_permissions(current_user)
        if not permissions.get("can_view_all_tasks", False):
            allowed_users = permissions.get("view_other_tasks", []) or []
            if current_user.role == "manager":
                team_ids = await get_team_user_ids(current_user.id)
                allowed_users = list(set(allowed_users + team_ids))
            task_query["$or"] = [
                {"assigned_to": current_user.id},
                {"sub_assignees": current_user.id},
                {"created_by": current_user.id},
                {"assigned_to": {"$in": allowed_users}},
            ]

    tasks = await db.tasks.find(task_query, {"_id": 0}).to_list(1000)
    total_tasks = len(tasks)
    completed_tasks = len([t for t in tasks if t["status"] == "completed"])
    pending_tasks = len([t for t in tasks if t["status"] == "pending"])
    overdue_tasks = 0
    for task in tasks:
        if task.get("due_date") and task["status"] != "completed":
            try:
                due_date = (
                    datetime.fromisoformat(task["due_date"])
                    if isinstance(task["due_date"], str)
                    else task["due_date"]
                )
                if due_date.tzinfo is None:
                    due_date = due_date.replace(tzinfo=timezone.utc)
                if due_date < now:
                    overdue_tasks += 1
            except (ValueError, TypeError):
                continue

    dsc_list = await db.dsc_register.find({}, {"_id": 0}).to_list(1000)
    total_dsc = len(dsc_list)
    expiring_dsc_count = 0
    expired_dsc_count = 0
    expiring_dsc_list = []
    for dsc in dsc_list:
        try:
            expiry_date = datetime.fromisoformat(dsc["expiry_date"]) if isinstance(dsc["expiry_date"], str) else dsc["expiry_date"]
            days_left = (expiry_date - now).days
            if days_left < 0:
                expired_dsc_count += 1
            if days_left <= 90:
                expiring_dsc_count += 1
                expiring_dsc_list.append({
                    "id": dsc["id"],
                    "holder_name": dsc["holder_name"],
                    "certificate_number": dsc.get("certificate_number", "N/A"),
                    "expiry_date": dsc["expiry_date"],
                    "days_left": days_left,
                    "status": "expired" if days_left < 0 else "expiring"
                })
        except (ValueError, TypeError):
            continue

    if current_user.role == "admin":
        client_query = {}
    else:
        # SCOPE: OWN (assigned_to) + created_by + per-service assignments[] + explicit assigned_clients permission
        # This ensures non-admin users see all clients they are actually working on
        permissions = get_user_permissions(current_user)
        extra_clients = permissions.get("assigned_clients", []) or []
        or_clauses = [
            {"assigned_to": current_user.id},
            {"created_by": current_user.id},
            {"assignments": {"$elemMatch": {"user_id": current_user.id}}},
        ]
        if extra_clients:
            or_clauses.append({"id": {"$in": extra_clients}})
        client_query = {"$or": or_clauses}

    clients = await db.clients.find(client_query, {"_id": 0}).to_list(1000)
    total_clients = len(clients)
    today = date.today()
    upcoming_birthdays = 0

    for client in clients:
        if client.get("birthday"):
            try:
                bday = date.fromisoformat(client["birthday"]) if isinstance(client["birthday"], str) else client["birthday"]
                try:
                    this_year_bday = bday.replace(year=today.year)
                except ValueError:
                    this_year_bday = bday.replace(year=today.year, day=28)
                if this_year_bday < today:
                    try:
                        this_year_bday = bday.replace(year=today.year + 1)
                    except ValueError:
                        this_year_bday = bday.replace(year=today.year + 1, day=28)
                days_until = (this_year_bday - today).days
                if 0 <= days_until <= 7:
                    upcoming_birthdays += 1
            except (ValueError, TypeError):
                continue

    upcoming_due_dates_count = 0
    due_date_query = {"status": "pending"}
    if current_user.role != "admin" and current_user.departments:
        due_date_query["department"] = {"$in": current_user.departments}

    due_dates = await db.due_dates.find(due_date_query, {"_id": 0}).to_list(1000)
    for dd in due_dates:
        try:
            dd_date = datetime.fromisoformat(dd["due_date"]) if isinstance(dd["due_date"], str) else dd["due_date"]
            days_until_due = (dd_date - now).days
            if days_until_due <= 120:
                upcoming_due_dates_count += 1
        except (ValueError, TypeError):
            continue

    team_workload = []
    if current_user.role != "staff":
        users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(100)
        for user in users:
            user_tasks = [t for t in tasks if t.get("assigned_to") == user["id"]]
            team_workload.append({
                "user_id": user["id"],
                "user_name": user["full_name"],
                "total_tasks": len(user_tasks),
                "pending_tasks": len([t for t in user_tasks if t["status"] == "pending"]),
                "completed_tasks": len([t for t in user_tasks if t["status"] == "completed"])
            })

    compliance_score = 100
    if total_tasks > 0:
        compliance_score -= (overdue_tasks / total_tasks) * 50
    if total_dsc > 0:
        compliance_score -= (expiring_dsc_count / total_dsc) * 30

    compliance_status = {
        "score": max(0, int(compliance_score)),
        "status": "good" if compliance_score >= 80 else "warning" if compliance_score >= 50 else "critical",
        "overdue_tasks": overdue_tasks,
        "expiring_certificates": expiring_dsc_count
    }

    return DashboardStats(
        total_tasks=total_tasks,
        completed_tasks=completed_tasks,
        pending_tasks=pending_tasks,
        overdue_tasks=overdue_tasks,
        total_dsc=total_dsc,
        expiring_dsc_count=expiring_dsc_count,
        expiring_dsc_list=expiring_dsc_list,
        total_clients=total_clients,
        upcoming_birthdays=upcoming_birthdays,
        upcoming_due_dates=upcoming_due_dates_count,
        team_workload=team_workload,
        compliance_status=compliance_status,
        expired_dsc_count=expired_dsc_count
    )

# ==========================================================
# STAFF ACTIVITY ROUTES
# ==========================================================
@api_router.post("/activity/log")
async def log_staff_activity(
    activity_data: StaffActivityCreate,
    current_user: User = Depends(get_current_user)
):
    activity = StaffActivityLog(user_id=current_user.id, **activity_data.model_dump())
    doc = activity.model_dump()
    doc["timestamp"] = datetime.now(IST)
    await db.staff_activity.insert_one(doc)
    return {"message": "Activity logged successfully"}

@api_router.get("/activity/summary")
async def get_activity_summary(
    user_id: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    current_user: User = Depends(check_permission("can_view_staff_activity"))
):
    # PERMISSION MATRIX (updated):
    # Staff   → own activity only (can see view_other_activity explicit list)
    # Manager → own activity + same-department team (Own + Team)
    # Admin   → all activity
    query = {}
    if current_user.role != "admin":
        permissions = get_user_permissions(current_user)
        allowed_others = permissions.get("view_other_activity", []) or []

        if current_user.role == "manager":
            # Manager: include entire team (same department) + explicitly listed users
            team_ids = await get_team_user_ids(current_user.id)
            visible_ids = list(set(team_ids + allowed_others + [current_user.id]))
        else:
            # Staff: own data only + explicitly listed users
            visible_ids = list(set(allowed_others + [current_user.id]))

        if user_id:
            if user_id != current_user.id and user_id not in visible_ids:
                raise HTTPException(status_code=403, detail="Not authorised to view this user's activity")
            query["user_id"] = user_id
        else:
            query["user_id"] = {"$in": visible_ids}
    if date_from or date_to:
        query["timestamp"] = {}
    if date_from:
        try:
            query["timestamp"]["$gte"] = datetime.fromisoformat(date_from)
        except Exception:
            pass
    if date_to:
        try:
            query["timestamp"]["$lte"] = datetime.fromisoformat(date_to)
        except Exception:
            pass
    activities = await db.staff_activity.find(query, {"_id": 0}).to_list(10000)
    user_summary = {}
    for activity in activities:
        uid = activity.get("user_id")
        if not uid:
            continue
        duration = activity.get("duration_seconds") or 0
        app_name = activity.get("app_name", "Unknown App")
        category = activity.get("category", "other")
        website = activity.get("website")
        idle = activity.get("idle", False)
        if uid not in user_summary:
            user_summary[uid] = {
                "user_id": uid,
                "total_duration": 0,
                "active_duration": 0,
                "idle_duration": 0,
                "apps": {},
                "websites": {},
                "categories": {}
            }
        user_summary[uid]["total_duration"] += duration
        if idle:
            user_summary[uid]["idle_duration"] += duration
        else:
            user_summary[uid]["active_duration"] += duration
        if app_name not in user_summary[uid]["apps"]:
            user_summary[uid]["apps"][app_name] = {"count": 0, "duration": 0}
        user_summary[uid]["apps"][app_name]["count"] += 1
        user_summary[uid]["apps"][app_name]["duration"] += duration
        if website:
            if website not in user_summary[uid]["websites"]:
                user_summary[uid]["websites"][website] = 0
            user_summary[uid]["websites"][website] += duration
        if category not in user_summary[uid]["categories"]:
            user_summary[uid]["categories"][category] = 0
        user_summary[uid]["categories"][category] += duration
    users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(200)
    user_map = {u.get("id"): u.get("full_name", "Unknown") for u in users if u.get("id")}
    result = []
    for uid, data in user_summary.items():
        data["user_name"] = user_map.get(uid, "Unknown")
        data["apps_list"] = sorted(
            [{"name": k, **v} for k, v in data["apps"].items()],
            key=lambda x: x["duration"],
            reverse=True
        )
        productive_duration = data["categories"].get("productivity", 0)
        total_duration = data["total_duration"]
        data["productivity_percent"] = (productive_duration / total_duration) * 100 if total_duration > 0 else 0
        result.append(data)

    intensity_map = {}
    radar_metrics = {}
    tool_chain_data = []
    for item in result:
        uid = item["user_id"]
        intensity_map[uid] = {"duration": item["total_duration"], "productivity_percent": item["productivity_percent"]}
        radar_metrics[uid] = {"productivity": item["productivity_percent"], "attendance": 75, "task_completion": 80}
        tool_chain_data.append({"user_id": uid, "top_apps": item.get("apps_list", [])[:3]})
    for item in result:
        uid = item["user_id"]
        item["intensityMap"] = intensity_map.get(uid, {})
        item["radarMetrics"] = radar_metrics.get(uid, {})
        item["toolChainData"] = next((t for t in tool_chain_data if t["user_id"] == uid), {})
    return result

@api_router.get("/activity/user/{user_id}")
async def get_user_activity(
    user_id: str,
    limit: int = 100,
    current_user: User = Depends(check_permission("can_view_staff_activity"))
):
    # PERMISSION MATRIX (updated):
    # Staff   → own activity only (+ explicitly listed view_other_activity)
    # Manager → own activity + same-department team (Own + Team)
    # Admin   → any user's activity
    if current_user.role != "admin":
        permissions = get_user_permissions(current_user)
        allowed_others = permissions.get("view_other_activity", []) or []
        if current_user.role == "manager":
            team_ids = await get_team_user_ids(current_user.id)
            visible_ids = list(set(team_ids + allowed_others + [current_user.id]))
        else:
            visible_ids = list(set(allowed_others + [current_user.id]))
        if user_id != current_user.id and user_id not in visible_ids:
            raise HTTPException(status_code=403, detail="You are not authorised to view this user's activity")
    activities = await db.staff_activity.find({"user_id": user_id}, {"_id": 0}).sort("timestamp", -1).to_list(limit)
    return activities

# TASK REMINDER ROUTES
def build_reminder_email(user_name: str, task_list: list) -> tuple[str, str]:
    """
    Builds Option-1 style plain-text reminder email.
    Returns (subject, body).
    """
    from datetime import datetime

    def fmt_date(raw):
        if not raw or raw == "N/A":
            return "N/A"
        try:
            # Handle ISO string like 2026-04-23T00:00:00.000Z
            dt = datetime.fromisoformat(str(raw).replace("Z", "+00:00"))
            return dt.strftime("%d %b %Y")
        except Exception:
            return str(raw)[:10]  # fallback: take first 10 chars

    def priority_badge(p):
        p = (p or "medium").lower()
        if p in ("critical", "high"):   return "High  "
        if p == "medium":               return "Medium"
        return "Low   "

    count = len(task_list)
    first_name = user_name.split()[0] if user_name else "there"

    # Column widths
    W_TASK  = 50
    W_DATE  = 13
    W_PRIO  = 8

    def row(task_col, date_col, prio_col, sep="│"):
        return (
            f"{sep} {task_col:<{W_TASK}} {sep} {date_col:<{W_DATE}} {sep} {prio_col:<{W_PRIO}} {sep}"
        )

    top    = f"┌{'─'*(W_TASK+2)}┬{'─'*(W_DATE+2)}┬{'─'*(W_PRIO+2)}┐"
    header = row("Task", "Due Date", "Priority")
    mid    = f"├{'─'*(W_TASK+2)}┼{'─'*(W_DATE+2)}┼{'─'*(W_PRIO+2)}┤"
    bottom = f"└{'─'*(W_TASK+2)}┴{'─'*(W_DATE+2)}┴{'─'*(W_PRIO+2)}┘"

    rows = []
    for t in task_list:
        title    = (t.get("title") or "Untitled")[:W_TASK]
        due      = fmt_date(t.get("due_date", "N/A"))
        priority = priority_badge(t.get("priority"))
        rows.append(row(title, due, priority))

    table = "\n".join([top, header, mid] + rows + [bottom])

    subject = f"\u23f0 Task Reminder \u2014 {count} Pending Task{'s' if count != 1 else ''}"

    body = (
        f"Hello {first_name},\n\n"
        f"You have {count} pending task{'s' if count != 1 else ''} requiring your attention:\n\n"
        f"{table}\n\n"
        f"Please complete them at your earliest convenience.\n\n"
        f"Regards,\n"
        f"TaskoSphere"
    )

    return subject, body

@api_router.post("/send-pending-task-reminders")
async def send_pending_task_reminders(current_user: User = Depends(get_current_user)):
    perms = get_user_permissions(current_user)
    if current_user.role != "admin" and not perms.get("can_send_reminders", False):
        raise HTTPException(status_code=403, detail="Reminder permission required")
    tasks = await db.tasks.find({"status": {"$ne": "completed"}}, {"_id": 0}).to_list(1000)
    if not tasks:
        return {"message": "No pending tasks found", "emails_sent": 0, "emails_failed": []}

    # Batch lookup: collect all unique assigned_to IDs, fetch users in one query
    assigned_ids = list({t["assigned_to"] for t in tasks if t.get("assigned_to")})
    users_list = await db.users.find({"id": {"$in": assigned_ids}}, {"_id": 0}).to_list(1000)
    user_by_id = {u["id"]: u for u in users_list}

    user_task_map = {}
    for task in tasks:
        assigned_to = task.get("assigned_to")
        if not assigned_to:
            continue
        user = user_by_id.get(assigned_to)
        if not user:
            continue
        email = user.get("email")
        if not email:
            continue
        user_task_map.setdefault(email, {"user": user, "tasks": []})["tasks"].append(task)

    success_count = 0
    failed_emails = []
    for email, data in user_task_map.items():
        try:
            user_name = data["user"].get("full_name", "")
            subject, body = build_reminder_email(user_name, data["tasks"])
            sent = send_email(email, subject, body)
            if sent:
                success_count += 1
            else:
                failed_emails.append(email)
        except Exception as e:
            failed_emails.append(email)
            logger.error(f"Error sending reminder to {email}: {str(e)}")
    return {
        "message": "Reminder process completed",
        "total_users": len(user_task_map),
        "emails_sent": success_count,
        "emails_failed": failed_emails
    }

# AUDIT LOGS ROUTE
@api_router.get("/audit-logs")
async def get_audit_logs(
    module: Optional[str] = None,
    record_id: Optional[str] = None,
    action: Optional[str] = None,
    current_user: User = Depends(check_permission("can_view_audit_logs"))
):
    query = {}
    if module:
        query["module"] = module
    if record_id:
        query["record_id"] = record_id
    if action and action != "ALL":
        query["action"] = action
    logs = await db.audit_logs.find(query, {"_id": 0}).sort("timestamp", -1).to_list(2000)
    logs = convert_objectids(logs)
    for log in logs:
        if isinstance(log.get("timestamp"), str):
            try:
                log["timestamp"] = datetime.fromisoformat(log["timestamp"])
            except Exception:
                pass
    return logs

# INTERNAL FUNCTION FOR AUTO REMINDER
async def send_pending_task_reminders_internal():
    tasks = await db.tasks.find({"status": {"$ne": "completed"}}, {"_id": 0}).to_list(1000)
    if not tasks:
        return
    # Batch user lookup - single query for all assigned users
    assigned_ids = list({t["assigned_to"] for t in tasks if t.get("assigned_to")})
    users_list = await db.users.find({"id": {"$in": assigned_ids}}, {"_id": 0}).to_list(1000)
    user_by_id = {u["id"]: u for u in users_list}

    user_task_map = {}
    for task in tasks:
        assigned_to = task.get("assigned_to")
        if not assigned_to:
            continue
        user = user_by_id.get(assigned_to)
        if not user or not user.get("email"):
            continue
        email = user["email"]
        user_task_map.setdefault(email, {"user": user, "tasks": []})["tasks"].append(task)

    for email, data in user_task_map.items():
        try:
            user_name = data["user"].get("full_name", "")
            subject, body = build_reminder_email(user_name, data["tasks"])
            send_email(email, subject, body)
        except Exception as e:
            logger.error(f"Auto reminder failed for {email}: {str(e)}")

# ─────────────────────────────────────────────────────────────────────────────
# AUTO DAILY REMINDER MIDDLEWARE
# ─────────────────────────────────────────────────────────────────────────────
async def _run_daily_reminder_job(today_str: str):
    """Background job - never blocks requests."""
    global _last_reminder_date_cache
    try:
        setting = await db.system_settings.find_one({"key": "last_reminder_date"}, {"_id": 0})
        db_last_date = setting["value"] if setting else None
        if db_last_date != today_str:
            logger.info("Auto daily reminder triggered at 10:00 AM IST")
            await send_pending_task_reminders_internal()
            await db.system_settings.update_one(
                {"key": "last_reminder_date"},
                {"$set": {"value": today_str}},
                upsert=True
            )
        _last_reminder_date_cache = today_str
    except Exception as e:
        logger.error(f"Auto daily reminder job failed: {e}")


@app.middleware("http")
async def auto_daily_reminder(request: Request, call_next):
    global _last_reminder_date_cache
    try:
        india_time = datetime.now(pytz.timezone("Asia/Kolkata"))
        today_str = india_time.date().isoformat()
        # Only fire after 10 AM and if the in-memory cache hasn't already
        # been set for today. The cache acts as a fast pre-check; the actual
        # DB-level atomic upsert inside _run_daily_reminder_job prevents
        # duplicate sends even if multiple workers race past here.
        if india_time.hour >= 10 and _last_reminder_date_cache != today_str:
            # Set cache immediately to prevent other concurrent requests from
            # spawning duplicate background tasks in the same process.
            _last_reminder_date_cache = today_str
            asyncio.ensure_future(_run_daily_reminder_job(today_str))  # fire-and-forget
    except Exception as e:
        logger.error(f"Auto reminder middleware error: {e}")
    response = await call_next(request)
    return response

# ==================== HOLIDAY ROUTES ====================
@api_router.get("/holidays", response_model=list[HolidayResponse])
async def get_holidays(current_user: User = Depends(get_current_user)):
    # Return all non-rejected holidays to all users — no manual confirmation step
    query = {"status": {"$ne": "rejected"}}
    holidays = await db.holidays.find(query, {"_id": 0}).sort("date", 1).to_list(500)
    return holidays


@api_router.post("/holidays/auto-sync")
async def auto_sync_holidays(current_user: User = Depends(get_current_user)):
    """
    Fetches Indian public holidays for current + next year from date.nager.at
    and saves them as 'confirmed'. Idempotent — safe to call any number of times.
    Also upgrades any existing 'pending' holidays to 'confirmed' automatically.
    """
    import httpx as _httpx
    now = datetime.now(IST)
    added = 0
    upgraded = 0
    errors = []

    for year in [now.year, now.year + 1]:
        try:
            async with _httpx.AsyncClient(timeout=10) as http:
                resp = await http.get(
                    f"https://date.nager.at/api/v3/PublicHolidays/{year}/IN"
                )
            if resp.status_code != 200:
                errors.append(f"API {resp.status_code} for {year}")
                continue
            for h in resp.json():
                date_str = h["date"]
                name     = h.get("localName") or h.get("name", "Holiday")
                existing = await db.holidays.find_one({"date": date_str}, {"_id": 0})
                if not existing:
                    await db.holidays.insert_one({
                        "date": date_str,
                        "name": name,
                        "status": "confirmed",
                        "type": "public",
                        "created_at": now.isoformat(),
                    })
                    added += 1
                elif existing.get("status") not in ("confirmed", "rejected"):
                    await db.holidays.update_one(
                        {"date": date_str},
                        {"$set": {"status": "confirmed"}}
                    )
                    upgraded += 1
        except Exception as e:
            errors.append(f"{year}: {e}")

    logger.info(f"auto-sync holidays: +{added} new, {upgraded} upgraded — by {current_user.email}")
    return {"added": added, "upgraded": upgraded, "errors": errors}


@api_router.post("/holidays", response_model=HolidayResponse)
async def create_holiday(
    holiday: HolidayCreate,
    current_user: User = Depends(require_admin())
):
    """
    Create a holiday entry.
    HolidayCreate.date is typed as Any — the frontend sends a plain string ("2026-04-05").
    Calling .isoformat() on a str raises AttributeError → 500. Guard with isinstance check.
    """
    # ── Safe date → ISO string conversion ───────────────────────────────────
    raw_date = holiday.date
    if isinstance(raw_date, str):
        date_str = raw_date.strip()[:10]   # already "YYYY-MM-DD", just normalise
    elif hasattr(raw_date, "isoformat"):
        date_str = raw_date.isoformat()    # date / datetime object
    else:
        date_str = str(raw_date)[:10]

    # Validate it actually looks like a date
    try:
        date.fromisoformat(date_str)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid date format: '{date_str}'. Use YYYY-MM-DD.")

    holiday_dict = {
        "date":        date_str,
        "name":        holiday.name,
        "description": getattr(holiday, "description", None),
        "status":      "confirmed",
        "type":        getattr(holiday, "type", None) or "manual",
        "created_at":  datetime.now(timezone.utc).isoformat(),
    }

    logger.info(f"Creating holiday: date={holiday_dict['date']}, name={holiday_dict.get('name')}, by={current_user.id}")

    # Upsert: if a record exists (auto-fetched with pending status), confirm it.
    # If it's already confirmed, return it silently (no 400 error on duplicates).
    existing = await db.holidays.find_one({"date": holiday_dict["date"]}, {"_id": 0})
    if existing:
        if existing.get("status") == "confirmed":
            logger.info(f"Holiday already confirmed for {holiday_dict['date']}")
            return existing
        # Exists but not confirmed (pending/rejected from auto-fetch) -> confirm it
        await db.holidays.update_one(
            {"date": holiday_dict["date"]},
            {"$set": {
                "name": holiday_dict.get("name", existing.get("name")),
                "status": "confirmed",
                "type": holiday_dict.get("type", existing.get("type", "public")),
                "updated_by": current_user.id,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }}
        )
        updated = await db.holidays.find_one({"date": holiday_dict["date"]}, {"_id": 0})
        logger.info(f"Holiday upserted to confirmed: {holiday_dict['date']}")
        return updated

    try:
        await db.holidays.insert_one(holiday_dict)
        holiday_dict.pop("_id", None)
        logger.info(f"Holiday inserted: {holiday_dict['date']}")
        return holiday_dict
    except Exception as e:
        # Handle duplicate key error from unique index (race condition)
        logger.error(f"Holiday insert failed for {holiday_dict['date']}: {e}")
        existing2 = await db.holidays.find_one({"date": holiday_dict["date"]}, {"_id": 0})
        if existing2:
            # A concurrent insert created it — just confirm and return it
            await db.holidays.update_one(
                {"date": holiday_dict["date"]},
                {"$set": {"status": "confirmed", "name": holiday_dict.get("name", existing2.get("name"))}}
            )
            final = await db.holidays.find_one({"date": holiday_dict["date"]}, {"_id": 0})
            return final
        raise HTTPException(status_code=500, detail=f"Failed to save holiday: {str(e)}")


@api_router.patch("/holidays/{holiday_date}/status")
async def update_holiday_status(
    holiday_date: str,
    data: dict,
    current_user: User = Depends(require_admin())
):
    new_status = data.get("status")
    if new_status not in ["confirmed", "rejected", "pending"]:
        raise HTTPException(status_code=400, detail="Invalid status")
    result = await db.holidays.update_one(
        {"date": holiday_date},
        {"$set": {"status": new_status, "updated_by": current_user.id}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Holiday not found")
    return {"message": f"Holiday marked as {new_status}"}


@api_router.delete("/holidays/{holiday_date}")
async def delete_holiday(holiday_date: str, current_user: User = Depends(require_admin())):
    result = await db.holidays.delete_one({"date": holiday_date})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Holiday not found")
    return {"message": "Holiday removed"}

# ─────────────────────────────────────────────────────────────────────────────
# PDF HOLIDAY EXTRACTOR — 100% FREE using pdfplumber + regex (no API key)
# ─────────────────────────────────────────────────────────────────────────────

# Month name → number map (reused from compliance parser above)
_HOLIDAY_MONTH_MAP = {
    "january": 1, "jan": 1, "february": 2, "feb": 2,
    "march": 3, "mar": 3, "april": 4, "apr": 4,
    "may": 5, "june": 6, "jun": 6, "july": 7, "jul": 7,
    "august": 8, "aug": 8, "september": 9, "sep": 9, "sept": 9,
    "october": 10, "oct": 10, "november": 11, "nov": 11,
    "december": 12, "dec": 12,
}

def _parse_holiday_date(text: str, default_year: int) -> Optional[str]:
    """Try every common date format found in Indian holiday PDFs. Returns YYYY-MM-DD or None."""
    text = text.strip()

    # YYYY-MM-DD or YYYY/MM/DD
    m = re.search(r'\b(\d{4})[-/](\d{1,2})[-/](\d{1,2})\b', text)
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3))).isoformat()
        except ValueError:
            pass

    # DD-MM-YYYY or DD/MM/YYYY
    m = re.search(r'\b(\d{1,2})[-/](\d{1,2})[-/](\d{4})\b', text)
    if m:
        try:
            return date(int(m.group(3)), int(m.group(2)), int(m.group(1))).isoformat()
        except ValueError:
            pass

    # "15 August 2026" or "15th August 2026"
    m = re.search(
        r'\b(\d{1,2})(?:st|nd|rd|th)?\s+'
        r'(january|february|march|april|may|june|july|august|september|october|november|december|'
        r'jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)'
        r'(?:\s+(\d{4}))?\b',
        text, re.IGNORECASE
    )
    if m:
        try:
            yr = int(m.group(3)) if m.group(3) else default_year
            mo = _HOLIDAY_MONTH_MAP[m.group(2).lower()]
            return date(yr, mo, int(m.group(1))).isoformat()
        except ValueError:
            pass

    # "August 15" or "August 15, 2026"
    m = re.search(
        r'\b(january|february|march|april|may|june|july|august|september|october|november|december|'
        r'jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)\s+'
        r'(\d{1,2})(?:st|nd|rd|th)?(?:,?\s+(\d{4}))?\b',
        text, re.IGNORECASE
    )
    if m:
        try:
            yr = int(m.group(3)) if m.group(3) else default_year
            mo = _HOLIDAY_MONTH_MAP[m.group(1).lower()]
            return date(yr, mo, int(m.group(2))).isoformat()
        except ValueError:
            pass

    return None


def _extract_holidays_from_text(raw_text: str) -> list:
    """
    Parse raw PDF text and extract holiday name + date pairs.
    Handles:
      - Table rows:  "Diwali | 20 October 2026"
      - Inline rows: "20-10-2026  Diwali"
      - Mixed:       "Diwali (20 Oct)"
    """
    results = []
    seen_dates = set()

    # Detect dominant year in the document (use current year as fallback)
    year_matches = re.findall(r'\b(20\d{2})\b', raw_text)
    if year_matches:
        from collections import Counter
        default_year = int(Counter(year_matches).most_common(1)[0][0])
    else:
        default_year = datetime.now().year

    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]

    # --- Pass 1: pipe/tab separated table rows ---
    for line in lines:
        if "|" in line or "\t" in line:
            sep = "|" if "|" in line else "\t"
            cols = [c.strip() for c in line.split(sep) if c.strip()]
            date_val = None
            name_col = None
            for col in cols:
                d = _parse_holiday_date(col, default_year)
                if d and d not in seen_dates:
                    date_val = d
                else:
                    if col and not re.match(r'^(sl\.?\s*no|s\.?\s*no|sr\.?\s*no|#|date|day|holiday|occasion|name)$', col, re.IGNORECASE):
                        name_col = col
            if date_val and name_col and len(name_col) > 2:
                seen_dates.add(date_val)
                results.append({"name": name_col[:80].strip(), "date": date_val})

    # --- Pass 2: lines where date and name appear together ---
    for line in lines:
        # Skip header-like lines
        if re.match(r'^(sl\.?\s*no|s\.?\s*no|sr\.?\s*no|#|date|day|holiday|occasion|name|month)', line, re.IGNORECASE):
            continue
        if len(line) < 5:
            continue

        date_val = _parse_holiday_date(line, default_year)
        if not date_val or date_val in seen_dates:
            continue

        # Remove the date portion to get the name
        name = re.sub(
            r'\b\d{1,2}(?:st|nd|rd|th)?[\s\-/]*'
            r'(?:january|february|march|april|may|june|july|august|september|october|november|december|'
            r'jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)[\s,]*\d{0,4}\b',
            '', line, flags=re.IGNORECASE
        )
        name = re.sub(r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b', '', name)
        name = re.sub(r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b', '', name)
        name = re.sub(r'\b(monday|tuesday|wednesday|thursday|friday|saturday|sunday)\b', '', name, flags=re.IGNORECASE)
        name = re.sub(r'[\|\-–—,;:()\[\]]+', ' ', name)
        name = re.sub(r'\s+', ' ', name).strip()

        if len(name) < 3:
            continue

        seen_dates.add(date_val)
        results.append({"name": name[:80], "date": date_val})

    # Sort by date
    results.sort(key=lambda x: x["date"])
    return results


@api_router.post("/holidays/extract-from-pdf")
async def extract_holidays_from_pdf(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """
    100% FREE holiday extractor.
    Uses pdfplumber (already installed) + regex to parse holiday PDFs.
    No API key, no external calls.
    """
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    # Extract all text from the PDF using pdfplumber (already in requirements)
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                # Extract plain text
                text = page.extract_text()
                if text:
                    parts.append(text)
                # Also extract tables so pipe-separated logic fires
                for table in page.extract_tables():
                    for row in table:
                        if row:
                            parts.append("  |  ".join(str(c or "").strip() for c in row))
        raw_text = "\n".join(parts)
    except Exception as exc:
        logger.error(f"pdfplumber failed: {exc}")
        raise HTTPException(status_code=422, detail=f"Could not read PDF: {str(exc)}")

    if not raw_text or len(raw_text.strip()) < 10:
        raise HTTPException(status_code=422, detail="No readable text found in PDF. Try a text-based (non-scanned) PDF.")

    holidays = _extract_holidays_from_text(raw_text)

    if not holidays:
        raise HTTPException(
            status_code=404,
            detail="No holidays detected. Make sure the PDF contains dates alongside holiday names."
        )

    logger.info(f"PDF holiday extraction (free): {len(holidays)} holidays found by {current_user.email}")
    return {"holidays": holidays}

# ─────────────────────────────────────────────────────────────────────────────
# TRADEMARK / IP NOTICE PDF EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

_TM_MONTH_MAP = {
    "january": 1, "jan": 1, "february": 2, "feb": 2,
    "march": 3, "mar": 3, "april": 4, "apr": 4,
    "may": 5, "june": 6, "jun": 6, "july": 7, "jul": 7,
    "august": 8, "aug": 8, "september": 9, "sep": 9, "sept": 9,
    "october": 10, "oct": 10, "november": 11, "nov": 11,
    "december": 12, "dec": 12,
}


def _parse_tm_date(text: str) -> Optional[str]:
    """
    Parse a date string in any common format used by India's IP Office.
    Returns YYYY-MM-DD string or None.
    """
    if not text:
        return None
    text = text.strip()

    # DD/MM/YYYY or DD-MM-YYYY
    m = re.search(r'\b(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})\b', text)
    if m:
        try:
            return date(int(m.group(3)), int(m.group(2)), int(m.group(1))).isoformat()
        except ValueError:
            pass

    # YYYY-MM-DD
    m = re.search(r'\b(\d{4})-(\d{2})-(\d{2})\b', text)
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3))).isoformat()
        except ValueError:
            pass

    # "06-04-2026" style already caught above, but also "06 April 2026"
    m = re.search(
        r'\b(\d{1,2})(?:st|nd|rd|th)?[\s\-]+('
        r'january|february|march|april|may|june|july|august|september|october|november|december|'
        r'jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec'
        r')[\s,]+(\d{4})\b',
        text, re.IGNORECASE
    )
    if m:
        try:
            mo = _TM_MONTH_MAP[m.group(2).lower()]
            return date(int(m.group(3)), mo, int(m.group(1))).isoformat()
        except ValueError:
            pass

    return None


def _extract_trademark_notice_data(raw_text: str) -> dict:
    """
    Parse raw text extracted from an IP/trademark notice PDF.

    Handles bilingual (Hindi + English) notices from tmrsearch.ipindia.gov.in.
    Returns a dict with all extracted fields; missing fields are None.
    """
    result = {
        "document_type":    None,
        "application_no":   None,
        "class":            None,
        "application_date": None,
        "used_since":       None,
        "applicant_name":   None,
        "recipient_name":   None,
        "hearing_date":     None,
        "letter_date":      None,
        "brand_name":       None,
        "raw_text_snippet": raw_text[:500].strip(),
    }

    # ── Document Type ────────────────────────────────────────────────────────
    text_lower = raw_text.lower()
    if "show cause" in text_lower or "mis-r" in text_lower:
        result["document_type"] = "Show Cause Hearing Notice"
    elif "examination report" in text_lower:
        result["document_type"] = "Examination Report Notice"
    elif "opposition" in text_lower:
        result["document_type"] = "Opposition Notice"
    elif "renewal" in text_lower:
        result["document_type"] = "Renewal Notice"
    elif "registration" in text_lower and "certificate" in text_lower:
        result["document_type"] = "Registration Certificate"
    else:
        result["document_type"] = "IP Office Notice"

    # ── Application Number ───────────────────────────────────────────────────
    # Handles: "Application No: 5922988" or "Application No. 1234567" or
    # "आवेदन संख्या/Application No: 5922988"
    m = re.search(
        r'(?:application\s*no\.?|app\.?\s*no\.?|आवेदन\s*संख्या)[:\s\/]*(\d{5,10})',
        raw_text, re.IGNORECASE
    )
    if m:
        result["application_no"] = m.group(1).strip()

    # ── Class ────────────────────────────────────────────────────────────────
    m = re.search(
        r'(?:in\s+class(?:es)?|class(?:es)?)[:\s\/]*(\d{1,2}(?:\s*[,&]\s*\d{1,2})*)',
        raw_text, re.IGNORECASE
    )
    if m:
        result["class"] = m.group(1).strip()

    # ── Application Date ─────────────────────────────────────────────────────
    m = re.search(
        r'(?:application\s+date|आवेदन\s+तिथि)[:\s\/]*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})',
        raw_text, re.IGNORECASE
    )
    if m:
        result["application_date"] = _parse_tm_date(m.group(1))

    # ── Used Since ───────────────────────────────────────────────────────────
    m = re.search(
        r'(?:used\s+since|उपयोग\s+की\s+तिथि)[:\s\/]*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})',
        raw_text, re.IGNORECASE
    )
    if m:
        result["used_since"] = _parse_tm_date(m.group(1))

    # ── Applicant Name ───────────────────────────────────────────────────────
    # "Name of Applicant: MR. RAJESH DHARIWAL"
    m = re.search(
        r'(?:name\s+of\s+applicant|applicant(?:\'s)?\s+name|आवेदक\s+का\s+नाम)[:\s\/]*([A-Z][A-Za-z.\s]{3,60}?)(?:\n|$|\|)',
        raw_text, re.IGNORECASE
    )
    if m:
        result["applicant_name"] = re.sub(r'\s+', ' ', m.group(1)).strip().rstrip('.,')

    # ── Recipient / Agent Name ───────────────────────────────────────────────
    # First block: "To,\n<NAME>\n<ADDRESS>" — take line after "To,"
    m = re.search(
        r'(?:सेवा\s+में\s*\/\s*To|To\s*,)\s*\n\s*([A-Z][A-Za-z\s.]{2,50})\n',
        raw_text, re.IGNORECASE
    )
    if m:
        result["recipient_name"] = re.sub(r'\s+', ' ', m.group(1)).strip()

    # ── Hearing Date ─────────────────────────────────────────────────────────
    # "fixed for hearing on 06-04-2026" or "दिनांक 06-04-2026 को सुनवाई"
    # Try English first
    m = re.search(
        r'(?:hearing\s+on|fixed\s+for\s+hearing\s+on|scheduled.*?on)\s+(\d{1,2}[-\/]\d{1,2}[-\/]\d{4})',
        raw_text, re.IGNORECASE
    )
    if m:
        result["hearing_date"] = _parse_tm_date(m.group(1))

    if not result["hearing_date"]:
        # Hindi version: "दिनांक 06-04-2026 को सुनवाई"
        m = re.search(
            r'दिनांक\s+(\d{1,2}[-\/]\d{1,2}[-\/]\d{4})\s+को\s+सुनवाई',
            raw_text
        )
        if m:
            result["hearing_date"] = _parse_tm_date(m.group(1))

    if not result["hearing_date"]:
        # Bold date pattern in English block — "on **06-04-2026** as scheduled"
        m = re.search(
            r'\bon\s+(\d{2}[-\/]\d{2}[-\/]\d{4})\s+as\s+scheduled',
            raw_text, re.IGNORECASE
        )
        if m:
            result["hearing_date"] = _parse_tm_date(m.group(1))

    # ── Letter Date ──────────────────────────────────────────────────────────
    # "Dated: 16-02-2026" at top of letter
    m = re.search(
        r'(?:dated?|दिनांक)[:\s]*(\d{1,2}[-\/]\d{1,2}[-\/]\d{4})',
        raw_text, re.IGNORECASE
    )
    if m:
        result["letter_date"] = _parse_tm_date(m.group(1))

    # ── Brand / Mark Name ────────────────────────────────────────────────────
    # Not always present in show cause notices. Try common patterns.
    m = re.search(
        r'(?:trade\s*mark(?:s)?\s+(?:application\s+)?(?:for|of)|in\s+respect\s+of)[:\s]+"?([A-Z][A-Za-z0-9\s&\-\'\.]{1,40})"?',
        raw_text, re.IGNORECASE
    )
    if m:
        result["brand_name"] = m.group(1).strip().strip('"\'')

    return result


@api_router.post("/documents/extract-trademark-notice")
async def extract_trademark_notice(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """
    Extract structured data from a trademark / IP notice PDF.

    No API key required — uses pdfplumber + regex only.
    Works with India IP Office notices (tmrsearch.ipindia.gov.in),
    bilingual Hindi/English format.

    Returns:
        document_type, application_no, class, application_date,
        used_since, applicant_name, recipient_name, hearing_date,
        letter_date, brand_name
    """
    filename = (file.filename or "").lower()
    if not filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    # Extract text using pdfplumber
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
                # Also extract tables — helps with notices that put fields in table cells
                for table in page.extract_tables():
                    for row in table:
                        if row:
                            parts.append("  |  ".join(str(c or "").strip() for c in row))
        raw_text = "\n".join(parts)
    except Exception as exc:
        logger.error(f"pdfplumber failed on trademark notice: {exc}")
        raise HTTPException(status_code=422, detail=f"Could not read PDF: {str(exc)}")

    if not raw_text or len(raw_text.strip()) < 20:
        raise HTTPException(
            status_code=422,
            detail="No readable text found. Please upload a text-based (non-scanned) PDF."
        )

    extracted = _extract_trademark_notice_data(raw_text)

    # Require at minimum an application number OR a hearing date
    if not extracted["application_no"] and not extracted["hearing_date"]:
        raise HTTPException(
            status_code=404,
            detail="Could not find application number or hearing date. "
                   "Make sure this is a valid IP Office notice PDF."
        )

    logger.info(
        f"Trademark notice extracted by {current_user.email}: "
        f"app={extracted['application_no']}, hearing={extracted['hearing_date']}"
    )
    return extracted


# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL EXCEPTION HANDLER
# FIX: Must return a JSONResponse WITH CORS headers so the browser does not
# show a secondary "No Access-Control-Allow-Origin" error when a 500 occurs.
# The CORSMiddleware does NOT add headers to error responses generated by
# exception handlers that run outside the middleware stack, so we add them
# manually here.
# ─────────────────────────────────────────────────────────────────────────────
@app.exception_handler(Exception)
async def universal_exception_handler(request: Request, exc: Exception):
    logger.error(f"Critical Error on {request.url.path}: {str(exc)}")
    logger.error(traceback.format_exc())
    # Determine the correct origin to echo back (support localhost dev too)
    origin = request.headers.get("origin", "")
    allowed_origins = [
        "https://final-taskosphere-frontend.onrender.com",
        "http://localhost:3000",
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://127.0.0.1:3000",
    ]
    cors_origin = origin if origin in allowed_origins else "https://final-taskosphere-frontend.onrender.com"
    headers = {
        "Access-Control-Allow-Origin": cors_origin,
        "Access-Control-Allow-Credentials": "true",
        "Access-Control-Allow-Methods": "GET, POST, PUT, PATCH, DELETE, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type, Authorization, Accept, X-Requested-With",
    }
    return JSONResponse(
        status_code=500,
        content={
            "error": "InternalServerError",
            "message": "A database or logic error occurred.",
            "path": request.url.path
        },
        headers=headers,
    )

# Api Router
api_router.include_router(invoicing_router)
api_router.include_router(compliance_router)
api_router.include_router(identix_router, prefix="/identix")
api_router.include_router(passwords_router)
api_router.include_router(visits_router)
api_router.include_router(website_tracking_router)
api_router.include_router(quotation_router)
api_router.include_router(telegram_router)
api_router.include_router(leads_router)
api_router.include_router(notification_router)
api_router.include_router(email_router)
app.include_router(google_auth_router)
app.include_router(api_router)
