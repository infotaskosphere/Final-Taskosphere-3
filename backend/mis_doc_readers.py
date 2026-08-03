"""
MIS Document Readers
====================
Universal "read anything the client sends us" layer for the MIS Report module.

Supported inputs
----------------
    * Excel      .xlsx / .xlsm / .xls      (every sheet, header row auto-detected)
    * CSV / TSV  .csv / .tsv / .txt        (delimiter sniffing)
    * PDF        .pdf                      (ruled tables, column-position tables,
                                            and plain text — works for bank
                                            statements, GST portal PDFs and
                                            balance sheets)
    * Word       .docx                     (all tables + all paragraph text)

Everything is normalised into a single `ParsedDoc`:

    ParsedDoc.tables -> List[pandas.DataFrame]   (headers already promoted)
    ParsedDoc.text   -> str                      (full plain text, layout kept)

Downstream parsers in `backend/mis_report.py` try each table until one maps to
the expected register/bank/GST shape, then fall back to text parsing.
"""

from __future__ import annotations

import io
import re
import csv
import logging
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any

import pandas as pd

logger = logging.getLogger(__name__)

EXCEL_EXT = (".xlsx", ".xlsm", ".xltx", ".xls")
CSV_EXT = (".csv", ".tsv", ".txt")
PDF_EXT = (".pdf",)
WORD_EXT = (".docx", ".doc")


@dataclass
class ParsedDoc:
    kind: str = "unknown"                       # excel | csv | pdf | word
    tables: List[pd.DataFrame] = field(default_factory=list)
    text: str = ""
    meta: Dict[str, Any] = field(default_factory=dict)

    @property
    def has_tables(self) -> bool:
        return any(t is not None and not t.empty for t in self.tables)


# Backward/forward-compatible alias — mis_report.py imports this name.
ParsedDocument = ParsedDoc


# ══════════════════════════════════════════════════════════════════════════
# GENERIC TABLE CLEANUP
# ══════════════════════════════════════════════════════════════════════════

def _clean_cell(v) -> str:
    if v is None:
        return ""
    s = str(v).replace("\n", " ").replace("\r", " ")
    return re.sub(r"\s+", " ", s).strip()


def _looks_like_header(cells: List[str]) -> int:
    """Score a row on how much it looks like a header row."""
    filled = [c for c in cells if c]
    if len(filled) < 2:
        return 0
    score = 0
    for c in filled:
        if re.fullmatch(r"[-+(]?[\d,.\s]+[)%]?", c):
            score -= 1                     # numbers in a header are unusual
        elif len(c) <= 40:
            score += 1
    return score


def _promote_header(df: pd.DataFrame, max_scan: int = 12) -> pd.DataFrame:
    """Find the real header row inside a raw grid and promote it."""
    if df is None or df.empty:
        return pd.DataFrame()
    best_idx, best_score = None, 0
    for i in range(min(max_scan, len(df))):
        cells = [_clean_cell(v) for v in df.iloc[i].tolist()]
        s = _looks_like_header(cells)
        if s > best_score:
            best_idx, best_score = i, s
    if best_idx is None:
        return df
    header = [_clean_cell(v) for v in df.iloc[best_idx].tolist()]
    header = [h if h else f"col_{i}" for i, h in enumerate(header)]
    out = df.iloc[best_idx + 1:].copy()
    out.columns = _dedupe(header)
    out = out.dropna(axis=0, how="all").dropna(axis=1, how="all")
    return out.reset_index(drop=True)


def _dedupe(names: List[str]) -> List[str]:
    seen: Dict[str, int] = {}
    out = []
    for n in names:
        if n in seen:
            seen[n] += 1
            out.append(f"{n}_{seen[n]}")
        else:
            seen[n] = 0
            out.append(n)
    return out


def _grid_to_df(rows: List[List[Any]]) -> pd.DataFrame:
    rows = [[_clean_cell(c) for c in r] for r in rows if r is not None]
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return pd.DataFrame()
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    return _promote_header(pd.DataFrame(rows))


# ══════════════════════════════════════════════════════════════════════════
# EXCEL / CSV
# ══════════════════════════════════════════════════════════════════════════

def _read_excel(file_bytes: bytes, filename: str) -> ParsedDoc:
    doc = ParsedDoc(kind="excel")
    engine = "xlrd" if filename.lower().endswith(".xls") else "openpyxl"
    try:
        xl = pd.ExcelFile(io.BytesIO(file_bytes), engine=engine)
    except Exception:
        xl = pd.ExcelFile(io.BytesIO(file_bytes))
    texts = []
    for sheet in xl.sheet_names:
        try:
            raw = pd.read_excel(xl, sheet_name=sheet, header=None, dtype=str)
        except Exception as e:      # a single broken sheet must not kill the upload
            logger.warning("MIS: sheet '%s' unreadable: %s", sheet, e)
            continue
        df = _promote_header(raw)
        if df is not None and not df.empty:
            df.attrs["sheet_name"] = sheet
            doc.tables.append(df)
            texts.append(f"--- {sheet} ---\n" + raw.fillna("").astype(str).to_string(index=False))
    doc.text = "\n".join(texts)
    doc.meta["sheets"] = xl.sheet_names
    return doc


def _read_csv(file_bytes: bytes, filename: str) -> ParsedDoc:
    doc = ParsedDoc(kind="csv")
    text = None
    for enc in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            text = file_bytes.decode(enc)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    if text is None:
        text = file_bytes.decode("utf-8", errors="ignore")
    sep = ","
    try:
        sep = csv.Sniffer().sniff(text[:8000], delimiters=",;\t|").delimiter
    except Exception:
        if filename.lower().endswith(".tsv"):
            sep = "\t"
    raw = pd.read_csv(io.StringIO(text), sep=sep, engine="python", header=None,
                      dtype=str, on_bad_lines="skip")
    df = _promote_header(raw)
    if df is not None and not df.empty:
        doc.tables.append(df)
    doc.text = text
    return doc


# ══════════════════════════════════════════════════════════════════════════
# WORD
# ══════════════════════════════════════════════════════════════════════════

def _read_word(file_bytes: bytes) -> ParsedDoc:
    doc = ParsedDoc(kind="word")
    try:
        import docx  # python-docx
    except ImportError:
        raise RuntimeError("python-docx is not installed — cannot read Word files.")
    d = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for t in d.tables:
        grid = []
        for row in t.rows:
            grid.append([c.text for c in row.cells])
        df = _grid_to_df(grid)
        if not df.empty:
            doc.tables.append(df)
        paragraphs.extend(" | ".join(_clean_cell(c) for c in r) for r in grid)
    doc.text = "\n".join(paragraphs)
    return doc


# ══════════════════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════════════════

_DATE_START = re.compile(r"^\d{1,2}[-/][A-Za-z0-9]{2,9}[-/]\d{2,4}")


def _page_lines(page) -> List[List[dict]]:
    try:
        words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
    except Exception:
        return []
    if not words:
        return []
    lines: Dict[int, List[dict]] = {}
    for w in words:
        key = round(w["top"] / 4.0)
        lines.setdefault(key, []).append(w)
    return [sorted(ws, key=lambda w: w["x0"]) for _, ws in sorted(lines.items())]


_HEADER_TOKENS = ("date", "narration", "particulars", "description", "withdrawal",
                  "deposit", "debit", "credit", "balance", "amount", "chq")


def _detect_columns(ordered: List[List[dict]]):
    """Return (header_index, columns) for the first line that looks like a header."""
    for i, ws in enumerate(ordered):
        text = " ".join(w["text"] for w in ws).lower()
        if sum(1 for t in _HEADER_TOKENS if t in text) < 3:
            continue
        cols: List[Dict[str, Any]] = []
        for w in ws:
            if cols and w["x0"] - cols[-1]["x1"] < 6:
                cols[-1]["label"] += " " + w["text"]
                cols[-1]["x1"] = w["x1"]
            else:
                cols.append({"label": w["text"], "x0": w["x0"], "x1": w["x1"]})
        if len(cols) >= 3:
            return i, cols
    return None, None


def _words_to_column_table(page, columns=None) -> Optional[List[List[str]]]:
    """
    Reconstruct a table from word positions — the only reliable way to read
    bank-statement PDFs, whose columns have no ruling lines and whose amounts
    are vertically offset from the row they belong to.

    `columns` carries the header layout detected on an earlier page so that
    continuation pages (which repeat no header) are still parsed.
    """
    ordered = _page_lines(page)
    if not ordered:
        return None

    header_idx, cols = _detect_columns(ordered)
    if cols is None:
        if not columns:
            return None
        cols, header_idx = columns, -1

    # column boundaries = midpoints between neighbouring header spans
    bounds = []
    for i, c in enumerate(cols):
        left = 0 if i == 0 else (cols[i - 1]["x1"] + c["x0"]) / 2
        right = 10_000 if i == len(cols) - 1 else (c["x1"] + cols[i + 1]["x0"]) / 2
        bounds.append((left, right))

    def bucket(w) -> int:
        mid = (w["x0"] + w["x1"]) / 2
        for i, (l, r) in enumerate(bounds):
            if l <= mid < r:
                return i
        return len(cols) - 1

    grid: List[List[str]] = [[c["label"] for c in cols]]
    current: Optional[List[str]] = None
    for ws in ordered[header_idx + 1:]:
        cells = [""] * len(cols)
        for w in ws:
            i = bucket(w)
            cells[i] = (cells[i] + " " + w["text"]).strip()
        if not any(cells):
            continue
        first = cells[0].strip()
        if _DATE_START.match(first):
            if current:
                grid.append(current)
            current = cells
        elif current:
            # continuation line: append text, fill blank numeric cells
            for i, v in enumerate(cells):
                if not v:
                    continue
                if current[i]:
                    current[i] = f"{current[i]} {v}".strip()
                else:
                    current[i] = v
    if current:
        grid.append(current)
    return grid if len(grid) > 1 else None



def _read_pdf(file_bytes: bytes) -> ParsedDoc:
    doc = ParsedDoc(kind="pdf")
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber is not installed — cannot read PDF files.")

    texts: List[str] = []
    column_grids: List[List[List[str]]] = []
    ruled_grids: List[List[List[str]]] = []

    last_cols = None
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            try:
                texts.append(page.extract_text(layout=True) or page.extract_text() or "")
            except Exception:
                try:
                    texts.append(page.extract_text() or "")
                except Exception:
                    texts.append("")
            # 1. ruled tables (GST portal PDFs, balance sheets with borders)
            try:
                for tbl in page.extract_tables() or []:
                    if tbl and len(tbl) > 1 and max(len(r) for r in tbl) > 1:
                        ruled_grids.append(tbl)
            except Exception:
                pass
            # 2. positional tables (bank statements), header carried across pages
            try:
                _, cols = _detect_columns(_page_lines(page))
                if cols:
                    last_cols = cols
                g = _words_to_column_table(page, last_cols)
                if g:
                    column_grids.append(g)
            except Exception:
                pass


    doc.text = "\n".join(texts)
    doc.meta["page_count"] = len(texts)

    # merge multi-page positional tables that share the same header
    if column_grids:
        merged: Dict[str, List[List[str]]] = {}
        for g in column_grids:
            key = "|".join(g[0]).lower()
            if key in merged:
                merged[key].extend(g[1:])
            else:
                merged[key] = list(g)
        for g in merged.values():
            df = _grid_to_df(g)
            if not df.empty:
                doc.tables.append(df)

    # merge ruled tables with identical headers too
    if ruled_grids:
        merged_r: Dict[str, List[List[Any]]] = {}
        for g in ruled_grids:
            key = "|".join(_clean_cell(c) for c in g[0]).lower()
            if key in merged_r:
                merged_r[key].extend(g[1:])
            else:
                merged_r[key] = list(g)
        for g in merged_r.values():
            df = _grid_to_df(g)
            if not df.empty:
                doc.tables.append(df)

    return doc


# ══════════════════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

def read_document(file_bytes: bytes, filename: str) -> ParsedDoc:
    name = (filename or "").lower().strip()
    if name.endswith(EXCEL_EXT):
        return _read_excel(file_bytes, name)
    if name.endswith(CSV_EXT):
        return _read_csv(file_bytes, name)
    if name.endswith(PDF_EXT):
        return _read_pdf(file_bytes)
    if name.endswith(".docx"):
        return _read_word(file_bytes)
    if name.endswith(".doc"):
        raise RuntimeError("Legacy .doc files are not supported — please save as .docx or PDF.")
    # unknown extension: sniff the magic bytes
    if file_bytes[:4] == b"%PDF":
        return _read_pdf(file_bytes)
    if file_bytes[:2] == b"PK":
        try:
            return _read_excel(file_bytes, name + ".xlsx")
        except Exception:
            return _read_word(file_bytes)
    return _read_csv(file_bytes, name)
